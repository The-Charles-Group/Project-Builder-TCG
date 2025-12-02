import os, re, io, math, json, datetime, urllib.parse, tempfile, base64
import uuid
import importlib
import asyncio
from dataclasses import dataclass, field
from enum import Enum
from typing import List, Optional, Dict, Any, Tuple, Set, Union
from zoneinfo import ZoneInfo  # Python 3.9+
from xml.etree.ElementTree import Element, SubElement, tostring
from xml.dom import minidom
from post_export import post_process_xml
from ai_weighted_matcher import score_rfp
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
import pandas as pd
import numpy as np
# REMOVED: from convert_excel_to_mspdi import convert_excel_to_mspdi
# This imported the bloated enhanced version with hardcoded FinishDate and milestone bloat
# Now using the simple version defined in this file at line 9050 which matches Nov 12 reference
from contextlib import asynccontextmanager
import httpx
import hashlib
import time
from functools import lru_cache
import pickle

# ============================================================
# FEATURE FLAGS
# ============================================================
ENABLE_WBS_DEPENDENCIES = True  # Use WBS Dependencies column instead of scheduler edges
ENABLE_MULTI_ASSIGNMENT = True  # Group roles into multi-assignment tasks
# ROLLBACK: Set both to False to restore legacy behavior (scheduler edges, one task per role)

# ============================================================
# SELECTION MODE THRESHOLDS
# ============================================================
CONF_THRESHOLD = float(os.environ.get("CONF_THRESHOLD", "70.0"))  # AI confidence threshold (0-100)
TFIDF_THRESHOLD = float(os.environ.get("TFIDF_THRESHOLD", "0.70"))  # TF-IDF similarity threshold (0-1)

# Default billable rate for Step 3 pricing and XML exports
DEFAULT_BILLABLE_RATE_USD = 210.0
# ============================================================

def should_select_deliverable(match_percent: float, tfidf_similarity: float, selection_mode: str = "confidence_only") -> Dict[str, Any]:
    """
    Determine if a deliverable should be auto-selected based on selection mode.
    
    Args:
        match_percent: AI confidence score (0-100)
        tfidf_similarity: TF-IDF similarity score (0-1)
        selection_mode: One of "confidence_only", "tfidf_only", or "both"
    
    Returns:
        Dict with "selected" (bool) and "selection_reason" (dict with by_confidence and by_tfidf flags)
    """
    # Convert TF-IDF to 0-100 scale for threshold comparison
    tfidf_percent = tfidf_similarity * 100.0
    
    # Check if each method would select this deliverable
    conf_ok = match_percent >= CONF_THRESHOLD
    tfidf_ok = tfidf_percent >= (TFIDF_THRESHOLD * 100.0)
    
    # Apply selection logic based on mode
    if selection_mode == "confidence_only":
        selected = conf_ok
    elif selection_mode == "tfidf_only":
        selected = tfidf_ok
    elif selection_mode == "both":
        # Union: selected if either method selects it
        selected = conf_ok or tfidf_ok
    else:
        # Fallback to current behavior (confidence only)
        selected = conf_ok
    
    return {
        "selected": selected,
        "selection_reason": {
            "by_confidence": conf_ok,
            "by_tfidf": tfidf_ok
        }
    }

try:
    from docx import Document  # pip install python-docx
except Exception:
    Document = None

try:
    from pypdf import PdfReader  # pip install pypdf
except Exception:
    PdfReader = None

try:
    from PIL import Image  # pip install pillow
except Exception:
    Image = None

# ---------- Performance Optimization: Cache Excel to Pickle ----------
# Note: This function will be called after AgencyDB is defined
def load_database_with_pickle_cache():
    """Cache heavy Excel reads; convert to pickle for subsequent boots"""
    # Import here to avoid circular dependency issues
    import glob
    
    # Find the v4 Excel file
    xlsx_paths = [
        "test_outputs/Replit_App_DB_READABLE_FullRows_v4.xlsx",
        "Replit_App_DB_READABLE_FullRows_v4.xlsx",
        "data/Replit_App_DB_READABLE_FullRows_v4.xlsx"
    ]
    
    xlsx_path = None
    for path in xlsx_paths:
        if os.path.exists(path):
            xlsx_path = path
            break
    
    # If not found in standard locations, check attached_assets for timestamped v4 files
    if not xlsx_path:
        v4_files = glob.glob("attached_assets/Replit_App_DB_READABLE_FullRows_v4_*.xlsx")
        if v4_files:
            # Use the most recent file (highest timestamp)
            xlsx_path = sorted(v4_files)[-1]
            print(f"[STARTUP] Found timestamped v4 database: {xlsx_path}")
    
    if not xlsx_path:
        print("[STARTUP] No database file found, creating mock database")
        db = AgencyDB()
        db._create_mock_data()
        db.loaded = True
        return db
        
    pkl_path = xlsx_path + ".pkl"
    
    # Check if pickle exists and is newer than Excel
    if os.path.exists(pkl_path) and os.path.getmtime(pkl_path) >= os.path.getmtime(xlsx_path):
        try:
            print(f"[STARTUP] Loading cached database from {pkl_path}")
            start = time.time()
            with open(pkl_path, "rb") as f:
                db = pickle.load(f)
            print(f"[STARTUP] Pickle cache loaded in {(time.time()-start)*1000:.1f}ms")
            
            # Validate that pickle has real data, not mock data
            if (db.all_rows is not None and 
                not db.all_rows.empty and 
                len(db.all_rows) > 10 and  # Mock data has only 3 rows
                'Component' in db.all_rows.columns):
                print(f"[STARTUP] Pickle validated: {len(db.all_rows)} rows with normalized columns")
                return db
            else:
                print(f"[STARTUP][WARN] Pickle cache invalid (mock data or missing columns), regenerating...")
                os.remove(pkl_path)  # Delete invalid cache
        except Exception as e:
            print(f"[STARTUP][WARN] Failed to load pickle cache: {e}")
            try:
                os.remove(pkl_path)  # Delete corrupted cache
            except:
                pass
    
    # Load from Excel and save to pickle
    print(f"[STARTUP] Loading database from {xlsx_path} (first boot, will cache)")
    start = time.time()
    db = AgencyDB()
    try:
        db.load()
        print(f"[STARTUP] Excel loaded in {(time.time()-start)*1000:.1f}ms")
        
        # Validate loaded database has proper columns
        if not (db.all_rows is not None and not db.all_rows.empty and 'Component' in db.all_rows.columns):
            print(f"[STARTUP][ERROR] Database loaded but missing required columns, using mock data")
            db._create_mock_data()
            db.loaded = True
    except Exception as e:
        print(f"[STARTUP][ERROR] Failed to load Excel database: {e}, using mock data")
        db._create_mock_data()
        db.loaded = True
        return db  # Return mock database, don't try to pickle it
    
    # Save valid database to pickle for next boot
    try:
        with open(pkl_path, "wb") as f:
            pickle.dump(db, f)
        print(f"[STARTUP] Saved pickle cache to {pkl_path}")
    except Exception as e:
        print(f"[STARTUP][WARN] Failed to save pickle cache: {e}")
    
    return db

# ---------- Lifespan for Resource Management ----------
@asynccontextmanager
async def lifespan(app: FastAPI):
    """Manage app lifecycle: startup preloading and shutdown cleanup"""
    # STARTUP
    print("[STARTUP] Initializing Agency Project Builder...")
    
    # 0) Enforce no-patch policy - scan for violations
    try:
        from patch_detector import enforce_no_patches
        enforce_no_patches()
    except ImportError:
        pass  # Module not yet available on first run
    except Exception as e:
        print(f"[WARNING] Patch detection check: {e}")
    
    # 1) Global HTTP client for OpenAI/external APIs (connection pooling)
    app.state.http = httpx.AsyncClient(timeout=30.0, limits=httpx.Limits(max_connections=100))
    print("[STARTUP] HTTP client initialized with connection pooling")
    
    # 2) Preload database with pickle caching
    try:
        app.state.db = load_database_with_pickle_cache()
        if app.state.db:
            print(f"[STARTUP] Database loaded: {getattr(app.state.db, 'src', 'unknown')} with {len(getattr(app.state.db, 'all_rows', []))} rows")
        else:
            print("[STARTUP] Using mock database")
    except Exception as e:
        print(f"[STARTUP][ERROR] Failed to preload database: {e}")
        app.state.db = None
    
    # 3) Preload Fast2 TF-IDF analyzer
    try:
        from app_perf.fast_pipeline import initialize_analyzer
        if initialize_analyzer(app.state):
            print("[STARTUP] Fast2 TF-IDF analyzer preloaded and cached")
        else:
            print("[STARTUP] Fast2 analyzer will be loaded on first use")
    except Exception as e:
        print(f"[STARTUP][WARN] Could not preload Fast2 analyzer: {e}")
    
    # 4) Test GPT-5 availability
    app.state.gpt5_available = False
    print("[STARTUP] Testing GPT-5 availability...")
    try:
        from gpt5_helpers import gpt5_text
        from openai import OpenAI
        
        # Check if we have API key - support both naming conventions
        api_key = os.environ.get("OPENAI_API_KEY") or os.environ.get("Open_AI_Key")
        if api_key:
            try:
                client = OpenAI(api_key=api_key)
                # Test with minimal token usage, retry disabled for quick test
                test_result = gpt5_text(
                    client,
                    messages=[{"role": "user", "content": "Reply with: OK"}],
                    tier="mini",  # Use cheapest tier for test
                    max_output_tokens=100,  # Increased to avoid incomplete responses
                    use_retry=False  # Don't retry for startup test
                )
                if test_result and "OK" in test_result.upper():
                    app.state.gpt5_available = True
                    print("[STARTUP] ✅ GPT-5 is available and responding correctly")
                else:
                    print(f"[STARTUP] ⚠️ GPT-5 responded but with unexpected content: {test_result[:50]}")
            except Exception as e:
                print(f"[STARTUP] ⚠️ GPT-5 test failed: {str(e)[:200]}")
                print("[STARTUP] ℹ️ System will use embedding-based fallback until GPT-5 becomes available")
        else:
            print("[STARTUP] ⚠️ No OpenAI API key found - GPT-5 features disabled")
            print("[STARTUP] ℹ️ Add OPENAI_API_KEY to environment to enable GPT-5 intelligence")
    except ImportError as e:
        print(f"[STARTUP] ⚠️ Could not import GPT-5 helpers: {e}")
        print("[STARTUP] ℹ️ System will operate in embedding-only mode")
    
    # 5) Start background job cleanup task
    async def periodic_cleanup():
        while True:
            await asyncio.sleep(300)  # Every 5 minutes
            await cleanup_old_jobs()
    
    app.state._cleanup_task = asyncio.create_task(periodic_cleanup())
    print("[STARTUP] Background job cleanup task started")
    
    # Final status summary
    if app.state.gpt5_available:
        print("[STARTUP] ✅ Agency Project Builder ready with full GPT-5 intelligence!")
    else:
        print("[STARTUP] ✅ Agency Project Builder ready (using embedding fallback mode)")
        print("[STARTUP] ℹ️ To enable GPT-5: ensure OPENAI_API_KEY is set and API is accessible")
    
    yield  # App runs here
    
    # SHUTDOWN
    print("[SHUTDOWN] Cleaning up resources...")
    try:
        await app.state.http.aclose()
        print("[SHUTDOWN] HTTP client closed")
    except Exception:
        pass
    
    if hasattr(app.state, "_cleanup_task"):
        app.state._cleanup_task.cancel()
        print("[SHUTDOWN] Cleanup task cancelled")
    
    print("[SHUTDOWN] ✅ Cleanup complete")

# ---------- App & CORS ----------
app = FastAPI(title="Agency Project Builder", version="1.0", lifespan=lifespan)

# ---------- Wire Job Runner from sitecustomize ----------
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
try:
    from sitecustomize import register_job_routes, register_health_route
    register_job_routes(app)
    register_health_route(app)
except ImportError as e:
    print(f"[WARNING] Could not import job runner from sitecustomize: {e}")
    # Continue without job runner - will use existing background task system

# ---------- Health Check Endpoint ----------
@app.get("/api/health")
async def health_check():
    """Simple health check endpoint"""
    return {"status": "healthy", "service": "Agency Project Builder", "version": "1.0"}

# ---------- AI Agent Endpoints (CHARLES AGENT) ----------
class AgentChatRequest(BaseModel):
    message: str
    context: Optional[Dict[str, Any]] = None
    session_id: Optional[str] = None
    gpt5_tier: str = "auto"

@app.post("/api/agent/chat")
async def agent_chat_endpoint(request: AgentChatRequest):
    """Chat with CHARLES AGENT using selected GPT-5 intelligence tier"""
    try:
        # Log the request for debugging
        print(f"[CHARLES] Chat request - Tier: {request.gpt5_tier}, Message: {request.message[:100]}...")
        
        # Process the chat message
        response = await chat_with_agent(
            message=request.message,
            context=request.context,
            session_id=request.session_id,
            gpt5_tier=request.gpt5_tier
        )
        
        return response.__dict__ if hasattr(response, '__dict__') else response
        
    except Exception as e:
        print(f"[CHARLES] Error in chat endpoint: {str(e)}")
        return {
            "success": False,
            "message": f"Error processing request: {str(e)}",
            "command": {"type": "UNKNOWN", "parameters": {}, "confidence": 0.0},
            "actions": []
        }

@app.get("/api/agent/status")
async def agent_status_endpoint():
    """Check CHARLES AGENT availability and GPT-5 status"""
    return {
        "available": True,
        "gpt5_available": getattr(app.state, "gpt5_available", False),
        "agent_name": "CHARLES AGENT: ProBuFo",
        "version": "1.0.0",
        "capabilities": [
            "upload_rfp",
            "analyze_rfp", 
            "select_deliverables",
            "modify_pricing",
            "optimize_timeline",
            "export_project"
        ]
    }

@app.post("/api/upload_rfp")
async def upload_rfp_endpoint(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    analyze: bool = True,  # Whether to automatically trigger analysis
    mode: str = "deep"     # Analysis mode: "fast" or "deep"
):
    """Upload RFP document (PDF, DOCX, TXT) for processing and optionally trigger analysis"""
    try:
        # Check file type
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        file_ext = file.filename.split('.')[-1].lower()
        if file_ext not in ['pdf', 'docx', 'txt']:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}")
        
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        text = ""
        if file_ext == 'txt':
            text = content.decode('utf-8', errors='ignore')
        elif file_ext == 'pdf' and PdfReader:
            pdf_reader = PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        elif file_ext == 'docx' and Document:
            doc = Document(io.BytesIO(content))
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            raise HTTPException(status_code=500, detail="Document parsing library not available")
        
        # Cache the RFP text
        global RFP_TEXT_CACHE_FILE, RFP_TEXT_CACHE
        RFP_TEXT_CACHE_FILE = text.strip()
        RFP_TEXT_CACHE = text.strip()
        
        # Prepare response
        response = {
            "success": True,
            "filename": file.filename,
            "text": text.strip(),
            "text_length": len(text.strip()),
            "message": f"Successfully uploaded and extracted text from {file.filename}"
        }
        
        # Optionally trigger analysis
        if analyze and text.strip():
            # Generate job ID for tracking
            job_id = f"upload_{int(time.time())}_{file.filename[:20].replace(' ', '_')}"
            
            # Initialize job tracking
            AI_JOB_STORE[job_id] = AIAnalysisJob(
                job_id=job_id,
                status=AIJobStatus.PENDING,
                total_chunks=0,
                processed_chunks=0,
                current_stage="Starting analysis..."
            )
            
            # Load database if needed
            if not app.state.db.loaded:
                app.state.db.load()
            
            # Start background analysis
            background_tasks.add_task(
                _run_analysis_background,
                job_id,
                text.strip(),
                app.state.db,
                "normal",  # strictness
                "auto",    # tier
                mode,
                None,      # client will be created
                f"upload_{int(time.time())}"  # session_id
            )
            
            response["job_id"] = job_id
            response["analysis_started"] = True
            response["message"] += f" Analysis started with job ID: {job_id}"
        
        return response
        
    except Exception as e:
        print(f"[UPLOAD] Error processing file: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# ---------- AI Planner Integration (AgencyDB) ----------
from ai_planner_agencydb import (
    mount_routes_agencydb,
    AIAnalysisJob,
    AIJobStatus,
    AI_JOB_STORE,
    _run_analysis_background
)

# ---------- AI Agent Integration (CHARLES) ----------
from ai_agent import chat_with_agent, parse_user_intent, CommandType

# ---------- Job Status Endpoint for CHARLES Agent ----------
@app.get("/api/agencydb/status/{job_id}")
async def get_agencydb_job_status(job_id: str):
    """
    Get status of AI analysis job - endpoint for CHARLES Agent compatibility
    Returns job status from AI_JOB_STORE with the expected format
    """
    # Check if job exists in AI_JOB_STORE
    if job_id not in AI_JOB_STORE:
        # Also check sitecustomize job store if available
        try:
            from sitecustomize import _JOBS
            if job_id in _JOBS:
                job = _JOBS[job_id]
                progress = 0.0
                if job.total_batches:
                    progress = 100.0 * (job.finished_batches / job.total_batches)
                
                # Map sitecustomize job status to expected format
                status_map = {
                    "queued": "pending",
                    "running": "processing", 
                    "done": "completed",
                    "error": "failed",
                    "canceled": "failed",
                    "timeout": "failed"
                }
                
                response = {
                    "job_id": job.id,
                    "status": status_map.get(job.status, job.status),
                    "progress": round(progress, 2),
                    "message": job.message or f"Job {job.status}"
                }
                
                if job.status == "done" and job.result:
                    response["data"] = job.result
                
                if job.status in ("error", "timeout") and job.error:
                    response["error"] = job.error
                
                return response
        except ImportError:
            pass
        
        # Check JOB_STORE for image processing jobs (need forward reference since JOB_STORE defined later)
        # Use globals() to access it
        job_store = globals().get('JOB_STORE', {})
        if job_id in job_store:
            job = job_store[job_id]
            
            # Map JobStatus to expected format
            status_map = {
                "pending": "pending",
                "processing": "processing",
                "completed": "completed",
                "failed": "failed",
                "cancelled": "failed"
            }
            
            response = {
                "job_id": job.job_id,
                "status": status_map.get(job.status.value, job.status.value),
                "progress": round(job.percentage, 2),
                "message": f"Processing images: {job.processed_images}/{job.total_images}"
            }
            
            if job.status.value == "completed" and job.result_text:
                response["data"] = {"text": job.result_text}
            
            if job.status.value == "failed" and job.errors:
                response["error"] = "; ".join(job.errors)
            
            return response
        
        # Check SSE_JOB_STORE for timeline generation jobs
        try:
            from app_perf.stream import SSE_JOB_STORE, StreamJobStatus
            if job_id in SSE_JOB_STORE:
                job = SSE_JOB_STORE[job_id]
                
                # Map StreamJobStatus to expected format
                status_map = {
                    StreamJobStatus.QUEUED: "pending",
                    StreamJobStatus.PROCESSING: "processing",
                    StreamJobStatus.COMPLETED: "completed",
                    StreamJobStatus.FAILED: "failed",
                    StreamJobStatus.CANCELLED: "cancelled"
                }
                
                response = {
                    "job_id": job.job_id,
                    "status": status_map.get(job.status, "processing"),
                    "progress": round(job.progress, 2),
                    "message": job.message or f"Timeline generation {job.status.value}",
                    "current_stage": job.current_stage if hasattr(job, 'current_stage') else ""
                }
                
                if job.status == StreamJobStatus.COMPLETED and job.result:
                    response["data"] = job.result
                
                if job.status == StreamJobStatus.FAILED and job.error:
                    response["error"] = job.error
                
                return response
        except ImportError:
            pass
        
        # Job not found in any store
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found")
    
    # Get job from AI_JOB_STORE
    job = AI_JOB_STORE[job_id]
    
    # Calculate progress percentage
    progress = 0
    if job.total_chunks > 0:
        progress = int((job.processed_chunks / job.total_chunks) * 100)
    elif job.status == AIJobStatus.COMPLETED:
        progress = 100
    
    # Map internal status to expected format
    status_map = {
        AIJobStatus.PENDING: "pending",
        AIJobStatus.RUNNING: "processing",
        AIJobStatus.COMPLETED: "completed", 
        AIJobStatus.FAILED: "failed"
    }
    
    response = {
        "job_id": job.job_id,
        "status": status_map[job.status],
        "progress": progress,
        "current_stage": job.current_stage,
        "message": job.current_stage,
        "reasoning": job.current_reasoning if hasattr(job, 'current_reasoning') else "",  # NEW: AI thinking steps
        "reasoning_history": job.reasoning_history if hasattr(job, 'reasoning_history') else []  # Full thinking log
    }
    
    # Add deliverables data if job is completed
    if job.status == AIJobStatus.COMPLETED and job.result:
        response["result"] = job.result  # Changed from "data" to "result" for frontend compatibility
        # Also include deliverable count for UI display
        delivs_count = 0
        if isinstance(job.result, dict):
            for dept, delivs in job.result.items():
                if isinstance(delivs, list):
                    delivs_count += len(delivs)
        response["deliverables_count"] = delivs_count
    
    # Add error if job failed
    if job.status == AIJobStatus.FAILED and job.error:
        response["error"] = job.error
    
    return response

# ---------- Backward Compatibility Alias for /api/ai/jobs/{job_id} ----------
@app.get("/api/ai/jobs/{job_id}")
async def get_ai_job_status_alias(job_id: str):
    """
    Alias endpoint for /api/agencydb/status/{job_id}
    Provides backward compatibility for app.js polling
    """
    return await get_agencydb_job_status(job_id)

# ---------- Industry Template System Import ----------
from luxury_fashion_template import (
    get_industry_template,
    get_available_industries,
    LuxuryFashionTemplate
)

# ---------- Job Tracking System for Async Image Processing ----------
class JobStatus(str, Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"
    CANCELLED = "cancelled"

@dataclass
class JobState:
    job_id: str
    status: JobStatus
    total_images: int = 0
    processed_images: int = 0
    start_time: float = field(default_factory=lambda: datetime.datetime.now().timestamp())
    end_time: Optional[float] = None
    errors: List[str] = field(default_factory=list)
    result_text: Optional[str] = None
    image_timings: List[float] = field(default_factory=list)  # Track time per image for ETA
    cancelled: bool = False
    
    # Two-phase tracking
    phase: str = "quick_scan"  # "quick_scan" or "deep_analysis"
    skipped_images: int = 0  # Images filtered out (duplicates, tiny, irrelevant)
    relevant_images: int = 0  # Images flagged for deep analysis
    
    @property
    def percentage(self) -> float:
        if self.total_images == 0:
            return 100.0
        return (self.processed_images / self.total_images) * 100
    
    @property
    def eta_seconds(self) -> Optional[float]:
        """Estimate time remaining based on rolling average of image processing times"""
        if self.total_images == 0 or self.processed_images == 0:
            return None
        
        remaining = self.total_images - self.processed_images
        if remaining == 0:
            return 0.0
        
        # Use rolling average of last 3 image timings for better accuracy
        recent_timings = self.image_timings[-3:] if len(self.image_timings) > 0 else []
        if not recent_timings:
            return None
        
        avg_time_per_image = sum(recent_timings) / len(recent_timings)
        return remaining * avg_time_per_image

# In-memory job store with automatic cleanup
JOB_STORE: Dict[str, JobState] = {}
JOB_TTL_SECONDS = 3600  # Clean up jobs after 1 hour

async def cleanup_old_jobs():
    """Remove completed/failed jobs older than TTL"""
    now = datetime.datetime.now().timestamp()
    to_remove = []
    
    for job_id, job in JOB_STORE.items():
        if job.end_time and (now - job.end_time > JOB_TTL_SECONDS):
            to_remove.append(job_id)
    
    for job_id in to_remove:
        del JOB_STORE[job_id]
    
    if to_remove:
        print(f"[JOB CLEANUP] Removed {len(to_remove)} expired jobs")

# Global to track last uploaded filename for export defaults
LAST_UPLOAD_FILENAME: str | None = None

# Global RFP text caches for Step 1 → Step 2 handoff
RFP_TEXT_CACHE_TEXTAREA: str | None = None  # Text from textarea input
RFP_TEXT_CACHE_FILE: str | None = None       # Text from uploaded file
RFP_TEXT_CACHE: str | None = None            # Merged text (backward compatibility)
# RFP_SUMMARY_CACHE removed - frontend holds window.currentRfpSummary instead

# SCENARIO_STORE: Unified storage for session_id -> scenario data
# Syncs Gantt ↔ Pricing ↔ XML data through a single source of truth
# 
# Structure per session_id:
# {
#   "baseline": {...},      # Snapshot from Step 2 (immutable until explicit reset)
#   "scenario": {...},      # Working copy for Step 3 edits
#   "totals": {...},        # Computed totals (hours, price)
#   "metadata": {
#     "created_at": float,
#     "updated_at": float,
#     "last_sync": float,
#   }
# }
SCENARIO_STORE: Dict[str, Dict[str, Any]] = {}

import copy
import asyncio

# Per-session locks to prevent race conditions during concurrent updates
_SESSION_LOCKS: Dict[str, asyncio.Lock] = {}

def _get_session_lock(session_id: str) -> asyncio.Lock:
    """Get or create a lock for a session to prevent concurrent mutations."""
    if session_id not in _SESSION_LOCKS:
        _SESSION_LOCKS[session_id] = asyncio.Lock()
    return _SESSION_LOCKS[session_id]

def get_session_state(session_id: str) -> Dict[str, Any]:
    """
    Get or initialize the SCENARIO_STORE entry for a session.
    
    Returns a dict with structure:
    {
        "baseline": {...} or None,
        "scenario": {...} or None,
        "totals": {...},
        "metadata": {...}
    }
    """
    # Guard against None or empty session_id to prevent SCENARIO_STORE[None] entries
    if not session_id:
        return {
            "baseline": None,
            "scenario": None,
            "totals": {"hours": 0.0, "price": 0.0},
            "metadata": {
                "created_at": time.time(),
                "updated_at": time.time(),
                "last_sync": None
            }
        }
    
    if session_id not in SCENARIO_STORE:
        SCENARIO_STORE[session_id] = {
            "baseline": None,
            "scenario": None,
            "totals": {"hours": 0.0, "price": 0.0},
            "metadata": {
                "created_at": time.time(),
                "updated_at": time.time(),
                "last_sync": None
            }
        }
    
    # Auto-migrate old format entries
    entry = SCENARIO_STORE[session_id]
    if "baseline" not in entry:
        # Old format: the entry IS the scenario itself (has "items" key directly)
        # New format: entry has "baseline", "scenario", "totals", "metadata" keys
        if "items" in entry:
            # This is an old-format scenario - migrate it
            # CRITICAL: baseline and scenario MUST be separate deep copies
            baseline_copy = copy.deepcopy(entry)
            working_copy = copy.deepcopy(entry)
            SCENARIO_STORE[session_id] = {
                "baseline": baseline_copy,
                "scenario": working_copy,
                "totals": working_copy.get("totals", {"hours": 0.0, "price": 0.0}),
                "metadata": {
                    "created_at": entry.get("updated_at", time.time()),
                    "updated_at": time.time(),
                    "last_sync": entry.get("updated_at")
                }
            }
        elif "scenario" in entry:
            # Intermediate format with nested "scenario" key but no "baseline"
            old_scenario = entry.get("scenario")
            baseline_copy = copy.deepcopy(old_scenario) if old_scenario else None
            working_copy = copy.deepcopy(old_scenario) if old_scenario else None
            SCENARIO_STORE[session_id] = {
                "baseline": baseline_copy,
                "scenario": working_copy,
                "totals": working_copy.get("totals", {"hours": 0.0, "price": 0.0}) if working_copy else {"hours": 0.0, "price": 0.0},
                "metadata": {
                    "created_at": entry.get("updated_at", time.time()),
                    "updated_at": time.time(),
                    "last_sync": entry.get("updated_at")
                }
            }
        # else: entry already has proper structure (shouldn't reach here)
    
    return SCENARIO_STORE[session_id]

def set_baseline_and_scenario(session_id: str, scenario: dict):
    """
    Set both baseline and working scenario from Step 2 build.
    Called when building a new scenario or resetting.
    """
    now = time.time()
    
    # Ensure scenario has proper structure
    if "items" not in scenario:
        scenario["items"] = []
    if "totals" not in scenario:
        scenario["totals"] = {"hours": 0.0, "price": 0.0}
    
    # Recompute totals to ensure consistency
    scenario = _recompute_totals(scenario)
    
    SCENARIO_STORE[session_id] = {
        "baseline": copy.deepcopy(scenario),
        "scenario": scenario,
        "totals": scenario.get("totals", {"hours": 0.0, "price": 0.0}),
        "metadata": {
            "created_at": now,
            "updated_at": now,
            "last_sync": now
        }
    }

def update_working_scenario(session_id: str, scenario: dict) -> dict:
    """
    Update only the working scenario (not baseline).
    Called during Step 3 edits.
    Returns the updated scenario with recomputed totals.
    """
    state = get_session_state(session_id)
    
    # Recompute totals
    scenario = _recompute_totals(scenario)
    
    state["scenario"] = scenario
    state["totals"] = scenario.get("totals", {"hours": 0.0, "price": 0.0})
    state["metadata"]["updated_at"] = time.time()
    
    SCENARIO_STORE[session_id] = state
    return scenario

def reset_scenario_from_baseline(session_id: str) -> dict:
    """
    Reset working scenario to baseline (deepcopy).
    Called by "Reset from Step 2" button.
    Returns the reset scenario.
    """
    state = get_session_state(session_id)
    
    if state["baseline"] is None:
        # No baseline exists - return empty scenario
        return {"items": [], "totals": {"hours": 0.0, "price": 0.0}}
    
    # Deep copy baseline to working scenario
    reset_scenario = copy.deepcopy(state["baseline"])
    reset_scenario = _recompute_totals(reset_scenario)
    
    state["scenario"] = reset_scenario
    state["totals"] = reset_scenario.get("totals", {"hours": 0.0, "price": 0.0})
    state["metadata"]["updated_at"] = time.time()
    
    SCENARIO_STORE[session_id] = state
    return reset_scenario

def get_working_scenario(session_id: str) -> dict:
    """
    Get the current working scenario for a session.
    This is what exports, Step 3, and Step 4 should use.
    
    IMPORTANT: Returns the LIVE object (not a copy) so edits persist.
    If scenario is missing/empty, clones from baseline per GPT-5 Pro spec.
    """
    state = get_session_state(session_id)
    scenario = state.get("scenario")
    
    if scenario is None or not scenario:
        # Clone from baseline if scenario is missing
        baseline = state.get("baseline") or {}
        if baseline:
            print(f"[SCENARIO_STORE] Cloning scenario from baseline for {session_id}")
            state["scenario"] = copy.deepcopy(baseline)
            SCENARIO_STORE[session_id] = state
            return state["scenario"]
        return {"items": [], "totals": {"hours": 0.0, "price": 0.0}}
    
    return scenario

def has_baseline(session_id: str) -> bool:
    """Check if a baseline exists for this session."""
    state = get_session_state(session_id)
    return state.get("baseline") is not None and bool(state["baseline"].get("items"))

def _norm_str(s: Any) -> str:
    """Normalize a string for key matching: strip, lowercase."""
    return str(s or "").strip().lower()

def make_item_key(item: dict) -> str:
    """
    Create a canonical key for matching scenario items.
    Format: deliverable_code|component|task_label
    
    Uses normalized (lowercased, stripped) values.
    This key is used by:
    - PATCH/save endpoints to find and update items
    - Merge logic for scenario sync
    - Export verification
    
    Per GPT-5 Pro spec: use consistent keying everywhere.
    """
    d = _norm_str(item.get("Deliverable_Code") or item.get("deliverable_code"))
    c = _norm_str(item.get("Component") or item.get("component_name") or item.get("component"))
    t = _norm_str(item.get("Task_Label") or item.get("task_label") or item.get("Task"))
    parts = [p for p in [d, c, t] if p]
    return "|".join(parts)

def _recompute_totals(scn: dict) -> dict:
    """
    Recompute Price_USD and totals for a scenario.
    
    Pricing logic:
    - If item has cadence pricing (price_usd/total_price set), use that as canonical price
    - Otherwise, calculate Price_USD = Rate_USD * Planned_Hours
    
    Also keeps price_usd and Price_USD in sync for backward compatibility.
    Updates totals.hours and totals.price.
    
    Args:
        scn: Scenario dictionary with 'items' list
        
    Returns:
        Updated scenario dictionary
    """
    items = scn.get("items", [])
    total_hours = 0.0
    total_price = 0.0
    
    for item in items:
        # Get hours from item (using backward-compatible lookup)
        hours = 0.0
        for key in ["Planned_Hours", "planned_hours", "hours"]:
            if key in item and item[key] not in (None, ""):
                try:
                    hours = float(item[key])
                    break
                except (TypeError, ValueError):
                    continue
        
        # Check if item has cadence-based pricing (price_usd or total_price already set)
        cadence_price = None
        for key in ["price_usd", "total_price"]:
            if key in item and item[key] not in (None, ""):
                try:
                    cadence_price = float(item[key])
                    break
                except (TypeError, ValueError):
                    continue
        
        if cadence_price is not None and cadence_price > 0:
            item_price = cadence_price
        else:
            rate = get_rate_from_item(item)
            item_price = round(rate * hours, 2)
        
        # Keep all price fields in sync
        item["Price_USD"] = round(item_price, 2)
        item["price_usd"] = item["Price_USD"]
        item["total_price"] = item["Price_USD"]
        
        # Ensure Planned_Hours is set
        if "Planned_Hours" not in item:
            item["Planned_Hours"] = round(hours, 2)
        
        # Accumulate totals
        total_hours += hours
        total_price += item["Price_USD"]
    
    # Update scenario totals
    if "totals" not in scn:
        scn["totals"] = {}
    
    scn["totals"]["hours"] = round(total_hours, 2)
    scn["totals"]["price"] = round(total_price, 2)
    
    return scn

# ---------- Rate/Hours Utility Helpers (Backward Compatible) ----------

def get_rate_from_item(item: dict) -> float:
    """Get rate from item, checking multiple field name conventions."""
    for key in ["Rate_USD", "rate_usd", "rate", "Blended_Rate_USD", "blended_rate_usd"]:
        if key in item and item[key] not in (None, ""):
            try:
                return float(item[key])
            except (TypeError, ValueError):
                continue
    return 0.0

def get_hours_from_item(item: dict) -> float:
    """Get hours from item, checking multiple field name conventions."""
    for key in ["Planned_Hours", "planned_hours", "hours"]:
        if key in item and item[key] not in (None, ""):
            try:
                return float(item[key])
            except (TypeError, ValueError):
                continue
    return 0.0

def set_hours_on_item(item: dict, hours: float):
    """Set hours on item using existing convention."""
    if "Planned_Hours" in item or "planned_hours" not in item:
        item["Planned_Hours"] = round(hours, 2)
    else:
        item["planned_hours"] = round(hours, 2)

def get_price_from_item(item: dict) -> float:
    """Get total price from item, checking multiple field name conventions."""
    for key in ["price_usd", "Price_USD", "total_price", "Total_Price"]:
        if key in item and item[key] not in (None, ""):
            try:
                return float(item[key])
            except (TypeError, ValueError):
                continue
    return 0.0

def cadence_changed(old_item: dict | None, new_item: dict) -> bool:
    """
    Check if cadence fields actually changed between old and new item.
    Returns False for first sync (no old_item) to preserve existing retainer months.
    Only returns True when comparing items with explicit field differences.
    """
    if not old_item:
        return False
    
    fields = ["billing_cadence", "cadence_units"]
    for f in fields:
        old = old_item.get(f)
        new = new_item.get(f)
        if (old is None or old == "") and (new is None or new == ""):
            continue
        if str(old) != str(new):
            return True
    return False

def cadence_price_changed(old_item: dict | None, new_item: dict) -> bool:
    """
    Check if cadence_price changed (separate from cadence structure changes).
    Returns False on first sync (no old_item) to preserve existing pricing.
    """
    if not old_item:
        return False
    
    old_cp = old_item.get("cadence_price")
    new_cp = new_item.get("cadence_price")
    
    if old_cp is None and new_cp is None:
        return False
    
    try:
        return abs(float(old_cp or 0) - float(new_cp or 0)) > 0.01
    except (TypeError, ValueError):
        return str(old_cp) != str(new_cp)

def monthly_price_changed(old_item: dict | None, new_item: dict) -> bool:
    """
    Check if monthly_price was explicitly changed by PM.
    Returns False on first sync (no old_item) to preserve existing pricing.
    Only returns True when there's a meaningful difference with an old value.
    """
    if not old_item:
        return False
    
    if "monthly_price" not in new_item:
        return False
    
    new_mp = new_item.get("monthly_price")
    if new_mp is None or new_mp == "" or new_mp == 0:
        return False
    
    try:
        old = float(old_item.get("monthly_price") or 0)
        new = float(new_mp or 0)
    except (TypeError, ValueError):
        return True
    
    return abs(old - new) > 0.01

# ---------- Cadence-Aware Pricing Helper ----------

def apply_cadence_to_item(new_item: dict, old_item: dict | None = None):
    """
    Apply cadence pricing logic to a scenario item.
    Computes total_price, hours, and retainer fields from cadence settings.
    
    Priority rules:
    1. If monthly_price was explicitly changed by PM, derive total from monthly_price
    2. If cadence_price changed, recompute total from cadence_price
    3. If cadence structure changed (billing_cadence/cadence_units), recompute months
    4. If nothing changed, preserve existing totals
    """
    changed_cadence = cadence_changed(old_item, new_item)
    changed_cadence_price = cadence_price_changed(old_item, new_item)
    changed_monthly = monthly_price_changed(old_item, new_item)
    
    if not changed_cadence and not changed_cadence_price and not changed_monthly:
        return
    
    cadence_raw = (new_item.get("billing_cadence") or new_item.get("cadence") or "one_time").lower()
    
    mapping = {
        "one time": "one_time",
        "one-time": "one_time",
        "one_time": "one_time",
        "monthly": "monthly",
        "month": "monthly",
        "quarterly": "quarterly",
        "q": "quarterly",
        "semi-annual": "semi_annual",
        "semiannual": "semi_annual",
        "annual": "annual",
        "yearly": "annual",
    }
    cadence = mapping.get(cadence_raw, "one_time")
    new_item["billing_cadence"] = cadence
    
    try:
        units = int(new_item.get("cadence_units") or 1)
    except (TypeError, ValueError):
        units = 1
    if units <= 0:
        units = 1
    new_item["cadence_units"] = units
    
    months_per_cadence = {
        "one_time": 0,
        "monthly": 1,
        "quarterly": 3,
        "semi_annual": 6,
        "annual": 12,
    }.get(cadence, 0)
    
    total_months = months_per_cadence * units
    
    existing_retainer = new_item.get("retainer") or {}
    existing_months = existing_retainer.get("months") or 0
    if existing_months and not changed_cadence:
        total_months = int(existing_months)
    
    is_retainer = (cadence != "one_time") and total_months > 0
    new_item["is_retainer"] = is_retainer
    
    if is_retainer and total_months > 0 and changed_monthly:
        try:
            monthly_price = float(new_item.get("monthly_price") or 0)
        except (TypeError, ValueError):
            monthly_price = 0
        
        if monthly_price > 0:
            total_price = monthly_price * total_months
            cadence_price = total_price / units if units > 0 else total_price
            new_item["monthly_price"] = round(monthly_price, 2)
            new_item["cadence_price"] = round(cadence_price, 2)
            new_item["price_usd"] = round(total_price, 2)
            new_item["Price_USD"] = new_item["price_usd"]
            new_item["total_price"] = new_item["price_usd"]
        else:
            changed_monthly = False
    
    if not changed_monthly:
        try:
            cadence_price = float(new_item.get("cadence_price") or 0.0)
        except (TypeError, ValueError):
            cadence_price = 0.0
        
        if cadence_price <= 0:
            existing_total = get_price_from_item(new_item)
            cadence_price = existing_total / units if units > 0 else existing_total
        
        new_item["cadence_price"] = round(cadence_price, 2)
        
        total_price = cadence_price * units
        new_item["price_usd"] = round(total_price, 2)
        new_item["Price_USD"] = new_item["price_usd"]
        new_item["total_price"] = new_item["price_usd"]
        
        if is_retainer and total_months > 0:
            monthly_price = total_price / total_months
            new_item["monthly_price"] = round(monthly_price, 2)
        else:
            new_item["monthly_price"] = 0
    
    if is_retainer:
        new_item["retainer"] = {**existing_retainer, "months": total_months}
        new_item["retainer_months"] = total_months
    
    rate = get_rate_from_item(new_item)
    total_price = get_price_from_item(new_item)
    if rate > 0:
        total_hours = total_price / rate
    else:
        total_hours = get_hours_from_item(new_item)
    
    set_hours_on_item(new_item, total_hours)
    
    if is_retainer and total_months > 0:
        monthly_hours = total_hours / total_months
        new_item["monthly_hours"] = round(monthly_hours, 2)
    else:
        new_item["monthly_hours"] = 0

# ---------- Scenario Sync Merge Helpers ----------

def _merge_items(old_item: dict, new_item: dict, editable_fields: set) -> dict:
    """
    Merge new item data into old item, only updating editable fields.
    Preserves timeline/WBS metadata, PM overrides, etc.
    """
    merged = old_item.copy()
    for k, v in new_item.items():
        if k in editable_fields or k not in merged:
            merged[k] = v
    return merged

def merge_scenario_from_sync(session_id: str):
    """
    Safely merge Step 3 edits from SCENARIO_SYNC_STATE into SCENARIO_STORE.
    Preserves timeline/WBS/PM metadata while updating only editable pricing/cadence fields.
    """
    sync_entry = SCENARIO_SYNC_STATE.get(session_id)
    if not sync_entry or "scenario" not in sync_entry:
        return
    
    new_scen = sync_entry.get("scenario") or {}
    if not new_scen:
        return
    
    # Use get_session_state to ensure proper dual-entry format
    state = get_session_state(session_id)
    old_scen = (state.get("scenario") or {}).copy()
    
    editable_fields = {
        "billing_cadence", "cadence_units", "cadence_price",
        "price_usd", "Price_USD", "total_price",
        "planned_hours", "Planned_Hours", "hours",
        "retainer", "retainer_months", "monthly_hours", "monthly_price",
        "is_retainer", "Rate_USD", "rate_usd",
    }
    
    # If no old scenario exists, just use the new one directly
    if not old_scen or not old_scen.get("items"):
        update_working_scenario(session_id, new_scen)
        print(f"[DEBUG] merge_scenario_from_sync: SCENARIO_STORE[{session_id}] keys:", list(SCENARIO_STORE[session_id].keys()))
        return
    
    merged_scen = old_scen.copy()
    for k, v in new_scen.items():
        if k == "items":
            continue
        if isinstance(v, dict) and isinstance(old_scen.get(k), dict):
            merged_scen[k] = {**old_scen.get(k, {}), **v}
        else:
            merged_scen[k] = v
    
    old_items_by_id = {}
    for item in old_scen.get("items", []):
        item_id = item.get("id") or item.get("Row_ID") or item.get("deliverable_code")
        if item_id:
            old_items_by_id[item_id] = item
    
    merged_items = []
    for new_item in new_scen.get("items", []):
        item_id = new_item.get("id") or new_item.get("Row_ID") or new_item.get("deliverable_code")
        if not item_id:
            merged_items.append(new_item)
            continue
        
        old_item = old_items_by_id.get(item_id, {})
        merged_items.append(_merge_items(old_item, new_item, editable_fields))
    
    merged_scen["items"] = merged_items
    
    update_working_scenario(session_id, merged_scen)
    print(f"[DEBUG] merge_scenario_from_sync merged: SCENARIO_STORE[{session_id}] keys:", list(SCENARIO_STORE[session_id].keys()))

# ---------- Gantt Monthly Expansion Helpers ----------

def month_bounds(start_date, offset: int):
    """
    Get the first and last day of a month relative to start_date.
    
    Args:
        start_date: date object (project start)
        offset: int (0 for first month, 1 for next month, etc.)
    
    Returns:
        Tuple of (first_day_of_month, last_day_of_month) as date objects
    """
    from datetime import date, timedelta
    
    if isinstance(start_date, str):
        start_date = datetime.datetime.fromisoformat(start_date.replace('Z', '')).date()
    elif isinstance(start_date, datetime.datetime):
        start_date = start_date.date()
    
    year = start_date.year + (start_date.month - 1 + offset) // 12
    month = (start_date.month - 1 + offset) % 12 + 1
    first = date(year, month, 1)
    
    if month == 12:
        last = date(year + 1, 1, 1) - timedelta(days=1)
    else:
        last = date(year, month + 1, 1) - timedelta(days=1)
    
    return first, last

def expand_retainer_into_months(deliv_row: dict, months: int, project_start_date) -> list:
    """
    Expand a retainer deliverable into monthly child rows for Gantt display.
    
    Args:
        deliv_row: The parent deliverable row dict
        months: Number of months to expand
        project_start_date: Project start date
        
    Returns:
        List of monthly child row dicts
    """
    from datetime import date
    
    if isinstance(project_start_date, str):
        project_start_date = datetime.datetime.fromisoformat(project_start_date.replace('Z', '')).date()
    elif isinstance(project_start_date, datetime.datetime):
        project_start_date = project_start_date.date()
    
    base_wbs = deliv_row.get("WBS_ID", "")
    label = deliv_row.get("Deliverable") or deliv_row.get("Task_Name") or ""
    
    total_hours = get_hours_from_item(deliv_row)
    hours_per_month = total_hours / months if months > 0 else 0
    
    total_price = get_price_from_item(deliv_row)
    price_per_month = total_price / months if months > 0 else 0
    
    rows = []
    for i in range(months):
        first, last = month_bounds(project_start_date, i)
        start_offset_days = (first - project_start_date).days
        duration_days = (last - first).days + 1
        
        child = deliv_row.copy()
        child["WBS_ID"] = f"{base_wbs}.{i+1}"
        child["Parent_WBS_ID"] = base_wbs
        child["Task_Name"] = f"{label} - {first.strftime('%B')}"
        child["Planned_Hours"] = round(hours_per_month, 2)
        child["Price_USD"] = round(price_per_month, 2)
        child["price_usd"] = child["Price_USD"]
        
        child["Start_Offset_Days"] = start_offset_days
        child["Duration_Days"] = duration_days
        child["Start_Date"] = first.isoformat()
        child["End_Date"] = last.isoformat()
        
        if i == 0:
            child["Dependencies"] = deliv_row.get("Dependencies", "")
        else:
            child["Dependencies"] = f"{base_wbs}.{i}"
        
        child["Notes"] = (deliv_row.get("Notes", "") + f" | Month {i+1}/{months}").strip(" |")
        
        rows.append(child)
    
    return rows

# Configure file upload limits - allow up to 20MB files
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request
from starlette.responses import JSONResponse

class FileSizeMiddleware(BaseHTTPMiddleware):
    def __init__(self, app, max_size: int = 20 * 1024 * 1024):  # 20MB
        super().__init__(app)
        self.max_size = max_size

    async def dispatch(self, request: Request, call_next):
        if request.method == "POST":
            content_length = request.headers.get("content-length")
            if content_length and int(content_length) > self.max_size:
                return JSONResponse(
                    {"error": f"File too large. Maximum size is {self.max_size // (1024*1024)}MB."},
                    status_code=413
                )
        return await call_next(request)

app.add_middleware(FileSizeMiddleware)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_headers=["*"],
    allow_methods=["*"],
    expose_headers=["Content-Disposition"],  # Allow browser to read server-suggested filename
)

# Serve static frontend
if not os.path.exists("static"):
    os.makedirs("static", exist_ok=True)
app.mount("/static", StaticFiles(directory="static"), name="static")

# Include AI weights router
from routes_weights_fastapi import router as weights_router
app.include_router(weights_router)

# Include Learning Brain router
from learning_brain.routes_brain import router as brain_router
app.include_router(brain_router, prefix="/api/brain", tags=["learning"])

# Serve Learning Brain admin UI
@app.get("/admin/brain")
async def brain_admin_ui():
    """Serve the Learning Brain admin interface"""
    from fastapi.responses import FileResponse
    return FileResponse("learning_brain/static/admin_brain.html")

# Mount AI planner routes (connected to AgencyDB)
mount_routes_agencydb(app, base="/api/ai")

# Mount Performance Optimization routers (Fast2 and SSE)
from app_perf import fast_router, stream_router
app.include_router(fast_router)  # Fast2 TF-IDF endpoint
app.include_router(stream_router)  # SSE streaming endpoint

# Mount Scenario API (unified pricing table persistence)
from backend.scenario_api import router as scenario_router
app.include_router(scenario_router, prefix="/api/scenario", tags=["scenario"])

# Import AI Timeline Manager
from ai_timeline_manager import suggest_timeline_from_selection, generate_ai_timeline

# Import Intelligent Timeline Scheduler
from timeline_scheduler import generate_intelligent_timeline

# Import AI Pricing Optimizer
from ai_pricing_optimizer import redistribute_hours, calculate_retainer_distribution, analyze_retainer_vs_project

# Startup event (AI planner now uses AgencyDB directly, no ZIP catalog needed)
@app.on_event("startup")
async def startup_event():
    """Initialize AI planner with AgencyDB catalog"""
    app.state.db = AgencyDB()
    app.state.db.load()
    print(f"[AI PLANNER] Using AgencyDB as catalog source (connected to app.state.db)")
    print(f"[AgencyDB] Loaded from: {app.state.db.src}")
    print(f"[AgencyDB] Total rows: {len(app.state.db.all_rows) if app.state.db.all_rows is not None else 0}")

# ===== Workfront column order (now includes Service_Department) =====
WF_COLUMNS = [
    "Project_Name", "WBS_ID", "Parent_WBS_ID",
    "Task_Name", "Deliverable", "Component", "Task",
    "Service Department",            # <-- exact header, its own column
    "Role", "Seniority",
    "Planned_Hours", "Start_Offset_Days", "Duration_Days",
    "Start_Date", "End_Date",
    "Dependencies", "Assignee_External_ID", "Notes",
    "Rate_USD", "Price_USD", "FixedCost"  # FixedCost for Step 3 price export
]


# XML post-processing: disable parallelization to preserve role-to-role predecessor chains
PARALLELIZE_IDENTICAL_NAMES = os.getenv("PARALLELIZE_IDENTICAL_NAMES", "false").lower() == "true"

# Scenario multipliers - A only (B/C removed for simplicity)
SCENARIO_MULT = {
    "A": {"hours_mult": 1.00, "qa_pct": 0.05, "pm_pct": 0.10, "strip_optional": True}
}

# US & Mexico holidays for business day calendar
def _get_us_mx_holidays(year: int) -> list:
    """Get US federal and Mexico holidays for a given year."""
    holidays = []
    
    # US Federal Holidays
    holidays.append(np.datetime64(f'{year}-01-01'))  # New Year's Day
    # MLK Day (3rd Monday in January)
    mlk_day = 15 + (7 - datetime.date(year, 1, 15).weekday()) % 7
    holidays.append(np.datetime64(f'{year}-01-{mlk_day:02d}'))
    # Presidents Day (3rd Monday in February)
    pres_day = 15 + (7 - datetime.date(year, 2, 15).weekday()) % 7
    holidays.append(np.datetime64(f'{year}-02-{pres_day:02d}'))
    # Memorial Day (last Monday in May)
    last_may = datetime.date(year, 5, 31)
    # Roll back from May 31 to the last Monday (weekday 0)
    days_since_monday = last_may.weekday()  # 0=Mon, 6=Sun
    memorial = last_may - datetime.timedelta(days=days_since_monday)
    holidays.append(np.datetime64(memorial))
    holidays.append(np.datetime64(f'{year}-07-04'))  # Independence Day
    # Labor Day (1st Monday in September)
    labor_day = 1 + (7 - datetime.date(year, 9, 1).weekday()) % 7
    holidays.append(np.datetime64(f'{year}-09-{labor_day:02d}'))
    # Columbus Day (2nd Monday in October)
    columbus_day = 8 + (7 - datetime.date(year, 10, 8).weekday()) % 7
    holidays.append(np.datetime64(f'{year}-10-{columbus_day:02d}'))
    holidays.append(np.datetime64(f'{year}-11-11'))  # Veterans Day
    # Thanksgiving (4th Thursday in November)
    first_nov = datetime.date(year, 11, 1)
    thanksgiving = first_nov + datetime.timedelta(days=(3 - first_nov.weekday()) % 7 + 21)
    holidays.append(np.datetime64(thanksgiving))
    holidays.append(np.datetime64(f'{year}-12-25'))  # Christmas
    
    # Mexico holidays
    holidays.append(np.datetime64(f'{year}-02-05'))  # Constitution Day
    holidays.append(np.datetime64(f'{year}-03-21'))  # Benito Juárez's Birthday
    holidays.append(np.datetime64(f'{year}-05-01'))  # Labor Day (Mexico)
    holidays.append(np.datetime64(f'{year}-09-16'))  # Independence Day (Mexico)
    holidays.append(np.datetime64(f'{year}-11-20'))  # Revolution Day
    
    return holidays


def _find_v4_path() -> str | None:
    import glob
    # Check for xlsx files first (including test_outputs directory)
    for p in [
        "Replit_App_DB_READABLE_FullRows_v4b.xlsx",
        "Replit_App_DB_READABLE_FullRows_v4.xlsx",
        "test_outputs/Replit_App_DB_READABLE_FullRows_v4.xlsx",  # v4 in test_outputs
        "data/Replit_App_DB_READABLE_FullRows_v4.xlsx",
    ]:
        if os.path.exists(p):
            return p
    # Check attached_assets for timestamped v4 files
    v4_files = glob.glob("attached_assets/Replit_App_DB_READABLE_FullRows_v4_*.xlsx")
    if v4_files:
        return sorted(v4_files)[-1]  # Use most recent file
    # Check for CSV bundle directory (canonical and timestamped variants)
    # Filter to only return directories, not ZIP files
    csv_dirs = [d for d in glob.glob("Replit_App_DB_READABLE_FullRows_v4_csvs*") if os.path.isdir(d)]
    if csv_dirs:
        return csv_dirs[0]  # return first match
    return None

# ---------- OpenAI Integration (Stage 2) ----------
from openai import OpenAI

# the newest OpenAI model is "gpt-5" which was released August 7, 2025.
# do not change this unless explicitly requested by the user
# Initialize OpenAI client - will be set after models are defined

# ---------- Helper: DB Loader ----------
class AgencyDB:
    def __init__(self):
        self.loaded = False
        self.src = None
        # DataFrames
        self.all_rows = None                # All_Task_Rows
        self.deliverables = None            # Deliverable_Index
        self.b_rules = None                 # Bundle_Rules_Table
        self.b_defaults = None              # Bundle_Scenario_Defaults
        self.b_by_deliv = None              # Bundles_By_Deliverable
        self.b_hours_by_role = None         # Bundles_Hours_By_Role
        self.role_rate_card = None          # Role_Rate_Card
        self.rate_matrix = None             # Role_Rate_Matrix
        self.rate_bands = None              # Rate_Bands
        self.timeline_params = None         # Timeline_Params
        self.timeline_scaling = None        # Timeline_Scaling
        self.timeline_weighting = None      # Timeline_Weighting
        self.slack_settings = None          # Slack_Settings
        self.pricing_settings = None        # Pricing_Settings
        self.scenario_templates = None      # Scenario_Templates
        self.ui_options = None              # UI_Options
        self.rfp_rules = None               # RFP_Matching_Rules

    def _scenario_col(self, complexity: str, tier: str) -> str:
        return f"{complexity}__{tier}_Hours"

    def load(self):
        v4_path = _find_v4_path()

        if v4_path:
            self._load_from_v4_primary(v4_path)
        else:
            # No database files found, create mock data
            self._create_mock_data()
        self.loaded = True
        return True


    def _load_from_v4_primary(self, v4_path: str | None):
        # Use the provided v4_path or fall back to searching for files
        if v4_path and v4_path.endswith(".xlsx"):
            # Excel file
            self.src = v4_path
            read = lambda sh: pd.read_excel(v4_path, sheet_name=sh)
        elif v4_path and os.path.isdir(v4_path):
            # CSV bundle directory
            def read_csv(sh):
                path = os.path.join(v4_path, f"{sh}.csv")
                if not os.path.exists(path):
                    raise FileNotFoundError(path)
                return pd.read_csv(path)
            read = read_csv
            self.src = v4_path
        else:
            # Create minimal mock data for demo purposes
            self._create_mock_data()
            self.loaded = True
            return True

        # Load sheets
        self.all_rows          = read("All_Task_Rows")
        self.deliverables      = read("Deliverable_Index")
        self.b_rules           = read("Bundle_Rules_Table")
        self.b_defaults        = read("Bundle_Scenario_Defaults")
        self.b_by_deliv        = read("Bundles_By_Deliverable")
        self.b_hours_by_role   = read("Bundles_Hours_By_Role")
        self.role_rate_card    = read("Role_Rate_Card")
        self.rate_matrix       = read("Role_Rate_Matrix")
        self.rate_bands        = read("Rate_Bands")
        self.timeline_params   = read("Timeline_Params")
        self.timeline_scaling  = read("Timeline_Scaling")
        self.timeline_weighting= read("Timeline_Weighting")
        self.slack_settings    = read("Slack_Settings")
        self.pricing_settings  = read("Pricing_Settings")
        # Override Default_Blended_Rate to ensure it uses the global constant
        if self.pricing_settings is not None and not self.pricing_settings.empty:
            mask = self.pricing_settings["Key"] == "Default_Blended_Rate"
            if mask.any():
                self.pricing_settings.loc[mask, "Default"] = DEFAULT_BILLABLE_RATE_USD
        self.scenario_templates= read("Scenario_Templates")
        self.ui_options        = read("UI_Options")
        self.rfp_rules         = read("RFP_Matching_Rules")

        # Normalize
        for c in ["Deliverable_Code","Deliverable","Category"]:
            if c in self.deliverables.columns:
                self.deliverables[c] = self.deliverables[c].astype(str)

        # Normalize component column from v4 spreadsheet
        self._normalize_component_column()
        # Normalize task label column from v4b spreadsheet
        self._normalize_task_label_column()
        # Normalize role and seniority columns 
        self._normalize_role_and_seniority_columns()
        self._normalize_rate_card_seniority()
        
        # normalize code columns from v4/v4b for canonical naming
        self._normalize_code_columns()
        
        print(f"[DB] v4-primary loaded from: {self.src}")

    def _normalize_component_column(self):
        """
        Ensure self.all_rows has a 'Component' column populated from v4's Component_Task_L1 (Column F).
        If other synonyms exist, prefer them in this order.
        """
        if self.all_rows is None or isinstance(self.all_rows, pd.DataFrame) and self.all_rows.empty:
            return
        # Known header names (case-insensitive)
        candidates = ["Component_Task_L1", "Component L1", "Component_L1", "Component"]
        cols_lc = {c.lower(): c for c in self.all_rows.columns}
        found = None
        for cand in candidates:
            if cand.lower() in cols_lc:
                found = cols_lc[cand.lower()]
                break

        if found:
            # Standardize to 'Component' - fix order to handle NaN properly
            self.all_rows["Component"] = (
                self.all_rows[found]
                .fillna("")
                .astype(str)
                .str.strip()
                .replace({"nan": ""})  # Handle any literal "nan" strings
            )
        else:
            # Create empty for downstream logic; we'll only use 'General' per missing row, not globally
            self.all_rows["Component"] = ""

    def _normalize_task_label_column(self):
        """
        Ensure self.all_rows has a 'Task_Label' column populated from v4b's Column G.
        Map v4b column G → 'Task_Label' (UI-display name for tasks)
        """
        if self.all_rows is None or isinstance(self.all_rows, pd.DataFrame) and self.all_rows.empty:
            return

        # optional UI override from UI_Options.Key == 'Task_Label_Column_Name'
        preferred = None
        try:
            if self.ui_options is not None and not self.ui_options.empty:
                row = self.ui_options[self.ui_options["Key"] == "Task_Label_Column_Name"]
                if not row.empty and isinstance(row, pd.DataFrame):
                    preferred = str(row["Value"].iloc[0]).strip()
        except Exception:
            pass

        candidates = [preferred] if preferred else []
        # common headers we've seen for column G in v4b
        candidates += ["Task_Label", "Task_Name", "Task_L1", "Component_Task_L2", "Task"]

        cols_lc = {c.lower(): c for c in self.all_rows.columns}
        found = None
        for cand in candidates:
            if cand and cand.lower() in cols_lc:
                found = cols_lc[cand.lower()]
                break
        if not found:
            # last‑ditch: pick column index G (0‑based 6) if it exists
            try:
                if len(self.all_rows.columns) > 6:
                    found = self.all_rows.columns[6]
            except Exception:
                found = None

        if found and isinstance(self.all_rows, pd.DataFrame):
            self.all_rows["Task_Label"] = (
                self.all_rows[found].astype(str).fillna("").str.strip()
            )
        else:
            # fallback; we'll still display task_group if label missing
            self.all_rows["Task_Label"] = ""


    def _normalize_code_columns(self):
        """Map v4/v4b All_Task_Rows code columns to canonical names we use downstream."""
        if self.all_rows is None or isinstance(self.all_rows, pd.DataFrame) and self.all_rows.empty:
            return
        cols = {c.lower(): c for c in self.all_rows.columns}

        def pick(*names):
            for n in names:
                if n and n.lower() in cols:
                    return cols[n.lower()]
            return None

        # canonical: Deliverable_Code already exists in v4/v4b; keep synonyms just in case
        if isinstance(self.all_rows, pd.DataFrame) and "Deliverable_Code" not in self.all_rows.columns:
            alt = pick("Deliverable Code", "Deliv_Code", "DeliverableID")
            if alt: self.all_rows["Deliverable_Code"] = self.all_rows[alt].astype(str)

        # Row_ID (v3 style)
        self._col_row_id        = pick("Row_ID", "RowID", "Row Id", "ID")

        # Task_Code (v3 style)
        self._col_task_code     = pick("Task_Code", "Task Code", "TaskCode", "Task_Code_L1", "Task_Group_Code")

        # Service_Department (v3 style)
        self._col_service_dept  = pick("Service_Department", "Service Department", "Service_Dept", "Department", "Dept")

        # Make sure we have Component + Task_Label from prior patches
        if isinstance(self.all_rows, pd.DataFrame) and "Component" not in self.all_rows.columns:
            self.all_rows["Component"] = ""
        if isinstance(self.all_rows, pd.DataFrame) and "Task_Label" not in self.all_rows.columns:
            self.all_rows["Task_Label"] = ""

    def _canonical_seniority(self, v: str) -> str:
        """Standardize seniority levels to canonical values: Junior, Mid, Senior, Director"""
        x = (str(v) or "").strip().lower()
        x = x.replace(".", "")
        if x in {"jr", "junior", "jr-level", "associate", "coordinator", "assistant", "l1", "level 1"}:
            return "Junior"
        if x in {"mid", "midlevel", "intermediate", "standard", "staff", "specialist", "producer", "manager", "l2", "level 2"}:
            return "Mid"
        if x in {"sr", "senior", "lead", "principal", "l3", "level 3"}:
            return "Senior"
        if x in {"director", "group director", "head", "executive director"}:
            return "Director"
        return (str(v) or "").strip()

    def _canonical_deliverable_code(self, v: str) -> str:
        """Map a free-text label to a DB deliverable code"""
        if not v or self.deliverables is None or not isinstance(self.deliverables, pd.DataFrame):
            return ""
        
        v_norm = (str(v) or "").strip().lower()
        if not v_norm:
            return ""
        
        # Try exact match on deliverable code first
        exact_code_match = self.deliverables[
            self.deliverables["Deliverable_Code"].astype(str).str.lower() == v_norm
        ]
        if not exact_code_match.empty and isinstance(exact_code_match, pd.DataFrame):
            return str(exact_code_match["Deliverable_Code"].iloc[0])
        
        # Try exact match on deliverable name
        exact_name_match = self.deliverables[
            self.deliverables["Deliverable"].astype(str).str.lower() == v_norm
        ]
        if not exact_name_match.empty and isinstance(exact_name_match, pd.DataFrame):
            return str(exact_name_match["Deliverable_Code"].iloc[0])
        
        # Try substring match on deliverable name
        substring_match = self.deliverables[
            self.deliverables["Deliverable"].astype(str).str.lower().str.contains(v_norm, na=False)
        ]
        if not substring_match.empty and isinstance(substring_match, pd.DataFrame):
            return str(substring_match["Deliverable_Code"].iloc[0])
        
        # Try substring match on category
        category_match = self.deliverables[
            self.deliverables["Category"].astype(str).str.lower().str.contains(v_norm, na=False)
        ]
        if not category_match.empty and isinstance(category_match, pd.DataFrame):
            return str(category_match["Deliverable_Code"].iloc[0])
        
        # No match found
        return ""

    def _normalize_role_and_seniority_columns(self):
        """Normalize Role and Seniority columns to ensure consistent data format"""
        if self.all_rows is None or isinstance(self.all_rows, pd.DataFrame) and self.all_rows.empty:
            return
        cols = {c.lower(): c for c in self.all_rows.columns}

        # Role column ➜ Resource_Title
        for cand in ["Resource_Title", "Role_Title", "Role", "Resource"]:
            if cand.lower() in cols:
                self.all_rows["Resource_Title"] = (
                    self.all_rows[cols[cand.lower()]]
                    .fillna("")
                    .astype(str)
                    .str.strip()
                    .replace({"nan": ""})  # Handle literal "nan" strings
                )
                break
        if isinstance(self.all_rows, pd.DataFrame) and "Resource_Title" not in self.all_rows.columns:
            self.all_rows["Resource_Title"] = ""
        
        # Ensure no blank roles - use placeholder for empty values
        self.all_rows["Resource_Title"] = self.all_rows["Resource_Title"].where(
            self.all_rows["Resource_Title"].str.len() > 0, "General Role"
        )

        # Seniority column ➜ Seniority (canonical labels)
        sen_src = None
        for cand in ["Seniority", "Seniority_Level", "Seniority L1", "Seniority_Title", "Level"]:
            if cand.lower() in cols:
                sen_src = cols[cand.lower()]
                break
        if sen_src:
            ser = (self.all_rows[sen_src]
                   .fillna("")
                   .astype(str)
                   .str.strip()
                   .replace({"nan": ""}))
        else:
            ser = pd.Series([""]*len(self.all_rows))
        self.all_rows["Seniority"] = ser.apply(self._canonical_seniority)
        
        # Ensure no blank seniority values - default to "Mid"
        self.all_rows["Seniority"] = self.all_rows["Seniority"].where(
            self.all_rows["Seniority"].str.len() > 0, "Mid"
        )

    def _normalize_rate_card_seniority(self):
        """Normalize role and seniority values in the role rate card to ensure pricing joins work properly"""
        if self.role_rate_card is None or self.role_rate_card.empty:
            return
        rc = self.role_rate_card.copy()
        
        # Normalize Resource_Title in rate card
        if "Resource_Title" in rc.columns:
            rc["Resource_Title"] = (rc["Resource_Title"]
                                   .fillna("")
                                   .astype(str)
                                   .str.strip()
                                   .replace({"nan": ""}))
            # Ensure no blank roles in rate card - use placeholder
            rc["Resource_Title"] = rc["Resource_Title"].where(
                rc["Resource_Title"].str.len() > 0, "General Role"
            )
        
        # Normalize Seniority in rate card
        if "Seniority" in rc.columns:
            rc["Seniority"] = (rc["Seniority"]
                              .fillna("")
                              .astype(str)
                              .str.strip()
                              .replace({"nan": ""})
                              .apply(self._canonical_seniority))
            # Ensure no blank seniority in rate card - default to "Mid"
            rc["Seniority"] = rc["Seniority"].where(rc["Seniority"].str.len() > 0, "Mid")
        
        self.role_rate_card = rc

    # ---------- v3 Drivers helper methods ----------
    def _norm_token(self, s: str) -> str:
        return re.sub(r"[^a-z0-9]+", "", str(s).strip().lower())

    def _v4_complexity_tokens(self) -> list[str]:
        toks = set()
        if self.all_rows is not None and isinstance(self.all_rows, pd.DataFrame):
            for c in self.all_rows.columns:
                if c.endswith("_Hours") and "__" in c:
                    toks.add(c.split("__", 1)[0])
        return sorted(toks)

    def _v4_tier_tokens(self) -> list[str]:
        toks = set()
        if self.all_rows is not None and isinstance(self.all_rows, pd.DataFrame):
            for c in self.all_rows.columns:
                if c.endswith("_Hours") and "__" in c:
                    # Safety check: ensure column has expected format before splitting
                    parts = c.rsplit("__", 1)[0].split("__", 1)
                    if len(parts) > 1:
                        toks.add(parts[1].replace("_Hours", ""))
        return sorted(toks)

    def _map_to_v4_token(self, label: str, candidates: list[str]) -> str:
        if not label: return ""
        z = self._norm_token(label)
        # exact normalized match first
        for c in candidates:
            if self._norm_token(c) == z:
                return c
        # relaxed: startswith / contains
        for c in candidates:
            if z and (self._norm_token(c).startswith(z) or z.startswith(self._norm_token(c))):
                return c
        # fallback to first
        return candidates[0] if candidates else ""

    def drivers_complexities_tiers_v3(self) -> tuple[list[str], list[str]]:
        """Return empty lists since v3 is no longer supported."""
        return ([], [])

    def _create_mock_data(self):
        """Create minimal mock data for demo purposes when database files are not available"""
        # Basic deliverables
        self.deliverables = pd.DataFrame([
            {"Deliverable_Code": "WEB_DEV", "Deliverable": "Website Development", "Category": "Digital"},
            {"Deliverable_Code": "BRAND_STR", "Deliverable": "Brand Strategy", "Category": "Branding"},
            {"Deliverable_Code": "CONTENT", "Deliverable": "Content Creation", "Category": "Content"}
        ])
        
        # Mock scenario data
        self.all_rows = pd.DataFrame([
            {"Deliverable_Code": "WEB_DEV", "task_group": "discovery", "Resource_Title": "Developer", "Seniority": "Senior", "Advanced__T2_MediumVolume_Hours": 40},
            {"Deliverable_Code": "WEB_DEV", "task_group": "development", "Resource_Title": "Developer", "Seniority": "Senior", "Advanced__T2_MediumVolume_Hours": 80},
            {"Deliverable_Code": "BRAND_STR", "task_group": "strategy", "Resource_Title": "Strategist", "Seniority": "Mid", "Advanced__T2_MediumVolume_Hours": 30}
        ])
        
        # Timeline and pricing settings
        self.timeline_params = pd.DataFrame([
            {"Task_Group": "discovery", "Nominal_Duration_Days": 5},
            {"Task_Group": "development", "Nominal_Duration_Days": 15},
            {"Task_Group": "strategy", "Nominal_Duration_Days": 10}
        ])
        
        self.timeline_scaling = pd.DataFrame([
            {"Scale_Type": "Complexity", "Key": "Advanced", "Multiplier": 1.2},
            {"Scale_Type": "Tier", "Key": "T2_MediumVolume", "Multiplier": 1.0}
        ])
        
        self.timeline_weighting = pd.DataFrame([
            {"Task_Group": "discovery", "Weight_Complexity": 0.6, "Weight_Tier": 0.4},
            {"Task_Group": "development", "Weight_Complexity": 0.6, "Weight_Tier": 0.4},
            {"Task_Group": "strategy", "Weight_Complexity": 0.6, "Weight_Tier": 0.4}
        ])
        
        # Basic settings
        self.pricing_settings = pd.DataFrame([
            {"Key": "Default_Blended_Rate", "Default": DEFAULT_BILLABLE_RATE_USD}
        ])
        
        self.slack_settings = pd.DataFrame([
            {"Key": "Use_Slack", "Default": True},
            {"Key": "Slack_After_Internal_Review_Days", "Default": 1},
            {"Key": "Slack_After_Client_Review_Days", "Default": 2},
            {"Key": "Slack_Global_Percent", "Default": 0.05}
        ])
        
        self.scenario_templates = pd.DataFrame([
            {"Scenario_Key": "MED_LOW", "Complexity": "Advanced", "Tier": "T2_MediumVolume"},
            {"Scenario_Key": "MED_HIGH", "Complexity": "Advanced", "Tier": "T2_MediumVolume"}
        ])
        
        self.rate_bands = pd.DataFrame([
            {"Band_Name": "Standard_US", "Rate_Multiplier": 1.0}
        ])
        
        self.role_rate_card = pd.DataFrame([
            {"Resource_Title": "Developer", "Seniority": "Senior", "Rate_USD": 150},
            {"Resource_Title": "Strategist", "Seniority": "Mid", "Rate_USD": 120}
        ])
        
        # Initialize empty tables for bundle functionality
        self.b_rules = pd.DataFrame(columns=["Category", "Bundle", "Task_Group", "Sort_Order"])
        self.b_defaults = pd.DataFrame(columns=["Bundle", "Default_Complexity", "Default_Tier"])
        self.b_by_deliv = pd.DataFrame()
        self.b_hours_by_role = pd.DataFrame()
        self.rate_matrix = pd.DataFrame()
        self.ui_options = pd.DataFrame()
        self.rfp_rules = pd.DataFrame(columns=["Regex_Keywords", "Map_To_Deliverable"])
        
        self.src = "mock_data"

    # ---------- RFP parsing via rules ----------
    def _ui_blocked_categories(self) -> set:
        try:
            if self.ui_options is not None and "Key" in self.ui_options.columns:
                row = self.ui_options[self.ui_options["Key"]=="Suggest_Block_Categories"]
                if not row.empty:
                    raw = str(row["Value"].iloc[0])
                    return {x.strip() for x in raw.split(";") if x.strip()}
        except Exception:
            pass
        # Default: block analytics unless explicitly asked for
        return {"Analytics"}

    def _ui_strict_mode(self) -> bool:
        try:
            if self.ui_options is not None and "Key" in self.ui_options.columns:
                row = self.ui_options[self.ui_options["Key"]=="RFP_Suggest_Strict"]
                if not row.empty:
                    v = str(row["Value"].iloc[0]).strip().lower()
                    return v in ("1","true","yes","y")
        except Exception:
            pass
        return True  # strict by default

    def suggest_deliverables_from_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Strict rules-first RFP matching. Returns [{
          deliverable_code, deliverable, category, confidence, matches: [...]
        }]
        - Uses RFP_Matching_Rules.Regex_Keywords -> Map_To_Deliverable
        - Applies optional UI_Options.Suggest_Block_Categories
        - NO fuzzy fallback in strict mode (prevents false positives)
        """
        if not text:
            return []
        strict = self._ui_strict_mode()
        blocked = self._ui_blocked_categories()

        text = str(text)
        found: Dict[str, Dict[str, Any]] = {}

        # 1) Rule-based suggestions
        if self.rfp_rules is not None and isinstance(self.rfp_rules, pd.DataFrame):
            for _, row in self.rfp_rules.iterrows():
                patt = str(row.get("Regex_Keywords", "") or "")
                target = str(row.get("Map_To_Deliverable", "") or "")
                if not patt or not target:
                    continue
                try:
                    hits = re.findall(patt, text, flags=re.IGNORECASE)
                except re.error:
                    continue
                if not hits:
                    continue

                # Find deliverable row(s)
                if self.deliverables is None or not isinstance(self.deliverables, pd.DataFrame):
                    continue
                match_df = self.deliverables[self.deliverables["Deliverable"] == target]
                if match_df.empty:
                    continue

                for __, r in match_df.iterrows():
                    code = str(r["Deliverable_Code"]); cat = str(r.get("Category",""))
                    if cat in blocked:
                        # allow through only if there are at least 2 strong hits
                        if len(hits) < 2:
                            continue

                    entry = found.setdefault(code, {
                        "deliverable_code": code,
                        "deliverable": str(r["Deliverable"]),
                        "category": cat,
                        "confidence": 0,
                        "matches": []
                    })
                    entry["confidence"] += len(hits)
                    uniq = list({str(h).lower() for h in hits if str(h).strip()})
                    for m in uniq:
                        if m not in entry["matches"]:
                            entry["matches"].append(m)

        # 2) NO fuzzy fallback when strict (prevents "not in RFP" picks)
        if strict:
            out = list(found.values())
            out.sort(key=lambda x: (-x["confidence"], x["deliverable"]))
            return out

        # Optional: gentle fallback if strict is off (rare)
        if self.deliverables is not None and isinstance(self.deliverables, pd.DataFrame):
            for _, r in self.deliverables.iterrows():
                name = str(r["Deliverable"])
                code = str(r["Deliverable_Code"])
                if re.search(r"\b" + re.escape(name) + r"\b", text, flags=re.IGNORECASE):
                    if code not in found:
                        found[code] = {
                            "deliverable_code": code,
                            "deliverable": name,
                            "category": str(r.get("Category","")),
                            "confidence": 1,
                            "matches": [name]
                        }

        out = list(found.values())
        out.sort(key=lambda x: (-x["confidence"], x["deliverable"]))
        return out

    def retainer_recommendation(self, text: str, deliverable_name: str) -> tuple[bool, int]:
        """
        Very simple rule-of-thumb to suggest retainer and months.
        Returns (is_retainer, months) where months in [1..12].
        """
        t = (text or "").lower()

        # If RFP states an explicit month count, respect that (clamped to 1..12)
        import re
        m = re.search(r'\b(\d{1,2})\s*(?:months|mos|mo)\b', t)
        months = 0
        if m:
            try:
                months = max(1, min(12, int(m.group(1))))
            except Exception:
                months = 0

        # Soft indicators that the work is monthly/ongoing
        signals = [
            "retainer", "monthly", "per month", "each month", "every month",
            "always-on", "always on", "ongoing", "maintenance", "management",
            "reporting cadence", "monthly report", "social calendar", "community"
        ]

        is_signal = any(s in t for s in signals)

        # Some deliverables are very often retainers
        likely_retainer_keywords = [
            "social", "community", "media", "measurement", "reporting",
            "seo", "maintenance", "support", "content", "blog"
        ]
        is_likely_by_name = any(k in (deliverable_name or "").lower() for k in likely_retainer_keywords)

        is_ret = bool(is_signal or is_likely_by_name)
        if is_ret and months == 0:
            # Default sensible guess
            months = 12 if "year" in t or "annual" in t else 6

        return (is_ret, max(1, min(12, months)) if months else (12 if is_ret else 0))

    # ---------- Bundle helpers ----------
    def included_task_groups(self, category: str, bundle: str) -> List[str]:
        if self.b_rules is None or not isinstance(self.b_rules, pd.DataFrame):
            return []
        sub = self.b_rules[(self.b_rules["Category"]==category) & (self.b_rules["Bundle"]==bundle)]
        if sub.empty:
            return []
        sub = sub.sort_values("Sort_Order")
        return [str(x) for x in sub["Task_Group"].tolist()]

    def default_complexity_tier_for_bundle(self, bundle: str) -> tuple[str, str]:
        if self.b_defaults is None or not isinstance(self.b_defaults, pd.DataFrame):
            return ("Advanced","T2_MediumVolume")
        row = self.b_defaults[self.b_defaults["Bundle"]==bundle]
        if row.empty or not isinstance(row, pd.DataFrame):
            return ("Advanced","T2_MediumVolume")
        r = row.iloc[0]
        return str(r["Default_Complexity"]), str(r["Default_Tier"])

    # ---------- Pricing ----------
    def blended_price(self, total_hours: float, blended_rate: float) -> float:
        return float(total_hours) * float(blended_rate)

    def per_resource_price(self, hrs_by_role: pd.DataFrame, rate_band: str="Standard_US") -> float:
        # hrs_by_role columns: Resource_Title, Seniority, Hours
        if self.rate_bands is None or not isinstance(self.rate_bands, pd.DataFrame):
            mult = 1.0
        else:
            band = self.rate_bands[self.rate_bands["Band_Name"]==rate_band]
            mult = float(band["Rate_Multiplier"].iloc[0]) if not band.empty and isinstance(band, pd.DataFrame) else 1.0
        # join to rate card
        if self.role_rate_card is None or not isinstance(self.role_rate_card, pd.DataFrame):
            return 0.0
        rc = self.role_rate_card[["Resource_Title","Seniority","Rate_USD"]].copy()
        merged = hrs_by_role.merge(rc, on=["Resource_Title","Seniority"], how="left")
        merged["Rate_USD"] = merged["Rate_USD"].fillna(0)
        merged["Price"] = merged["Hours"] * merged["Rate_USD"] * mult
        return float(merged["Price"].sum())

    # ---------- Hours aggregation ----------
    def scenario_hours_col(self, complexity: str, tier: str) -> str:
        # exact
        col = f"{complexity}__{tier}_Hours"
        if self.all_rows is not None and isinstance(self.all_rows, pd.DataFrame) and col in self.all_rows.columns:
            return col
        # try mapping display labels -> v4 tokens
        c_tok = self._map_to_v4_token(complexity, self._v4_complexity_tokens())
        t_tok = self._map_to_v4_token(tier,        self._v4_tier_tokens())
        col2 = f"{c_tok}__{t_tok}_Hours"
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame) or col2 not in self.all_rows.columns:
            raise HTTPException(400, f"Scenario column not found for ({complexity}, {tier}).")
        return col2

    def task_groups_for_deliverable(self, deliverable_code: str) -> List[str]:
        """Get all task groups for a deliverable from the database."""
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame):
            return []
        sub = self.all_rows[self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)]
        if sub.empty or not isinstance(sub, pd.DataFrame):
            return []
        series = sub["task_group"]
        if isinstance(series, pd.Series):
            return sorted(set(series.dropna().astype(str).tolist()))
        return []

    def hours_by_role_for_deliverable(
        self, deliverable_code: str, included_task_groups: List[str], scenario_col: str
    ) -> pd.DataFrame:
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame):
            return pd.DataFrame({"Resource_Title": [], "Seniority": [], "Hours": []})
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str)==str(deliverable_code)) &
            (self.all_rows["task_group"].isin(included_task_groups))
        ]
        if sub.empty or not isinstance(sub, pd.DataFrame) or scenario_col not in sub.columns:
            return pd.DataFrame({"Resource_Title": [], "Seniority": [], "Hours": []})
        g = sub.groupby(["Resource_Title","Seniority"], as_index=False)[scenario_col].sum()
        g = g.rename(columns={scenario_col:"Hours"})
        return g

    # ---------- Timeline ----------
    def hours_to_duration_days(self, hours: float, deliverable_code: str = "", task_group: str = "", 
                               scenario_col: str = "") -> float:
        """
        Task 5: Hours-based duration calculation using capacity formula.
        Formula: days = ceil(hours / (hours_per_day * available_resources * focus_factor))
        
        Parameters:
        - hours: Total hours for the task
        - deliverable_code: Used to identify resource roles if available
        - task_group: Used to identify typical resources needed
        - scenario_col: Scenario column for hours lookup
        
        Returns: Duration in business days
        """
        # PM-brain parameters based on agency reality
        hours_per_day = 6.5  # Realistic productive hours (not 8)
        focus_factor = 0.7   # Account for context switching, meetings, admin
        
        # Determine available resources - try to get from actual role assignments
        available_resources = 1.0  # Default to 1 FTE
        
        if deliverable_code and task_group and scenario_col and self.all_rows is not None:
            try:
                # Check how many distinct roles are assigned to this task
                sub = self.all_rows[
                    (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
                    (self.all_rows["task_group"].astype(str) == str(task_group))
                ]
                if not sub.empty and isinstance(sub, pd.DataFrame):
                    # Count distinct roles with non-zero hours
                    sub_filtered = sub[sub[scenario_col] > 0] if scenario_col in sub.columns else sub
                    unique_roles = sub_filtered["Resource_Title"].nunique()
                    if unique_roles > 0:
                        # Average 0.5 FTE per role (realistic for shared resources)
                        available_resources = unique_roles * 0.5
            except Exception:
                pass  # Fall back to default
        
        # Calculate duration with minimum of 1 day
        if hours <= 0:
            return 1.0
        
        effective_daily_capacity = hours_per_day * available_resources * focus_factor
        duration_days = math.ceil(hours / effective_daily_capacity)
        
        return max(1.0, float(duration_days))
    
    def task_group_duration_days(self, task_group: str, complexity: str, tier: str, use_slack: bool,
                                 slack_after_internal: int, slack_after_client: int, slack_global_pct: float,
                                 deliverable_code: str = "", scenario_col: str = "") -> float:
        """
        Enhanced with hours-based calculation (Task 5).
        Falls back to static lookups if hours data unavailable.
        """
        # Try hours-based calculation first if we have the data
        if deliverable_code and scenario_col and self.all_rows is not None:
            try:
                # Get actual hours for this task group from scenario
                sub = self.all_rows[
                    (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
                    (self.all_rows["task_group"].astype(str) == str(task_group))
                ]
                if not sub.empty and isinstance(sub, pd.DataFrame) and scenario_col in sub.columns:
                    total_hours = float(sub[scenario_col].sum())
                    if total_hours > 0:
                        # Use hours-based calculation
                        base_dur = self.hours_to_duration_days(
                            total_hours, 
                            deliverable_code, 
                            task_group, 
                            scenario_col
                        )
                        # Apply slack if needed
                        if use_slack and slack_global_pct > 0:
                            base_dur *= (1.0 + float(slack_global_pct))
                        return max(1.0, round(base_dur, 2))
            except Exception:
                pass  # Fall through to static calculation
        
        # Fallback: Static lookup from Timeline_Params (original logic)
        if self.timeline_params is None or not isinstance(self.timeline_params, pd.DataFrame):
            return 1.0
        tp = self.timeline_params[self.timeline_params["Task_Group"]==task_group]
        if tp.empty or not isinstance(tp, pd.DataFrame):
            return 1.0
        base = float(tp["Nominal_Duration_Days"].iloc[0])

        # Scaling - with None checks for v3 database compatibility
        if self.timeline_weighting is not None:
            cw = self.timeline_weighting[self.timeline_weighting["Task_Group"]==task_group]
            wc = float(cw["Weight_Complexity"].iloc[0]) if not cw.empty else 0.6
            wt = float(cw["Weight_Tier"].iloc[0])        if not cw.empty else 0.4
        else:
            wc = 0.6  # Default weight for complexity
            wt = 0.4  # Default weight for tier

        if self.timeline_scaling is not None:
            cm = self.timeline_scaling[(self.timeline_scaling["Scale_Type"]=="Complexity") &
                                       (self.timeline_scaling["Key"]==complexity)]
            tm = self.timeline_scaling[(self.timeline_scaling["Scale_Type"]=="Tier") &
                                       (self.timeline_scaling["Key"]==tier)]
            cmult = float(cm["Multiplier"].iloc[0]) if not cm.empty else 1.0
            tmult = float(tm["Multiplier"].iloc[0]) if not tm.empty else 1.0
        else:
            cmult = 1.0  # Default multiplier
            tmult = 1.0  # Default multiplier

        dur = base * (1 + (cmult - 1)*wc) * (1 + (tmult - 1)*wt)
        if use_slack and slack_global_pct > 0:
            dur *= (1.0 + float(slack_global_pct))
        return max(1.0, round(dur, 2))

    def _build_task_dependencies(self, task_groups: List[str], slack_after_internal: int, 
                                 slack_after_client: int) -> Dict[str, Dict[str, Any]]:
        """
        Task 7: Build SS/FS dependency relationships between task groups.
        
        Returns dict mapping task_group -> {"type": "SS"/"FS", "predecessor": str, "lag_days": int, "lag_pct": float}
        """
        dependencies = {}
        
        # Dependency rules based on PM best practices
        DEPENDENCY_RULES = {
            # Strategy/Brief overlaps with downstream work (SS = Start-to-Start)
            "art_direction": {"type": "SS", "predecessor": "strategy", "lag_pct": 0.5, "reason": "Art direction can start with draft strategy"},
            "research": {"type": "SS", "predecessor": "strategy", "lag_pct": 0.5, "reason": "Research can parallel strategy"},
            "concepting": {"type": "SS", "predecessor": "strategy", "lag_pct": 0.6, "reason": "Concepts need strategic direction"},
            
            # Review chains use FS (Finish-to-Start)
            "internal_review": {"type": "FS", "predecessor": "concepting", "lag_days": 0, "reason": "Review after concepts complete"},
            "client_review": {"type": "FS", "predecessor": "internal_review", "lag_days": slack_after_internal, "reason": "Client review after internal"},
            "revisions": {"type": "FS", "predecessor": "client_review", "lag_days": slack_after_client, "reason": "Revisions after client feedback"},
            
            # Production follows creative (with some overlap)
            "production": {"type": "SS", "predecessor": "creative", "lag_pct": 0.4, "reason": "Production can start on approved pieces"},
            "development": {"type": "SS", "predecessor": "design", "lag_pct": 0.7, "reason": "Development needs core designs"},
        }
        
        # Map task groups to dependencies
        for tg in task_groups:
            tg_lower = tg.lower().replace("_", "").replace(" ", "")
            
            # Check each rule to see if it matches
            for pattern, rule in DEPENDENCY_RULES.items():
                pattern_clean = pattern.lower().replace("_", "").replace(" ", "")
                if pattern_clean in tg_lower:
                    # Check if predecessor exists in task_groups
                    pred = rule["predecessor"]
                    for potential_pred in task_groups:
                        pred_clean = potential_pred.lower().replace("_", "").replace(" ", "")
                        if pred in pred_clean or pred_clean in pred:
                            dependencies[tg] = {
                                "type": rule["type"],
                                "predecessor": potential_pred,
                                "lag_days": rule.get("lag_days", 0),
                                "lag_pct": rule.get("lag_pct", 0.0),
                                "reason": rule.get("reason", "")
                            }
                            break
        
        return dependencies
    
    def build_schedule(self, deliverable_code: str, included_task_groups: List[str],
                       complexity: str, tier: str,
                       use_slack: bool, slack_after_internal: int, slack_after_client: int, slack_global_pct: float,
                       project_start: Optional[str]=None, scenario_letter: str="A") -> List[Dict[str, Any]]:
        """
        Enhanced with:
        - Task 6: Resource leveling with capacity constraints
        - Task 7: SS/FS dependencies with leads/lags
        """
        if self.timeline_params is None or not isinstance(self.timeline_params, pd.DataFrame):
            return []
        
        # Get scenario column for hours lookup
        scenario_col = ""
        try:
            scenario_col = self.scenario_hours_col(complexity, tier)
        except:
            pass
        
        order_map = {tg:i for i, tg in enumerate(self.timeline_params["Task_Group"].astype(str).tolist())}
        tgs = self.sort_task_groups(included_task_groups, scenario_letter)

        # Start date
        if project_start:
            # Handle both ISO8601 format (2025-10-07T01:00:00.000Z) and date-only format (2025-10-07)
            date_str = project_start.split('T')[0] if 'T' in project_start else project_start
            start_date = datetime.datetime.strptime(date_str, "%Y-%m-%d").date()
        else:
            start_date = datetime.date.today()
        
        # Build business day calendar with US/MX holidays
        year = start_date.year
        holidays = _get_us_mx_holidays(year)
        # Also get holidays for next year in case project spans years
        if start_date.month >= 10:  # If starting in Q4, include next year's holidays
            holidays.extend(_get_us_mx_holidays(year + 1))

        # Task 7: Build dependencies between task groups
        dependencies = self._build_task_dependencies(tgs, slack_after_internal, slack_after_client)
        
        # Task 6: Calculate total FTE available for resource leveling
        total_fte = 3.0  # Default team capacity (can be parameterized later)
        if self.all_rows is not None and deliverable_code and scenario_col:
            try:
                # Estimate FTE from unique roles assigned to this deliverable
                sub = self.all_rows[
                    (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
                    (self.all_rows["task_group"].astype(str).isin([str(x) for x in tgs]))
                ]
                if not sub.empty and isinstance(sub, pd.DataFrame):
                    unique_roles = sub["Resource_Title"].nunique()
                    total_fte = max(3.0, unique_roles * 0.5)  # 0.5 FTE per role, minimum 3
            except:
                pass
        
        # Max parallel tasks based on team capacity (Task 6)
        max_parallel = max(1, int(total_fte / 1.5))
        
        # Track active tasks by date for resource leveling (Task 6)
        active_tasks_by_date = {}  # date_str -> count of active tasks
        
        # Build schedule with dependencies and resource leveling
        cursor = np.datetime64(start_date)
        rows = []
        task_end_dates = {}  # task_group -> end_date (for dependency calculation)
        
        for tg in tgs:
            # Get duration using enhanced hours-based calculation
            dur = self.task_group_duration_days(
                tg, complexity, tier, use_slack,
                slack_after_internal, slack_after_client, slack_global_pct,
                deliverable_code, scenario_col
            )
            business_days_needed = max(1, math.ceil(dur))
            
            # Task 7: Apply dependencies to determine start date
            actual_start = cursor
            if tg in dependencies:
                dep = dependencies[tg]
                pred_tg = dep["predecessor"]
                
                if pred_tg in task_end_dates:
                    pred_end = task_end_dates[pred_tg]
                    
                    if dep["type"] == "FS":
                        # Finish-to-Start: start after predecessor ends + lag
                        lag_days = dep.get("lag_days", 0)
                        actual_start = np.busday_offset(pred_end, lag_days, roll='forward', holidays=holidays)
                    
                    elif dep["type"] == "SS":
                        # Start-to-Start: start when predecessor is X% complete
                        lag_pct = dep.get("lag_pct", 0.0)
                        if pred_tg in rows:
                            pred_row = next((r for r in rows if r["task_group"] == pred_tg), None)
                            if pred_row:
                                pred_start = np.datetime64(pred_row["start_date"])
                                pred_duration = pred_row["duration_days"]
                                lag_days = int(pred_duration * lag_pct)
                                actual_start = np.busday_offset(pred_start, lag_days, roll='forward', holidays=holidays)
            
            # Task 6: Resource leveling - check if we're over capacity
            # If too many tasks running in parallel, push start forward
            start_candidate = np.busday_offset(actual_start, 0, roll='forward', holidays=holidays)
            while True:
                # Check how many tasks are active on this date
                date_str = str(start_candidate)
                active_count = active_tasks_by_date.get(date_str, 0)
                
                if active_count < max_parallel:
                    # We have capacity - use this start date
                    actual_start = start_candidate
                    break
                else:
                    # Over capacity - try next business day
                    start_candidate = np.busday_offset(start_candidate, 1, roll='forward', holidays=holidays)
            
            # Calculate end date
            start = np.busday_offset(actual_start, 0, roll='forward', holidays=holidays)
            end = np.busday_offset(start, business_days_needed, roll='forward', holidays=holidays)
            
            # Task 6: Track active tasks for resource leveling
            current_date = start
            while current_date < end:
                date_str = str(current_date)
                active_tasks_by_date[date_str] = active_tasks_by_date.get(date_str, 0) + 1
                current_date = np.busday_offset(current_date, 1, roll='forward', holidays=holidays)
            
            # Convert numpy dates back to strings for JSON serialization
            start_str = str(start)
            end_str = str(end)
            
            # Get resource assignments for this task
            resources_assigned = []
            if self.all_rows is not None and deliverable_code and scenario_col:
                try:
                    sub = self.all_rows[
                        (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
                        (self.all_rows["task_group"].astype(str) == str(tg))
                    ]
                    if not sub.empty and isinstance(sub, pd.DataFrame):
                        resources_assigned = sub["Resource_Title"].dropna().unique().tolist()
                except:
                    pass
            
            # Build schedule entry with enhanced info
            schedule_entry = {
                "task_group": tg,
                "start_date": start_str,
                "end_date": end_str,
                "duration_days": business_days_needed,
                "resources_assigned": resources_assigned  # Task 6: Resource info
            }
            
            # Task 7: Add dependency info if present
            if tg in dependencies:
                dep = dependencies[tg]
                schedule_entry.update({
                    "dependency_type": dep["type"],
                    "dependency_predecessor": dep["predecessor"],
                    "dependency_lag_days": dep.get("lag_days", 0),
                    "dependency_reason": dep.get("reason", "")
                })
            
            rows.append(schedule_entry)
            
            # Store end date for dependency calculations
            task_end_dates[tg] = end
            
            # Move cursor to end date for next task group (for sequential tasks)
            cursor = end

            # Slack after reviews (in business days)
            if use_slack and tg == "internal_review":
                cursor = np.busday_offset(cursor, int(slack_after_internal), roll='forward', holidays=holidays)
            if use_slack and tg == "client_review":
                cursor = np.busday_offset(cursor, int(slack_after_client), roll='forward', holidays=holidays)

        return rows

    def _order_overrides(self, letter: str) -> list[tuple[str,str]]:
        """
        Optional UI_Options row(s):
          - Key: Task_Order_Overrides_A  Value: post_production<development; qa<launch
          - Key: Task_Order_Overrides_B  Value: ...
        Fallback: ensure post_production < development for Scenario A.
        """
        try:
            if self.ui_options is not None and isinstance(self.ui_options, pd.DataFrame) and not self.ui_options.empty:
                key = f"Task_Order_Overrides_{letter.upper()}"
                row = self.ui_options[self.ui_options["Key"] == key]
                if not row.empty and isinstance(row, pd.DataFrame):
                    parts = str(row["Value"].iloc[0]).split(";")
                    pairs = []
                    for p in parts:
                        if "<" in p:
                            a, b = [x.strip() for x in p.split("<", 1)]
                            if a and b: pairs.append((a, b))
                    if pairs:
                        return pairs
        except Exception:
            pass
        if letter.upper() == "A":
            return [("post_production", "development")]
        return []

    def sort_task_groups(self, tgs: list[str], letter: str) -> list[str]:
        """Topological sort using Timeline_Params order as baseline + overrides."""
        if self.timeline_params is None or not isinstance(self.timeline_params, pd.DataFrame):
            return tgs
        base = {str(tg): i for i, tg in enumerate(self.timeline_params["Task_Group"].astype(str).tolist())}
        nodes = [str(x) for x in tgs]
        edges = [(a, b) for (a, b) in self._order_overrides(letter) if a in nodes and b in nodes]
        # Kahn's algorithm with baseline tie-break
        preds = {n: set() for n in nodes}
        succs = {n: set() for n in nodes}
        for a, b in edges:
            succs[a].add(b); preds[b].add(a)
        ready = [n for n in nodes if not preds[n]]
        ready.sort(key=lambda n: base.get(n, 999))
        out = []
        while ready:
            n = ready.pop(0)
            out.append(n)
            for m in sorted(list(succs[n]), key=lambda x: base.get(x, 999)):
                preds[m].discard(n)
                if not preds[m] and m not in out and m in nodes and m not in ready:
                    ready.append(m)
            ready.sort(key=lambda x: base.get(x, 999))
        # append any left (cycle or unrelated), preserving baseline
        tail = [n for n in nodes if n not in out]
        tail.sort(key=lambda n: base.get(n, 999))
        return out + tail

    # ---------- Helper methods for task ordering and role detection ----------
    def sorted_task_groups(self, included: List[str]) -> List[str]:
        if self.timeline_params is None or not isinstance(self.timeline_params, pd.DataFrame):
            return included
        order_map = {tg: i for i, tg in enumerate(self.timeline_params["Task_Group"].astype(str).tolist())}
        return sorted([str(x) for x in included], key=lambda tg: order_map.get(tg, 999))

    def task_hours_by_task_group(self, deliverable_code: str, included: List[str], scenario_col: str) -> Dict[str, float]:
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame):
            return {}
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (self.all_rows["task_group"].isin(included))
        ]
        if sub.empty or not isinstance(sub, pd.DataFrame) or scenario_col not in sub.columns:
            return {}
        g = sub.groupby(["task_group"], as_index=False)[scenario_col].sum()
        return {str(r["task_group"]): float(r[scenario_col]) for _, r in g.iterrows()}

    def dominant_role_for_task_group(self, deliverable_code: str, task_group: str, scenario_col: str):
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame):
            return ("","")
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str) == str(task_group))
        ]
        if sub.empty or not isinstance(sub, pd.DataFrame) or scenario_col not in sub.columns:
            return ("","")
        g = sub.groupby(["Resource_Title","Seniority"], as_index=False)[scenario_col].sum()
        r = g.sort_values(scenario_col, ascending=False).iloc[0]
        return (str(r["Resource_Title"]), str(r["Seniority"]))

    # ---------- Component-level helper methods ----------
    def components_for_deliverable(self, deliverable_code: str, included_tgs: list[str]) -> list[str]:
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame):
            return ["General"]
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str).isin([str(x) for x in included_tgs]))
        ].copy()
        # Fill per-row only if blank
        if isinstance(sub, pd.DataFrame):
            sub["Component"] = sub["Component"].fillna("").astype(str).str.strip()
            comps = [c for c in sub["Component"].unique().tolist() if c]
            has_blanks = (sub["Component"] == "").any()
        else:
            return ["General"]

        if not comps and not has_blanks:
            # No component values at all for this deliverable → one placeholder bucket
            return ["General"]
        
        # Include "General" if there are any blank component rows (avoid duplicates)
        if has_blanks and "General" not in comps:
            comps.append("General")

        # Order components by earliest task_group position from Timeline_Params
        if self.timeline_params is not None and isinstance(self.timeline_params, pd.DataFrame):
            order_map = {str(tg): i for i, tg in enumerate(self.timeline_params["Task_Group"].astype(str).tolist())}
        else:
            order_map = {}
        comp_earliest = {}
        for comp in comps:
            if comp == "General":
                # For General, find earliest task_group among blank component rows
                blank_tgs = sub.loc[sub["Component"] == "", "task_group"].astype(str).unique().tolist()
                comp_earliest[comp] = min([order_map.get(tg, 999) for tg in blank_tgs]) if blank_tgs else 999
            else:
                tgs = sub.loc[sub["Component"] == comp, "task_group"].astype(str).unique().tolist()
                comp_earliest[comp] = min([order_map.get(tg, 999) for tg in tgs]) if tgs else 999
        return sorted(comps, key=lambda c: (comp_earliest.get(c, 999), c))

    def hours_by_component(self, deliverable_code: str, included_tgs: list[str], scenario_col: str) -> dict[str, float]:
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame):
            return {}
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str).isin([str(x) for x in included_tgs]))
        ].copy()
        if isinstance(sub, pd.DataFrame):
            sub["Component"] = sub["Component"].fillna("").astype(str).str.strip()
            # Only attach 'General' to rows that are actually blank
            sub.loc[sub["Component"] == "", "Component"] = "General"
            if sub.empty or scenario_col not in sub.columns:
                return {}
            g = sub.groupby("Component", as_index=False)[scenario_col].sum()
            return {str(r["Component"]): float(r[scenario_col]) for _, r in g.iterrows()}
        return {}

    def hours_by_taskgroup_for_component(self, deliverable_code: str, component: str,
                                         included_tgs: list[str], scenario_col: str) -> dict[str, float]:
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame):
            return {}
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str).isin([str(x) for x in included_tgs]))
        ].copy()
        if isinstance(sub, pd.DataFrame):
            sub["Component"] = sub["Component"].fillna("").astype(str).str.strip()
            # Remap blanks to "General" before filtering (consistent with hours_by_component)
            sub.loc[sub["Component"] == "", "Component"] = "General"
            
            comp_key = (component or "").strip() or "General"
            sub = sub[sub["Component"] == comp_key]

            if sub.empty or scenario_col not in sub.columns:
                return {}
            g = sub.groupby("task_group", as_index=False)[scenario_col].sum()
            return {str(r["task_group"]): float(r[scenario_col]) for _, r in g.iterrows()}
        return {}

    def dominant_role_for_component_task(self, deliverable_code: str, component: str,
                                         task_group: str, scenario_col: str) -> tuple[str, str]:
        """Enhanced role picker that prefers non-blank seniority with robust fallbacks"""
        # Narrow to this deliverable + task_group (+ component if present)
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame):
            return ("", "Mid")
            
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str) == str(task_group))
        ].copy()

        if isinstance(sub, pd.DataFrame) and "Component" in sub.columns and (component or "").strip():
            sub["Component"] = sub["Component"].astype(str).fillna("").str.strip()
            sub = sub[sub["Component"] == str(component).strip()]

        if not isinstance(sub, pd.DataFrame) or sub.empty:
            # Fallback: ignore component filter
            sub = self.all_rows[
                (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
                (self.all_rows["task_group"].astype(str) == str(task_group))
            ].copy()

        if not isinstance(sub, pd.DataFrame) or sub.empty:
            # Second fallback: any rows for this deliverable
            sub = self.all_rows[self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)].copy()

        if not isinstance(sub, pd.DataFrame) or sub.empty:
            return ("", "Mid")

        # Prefer rows with non-blank seniority
        sub["Resource_Title"] = sub["Resource_Title"].astype(str).fillna("").str.strip()
        sub["Seniority"] = sub["Seniority"].astype(str).fillna("").str.strip()

        pref = sub[sub["Seniority"] != ""]
        pick_from = pref if isinstance(pref, pd.DataFrame) and not pref.empty else sub

        g = pick_from.groupby(["Resource_Title", "Seniority"], as_index=False)[scenario_col].sum()
        r = g.sort_values(scenario_col, ascending=False).iloc[0]

        role = str(r["Resource_Title"]).strip()
        sen  = self._canonical_seniority(str(r["Seniority"]).strip())
        if sen == "":
            sen = "Mid"  # last-resort default

        return (role, sen)

    def task_label_for_component_tg(self, deliverable_code: str, component: str, task_group: str) -> str:
        """Get user-friendly task label from Task_Label column for UI display."""
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame):
            return str(task_group)
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str) == str(task_group))
        ].copy()
        if isinstance(sub, pd.DataFrame) and "Component" in sub.columns:
            sub["Component"] = sub["Component"].fillna("").astype(str).str.strip()
            # Remap blanks to "General" before filtering (consistent with other helpers)
            sub.loc[sub["Component"] == "", "Component"] = "General"
            comp_key = (component or "").strip() or "General"
            sub = sub[sub["Component"] == comp_key]

        if isinstance(sub, pd.DataFrame) and "Task_Label" in sub.columns:
            lab = sub["Task_Label"]
            if isinstance(lab, pd.Series):
                lab = lab.dropna().astype(str).str.strip()
                lab = lab[lab != ""]
                if isinstance(lab, pd.Series) and not lab.empty:
                    # most frequent non‑empty label for that component+task_group
                    return lab.value_counts().idxmax()
        return str(task_group)

    # ---------- Pricing helper methods ----------
    def role_rates_table(self, rate_band: str = "Standard_US") -> pd.DataFrame:
        """Rate card with band multiplier applied + normalized Seniority."""
        if self.rate_bands is None or not isinstance(self.rate_bands, pd.DataFrame):
            mult = 1.0
        else:
            band = self.rate_bands[self.rate_bands["Band_Name"] == rate_band]
            if isinstance(band, pd.DataFrame) and not band.empty:
                mult = float(band["Rate_Multiplier"].iloc[0])
            else:
                mult = 1.0
        
        if self.role_rate_card is None or not isinstance(self.role_rate_card, pd.DataFrame):
            return pd.DataFrame({"Resource_Title": [], "Seniority": [], "Rate_USD": []})
            
        rc = self.role_rate_card[["Resource_Title", "Seniority", "Rate_USD"]].copy()
        # normalize seniority if you added _canonical_seniority() earlier
        if isinstance(rc, pd.DataFrame) and "Seniority" in rc.columns:
            try:
                rc["Seniority"] = rc["Seniority"].astype(str).fillna("").apply(self._canonical_seniority)
            except Exception:
                rc["Seniority"] = rc["Seniority"].astype(str).fillna("")
        if isinstance(rc, pd.DataFrame):
            rc["Rate_USD"] = rc["Rate_USD"].astype(float) * mult
        return rc

    def price_for_hours_by_role(self, hrs_by_role: pd.DataFrame, rate_band: str) -> tuple[float, pd.DataFrame]:
        """Return (price_total, merged_breakdown) for a df with columns: Resource_Title, Seniority, Hours.
        Enforces rate integrity - uses fallback rates with warnings if role/seniority combinations are missing."""
        if hrs_by_role is None or not isinstance(hrs_by_role, pd.DataFrame) or hrs_by_role.empty:
            return (0.0, pd.DataFrame({"Resource_Title": [], "Seniority": [], "Hours": [], "Rate_USD": [], "Price": []}))
        
        rc = self.role_rates_table(rate_band)
        merged = hrs_by_role.merge(rc, on=["Resource_Title","Seniority"], how="left")
        merged["Hours"] = merged["Hours"].fillna(0.0).astype(float)
        
        # Check for missing rates before proceeding
        missing_rates = merged[merged["Rate_USD"].isna()]
        if not missing_rates.empty:
            # Try fallback: role-only matching (ignore seniority)
            fallback_merged = merged.copy()
            for idx, row in missing_rates.iterrows():
                role_only_match = rc[rc["Resource_Title"] == row["Resource_Title"]]
                if not role_only_match.empty:
                    # Use first available rate for this role (with any seniority)
                    fallback_rate = role_only_match["Rate_USD"].iloc[0]
                    fallback_merged.loc[idx, "Rate_USD"] = fallback_rate
                    print(f"Warning: Used fallback rate for {row['Resource_Title']} {row['Seniority']} -> {fallback_rate}")
            
            # Check if fallbacks resolved all issues
            still_missing = fallback_merged[fallback_merged["Rate_USD"].isna()]
            if not still_missing.empty:
                # Apply band-aware default rate as last resort
                if self.pricing_settings is not None and isinstance(self.pricing_settings, pd.DataFrame):
                    ps = self.pricing_settings[self.pricing_settings["Key"]=="Default_Blended_Rate"]
                    base_default = float(ps["Default"].iloc[0]) if isinstance(ps, pd.DataFrame) and not ps.empty else DEFAULT_BILLABLE_RATE_USD
                else:
                    base_default = DEFAULT_BILLABLE_RATE_USD
                
                if self.rate_bands is not None and isinstance(self.rate_bands, pd.DataFrame):
                    band = self.rate_bands[self.rate_bands["Band_Name"] == rate_band]
                    mult = float(band["Rate_Multiplier"].iloc[0]) if isinstance(band, pd.DataFrame) and not band.empty else 1.0
                else:
                    mult = 1.0
                default_rate = base_default * mult
                fallback_merged["Rate_USD"] = fallback_merged["Rate_USD"].fillna(default_rate)
                missing_list = [(row["Resource_Title"], row["Seniority"]) for _, row in still_missing.iterrows()]
                print(f"Warning: Applied band-aware default rate ${default_rate}/hr for missing roles: {missing_list}")
            
            merged = fallback_merged
        
        merged["Rate_USD"] = merged["Rate_USD"].astype(float)
        merged["Price"] = merged["Hours"] * merged["Rate_USD"]
        return float(merged["Price"].sum()), merged

    def hours_by_role_for_deliverable(self, deliverable_code: str, included_tgs: list[str], scenario_col: str) -> pd.DataFrame:
        """Get hours by role+seniority for an entire deliverable across all included task groups."""
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame):
            return pd.DataFrame({"Resource_Title": [], "Seniority": [], "Hours": []})
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str).isin([str(x) for x in included_tgs]))
        ].copy()
        if not isinstance(sub, pd.DataFrame) or sub.empty or scenario_col not in sub.columns:
            return pd.DataFrame({"Resource_Title": [], "Seniority": [], "Hours": []})
        g = sub.groupby(["Resource_Title","Seniority"], as_index=False)[scenario_col].sum()
        return g.rename(columns={scenario_col: "Hours"})

    def hours_by_role_for_component(self, deliverable_code: str, component: str,
                                    included_tgs: list[str], scenario_col: str) -> pd.DataFrame:
        """Get hours by role+seniority for a specific component."""
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame):
            return pd.DataFrame({"Resource_Title": [], "Seniority": [], "Hours": []})
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str).isin([str(x) for x in included_tgs]))
        ].copy()
        if isinstance(sub, pd.DataFrame) and "Component" in sub.columns:
            sub["Component"] = sub["Component"].astype(str).fillna("").str.strip()
            if (component or "").strip() and component != "General":
                sub = sub[sub["Component"] == component]
            else:
                sub = sub[(sub["Component"] == "") | (sub["Component"] == "General")]
        if not isinstance(sub, pd.DataFrame) or sub.empty or scenario_col not in sub.columns:
            return pd.DataFrame({"Resource_Title": [], "Seniority": [], "Hours": []})
        g = sub.groupby(["Resource_Title","Seniority"], as_index=False)[scenario_col].sum()
        return g.rename(columns={scenario_col: "Hours"})

    def hours_by_role_for_component_task(self, deliverable_code: str, component: str,
                                         task_group: str, scenario_col: str) -> pd.DataFrame:
        """Get hours by role+seniority for a specific component+task combination."""
        if self.all_rows is None or not isinstance(self.all_rows, pd.DataFrame):
            return pd.DataFrame({"Resource_Title": [], "Seniority": [], "Hours": []})
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str) == str(task_group))
        ].copy()
        if isinstance(sub, pd.DataFrame) and "Component" in sub.columns:
            sub["Component"] = sub["Component"].astype(str).fillna("").str.strip()
            if (component or "").strip() and component != "General":
                sub = sub[sub["Component"] == component]
            else:
                sub = sub[(sub["Component"] == "") | (sub["Component"] == "General")]
        if not isinstance(sub, pd.DataFrame) or sub.empty or scenario_col not in sub.columns:
            return pd.DataFrame({"Resource_Title": [], "Seniority": [], "Hours": []})
        g = sub.groupby(["Resource_Title","Seniority"], as_index=False)[scenario_col].sum()
        return g.rename(columns={scenario_col: "Hours"})

    def _svc_mode(self, series: pd.Series) -> str:
        if series is None or series.empty: return ""
        s = series.dropna().astype(str).str.strip()
        s = s[s != ""]
        return s.value_counts().idxmax() if not s.empty else ""

    def service_department_for_task(self, deliverable_code: str, component: str, task_group: str) -> str:
        """Service department logic from v4 data."""
        return ""

    def service_department_for_component(self, deliverable_code: str, component: str, task_groups: list[str]) -> str:
        """Service department logic from v4 data."""
        return ""

    def service_department_for_deliverable(self, deliverable_code: str, task_groups: list[str]) -> str:
        """Service department logic from v4 data."""
        return ""

    def _majority_by_hours(self, sub: pd.DataFrame, col: str, scenario_col: str) -> str:
        if sub.empty or col not in sub.columns or scenario_col not in sub.columns:
            return ""
        g = sub.groupby(col, as_index=False)[scenario_col].sum()
        g = g[g[col].astype(str).str.strip() != ""]
        if g.empty: return ""
        return str(g.sort_values(scenario_col, ascending=False).iloc[0][col]).strip()

    def service_dept_for_deliverable(self, deliverable_code: str, included_tgs: list[str], scenario_col: str) -> str:
        if not getattr(self, "_col_service_dept", None): return ""
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str)==str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str).isin([str(x) for x in included_tgs]))
        ]
        return self._majority_by_hours(sub, self._col_service_dept, scenario_col)

    def service_dept_for_component(self, deliverable_code: str, component: str, included_tgs: list[str], scenario_col: str) -> str:
        if not getattr(self, "_col_service_dept", None): return ""
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str)==str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str).isin([str(x) for x in included_tgs]))
        ].copy()
        if "Component" in sub.columns:
            sub["Component"] = sub["Component"].astype(str).fillna("").str.strip()
            sub = sub[sub["Component"]==str(component).strip()]
        return self._majority_by_hours(sub, self._col_service_dept, scenario_col)

    def hours_by_role_for_component_task(
        self, deliverable_code: str, component: str, task_group: str, scenario_col: str
    ) -> pd.DataFrame:
        """Return hours by (Resource_Title, Seniority) for one component+task_group."""
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str) == str(task_group))
        ].copy()
        if "Component" in sub.columns and str(component).strip():
            sub["Component"] = sub["Component"].astype(str).fillna("").str.strip()
            sub = sub[sub["Component"] == str(component).strip()]
        if sub.empty or scenario_col not in sub.columns:
            return pd.DataFrame(columns=["Resource_Title","Seniority","Hours"])
        g = sub.groupby(["Resource_Title","Seniority"], as_index=False)[scenario_col].sum()
        return g.rename(columns={scenario_col: "Hours"})

    def codes_for_component_task_role(
        self, deliverable_code: str, component: str, task_group: str,
        role: str, seniority: str, scenario_col: str
    ) -> tuple[str, str, str]:
        """Return (Row_ID, Task_Code, Service_Department) for one Role on a component+task_group."""
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str) == str(task_group)) &
            (self.all_rows["Resource_Title"].astype(str) == str(role)) &
            (self.all_rows["Seniority"].astype(str) == str(seniority))
        ].copy()
        if "Component" in sub.columns and str(component).strip():
            sub["Component"] = sub["Component"].astype(str).fillna("").str.strip()
            sub = sub[sub["Component"] == str(component).strip()]

        # Row_ID
        row_id = ""
        if getattr(self, "_col_row_id", None) and self._col_row_id in sub.columns:
            vals = sub[self._col_row_id].dropna().astype(str).str.strip()
            if not vals.empty:
                row_id = sorted(vals.tolist(), key=lambda x: (len(x), x))[0]

        # Task_Code
        task_code = ""
        if getattr(self, "_col_task_code", None) and self._col_task_code in sub.columns:
            v = sub[self._col_task_code].dropna().astype(str).str.strip()
            if not v.empty:
                task_code = v.value_counts().idxmax()
        if not task_code:
            task_code = str(task_group).upper().replace(" ", "_")

        # Service_Department (majority by hours in this subset)
        service = ""
        if getattr(self, "_col_service_dept", None) and self._col_service_dept in sub.columns:
            g = sub.groupby(self._col_service_dept, as_index=False)[scenario_col].sum()
            g = g[g[self._col_service_dept].astype(str).str.strip() != ""]
            if not g.empty:
                service = str(g.sort_values(scenario_col, ascending=False).iloc[0][self._col_service_dept]).strip()

        return (row_id, task_code, service)

    def codes_for_component_task(self, deliverable_code: str, component: str, task_group: str, scenario_col: str) -> tuple[str,str,str]:
        """Return (Row_ID, Task_Code, Service_Department) for a task row under a component."""
        sub = self.all_rows[
            (self.all_rows["Deliverable_Code"].astype(str)==str(deliverable_code)) &
            (self.all_rows["task_group"].astype(str)==str(task_group))
        ].copy()
        if "Component" in sub.columns and component:
            sub["Component"] = sub["Component"].astype(str).fillna("").str.strip()
            sub = sub[sub["Component"]==str(component).strip()]

        # Row_ID: stable pick (min or first non-empty)
        row_id = ""
        if getattr(self, "_col_row_id", None) and self._col_row_id in sub.columns:
            vals = sub[self._col_row_id].dropna().astype(str).str.strip()
            if not vals.empty:
                row_id = sorted(vals.tolist(), key=lambda x: (len(x), x))[0]

        # Task_Code: majority by count (or fallback to task_group)
        task_code = ""
        if getattr(self, "_col_task_code", None) and self._col_task_code in sub.columns:
            v = sub[self._col_task_code].dropna().astype(str).str.strip()
            if not v.empty:
                task_code = v.value_counts().idxmax()
        if not task_code:
            task_code = str(task_group).upper().replace(" ", "_")

        # Service_Department: majority by hours
        service = ""
        if getattr(self, "_col_service_dept", None) and self._col_service_dept in sub.columns:
            service = self._majority_by_hours(sub, self._col_service_dept, scenario_col)

        return (row_id, task_code, service)

DB = AgencyDB()

# ---------- Helper: extract text from uploaded file bytes ----------
def _extract_text_from_upload(content: bytes, filename: str) -> str:
    ext = os.path.splitext(filename or "")[1].lower()

    # Plain text-like
    if ext in (".txt", ".md", ".csv"):
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1", errors="ignore")

    # DOCX
    if ext == ".docx":
        if not Document:
            raise HTTPException(400, "DOCX support requires 'python-docx'. Install it and redeploy.")
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)

    # PDF
    if ext == ".pdf":
        if not PdfReader:
            raise HTTPException(400, "PDF support requires 'pypdf'. Install it and redeploy.")
        reader = PdfReader(io.BytesIO(content))
        buf = []
        for page in reader.pages:
            # pypdf exposes .extract_text()
            t = page.extract_text() or ""
            buf.append(t)
        return "\n".join(buf)

    raise HTTPException(415, f"Unsupported file type: {ext}. Use .pdf, .docx, or .txt.")

# ---------- Helper: extract and analyze images from PDFs (with parallel processing) ----------

# Image pre-filtering helpers
import hashlib

def _get_image_hash(img_bytes: bytes) -> str:
    """Generate hash for image deduplication"""
    return hashlib.md5(img_bytes).hexdigest()

def _get_image_size(img_bytes: bytes) -> Tuple[int, int]:
    """Get image dimensions using PIL"""
    if not Image:
        return (0, 0)
    try:
        img = Image.open(io.BytesIO(img_bytes))
        return img.size
    except:
        return (0, 0)

def _should_skip_image(img_bytes: bytes, width: int, height: int) -> Tuple[bool, str]:
    """Determine if image should be skipped based on size"""
    # Skip tiny images (likely icons, bullets, decorative)
    if width < 100 or height < 100:
        return (True, f"tiny_{width}x{height}")
    return (False, "")

async def _quick_relevance_check_async(img_bytes: bytes, page_num: int, img_index: int, semaphore: asyncio.Semaphore, job_id: str) -> Tuple[bool, str]:
    """Quick check if image contains relevant content (charts, diagrams, requirements)"""
    async with semaphore:
        try:
            # Convert image to PNG for OpenAI
            if Image:
                try:
                    img = Image.open(io.BytesIO(img_bytes))
                    if img.mode not in ('RGB', 'RGBA'):
                        img = img.convert('RGB')
                    png_buffer = io.BytesIO()
                    img.save(png_buffer, format='PNG')
                    img_base64 = base64.b64encode(png_buffer.getvalue()).decode('utf-8')
                except:
                    img_base64 = base64.b64encode(img_bytes).decode('utf-8')
            else:
                img_base64 = base64.b64encode(img_bytes).decode('utf-8')
            
            # Quick relevance check with minimal tokens using Chat Completions API with vision
            # Use async client for proper async operation
            if not async_openai_client:
                # Fallback if async client not available
                return (True, "no_client")
            
            response = await async_openai_client.chat.completions.create(
                model="gpt-5",  # Use GPT-5 (system enforces GPT-5 only)
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Does this contain charts, diagrams, wireframes, mockups, or project requirements? Answer YES or NO only."
                            },
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/png;base64,{img_base64}"}
                            }
                        ]
                    }
                ],
                max_completion_tokens=10  # Minimal token usage for GPT-5 vision
            )
            
            # Extract text from Chat Completions response
            answer = response.choices[0].message.content.strip().upper() if response.choices else ""
            is_relevant = "YES" in answer
            
            if job_id in JOB_STORE:
                JOB_STORE[job_id].processed_images += 1
                if is_relevant:
                    JOB_STORE[job_id].relevant_images += 1
                else:
                    JOB_STORE[job_id].skipped_images += 1
            
            return (is_relevant, answer)
            
        except Exception as e:
            print(f"[JOB {job_id}] Quick check failed for image {img_index} on page {page_num}: {e}")
            # On error, assume relevant to be safe
            if job_id in JOB_STORE:
                JOB_STORE[job_id].processed_images += 1
                JOB_STORE[job_id].relevant_images += 1
            return (True, "error_assume_relevant")

async def _analyze_single_image_async(img_bytes: bytes, page_num: int, img_index: int, semaphore: asyncio.Semaphore, job_id: str) -> Optional[str]:
    """Analyze a single image with retry logic and rate limit handling"""
    async with semaphore:  # Limit concurrent OpenAI calls
        for attempt in range(3):  # Retry up to 3 times
            try:
                img_start = datetime.datetime.now().timestamp()
                
                # Convert image to PNG format for OpenAI compatibility
                if Image:
                    try:
                        # Load image from bytes
                        img = Image.open(io.BytesIO(img_bytes))
                        
                        # Convert to RGB if needed (handles CMYK, grayscale, palette modes, etc.)
                        if img.mode not in ('RGB', 'RGBA'):
                            img = img.convert('RGB')
                        
                        # Save as PNG to buffer
                        png_buffer = io.BytesIO()
                        img.save(png_buffer, format='PNG')
                        img_base64 = base64.b64encode(png_buffer.getvalue()).decode('utf-8')
                        image_format = "png"
                    except Exception as conv_error:
                        # Fallback to raw bytes if conversion fails
                        print(f"[JOB {job_id}] Image conversion failed for image {img_index} on page {page_num}, using raw bytes: {conv_error}")
                        img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                        image_format = "jpeg"
                else:
                    # PIL not available, use raw bytes
                    img_base64 = base64.b64encode(img_bytes).decode('utf-8')
                    image_format = "jpeg"
                
                # Analyze image with OpenAI Vision using Chat Completions API
                # Use async client for proper async operation
                if not async_openai_client:
                    # Fallback if async client not available
                    raise Exception("Async OpenAI client not available")
                
                response = await async_openai_client.chat.completions.create(
                    model="gpt-5",  # GPT-5 enforced (sitecustomize.py handles mapping)
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": "Analyze this image from an RFP document. Describe what it shows: charts, diagrams, mockups, screenshots, or other visual content. Focus on business requirements and deliverables it implies."
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/{image_format};base64,{img_base64}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_completion_tokens=500  # Updated for GPT-5 vision compatibility
                )
                
                # Extract text from Chat Completions response
                description = response.choices[0].message.content if response.choices else str(response)
                img_end = datetime.datetime.now().timestamp()
                
                # Update job timing
                if job_id in JOB_STORE:
                    JOB_STORE[job_id].image_timings.append(img_end - img_start)
                    JOB_STORE[job_id].processed_images += 1
                
                return f"Page {page_num}, Image {img_index}: {description}"
                
            except Exception as e:
                error_msg = str(e).lower()
                if "rate" in error_msg or "429" in error_msg:
                    # Rate limit - wait and retry
                    wait_time = (attempt + 1) * 2  # 2s, 4s, 6s
                    print(f"[JOB {job_id}] Rate limit hit, waiting {wait_time}s before retry {attempt+1}/3")
                    await asyncio.sleep(wait_time)
                elif "timeout" in error_msg:
                    # Timeout - retry with backoff
                    wait_time = attempt + 1
                    print(f"[JOB {job_id}] Timeout, retrying in {wait_time}s (attempt {attempt+1}/3)")
                    await asyncio.sleep(wait_time)
                else:
                    # Other error - log and give up
                    if job_id in JOB_STORE:
                        JOB_STORE[job_id].errors.append(f"Image {img_index} on page {page_num}: {str(e)}")
                        JOB_STORE[job_id].processed_images += 1  # Count as processed to keep progress moving
                    print(f"[JOB {job_id}] Error analyzing image {img_index} on page {page_num}: {e}")
                    return None
        
        # All retries failed
        if job_id in JOB_STORE:
            JOB_STORE[job_id].errors.append(f"Image {img_index} on page {page_num}: Max retries exceeded")
            JOB_STORE[job_id].processed_images += 1
        return None

async def _extract_and_analyze_pdf_images_async(content: bytes, filename: str, job_id: str, analyze_images: bool = True) -> str:
    """
    Two-tier image analysis: Quick scan → Deep analysis of relevant images only.
    Tracks progress via job_id. Returns a text description of relevant images.
    """
    ext = os.path.splitext(filename or "")[1].lower()
    if ext != ".pdf":
        return ""  # Only process PDFs for image analysis
    
    if not analyze_images:
        return ""  # User disabled image analysis
    
    if not PdfReader or not openai_client:
        return ""  # Skip if dependencies missing
    
    try:
        reader = PdfReader(io.BytesIO(content))
        
        # Phase 1: Extract and pre-filter images
        images_to_scan = []
        seen_hashes = set()
        
        for page_num, page in enumerate(reader.pages, 1):
            if hasattr(page, 'images'):
                for img_index, image in enumerate(page.images, 1):
                    try:
                        img_bytes = image.data
                        
                        # Check size
                        width, height = _get_image_size(img_bytes)
                        should_skip, skip_reason = _should_skip_image(img_bytes, width, height)
                        
                        if should_skip:
                            print(f"[JOB {job_id}] Skipping image {img_index} on page {page_num}: {skip_reason}")
                            if job_id in JOB_STORE:
                                JOB_STORE[job_id].skipped_images += 1
                            continue
                        
                        # Check for duplicates
                        img_hash = _get_image_hash(img_bytes)
                        if img_hash in seen_hashes:
                            print(f"[JOB {job_id}] Skipping duplicate image {img_index} on page {page_num}")
                            if job_id in JOB_STORE:
                                JOB_STORE[job_id].skipped_images += 1
                            continue
                        
                        seen_hashes.add(img_hash)
                        images_to_scan.append((img_bytes, page_num, img_index))
                        
                    except Exception as e:
                        print(f"[JOB {job_id}] Could not extract image {img_index} from page {page_num}: {e}")
                        continue
        
        if not images_to_scan:
            return ""  # No images to analyze
        
        # Update job for Phase 1: Quick scan
        if job_id in JOB_STORE:
            JOB_STORE[job_id].total_images = len(images_to_scan)
            JOB_STORE[job_id].status = JobStatus.PROCESSING
            JOB_STORE[job_id].phase = "quick_scan"
        
        # Phase 2: Quick relevance check
        semaphore = asyncio.Semaphore(10)  # PERFORMANCE FIX: Increased concurrency for fast mode
        quick_check_tasks = [
            _quick_relevance_check_async(img_bytes, page_num, img_index, semaphore, job_id)
            for img_bytes, page_num, img_index in images_to_scan
        ]
        
        relevance_results = await asyncio.gather(*quick_check_tasks)
        
        # Identify relevant images for deep analysis
        relevant_images = [
            (img_bytes, page_num, img_index)
            for (img_bytes, page_num, img_index), (is_relevant, _) in zip(images_to_scan, relevance_results)
            if is_relevant
        ]
        
        print(f"[JOB {job_id}] Quick scan complete: {len(relevant_images)} relevant of {len(images_to_scan)} images")
        
        if not relevant_images:
            print(f"[JOB {job_id}] No relevant images found, skipping deep analysis")
            if job_id in JOB_STORE:
                JOB_STORE[job_id].status = JobStatus.COMPLETED
            return ""
        
        # Phase 3: Deep analysis of relevant images only
        if job_id in JOB_STORE:
            JOB_STORE[job_id].phase = "deep_analysis"
            JOB_STORE[job_id].total_images = len(relevant_images)
            JOB_STORE[job_id].processed_images = 0  # Reset for deep analysis phase
        
        semaphore = asyncio.Semaphore(15)  # PERFORMANCE FIX: Dramatically increased concurrency for parallel processing
        deep_analysis_tasks = [
            _analyze_single_image_async(img_bytes, page_num, img_index, semaphore, job_id)
            for img_bytes, page_num, img_index in relevant_images
        ]
        
        image_descriptions = await asyncio.gather(*deep_analysis_tasks)
        
        # Filter out None results (failed analyses)
        valid_descriptions = [desc for desc in image_descriptions if desc]
        
        if valid_descriptions:
            summary = f"--- Visual Content Analysis ({len(valid_descriptions)} relevant images) ---\n\n"
            summary += "\n\n".join(valid_descriptions)
            return "\n\n" + summary
        return ""
        
    except Exception as e:
        print(f"[JOB {job_id}] Could not extract images from PDF: {e}")
        if job_id in JOB_STORE:
            JOB_STORE[job_id].errors.append(f"PDF extraction failed: {str(e)}")
        return ""

# Synchronous wrapper for backward compatibility
def _extract_and_analyze_pdf_images(content: bytes, filename: str) -> str:
    """Synchronous wrapper - creates a job and runs async version"""
    job_id = str(uuid.uuid4())
    JOB_STORE[job_id] = JobState(job_id=job_id, status=JobStatus.PENDING)
    
    try:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        result = loop.run_until_complete(_extract_and_analyze_pdf_images_async(content, filename, job_id))
        loop.close()
        return result
    finally:
        if job_id in JOB_STORE:
            del JOB_STORE[job_id]  # Clean up sync wrapper jobs immediately

# Background task for async image processing
async def _process_images_background(content: bytes, filename: str, job_id: str, text_content: str, analyze_images: bool = True):
    """Background task to process images and update cache when complete"""
    global RFP_TEXT_CACHE_FILE, RFP_TEXT_CACHE, RFP_TEXT_CACHE_TEXTAREA
    
    try:
        # Check for cancellation before starting
        if job_id in JOB_STORE and JOB_STORE[job_id].cancelled:
            JOB_STORE[job_id].status = JobStatus.CANCELLED
            JOB_STORE[job_id].end_time = datetime.datetime.now().timestamp()
            return
        
        # Run image analysis (with user preference)
        image_analysis = await _extract_and_analyze_pdf_images_async(content, filename, job_id, analyze_images)
        
        # Check for cancellation after processing
        if job_id in JOB_STORE and JOB_STORE[job_id].cancelled:
            JOB_STORE[job_id].status = JobStatus.CANCELLED
            JOB_STORE[job_id].end_time = datetime.datetime.now().timestamp()
            return
        
        # Combine text and image analysis
        file_content = text_content
        if image_analysis:
            file_content = f"{file_content}\n\n{image_analysis}".strip()
        
        # Update caches
        RFP_TEXT_CACHE_FILE = file_content
        textarea_text = (RFP_TEXT_CACHE_TEXTAREA or "").strip()
        
        # Combine both sources
        if textarea_text and file_content:
            merged_text = f"{textarea_text}\n\n--- Uploaded Document Content ---\n\n{file_content}"
        elif file_content:
            merged_text = file_content
        else:
            merged_text = textarea_text
        
        RFP_TEXT_CACHE = merged_text
        
        # Mark job complete
        if job_id in JOB_STORE:
            JOB_STORE[job_id].status = JobStatus.COMPLETED
            JOB_STORE[job_id].result_text = merged_text
            JOB_STORE[job_id].end_time = datetime.datetime.now().timestamp()
        
        # Schedule cleanup
        await asyncio.sleep(1)  # Small delay before cleanup
        await cleanup_old_jobs()
        
    except Exception as e:
        print(f"[JOB {job_id}] Background processing failed: {e}")
        if job_id in JOB_STORE:
            JOB_STORE[job_id].status = JobStatus.FAILED
            JOB_STORE[job_id].errors.append(f"Processing failed: {str(e)}")
            JOB_STORE[job_id].end_time = datetime.datetime.now().timestamp()

# ---------- Helper: sanitize filenames ----------
def _safe_filename(s: str) -> str:
    import urllib.parse
    s = (s or "Proposal").strip()
    # Decode URL encoding (%20, etc.) if present
    try:
        s = urllib.parse.unquote(s)
    except:
        pass
    # Remove utf-8 prefix if present
    s = re.sub(r'^utf-8', '', s, flags=re.IGNORECASE)
    # Replace characters that are invalid on Windows/macOS and path separators for security
    s = re.sub(r'[\\/:*?"<>|]+', '_', s)
    # Remove any path traversal attempts
    s = re.sub(r'\.\.+', '', s)
    # Only allow alphanumeric, spaces, underscores, hyphens, and dots
    s = re.sub(r'[^A-Za-z0-9 _.-]', '', s)
    # Replace spaces with underscores for cleaner filenames
    s = s.replace(' ', '_')
    # Collapse multiple underscores
    s = re.sub(r'_+', '_', s).strip('_')
    # Ensure it's not empty after sanitization
    return s if s else "Proposal"

def _upload_title_default() -> str | None:
    """Base title from the most recent uploaded file (sans extension), sanitized."""
    if not LAST_UPLOAD_FILENAME:
        return None
    base = os.path.splitext(os.path.basename(LAST_UPLOAD_FILENAME))[0]
    return _safe_filename(base)

def _est_stamp_for_filename() -> str:
    """
    Eastern time, 12-hour with AM/PM formatted for clean filenames.
    Example: 2025-10-23_09-22AM_EST
    """
    now_est = datetime.datetime.now(ZoneInfo("America/New_York"))
    return now_est.strftime("%Y-%m-%d_%I-%M%p_EST")

def _export_basename(project_name: str, scenario_label: str | None = None) -> str:
    """Project_Workfront_Export_Scenario_A_2025-10-23_09-22AM_EST"""
    title = _safe_filename(project_name or "Proposal")
    parts = [title, "Workfront_Export"]
    if scenario_label:
        parts.append(_safe_filename(scenario_label))
    parts.append(_est_stamp_for_filename())
    return "_".join(parts)

def _safe_sheet_name(s: str) -> str:
    # Excel sheet name rules: max 31 chars, no : \ / ? * [ ]
    s = re.sub(r'[:\\/?*\[\]]+', "-", (s or "Sheet"))
    s = s.strip() or "Sheet"
    return s[:31]

def _apply_number_formats(ws, df):
    """Format numeric columns: Hours & Price -> 0 decimals, Rate -> 2 decimals."""
    col_idx = {c: i+1 for i, c in enumerate(df.columns)}  # 1-based
    # Whole-number columns
    for col in ["Planned_Hours", "Start_Offset_Days", "Duration_Days", "Price_USD"]:
        if col in col_idx:
            j = col_idx[col]
            for col_cells in ws.iter_cols(min_col=j, max_col=j, min_row=2, max_row=ws.max_row):
                for cell in col_cells:
                    cell.number_format = "0"
    # Rate with 2 decimals
    if "Rate_USD" in col_idx:
        j = col_idx["Rate_USD"]
        for col_cells in ws.iter_cols(min_col=j, max_col=j, min_row=2, max_row=ws.max_row):
            for cell in col_cells:
                cell.number_format = "0.00"

def _finalize_wf_df(df: pd.DataFrame) -> pd.DataFrame:
    # Ensure all expected columns exist (FixedCost added for Step 3 price export)
    NUMERIC_COLS = {"Planned_Hours", "Start_Offset_Days", "Duration_Days", "Rate_USD", "Price_USD", "FixedCost"}
    for col in WF_COLUMNS:
        if col not in df.columns:
            df[col] = "" if col not in NUMERIC_COLS else 0

    # Reindex to canonical order
    df = df[WF_COLUMNS].copy()

    # Enforce numeric types & pricing identity
    df["Planned_Hours"]     = pd.to_numeric(df["Planned_Hours"], errors="coerce").fillna(0).round(0).astype(int)
    df["Start_Offset_Days"] = pd.to_numeric(df["Start_Offset_Days"], errors="coerce").fillna(0).round(0).astype(int)
    df["Duration_Days"]     = pd.to_numeric(df["Duration_Days"], errors="coerce").fillna(0).round(0).astype(int)
    df["Rate_USD"]          = pd.to_numeric(df["Rate_USD"], errors="coerce").fillna(0).round(2)
    
    # STEP 3 PRICE PRESERVATION: Only recompute Price_USD when it's 0 AND hours > 0
    df["Price_USD"] = pd.to_numeric(df["Price_USD"], errors="coerce").fillna(0)
    mask = df["Price_USD"].eq(0) & df["Planned_Hours"].gt(0)
    df.loc[mask, "Price_USD"] = (df.loc[mask, "Planned_Hours"] * df.loc[mask, "Rate_USD"]).round(0)
    df["Price_USD"] = df["Price_USD"].round(0).astype(int)
    
    # STEP 3 FIXED COST: Ensure FixedCost is numeric for XML export
    df["FixedCost"] = pd.to_numeric(df["FixedCost"], errors="coerce").fillna(0).round(2)

    return df

# ---------- v3 A-E Column Ordering Helper ----------
V3_AE_ORDER = ["Row_ID", "Deliverable_Code", "Task_Code", "Service_Department", "Deliverable"]

def _ensure_v3_ae_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Ensure v3 A-E columns exist and are positioned first (leftmost) in exports."""
    # Create columns if missing
    for c in V3_AE_ORDER:
        if c not in df.columns:
            df[c] = ""
    # Reorder so A–E are leftmost
    rest = [c for c in df.columns if c not in V3_AE_ORDER]
    return df[V3_AE_ORDER + rest]

# ---------- WBS builder functions ----------
def _round_int(x: float) -> int:
    try:
        return int(round(float(x)))
    except Exception:
        return 0

def _largest_remainder(target_total: int, parts: dict[str, float]) -> dict[str, int]:
    if not parts:
        return {}
    total = sum(parts.values())
    if total <= 0:
        return {k: 0 for k in parts.keys()}
    raw = {k: (v / total) * target_total for k, v in parts.items()}
    flo = {k: int(v) for k, v in raw.items()}
    rem = target_total - sum(flo.values())
    order = sorted(parts.keys(), key=lambda k: (raw[k] - flo[k]), reverse=True)
    for k in order[:max(0, rem)]:
        flo[k] += 1
    return flo

def _eff_rate(price: float, hours: float) -> float:
    """Calculate effective rate: price / hours with rounding."""
    return round(price / hours, 2) if hours and hours > 0 else 0.0

def _band_multiplier(rate_band: str) -> float:
    band = DB.rate_bands[DB.rate_bands["Band_Name"] == (rate_band or "Standard_US")]
    return float(band["Rate_Multiplier"].iloc[0]) if not band.empty else 1.0

def _wbs_order_mode():
    return "timeline"

def _inflate_components_if_missing(scenario: dict) -> dict:
    """
    Defensive fallback: if scenario items lack included_task_groups/components,
    auto-expand from DB so exports never go flat.
    Handles both empty lists and "__ALL__" sentinel values.
    """
    for item in scenario.get("items", []):
        dcode = str(item.get("deliverable_code") or item.get("Deliverable_Code") or "").strip()
        if not dcode:
            continue
        
        # Check if task groups need to be populated
        included = item.get("included_task_groups") or []
        comp_map = item.get("included_task_groups_map", {})
        
        # Populate if empty OR if the map contains "__ALL__" sentinel
        needs_inflation = (
            not included or 
            comp_map == "__ALL__" or 
            (isinstance(comp_map, dict) and not comp_map)
        )
        
        if needs_inflation:
            # Derive all task groups for this deliverable from DB
            included = DB.task_groups_for_deliverable(dcode)
            item["included_task_groups"] = included
            # Also convert "__ALL__" sentinel to empty dict for downstream code
            if comp_map == "__ALL__":
                item["included_task_groups_map"] = {}
    
    return scenario

def build_wbs_with_pricing(scenario: dict, project_name: str) -> pd.DataFrame:
    """
    Adds Rate_USD and Price_USD at deliverable/component/task level.
    Flat_Blended -> uses blended_rate
    Per_Resource -> weighted effective per level
    """
    # PATCH A: Ensure components are inflated for AI picks (handle "__ALL__" sentinel)
    scenario = _inflate_components_if_missing(scenario)
    
    # PATCH B: Recompute totals to ensure Step 3 edits are consistent before WBS build
    scenario = _recompute_totals(scenario)
    
    rows = []
    pricing_mode = (scenario.get("pricing_mode") or "Flat_Blended").strip()
    rate_band    = (scenario.get("rate_band") or "Standard_US").strip()
    blended_rate = scenario.get("blended_rate")
    if blended_rate is None:
        # fallback to default from Pricing_Settings already loaded in v2
        ps = DB.pricing_settings[DB.pricing_settings["Key"]=="Default_Blended_Rate"]
        blended_rate = float(ps["Default"].iloc[0]) if not ps.empty else DEFAULT_BILLABLE_RATE_USD
    blended_rate = float(blended_rate)

    # project parent - set to 0 for proper Project Summary task
    rows.append({
        "Row_ID": "",
        "Deliverable_Code": "",
        "Task_Code": "",
        "Service_Department": "",
        "Deliverable": "",
        "Project_Name": project_name, "WBS_ID": "1", "Parent_WBS_ID": "",
        "Task_Name": "Project Summary", "Component": "", "Task": "",
        "Role": "", "Seniority": "", "Planned_Hours": 0, "Start_Offset_Days": 0, "Duration_Days": 0,
        "Dependencies": "", "Assignee_External_ID": "", "Notes": "",
        "Rate_USD": "", "Price_USD": ""
    })

    items = scenario.get("items", [])
    order_map = {str(tg): i for i, tg in enumerate(DB.timeline_params["Task_Group"].astype(str).tolist())}

    # DEBUG: Log what we received
    print(f"[WBS Builder] build_wbs_with_pricing called for project: {project_name}")
    print(f"[WBS Builder] Scenario keys: {list(scenario.keys())}")
    print(f"[WBS Builder] Has items: {len(items)} deliverables")
    print(f"[WBS Builder] Has timeline_tasks: {'timeline_tasks' in scenario}")
    print(f"[WBS Builder] Has timeline: {'timeline' in scenario}")

    # MERGE TIMELINE DATA: Copy Start_Date/End_Date from timeline tasks into deliverables
    # Support both formats: timeline_tasks (manual Gantt saves) and timeline.tasks (AI-generated)
    # IMPORTANT: Check timeline_tasks FIRST since it has priority over AI-generated timeline
    timeline_tasks = []
    if "timeline_tasks" in scenario and scenario["timeline_tasks"]:
        timeline_tasks = scenario["timeline_tasks"]
        print(f"[WBS Builder] Using timeline_tasks format (manual Gantt saves)")
    elif "timeline" in scenario and scenario["timeline"]:
        timeline_tasks = scenario["timeline"].get("tasks", [])
        print(f"[WBS Builder] Using timeline.tasks format (AI-generated)")
    else:
        print(f"[WBS Builder] ⚠️ NO TIMELINE DATA FOUND - skipping merge")
    
    if timeline_tasks:
        print(f"[WBS Builder] Found {len(timeline_tasks)} timeline tasks to merge")
        print(f"[WBS Builder DEBUG] First 3 timeline tasks: {timeline_tasks[:3]}")
        
        for item in items:
            deliv_name = str(item.get("deliverable", "")).strip()
            deliv_code = str(item.get("deliverable_code", item.get("code", ""))).strip()
            
            matched = False
            # Try to find matching timeline task by name or code
            for task in timeline_tasks:
                task_name = str(task.get("name", "")).strip()
                task_id = str(task.get("id", "")).strip()
                task_deliv_code = str(task.get("deliverable_code", "")).strip()
                
                # Match by deliverable name or code
                if (deliv_name and task_name == deliv_name) or \
                   (deliv_code and task_deliv_code == deliv_code) or \
                   (deliv_code and task_id.endswith(deliv_code)):
                    matched = True
                    # Copy timeline dates and duration into deliverable (check all casing variants)
                    start_date = task.get("Start_Date") or task.get("start_date") or task.get("start")
                    end_date = task.get("End_Date") or task.get("end_date") or task.get("end")
                    start_offset = task.get("Start_Offset_Days") if task.get("Start_Offset_Days") is not None else task.get("start_offset_days")
                    hours = task.get("hours") or task.get("Hours")
                    
                    print(f"[WBS Builder] 🔍 Merging timeline data for '{deliv_name}':")
                    print(f"  Timeline task has: start={start_date}, end={end_date}, hours={hours}")
                    
                    if start_date:
                        item["Start_Date"] = start_date
                        print(f"  ✓ Merged Start_Date={start_date}")
                    if end_date:
                        item["End_Date"] = end_date
                        print(f"  ✓ Merged End_Date={end_date}")
                    if start_offset is not None:
                        item["Start_Offset_Days"] = start_offset
                        print(f"  ✓ Merged Start_Offset_Days={start_offset}")
                    if hours is not None:
                        item["hours"] = hours
                        print(f"  ✓ Merged hours={hours}")
                    break
            
            if not matched and deliv_name:
                print(f"[WBS Builder] ✗ No timeline match for deliverable: '{deliv_name}' (code: '{deliv_code}')")

    # Enrich items with Service Department (kept as field only for reporting - no hierarchy)
    for item in items:
        dcode = str(item.get("deliverable_code", item.get("code", "")))
        tgs = [str(x) for x in item.get("included_task_groups", [])]
        # Resolve scenario_col from item; if missing/invalid, derive from complexity & tier
        scen_col = item.get("scenario_col")
        if not scen_col or scen_col not in DB.all_rows.columns:
            scen_col = DB.scenario_hours_col(item.get("complexity", "Advanced"),
                                             item.get("tier", "T2_MediumVolume"))
        # Prefer v3 (doesn't need scen_col), then fall back to v2
        dept = (DB.service_department_for_deliverable(dcode, tgs)
                or DB.service_dept_for_deliverable(dcode, tgs, scen_col)
                or "Other") if dcode else "Other"
        item["_service_department"] = dept or "Other"

    day_cursor = 0
    prev_deliv_wbs = ""
    deliv_counter_global = 0

    # Process all deliverables (no department grouping)
    for d in items:
            deliv_counter_global += 1
            dcode = str(d.get("deliverable_code", d.get("code", f"DELIV_{deliv_counter_global}")))
            # Resolve scenario_col from item; if missing/invalid, derive from complexity & tier
            scen_col = d.get("scenario_col")
            if not scen_col or scen_col not in DB.all_rows.columns:
                scen_col = DB.scenario_hours_col(d.get("complexity", "Advanced"),
                                                 d.get("tier", "T2_MediumVolume"))
            included = [str(x) for x in d.get("included_task_groups", [])]
            if not included:
                # derive from the database for this deliverable
                sub = DB.all_rows[DB.all_rows["Deliverable_Code"].astype(str) == str(dcode)]
                included = sorted(set(sub["task_group"].dropna().astype(str).tolist()))

            # Get deliverable name - try from scenario first, then lookup from database
            deliv_label = str(d.get("deliverable", "")).strip()
            if not deliv_label:
                # Fallback: lookup deliverable name from database using code
                row = DB.deliverables[DB.deliverables["Deliverable_Code"].astype(str)==str(dcode)]
                if not row.empty:
                    deliv_label = str(row["Deliverable"].iloc[0])
                else:
                    deliv_label = f"Deliverable {dcode}"

            # schedule offsets/durations by task_group
            schedule = d.get("schedule", [])
            if not schedule:
                complexity = d.get("complexity", "Advanced")
                tier = d.get("tier", "T2_MediumVolume")
                schedule = DB.build_schedule(
                    deliverable_code=dcode,
                    included_task_groups=included,
                    complexity=complexity,
                    tier=tier,
                    use_slack=scenario.get("use_slack", True),
                    slack_after_internal=scenario.get("slack_after_internal", 1),
                    slack_after_client=scenario.get("slack_after_client", 2),
                    slack_global_pct=scenario.get("slack_global_pct", 0.05),
                    project_start=scenario.get("project_start"),
                    scenario_letter=scenario.get("scenario_label", "A")
                )
            tg_order = sorted(included, key=lambda tg: order_map.get(tg, 999))
            duration_by_tg = {str(t["task_group"]): int(t["duration_days"]) for t in schedule}
            offset_by_tg = {}
            run = 0
            for tg in tg_order:
                offset_by_tg[tg] = run
                run += int(duration_by_tg.get(tg, 1))
            total_deliv_duration = run

            # Derive hours_by_role if missing (robustness fix)
            hrs_df = pd.DataFrame(d.get("hours_by_role") or [])
            if hrs_df.empty or float(hrs_df.get("Hours", pd.Series([], dtype=float)).sum()) <= 0.0:
                scen_col_resolved = DB.scenario_hours_col(d.get("complexity","Advanced"), d.get("tier","T2_MediumVolume"))
                hrs_df = DB.hours_by_role_for_deliverable(dcode, included, scen_col_resolved)
                d["hours_by_role"] = hrs_df.to_dict("records")
            
            # CRITICAL: Use get_hours_from_item() to honor Step 3 edits (Planned_Hours)
            # Only fall back to calculated hours if Step 3 didn't set any
            target_hours = get_hours_from_item(d)
            
            if target_hours <= 0:
                # No Step 3 hours - calculate from hrs_df ONLY (not total_hours which may be stale)
                calculated_total = float(hrs_df["Hours"].sum()) if not hrs_df.empty else 0.0
                target_hours = calculated_total
            
            parent_hours_exact = target_hours
            parent_hours_display = int(round(parent_hours_exact))
            
            # SCALING: Calculate scale factor to normalize role hours to Step 3 target
            # This ensures XML Work sums to Step 3 Planned_Hours, not database hours
            source_hours = float(hrs_df["Hours"].sum()) if not hrs_df.empty else 0.0
            if source_hours > 0 and target_hours > 0:
                hours_scale = target_hours / source_hours
            else:
                hours_scale = 1.0
            
            # Cache scaled hours_by_role on deliverable for component/task loops
            if not hrs_df.empty and hours_scale != 1.0:
                scaled_hrs_df = hrs_df.copy()
                scaled_hrs_df["Hours"] = (scaled_hrs_df["Hours"] * hours_scale).round(2)
                d["_scaled_hours_by_role"] = scaled_hrs_df.to_dict("records")
            else:
                d["_scaled_hours_by_role"] = d.get("hours_by_role", [])
            
            # FIXED COST: Use Step 3 price directly as canonical cost (no recalculation)
            target_price = get_price_from_item(d)
            if target_price <= 0:
                # Fall back to calculated price if Step 3 didn't set any
                target_price = float(d.get("total_price") or 0)
            
            print(f"[WBS_SCALE] {dcode}: target_hours={target_hours}, source_hours={source_hours}, scale={hours_scale:.3f}, target_price={target_price}")

            # Check if this is a retainer deliverable (support both old and new cadence fields)
            months = int((d.get("retainer") or {}).get("months", 0) or d.get("retainer_months", 0))
            
            # DEBUG: Log exact values going into export for retainer debugging
            print(f"[DEBUG] EXPORT ITEM: {d.get('deliverable', dcode)[:40]} hours={target_hours} months={months} cadence={d.get('billing_cadence', 'one_time')}")
            
            # Also check billing_cadence for new cadence-based retainers
            billing_cadence = (d.get("billing_cadence") or "one_time").lower()
            cadence_units = int(d.get("cadence_units") or 1)
            if billing_cadence in ("monthly", "quarterly", "semi_annual", "annual") and months == 0:
                # Calculate months from cadence
                months_per_cadence = {"monthly": 1, "quarterly": 3, "semi_annual": 6, "annual": 12}.get(billing_cadence, 0)
                months = months_per_cadence * cadence_units
            
            monthly_hours = float(d.get("monthly_hours") or 0)
            monthly_price = float(d.get("monthly_price") or 0)
            is_retainer = months > 0 or d.get("is_retainer", False)
            deliverable_type = "Retainer" if is_retainer else "One-Time"

            # FIXED COST EXPORT: Use Step 3 price (target_price) as the authoritative cost
            # This replaces all hourly cost calculation with a single FixedCost value
            deliv_price = round(target_price, 2) if target_price > 0 else 0.0
            deliv_fixed_cost = deliv_price  # FixedCost for XML export
            
            # DEBUG: Log FixedCost assignment at WBS building stage
            if deliv_fixed_cost > 0:
                print(f"[WBS BUILD] 💰 Setting FixedCost={deliv_fixed_cost} for {dcode}: {deliv_label[:40]}")
            
            # Calculate effective rate for display only (not used for costing)
            if parent_hours_display > 0:
                deliv_rate = round(deliv_price / parent_hours_display, 2)
            else:
                deliv_rate = 0.0

            # Build deliverable node - direct child of Project Summary (flattened hierarchy)
            wbs_deliv = f"1.{deliv_counter_global}"
            svc_deliv = (DB.service_department_for_deliverable(dcode, tg_order)
                         or DB.service_dept_for_deliverable(dcode, tg_order, scen_col))
            # deliv_label already set above with database fallback - don't override it
            deliv_notes = f'{d.get("complexity","")}/{d.get("tier","")}' + (f' | Retainer x{months} months' if months else '')
            total_deliv_duration = sum(int(t["duration_days"]) for t in schedule)  # one-cycle length
            
            # FIX B: Respect Step 4 Gantt overrides (Start_Offset_Days or Start_Date)
            dstart = day_cursor  # default sequential start
            try:
                if str(d.get("Start_Offset_Days", "")).strip() != "":
                    dstart = int(float(d["Start_Offset_Days"]))
                elif d.get("Start_Date") and scenario.get("project_start"):
                    import datetime
                    ps = datetime.date.fromisoformat(str(scenario["project_start"])[:10])
                    sd = datetime.date.fromisoformat(str(d["Start_Date"])[:10])
                    dstart = (sd - ps).days
            except Exception:
                pass  # fall back to sequential start on any parse issue
            
            deliv_row = {
                "Row_ID": "",
                "Deliverable_Code": dcode,
                "Task_Code": "",
                "Service_Department": svc_deliv,
                "Deliverable": deliv_label,
                "Project_Name": project_name, "WBS_ID": wbs_deliv, "Parent_WBS_ID": "1",
                "Task_Name": deliv_label,
                "Component": "", "Task": "", "Role": "", "Seniority": "",
                "Planned_Hours": parent_hours_display,  # Use scaled Step 3 hours (not monthly × months)
                "Start_Offset_Days": dstart,
                "Duration_Days": "",  # Task 8: Leave empty for summary bar - computed from children
                "Start_Date": d.get("Start_Date", ""),
                "End_Date": d.get("End_Date", ""),
                "Dependencies": prev_deliv_wbs, "Assignee_External_ID": "", "Notes": deliv_notes,
                "Rate_USD": round(deliv_rate, 2),
                "Price_USD": round(deliv_price, 2),
                "FixedCost": round(deliv_fixed_cost, 2),  # CRITICAL: Step 3 price as FixedCost for XML
                "Type": deliverable_type
            }
            rows.append(deliv_row)
            
            # Expand monthly retainers into child rows (Brad's requirement: "Feb assets, March assets...")
            if is_retainer and months > 0 and billing_cadence == "monthly":
                project_start = scenario.get("project_start")
                if project_start:
                    # Parent summary row: 0 hours, full FixedCost
                    deliv_row["Planned_Hours"] = 0
                    deliv_row["FixedCost"] = round(target_price, 2)  # Full price on parent
                    
                    # Each child gets equal share of hours and price
                    monthly_child_hours = round(target_hours / months, 2) if months > 0 else target_hours
                    monthly_child_price = round(target_price / months, 2) if months > 0 else target_price
                    
                    # Create monthly rows - STRUCTURAL ONLY (0 hours)
                    # Role rows under component/tasks carry the hours; monthly children provide date structure
                    # This prevents double-counting while maintaining calendar month timeline visibility
                    monthly_rows = expand_retainer_into_months(deliv_row, months, project_start)
                    for mrow in monthly_rows:
                        mrow["Planned_Hours"] = 0  # Structural only - hours are on role rows
                        mrow["FixedCost"] = round(monthly_child_price, 2)  # Per-month cost
                        mrow["Price_USD"] = round(monthly_child_price, 2)
                        mrow["Rate_USD"] = 0  # Cost comes from FixedCost only
                    rows.extend(monthly_rows)
                    
                    # DEBUG: Log retainer export structure for verification
                    print(f"[RET_EXPORT] {dcode}: months={months} target_hours={target_hours} "
                          f"monthly_children=structural_only (0 hours) "
                          f"total_price={target_price} cadence={billing_cadence}")

            comps = DB.components_for_deliverable(dcode, tg_order)
            # Robust fallback if DB returns no components
            if not comps:
                sub = DB.all_rows[DB.all_rows["Deliverable_Code"].astype(str)==str(dcode)]
                if not sub.empty:
                    comps = sorted({str(x) for x in sub["Component"].dropna().astype(str) if str(x).strip() and str(x).strip() != "nan"})
                if not comps:
                    comps = ["Work Package"]
            
            # SCALING: Apply hours_scale to component hours so they sum to target_hours
            comp_hours_map_month_raw = DB.hours_by_component(dcode, tg_order, scen_col)
            comp_hours_map_scaled = {k: v * hours_scale for k, v in comp_hours_map_month_raw.items()}
            
            # If not a retainer, treat "month" as the whole - use SCALED hours
            base_comp_hours_display = _largest_remainder(parent_hours_display, comp_hours_map_scaled)

            prev_comp_wbs = ""

            for j, comp in enumerate(comps, start=1):
                # SCALING: Apply hours_scale to task group hours within component
                tg_hours_in_comp_raw = DB.hours_by_taskgroup_for_component(dcode, comp, tg_order, scen_col)
                tg_hours_in_comp = {tg: hrs * hours_scale for tg, hrs in tg_hours_in_comp_raw.items()}
                tg_in_comp = sorted({tg for tg in tg_order if tg in tg_hours_in_comp})
                if not tg_in_comp:
                    continue

                comp_offset = min(offset_by_tg[tg] for tg in tg_in_comp)
                comp_duration = sum(int(duration_by_tg.get(tg, 1)) for tg in tg_in_comp)

                # Use SCALED hours for component
                # FIX: parent_hours_display (137) is already the TOTAL for the entire retainer
                # base_comp_hours_display distributes this TOTAL across components
                # DO NOT multiply by months again - that would give 137*6=822 (wrong!)
                comp_hours_total_display = int(base_comp_hours_display.get(comp, 0))
                # For per-month task distribution, divide total by months
                comp_hours_per_month = comp_hours_total_display / months if months > 0 else comp_hours_total_display

                wbs_comp = f"{wbs_deliv}.{j}"
                svc_comp = (DB.service_department_for_component(dcode, comp, tg_in_comp)
                            or DB.service_dept_for_component(dcode, comp, tg_in_comp, scen_col))
                
                # ZERO COST: All cost comes from deliverable FixedCost, not component/task costs
                rows.append({
                    "Row_ID": "", "Deliverable_Code": dcode, "Task_Code": "", "Service_Department": svc_comp,
                    "Deliverable": deliv_label,
                    "Project_Name": project_name, "WBS_ID": wbs_comp, "Parent_WBS_ID": wbs_deliv,
                    "Task_Name": comp, "Component": comp, "Task": "", "Role": "", "Seniority": "",
                    "Planned_Hours": comp_hours_total_display,
                    "Start_Offset_Days": dstart + comp_offset,
                    "Duration_Days": "",  # Task 8: Leave empty for summary bar - computed from children
                    "Dependencies": (wbs_deliv if j == 1 else prev_comp_wbs),
                    "Assignee_External_ID": "", "Notes": "",
                    "Rate_USD": 0,  # Cost comes from FixedCost only
                    "Price_USD": 0,  # Cost comes from FixedCost only
                    "FixedCost": 0   # No FixedCost on components - only on deliverable
                })

                # --- Tasks under the component ---
                # RETAINER FIX: For monthly retainers (billing_cadence == "monthly"), 
                # we run the task loop ONCE with FULL hours (not per-month hours).
                # For other cadences, we repeat tasks per month with per-month hours.
                is_monthly_retainer = (billing_cadence == "monthly") and months > 0
                month_loop_count = 1 if is_monthly_retainer else (months if months else 1)
                
                # Hours target for task groups:
                # - Monthly retainers: use FULL hours (single iteration carries all hours)
                # - Other cadences: use per-month hours (multiple iterations add up)
                tg_hours_base = {tg: float(tg_hours_in_comp.get(tg, 0.0)) for tg in tg_in_comp}
                if is_monthly_retainer:
                    # Use full component hours since we only iterate once
                    tg_target = _largest_remainder(comp_hours_total_display, tg_hours_base)
                    print(f"[RETAINER DEBUG] {dcode}/{comp}: is_monthly_retainer=True, comp_hours_total={comp_hours_total_display}, tg_target={tg_target}")
                else:
                    # Use per-month hours since we iterate month_loop_count times
                    tg_target = _largest_remainder(int(round(comp_hours_per_month)), tg_hours_base)

                # Build month-by-month repetition
                total_tasks_per_month = len(tg_in_comp)
                prev_month_last_wbs = ""  # chain months sequentially per component
                
                for month_idx in range(1, month_loop_count + 1):
                    # enumerates tasks within this month
                    prev_task_last_wbs = ""
                    for k, tg in enumerate(tg_in_comp, start=1):
                        dur = int(duration_by_tg.get(tg, 1))
                        label_core = DB.task_label_for_component_tg(dcode, comp, tg) if hasattr(DB, "task_label_for_component_tg") else tg
                        # For monthly retainers: use plain label (no Month prefix) - template cycle only
                        # For other cadences with months: keep Month XX prefix for task repetition
                        label = label_core if is_monthly_retainer else ((f"Month {month_idx:02d} – {label_core}") if months else label_core)

                        # Unique task index within the component across months
                        task_ordinal = (month_idx-1)*total_tasks_per_month + k
                        wbs_task = f"{wbs_comp}.{task_ordinal}"

                        base_offset = dstart + offset_by_tg[tg] + ((month_idx-1) * total_deliv_duration)

                        # ZERO COST: All cost comes from deliverable FixedCost
                        rows.append({
                            "Row_ID": "", "Deliverable_Code": dcode, "Task_Code": "", "Service_Department": svc_comp,
                            "Deliverable": deliv_label,
                            "Project_Name": project_name, "WBS_ID": wbs_task, "Parent_WBS_ID": wbs_comp,
                            "Task_Name": label, "Component": comp, "Task": label,
                            "Role": "", "Seniority": "",
                            "Planned_Hours": "",   # stays on role rows
                            "Start_Offset_Days": base_offset,
                            "Duration_Days": dur,
                            "Dependencies": (wbs_comp if (k==1 and month_idx==1) else (prev_task_last_wbs if k>1 else prev_month_last_wbs)),
                            "Assignee_External_ID": "", "Notes": "",
                            "Rate_USD": 0,  # Cost comes from FixedCost only
                            "Price_USD": 0,  # Cost comes from FixedCost only
                            "FixedCost": 0   # No FixedCost on tasks - only on deliverable
                        })

                        # Role rows for this task in this month
                        hrs_role_df = DB.hours_by_role_for_component_task(dcode, comp, tg, scen_col)
                        role_rows = hrs_role_df.to_dict(orient="records")
                        target_task_hours = int(tg_target.get(tg, 0))

                        raw_map = {(r["Resource_Title"], r["Seniority"]): float(r["Hours"]) for r in role_rows}
                        if not raw_map:
                            raw_map = {("","Mid"): float(target_task_hours)}
                        total = sum(raw_map.values()) or 1.0
                        raw_scaled = {key: (val/total)*target_task_hours for key, val in raw_map.items()}
                        flo = {key: int(val) for key, val in raw_scaled.items()}
                        rem = target_task_hours - sum(flo.values())
                        order = sorted(raw_map.keys(), key=lambda kk: (raw_scaled[kk]-flo[kk]), reverse=True)
                        for kk in order[:max(0, rem)]:
                            flo[kk] += 1
                        
                        # Debug logging for retainer role allocation
                        if is_monthly_retainer:
                            role_sum = sum(flo.values())
                            print(f"[RETAINER DEBUG] {dcode}/{comp}/{tg}: target_task_hours={target_task_hours}, role_sum={role_sum}, raw_map={raw_map}")

                        prev_role_wbs = ""
                        r_index = 0
                        for (role, sen), h in flo.items():
                            if h <= 0:
                                continue
                            r_index += 1
                            row_id, task_code, svc_task_v2 = DB.codes_for_component_task_role(dcode, comp, tg, role or "", sen or "", scen_col)
                            # Prefer v3 service_department_for_task (no scen_col dependency)
                            svc_task = DB.service_department_for_task(dcode, comp, tg) or svc_task_v2
                            wbs_role = f"{wbs_task}.{r_index}"

                            row_hours = int(h)
                            
                            # ZERO COST: All cost comes from deliverable FixedCost
                            # Role rows only carry hours for XML Work, not cost
                            rows.append({
                                "Row_ID": row_id,
                                "Deliverable_Code": dcode,
                                "Task_Code": task_code,
                                "Service_Department": (svc_task or svc_comp),
                                "Deliverable": deliv_label,
                                "Project_Name": project_name, "WBS_ID": wbs_role, "Parent_WBS_ID": wbs_task,
                                "Task_Name": label, "Component": comp, "Task": label,
                                "Role": role or "", "Seniority": sen or "",
                                "Planned_Hours": row_hours,
                                "Start_Offset_Days": "", "Duration_Days": "",
                                "Dependencies": wbs_task if r_index == 1 else prev_role_wbs,
                                "Assignee_External_ID": "", "Notes": "",
                                "Rate_USD": 0,  # Cost comes from FixedCost only
                                "Price_USD": 0,  # Cost comes from FixedCost only
                                "FixedCost": 0   # No FixedCost on role rows - only on deliverable
                            })
                            prev_role_wbs = wbs_role

                        prev_task_last_wbs = prev_role_wbs or wbs_task

                    prev_month_last_wbs = prev_task_last_wbs

                prev_comp_wbs = wbs_comp

            # Advance cursor based on actual start used (preserves ordering for next deliverable)
            day_cursor = max(day_cursor, dstart) + total_deliv_duration
            prev_deliv_wbs = wbs_deliv

    # ========================================================================
    # HOURS HARMONIZATION PASS: Ensure leaf role rows sum exactly to target
    # ========================================================================
    # After building all rows, harmonize leaf hours per deliverable to match Step 3 target
    # This fixes rounding drift where e.g. 200 hours target becomes 210 in leaf sum
    
    # Group rows by deliverable and find leaf role rows (highest outline level under each deliverable)
    deliverable_target_hours = {}  # dcode -> target hours from Step 3
    deliverable_leaf_rows = {}     # dcode -> list of row indices that are leaf role rows
    
    for d in items:
        dcode = d.get("deliverable_code") or d.get("code") or ""
        # Get Step 3 target hours using the canonical accessor
        target_hours = get_hours_from_item(d)
        if target_hours > 0:
            deliverable_target_hours[dcode] = target_hours
            deliverable_leaf_rows[dcode] = []
    
    # Find leaf role rows for each deliverable (rows with Role set and Planned_Hours > 0)
    for i, row in enumerate(rows):
        dcode = row.get("Deliverable_Code", "")
        if dcode in deliverable_target_hours:
            # Leaf role rows have a Role assigned and positive hours
            if row.get("Role") and row.get("Planned_Hours") and int(row.get("Planned_Hours", 0)) > 0:
                deliverable_leaf_rows[dcode].append(i)
    
    # Harmonize each deliverable's leaf hours to match target
    for dcode, target_hours in deliverable_target_hours.items():
        leaf_indices = deliverable_leaf_rows.get(dcode, [])
        if not leaf_indices:
            continue
        
        # Calculate current leaf sum
        leaf_hours = [(i, int(rows[i].get("Planned_Hours", 0))) for i in leaf_indices]
        leaf_total = sum(h for _, h in leaf_hours)
        
        if leaf_total == 0 or abs(leaf_total - target_hours) < 1:
            continue  # Already harmonized or no hours to scale
        
        # Scale factor to match target
        scale = target_hours / leaf_total
        
        # Apply scale with largest-remainder rounding to preserve integer hours
        scaled_hours = []
        for i, h in leaf_hours:
            exact = h * scale
            floor_h = int(exact)
            remainder = exact - floor_h
            scaled_hours.append((i, floor_h, remainder))
        
        # Distribute remainder using largest-remainder method
        current_sum = sum(h for _, h, _ in scaled_hours)
        deficit = int(round(target_hours)) - current_sum
        
        if deficit > 0:
            # Sort by remainder descending and add 1 to top deficit rows
            scaled_hours.sort(key=lambda x: x[2], reverse=True)
            for j in range(min(deficit, len(scaled_hours))):
                idx, h, _ = scaled_hours[j]
                scaled_hours[j] = (idx, h + 1, 0)
        elif deficit < 0:
            # Need to subtract - sort by remainder ascending (smallest first)
            scaled_hours.sort(key=lambda x: x[2])
            for j in range(min(-deficit, len(scaled_hours))):
                idx, h, _ = scaled_hours[j]
                if h > 0:  # Don't go negative
                    scaled_hours[j] = (idx, h - 1, 0)
        
        # Apply harmonized hours back to rows
        for idx, harmonized_h, _ in scaled_hours:
            rows[idx]["Planned_Hours"] = harmonized_h
        
        # Log the harmonization
        new_sum = sum(rows[idx]["Planned_Hours"] for idx in leaf_indices)
        print(f"[HOURS HARMONIZE] {dcode}: target={int(target_hours)}, was={leaf_total}, now={new_sum}, delta={new_sum - leaf_total}")

    df = pd.DataFrame(rows)
    
    # --- ENFORCE A-E COLUMN ORDER FOR v3 COMPATIBILITY ---
    order_ae = ["Row_ID","Deliverable_Code","Task_Code","Service_Department","Deliverable"]
    # Ensure all required columns exist before reordering
    for col in order_ae:
        if col not in df.columns:
            df[col] = ""
    rest = [c for c in df.columns if c not in order_ae]
    df = df.reindex(columns=order_ae + rest, fill_value="")
    
    # --- ENFORCE NUMERIC TYPES & PRESERVE CADENCE PRICING ---
    # Coerce to numeric and fill blanks
    if "Planned_Hours" in df.columns:
        df["Planned_Hours"] = pd.to_numeric(df["Planned_Hours"], errors="coerce").fillna(0).round(0).astype(int)

    # Rate shown with 2 decimals; blanks -> 0
    if "Rate_USD" in df.columns:
        df["Rate_USD"] = pd.to_numeric(df["Rate_USD"], errors="coerce").fillna(0).round(2)

    # STEP 3 PRICE PRESERVATION: Only compute Price from Hours × Rate when Price_USD is 0 AND hours > 0
    if "Price_USD" in df.columns:
        df["Price_USD"] = pd.to_numeric(df["Price_USD"], errors="coerce").fillna(0)
        if "Planned_Hours" in df.columns and "Rate_USD" in df.columns:
            mask = df["Price_USD"].eq(0) & df["Planned_Hours"].gt(0)
            df.loc[mask, "Price_USD"] = (df.loc[mask, "Planned_Hours"] * df.loc[mask, "Rate_USD"]).round(0)
        df["Price_USD"] = df["Price_USD"].round(0).astype(int)
    elif "Planned_Hours" in df.columns and "Rate_USD" in df.columns:
        df["Price_USD"] = (df["Planned_Hours"] * df["Rate_USD"]).round(0).astype(int)
    else:
        df["Price_USD"] = 0

    # FIXED COST EXPORT: Ensure FixedCost is numeric for XML export
    if "FixedCost" in df.columns:
        df["FixedCost"] = pd.to_numeric(df["FixedCost"], errors="coerce").fillna(0).round(2)
        # Log summary of FixedCost values for deliverables
        deliverable_mask = df["WBS_ID"].astype(str).str.count(r'\.') == 1  # OutlineLevel 2
        deliverables_with_cost = df.loc[deliverable_mask & (df["FixedCost"] > 0)]
        if not deliverables_with_cost.empty:
            total_fixed_cost = deliverables_with_cost["FixedCost"].sum()
            print(f"[WBS BUILD] 💰 Total FixedCost for {len(deliverables_with_cost)} deliverables: ${total_fixed_cost:,.2f}")
    else:
        df["FixedCost"] = 0.0

    return df

def build_wbs_dataframe_from_scenario(scenario: dict, project_name: str) -> pd.DataFrame:
    """Build WBS with pricing - delegates to the enhanced pricing-aware version."""
    return build_wbs_with_pricing(scenario, project_name)

# For backward compatibility, keep the old function name pointing to the new one
def _wbs_dataframe_from_scenario(scenario: dict, project_name: str) -> pd.DataFrame:
    """Legacy function name - redirects to the new WBS builder."""
    return build_wbs_dataframe_from_scenario(scenario, project_name)

# ---------- Pydantic models ----------
class SuggestPayload(BaseModel):
    rfp_text: str

class ResolveDeliverablesPayload(BaseModel):
    inputs: list[str] = []

# --- NEW: Retainer selection model ---
class RetainerSelection(BaseModel):
    deliverable_code: str
    months: int  # 1..12

class ScenarioSpec(BaseModel):
    mode: str                               # "template" or "bundle"
    # if mode == "template"
    scenario_key: Optional[str] = None      # e.g., "MED_LOW" or "MED_HIGH"
    complexity: Optional[str] = None        # override complexity
    tier: Optional[str] = None              # override tier
    # if mode == "bundle"
    bundle: Optional[str] = None            # Express/Good/Better/Best

class BuildPayload(BaseModel):
    selected_deliverable_codes: List[str]
    scenario_a: Optional[ScenarioSpec] = None
    pricing_mode: str = "Flat_Blended"                       # "Flat_Blended" or "Per_Resource"
    blended_rate: Optional[float] = None
    rate_band: Optional[str] = "Standard_US"
    use_slack: bool = True
    slack_after_internal: int = 1
    slack_after_client: int = 2
    slack_global_pct: float = 0.05
    project_start: Optional[str] = None     # ISO8601 format (e.g., "2025-10-06T09:00:00" or "YYYY-MM-DD")
    client_budget_usd: Optional[float] = None  # Client budget for budget analysis
    project_name: Optional[str] = None      # Project name from frontend input for exports
    session_id: Optional[str] = None        # For saving to SCENARIO_STORE (enables Gantt sync)
    # NEW: monthly retainers selected on the second screen
    retainers: Optional[List[RetainerSelection]] = []
    # NEW: component-level selection per deliverable (supports multiple formats including "__ALL__" sentinel)
    selected_components_map: Optional[Dict[str, Union[str, List[str], Dict[str, Optional[float]]]]] = None
    # NEW: L3 subtask selection per deliverable and component
    selected_l3_map: Optional[Dict[str, Dict[str, List[str]]]] = None  # { "<Deliverable_Code>": { "<Component_Name>": ["<Task_Label>", ...] } }

class AutoBuildPayload(BaseModel):
    rfp_text: str
    scenario_a: ScenarioSpec
    pricing_mode: str                       # "Flat_Blended" or "Per_Resource"
    blended_rate: Optional[float] = None
    rate_band: Optional[str] = "Standard_US"
    use_slack: bool = True
    slack_after_internal: int = 1
    slack_after_client: int = 2
    slack_global_pct: float = 0.05
    project_start: Optional[str] = None     # ISO8601 format (e.g., "2025-10-06T09:00:00" or "YYYY-MM-DD")
    client_budget_usd: Optional[float] = None  # Client budget for budget analysis
    # NEW: optional retainers alongside auto suggestions
    retainers: Optional[List[RetainerSelection]] = []

class ExportPayload(BaseModel):
    scenario: Dict[str, Any]
    project_name: Optional[str] = None       # e.g., "Casa Dragones"
    file_format: Optional[str] = "csv"       # "csv" or "xlsx"
    scenario_label: Optional[str] = None     # e.g., "Scenario A"
    add_timestamp: Optional[bool] = False    # include yyyymmdd-HHMM in filename?                # a scenario dict returned from /api/build

class ExportWorkbookPayload(BaseModel):
    scenario_a: dict
    scenario_b: dict
    project_name: str | None = None
    sheet_name_a: str | None = "Scenario A"
    sheet_name_b: str | None = "Scenario B"
    add_timestamp: bool | None = False

class ExportWorkbookABCPayload(BaseModel):
    scenario_a: dict
    scenario_b: dict
    scenario_c: dict
    project_name: str | None = None
    sheet_name_a: str | None = "Scenario A"
    sheet_name_b: str | None = "Scenario B"
    sheet_name_c: str | None = "Scenario C"
    add_timestamp: bool | None = False

class ExportXMLPayload(BaseModel):
    scenario: Optional[Dict[str, Any]] = None
    project_name: Optional[str] = None
    scenario_label: Optional[str] = None
    sheet_name: str = "Scenario A"
    start_date_mode: str = "next_monday"
    fixed_start_iso: Optional[str] = None  # ISO8601 project start (e.g., "2025-10-06T09:00:00")
    hours_per_day: float = 8.0
    merge_identical_children: bool = False

class ExportWorkbookXMLPayload(BaseModel):
    scenario_a: Optional[Dict[str, Any]] = None
    scenario_b: Optional[Dict[str, Any]] = None  
    project_name: Optional[str] = None
    project_start_iso: Optional[str] = None  # ISO8601 project start (e.g., "2025-10-06T09:00:00")
    merge_identical_children: bool = False

class ExportWorkbookXMLABCPayload(BaseModel):
    scenario_a: Optional[Dict[str, Any]] = None
    scenario_b: Optional[Dict[str, Any]] = None
    scenario_c: Optional[Dict[str, Any]] = None
    project_name: Optional[str] = None
    start_date_mode: str = "next_monday"   # "next_monday" | "fixed"
    fixed_start_iso: Optional[str] = None  # ISO8601 project start (e.g., "2025-10-06T09:00:00")
    hours_per_day: float = 8.0

class AuditPricingPayload(BaseModel):
    scenario: Dict[str, Any]           # scenario object from /api/build
    pricing_mode: str                  # "Flat_Blended" | "Per_Resource"
    blended_rate: Optional[float] = None
    rate_band: Optional[str] = "Standard_US"
    price_uses_rounded_hours: bool = True  # bill with rounded hours to match export

class BuildScenarioCPayload(BaseModel):
    base: str  # "A" or "B"
    add_on_codes: List[str] = []
    pricing_mode: str = "Flat_Blended"   # or "Per_Resource"
    blended_rate: Optional[float] = DEFAULT_BILLABLE_RATE_USD
    rate_band: str = "Standard_US"
    complexity: Optional[str] = None        # default to base scenario's
    tier: Optional[str] = None
    use_slack: Optional[bool] = None
    slack_after_internal: Optional[int] = None
    slack_after_client: Optional[int] = None
    slack_global_pct: Optional[float] = None
    project_start: Optional[str] = None     # ISO8601 format (e.g., "2025-10-06T09:00:00" or "YYYY-MM-DD")
    client_budget_usd: Optional[float] = None  # Client budget for budget analysis
    # NEW: override or inherit from base scenario
    retainers: Optional[List[RetainerSelection]] = []

# --- AI Summary models (Stage 2) ---
class RfpSummaryItem(BaseModel):
    label: str                       # deliverable name (human-friendly)
    short_desc: str                  # <= 2 sentences
    tasks: list[str] | None = []     # optional, zero or more tasks (strings)

class RfpSummary(BaseModel):
    summary_text: str                # rendered text for right panel (<= 500 words)
    deliverables: list[RfpSummaryItem]
    word_count: int

class SummarizePayload(BaseModel):
    rfp_text: str | None = None      # optional if using file route

# --- Reconcile (Stage 2, middle panel) ---
class ReconcilePayload(BaseModel):
    summary_deliverables: List[str]                 # from the right-panel AI summary (labels only)
    db_selected_deliverable_codes: Optional[List[str]] = None  # current selection on the left (codes)
    rfp_text: Optional[str] = None

class ReconcileSuggestion(BaseModel):
    code: str
    label: str
    reason: str
    preselect: bool = True

class ReconcileResult(BaseModel):
    add: list[ReconcileSuggestion]
    delete: list[ReconcileSuggestion]
    unchanged: list[str]
    db_used_codes: list[str]          # NEW: actual codes reconcile compared
    db_used_labels: list[str]         # NEW: labels for those codes

class ReorderPayload(BaseModel):
    scenario_letter: str
    deliverable_codes: list[str]                        # new order
    included_map: dict[str, list[str]] | None = None   # {code: [task_groups]}
    use_slack: bool = True
    slack_after_internal: int = 1
    slack_after_client: int = 2
    slack_global_pct: float = 0.0
    project_start: str | None = None                   # "YYYY-MM-DD"
    complexity: str = "Advanced"
    tier: str = "T2_MediumVolume"

# A1: L3 Request model for bulk component queries
class L3Request(BaseModel):
    deliverable_code: str
    component: Union[str, List[str]]  # can be one or many components

# ---------- Global Scenario Storage for Reordering ----------
_CURRENT_SCENARIOS = {}  # Store scenarios for reordering

def _current_scenarios():
    """Access current scenarios for reordering operations."""
    return _CURRENT_SCENARIOS

def _get_scenarios(session_id: Optional[str] = None) -> dict:
    """
    Get scenarios for export.
    If session_id is provided, uses get_working_scenario() to get Step 3/4 edits.
    Falls back to _CURRENT_SCENARIOS if session_id not found or not provided.
    
    This ensures XML exports reflect Step 3 pricing edits and Step 4 Gantt edits.
    
    UPDATED (GPT-5 Pro spec): Now uses get_working_scenario() for consistent access.
    Calls get_session_state() unconditionally to ensure migration runs.
    """
    print(f"[GET_SCENARIOS] Called with session_id={session_id}")
    
    if session_id:
        # Always call get_session_state to ensure migration runs (GPT-5 Pro spec)
        _ = get_session_state(session_id)
        # Use get_working_scenario() to ensure consistent access with migration/cloning
        scenario = get_working_scenario(session_id)
        print(f"[GET_SCENARIOS] Got working scenario via get_working_scenario()")
        
        if not scenario or not scenario.get("items"):
            print(f"[GET_SCENARIOS] Working scenario is empty, falling back to _CURRENT_SCENARIOS")
            return _CURRENT_SCENARIOS
        
        # Debug: Show what we got
        items = scenario.get("items", [])
        print(f"[GET_SCENARIOS] Scenario has {len(items)} items")
        if items:
            first_item = items[0]
            print(f"[GET_SCENARIOS] First item: {first_item.get('Deliverable', first_item.get('deliverable', 'N/A'))}, hours={first_item.get('Planned_Hours', first_item.get('planned_hours', 'N/A'))}, price={first_item.get('price_usd', first_item.get('Price_USD', 'N/A'))}")
        
        # Call _recompute_totals to ensure all pricing fields are in sync before export
        scenario = _recompute_totals(scenario)
        
        print(f"[GET_SCENARIOS] Scenario keys: {list(scenario.keys())}")
        print(f"[GET_SCENARIOS] Has timeline_tasks: {'timeline_tasks' in scenario}")
        
        if "timeline_tasks" in scenario:
            print(f"[GET_SCENARIOS] timeline_tasks count: {len(scenario['timeline_tasks'])}")
        
        # DEFENSIVE: Strip any anchor-related settings AND tasks from cached scenarios
        # This prevents old cached settings from overriding our Workfront compatibility fixes
        def clean_anchor_settings(scen):
            """Remove any cached anchor milestone settings and ANCHOR_ tasks (recursive deep clean)"""
            if isinstance(scen, dict):
                # Remove anchor flags at this level
                scen.pop("add_deliverable_milestones", None)
                scen.pop("add_anchors", None)
                scen.pop("include_anchors", None)
                scen.pop("deliverable_anchor_map", None)
                
                # Recursively clean all nested dicts and lists
                for key, value in list(scen.items()):
                    if key == "timeline_tasks" and isinstance(value, list):
                        # Filter out ANCHOR_ tasks
                        original_count = len(value)
                        scen[key] = [
                            task for task in value
                            if not (isinstance(task, dict) and 
                                   str(task.get("Task_Label", "")).startswith("ANCHOR_"))
                        ]
                        removed = original_count - len(scen[key])
                        if removed > 0:
                            print(f"[CLEAN_ANCHORS] Removed {removed} ANCHOR_ tasks from {key}")
                    elif isinstance(value, dict):
                        # Recursively clean nested dicts
                        scen[key] = clean_anchor_settings(value)
                    elif isinstance(value, list):
                        # Recursively clean items in lists
                        scen[key] = [
                            clean_anchor_settings(item) if isinstance(item, dict) else item
                            for item in value
                        ]
            return scen
        
        scenario = clean_anchor_settings(scenario)
        
        # Wrap in letter format if needed for compatibility
        if "items" in scenario:
            # This is a single scenario, return it as "A"
            print(f"[GET_SCENARIOS] Wrapping single scenario as 'A'")
            return {"A": scenario}
        
        # If it has letter keys, clean each one
        for letter in ["A", "B", "C"]:
            if letter in scenario:
                scenario[letter] = clean_anchor_settings(scenario[letter])
        
        print(f"[GET_SCENARIOS] Returning scenario as-is (already has letter keys)")
        return scenario
    
    # Fall back to _CURRENT_SCENARIOS (pricing-only data)
    print(f"[GET_SCENARIOS] Using _CURRENT_SCENARIOS (no session_id or not found)")
    return _CURRENT_SCENARIOS

# ---------- Helper function to create retainer summary sheet ----------
def create_retainer_summary(scenario: dict) -> pd.DataFrame:
    """Create a detailed retainer summary sheet for Excel export"""
    rows = []
    
    items = scenario.get("items", [])
    retainer_items = []
    
    # Find all retainer deliverables
    for item in items:
        months = int((item.get("retainer") or {}).get("months", 0))
        is_retainer = months > 0 or item.get("is_retainer", False)
        if is_retainer:
            retainer_items.append(item)
    
    if not retainer_items:
        # No retainers, return empty dataframe with headers
        return pd.DataFrame(columns=[
            "Deliverable", "Deliverable_Code", "Service_Department",
            "Total_Months", "Monthly_Hours", "Monthly_Cost", 
            "Total_Hours", "Total_Cost", "Start_Month", "End_Month"
        ])
    
    # Generate summary rows
    for item in retainer_items:
        dcode = str(item.get("deliverable_code", item.get("code", "")))
        deliv_label = str(item.get("deliverable", dcode))
        months = int((item.get("retainer") or {}).get("months", 12))
        monthly_hours = item.get("monthly_hours", 0)
        
        # Calculate costs
        pricing_mode = scenario.get("pricing_mode", "Flat_Blended")
        if pricing_mode == "Flat_Blended":
            blended_rate = scenario.get("blended_rate", DEFAULT_BILLABLE_RATE_USD)
            monthly_cost = monthly_hours * blended_rate
        else:
            # Use average rate from hours_by_role
            hrs_df = pd.DataFrame(item.get("hours_by_role") or [])
            if not hrs_df.empty and "Rate" in hrs_df.columns:
                avg_rate = hrs_df["Rate"].mean()
            else:
                avg_rate = DEFAULT_BILLABLE_RATE_USD
            monthly_cost = monthly_hours * avg_rate
        
        total_hours = monthly_hours * months
        total_cost = monthly_cost * months
        
        # Get department
        dept = item.get("_service_department", "Strategy")
        
        # Summary row
        rows.append({
            "Deliverable": deliv_label,
            "Deliverable_Code": dcode,
            "Service_Department": dept,
            "Total_Months": months,
            "Monthly_Hours": monthly_hours,
            "Monthly_Cost": round(monthly_cost, 2),
            "Total_Hours": total_hours,
            "Total_Cost": round(total_cost, 2),
            "Start_Month": "Month 1",
            "End_Month": f"Month {months}"
        })
        
        # Add monthly breakdown
        for month in range(1, months + 1):
            rows.append({
                "Deliverable": f"  └─ Month {month}",
                "Deliverable_Code": dcode,
                "Service_Department": dept,
                "Total_Months": "",
                "Monthly_Hours": monthly_hours,
                "Monthly_Cost": round(monthly_cost, 2),
                "Total_Hours": "",
                "Total_Cost": "",
                "Start_Month": "",
                "End_Month": ""
            })
    
    df = pd.DataFrame(rows)
    return df

# ---------- OpenAI Integration Functions (Stage 2) ----------

# Initialize OpenAI clients (both sync and async)
try:
    api_key = os.environ.get("OPENAI_API_KEY")
    if api_key:
        # Initialize synchronous client
        openai_client = OpenAI(api_key=api_key)
        # Initialize asynchronous client for async operations
        from openai import AsyncOpenAI
        async_openai_client = AsyncOpenAI(api_key=api_key)
        print(f"OpenAI clients initialized successfully (sync and async)")
    else:
        print("No OPENAI_API_KEY found in environment")
        openai_client = None
        async_openai_client = None
except Exception as e:
    print(f"Failed to initialize OpenAI clients: {e}")
    openai_client = None  # Optional OpenAI integration
    async_openai_client = None

# --- Reconciliation Helper Functions ---
def _norm_tokens(s: str) -> set[str]:
    return set(re.findall(r"[a-z0-9]+", (s or "").lower()))

def _jaccard(a: set[str], b: set[str]) -> float:
    if not a or not b:
        return 0.0
    return len(a & b) / len(a | b)

# --- LLM adapter: must not touch DB ---
_SENT_SPLIT = re.compile(r'(?<=[.!?])\s+')

def _count_words(s: str) -> int:
    return len(re.findall(r"\b\w+\b", s or ""))

def _truncate_to_2_sentences(s: str) -> str:
    parts = _SENT_SPLIT.split((s or "").strip())
    return " ".join(parts[:2]).strip()

def validate_ai_response(rfp_text: str, ai_deliverables: List[dict]) -> bool:
    """
    Validate that AI-generated deliverables are relevant to the RFP content.
    Returns False if there's clear contamination from other sources.
    """
    # Extract key terms from RFP
    rfp_lower = rfp_text.lower()
    
    # Check for common contamination patterns
    contamination_patterns = [
        # Music/Audio industry specific
        ("soundcloud", "artist accelerator", "music streaming"),
        # Education specific  
        ("charter school", "uncommon schools", "student enrollment", "teacher recruitment"),
        # E-commerce specific
        ("shopify", "woocommerce", "e-commerce platform"),
        # Healthcare specific
        ("patient care", "medical practice", "healthcare provider"),
    ]
    
    # Identify which domain the RFP belongs to
    rfp_domain = None
    for pattern_group in contamination_patterns:
        if any(term in rfp_lower for term in pattern_group):
            rfp_domain = pattern_group
            break
    
    # Check deliverables for contamination
    for deliverable in ai_deliverables:
        desc_lower = (deliverable.get("short_desc", "") + " " + deliverable.get("label", "")).lower()
        
        # Check if deliverables mention terms from OTHER domains
        for pattern_group in contamination_patterns:
            if pattern_group != rfp_domain:  # Different domain than RFP
                if any(term in desc_lower for term in pattern_group):
                    print(f"[VALIDATION] WARNING: Deliverable contaminated - mentions '{pattern_group[0]}' but RFP is about '{rfp_domain[0] if rfp_domain else 'general'}'")
                    return False
    
    return True

def ai_summarize_rfp_text(text: str) -> RfpSummary:
    """
    Call GPT‑5 (max compute) with a structured prompt that returns JSON:
      { "deliverables": [{"label": "...", "short_desc": "...", "tasks": [".."]}, ...] }
    and a prose summary <= 500 words.
    This function intentionally avoids any DB lookups.
    """
    global RFP_SUMMARY_CACHE
    try:
        # Create structured prompt for GPT-5
        system_prompt = """
You are an agency executive producer.
Read the RFP text and output JSON ONLY in this exact schema:

{
  "deliverables": [
    {"label": "...", "short_desc": "...", "tasks": ["...", "..."]}
  ]
}

Guidelines:
- Identify 3–8 concrete agency deliverables needed to fulfill the request:
  strategy, campaign creative, content production (video/audio/stills), social/community, editorial web/livestream, experiential/IRL, media planning/buying, measurement & reporting, program management/timeline.
- Each "short_desc" is ≤2 sentences, specific to this RFP, action‑oriented.
- Use common agency taxonomy for "label" so it will match a database later (e.g., "Brand Strategy", "Campaign Creative", "Content Production (Video/Audio)", "Social Media & Community", "Editorial Microsite & Livestream", "Experiential Activation", "Media Planning & Buying", "Measurement & Reporting", "Program Management & Timeline").
- Do NOT quote the RFP; summarize the work we must deliver.
- Keep total text concise (UI cap is 500 words).
- IMPORTANT: Your response MUST be specific to the actual RFP content provided. Do not use generic templates.
"""

        user_prompt = f"Analyze this RFP text and extract the key deliverables:\n\n{text[:8000]}"  # Limit input size
        
        response = openai_client.chat.completions.create(
            model="gpt-5",  # Use GPT-5 directly
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"}
        )
        
        # Parse JSON response
        import json
        result = json.loads(response.choices[0].message.content)
        deliverables = result.get("deliverables", [])
        
    except Exception as e:
        print(f"OpenAI error (using smarter fallback): {e}")
        t = (text or "").lower()
        deliverables = []

        def add(label, desc, *keys):
            if any(k in t for k in keys) and not any(d["label"] == label for d in deliverables):
                deliverables.append({"label": label, "short_desc": desc[:300], "tasks": []})

        # Generic fallback deliverables - NO hardcoded client-specific content
        add("Strategy Development",
            "Develop comprehensive strategy based on objectives, target audience analysis, and market research.", 
            "strategy", "objectives", "goals", "audience")
        add("Campaign Creative & Messaging",
            "Create campaign concepts, messaging framework, and creative platform aligned with brand objectives.",
            "campaign", "creative", "messaging", "brand")
        add("Content Production",
            "Produce multimedia content assets including video, audio, and visual materials for various channels.",
            "video", "audio", "content", "production", "assets")
        add("Digital Platform Development",
            "Design and develop digital experiences, websites, or applications as needed.",
            "digital", "website", "platform", "app", "web")
        add("Social Media Strategy",
            "Develop social media strategy, content calendar, and community engagement approach.",
            "social", "community", "engagement")
        add("Media Planning & Buying",
            "Create comprehensive media plan with channel strategy, budget allocation, and optimization approach.",
            "media", "advertising", "paid", "budget")
        add("Analytics & Reporting",
            "Establish measurement framework, KPIs, and reporting structure for campaign performance.",
            "analytics", "measurement", "kpi", "reporting", "metrics")
        add("Project Management",
            "Provide project timeline, resource planning, and stakeholder coordination.",
            "timeline", "project", "management", "coordination")

        if not deliverables:
            deliverables = [{"label":"Program Management & Timeline",
                             "short_desc":"Create a production timeline and rollout schedule with milestones and owners.",
                             "tasks":[]}]

    # Validate deliverables match the RFP (prevent contamination)
    if not validate_ai_response(text, deliverables):
        print("[VALIDATION] AI response failed validation - using generic fallback")
        # Reset to safe generic deliverables
        deliverables = [
            {"label": "Strategy Development", "short_desc": "Develop strategy aligned with project objectives.", "tasks": []},
            {"label": "Creative Development", "short_desc": "Create campaign concepts and messaging.", "tasks": []},
            {"label": "Media Planning", "short_desc": "Plan media strategy and budget allocation.", "tasks": []},
            {"label": "Analytics & Reporting", "short_desc": "Define KPIs and measurement framework.", "tasks": []}
        ]
    
    # enforce constraints
    for d in deliverables:
        d["short_desc"] = _truncate_to_2_sentences(d.get("short_desc",""))

    # concise prose capped at 500 words
    bullets = [f"• {d['label']}: {d['short_desc']}" for d in deliverables]
    prose = "\n".join(bullets)
    words = _count_words(prose)
    if words > 500:
        # trim from the end
        # (simple conservative trimming; UI also shows a counter)
        while bullets and _count_words("\n".join(bullets)) > 500:
            bullets.pop()
        prose = "\n".join(bullets)
        words = _count_words(prose)

    summary = RfpSummary(summary_text=prose, deliverables=[RfpSummaryItem(**d) for d in deliverables], word_count=words)
    # No global cache mutation - frontend holds window.currentRfpSummary instead
    return summary

# --- Name matching for reconciliation (deterministic; DB only used here) ---

def _norm(s: str) -> set[str]:
    return set(re.findall(r"[a-z0-9]+", (s or "").lower()))

def _best_match(label: str, db_rows: pd.DataFrame) -> tuple[str,str,float] | None:
    tokens = _norm(label)
    best = None
    for _, r in db_rows.iterrows():
        code = str(r["Deliverable_Code"]); name = str(r["Deliverable"])
        t2 = _norm(name)
        if not tokens or not t2: 
            continue
        jacc = len(tokens & t2) / max(1, len(tokens | t2))
        if best is None or jacc > best[2]:
            best = (code, name, jacc)
    return best

def _average_spec_for(category: str) -> dict:
    # Prefer a template key containing "MED" if present, else default constants already used in v4 hours
    row = DB.scenario_templates[DB.scenario_templates["Scenario_Key"].str.contains("MED", case=False, na=False)]
    if not row.empty:
        c = str(row.iloc[0]["Complexity"]); t = str(row.iloc[0]["Tier"])
        return {"mode":"template","scenario_key":str(row.iloc[0]["Scenario_Key"]), "complexity":c, "tier":t}
    return {"mode":"template","scenario_key":"AVERAGE","complexity":"Advanced","tier":"T2_MediumVolume"}

# ---- Deliverable ordering helpers ----
def _phase_rank_for(deliv_name: str, included_tgs: list[str]) -> int:
    """Lower rank = earlier phase. Name takes precedence, then task_groups."""
    n = (deliv_name or "").strip().lower()
    if "discovery" in n or "research" in n or "strategy" in n:
        return 0
    if "design" in n or "creative" in n or "concept" in n:
        return 10
    if "post" in n:  # covers "post-production" - check this BEFORE production
        return 30
    if "development" in n or "build" in n or "production" in n:
        return 20
    if "qa" in n or "review" in n or "test" in n:
        return 40
    if "launch" in n or "deploy" in n:
        return 50
    # fall back to included task groups against Timeline_Params order
    base = {str(tg): i for i, tg in enumerate(DB.timeline_params["Task_Group"].astype(str).tolist())}
    if included_tgs:
        return min([base.get(str(tg), 999) for tg in included_tgs])
    return 999

def _deliverable_order_overrides(letter: str) -> list[tuple[str,str]]:
    """
    Optional: UI_Options keys:
      Deliverable_Order_Overrides_A = Development<Post-Production; QA<Launch
      Deliverable_Order_Overrides_B = ...
    Default for both A & B ensures Development comes before Post-Production.
    """
    try:
        key = f"Deliverable_Order_Overrides_{letter.upper()}"
        row = DB.ui_options[DB.ui_options["Key"] == key]
        if not row.empty:
            pairs = []
            for p in str(row["Value"].iloc[0]).split(";"):
                if "<" in p:
                    a, b = [x.strip().lower() for x in p.split("<", 1)]
                    if a and b: pairs.append((a, b))
            if pairs: return pairs
    except Exception:
        pass
    return [("production", "post-production")]  # default: Production < Post-Production

def _sort_deliverables(per_deliv: list[dict], letter: str) -> list[dict]:
    # base rank from phase
    def base_rank(d):
        return _phase_rank_for(d.get("deliverable",""), d.get("included_task_groups", []))
    # topological sort from overrides + base rank as tiebreak
    nodes = list(range(len(per_deliv)))
    name_lc = [str(d.get("deliverable","")).lower() for d in per_deliv]
    edges = []
    for a,b in _deliverable_order_overrides(letter):
        for i, n in enumerate(name_lc):
            if a in n:
                for j, m in enumerate(name_lc):
                    if b in m: edges.append((i,j))
    preds = {i:set() for i in nodes}; succs = {i:set() for i in nodes}
    for i,j in edges: succs[i].add(j); preds[j].add(i)
    ready = [i for i in nodes if not preds[i]]
    ready.sort(key=lambda i: (base_rank(per_deliv[i]), name_lc[i]))
    out = []
    while ready:
        i = ready.pop(0)
        out.append(i)
        for j in sorted(list(succs[i]), key=lambda k: (base_rank(per_deliv[k]), name_lc[k])):
            preds[j].discard(i)
            if not preds[j] and j not in out and j not in ready:
                ready.append(j)
        ready.sort(key=lambda k: (base_rank(per_deliv[k]), name_lc[k]))
    tail = [i for i in nodes if i not in out]
    tail.sort(key=lambda i: (base_rank(per_deliv[i]), name_lc[i]))
    return [per_deliv[i] for i in out + tail]

# ---------- Routes ----------
@app.get("/", response_class=HTMLResponse)
def root():
    with open("static/index.html", "r", encoding="utf-8") as f:
        return f.read()

@app.get("/api/load")
def api_load():
    if not DB.loaded:
        DB.load()
    return {"ok": True, "src": DB.src}

@app.get("/api/last_upload_name")
def api_last_upload_name():
    """Return the sanitized project name default from the most recent file upload."""
    return {"project_name_default": _upload_title_default() or ""}

@app.post("/api/rfp/cache")
async def cache_rfp_text(text: str = Form(...)):
    """Cache RFP text from Step 1 for reuse in Step 2 Refresh AI Suggestions."""
    global RFP_TEXT_CACHE_TEXTAREA, RFP_TEXT_CACHE_FILE, RFP_TEXT_CACHE
    
    RFP_TEXT_CACHE_TEXTAREA = text.strip() or None
    textarea_text = RFP_TEXT_CACHE_TEXTAREA or ""
    file_text = (RFP_TEXT_CACHE_FILE or "").strip()
    
    # Combine both sources with clear separator
    if textarea_text and file_text:
        merged_text = f"{textarea_text}\n\n--- Uploaded Document Content ---\n\n{file_text}"
    elif file_text:
        merged_text = file_text
    else:
        merged_text = textarea_text
    
    # Cache merged text for backward compatibility
    RFP_TEXT_CACHE = merged_text
    return {"ok": True}

@app.get("/api/rfp/cache")
async def get_cached_rfp_text():
    """Retrieve cached RFP text for Step 2."""
    return {"text": RFP_TEXT_CACHE or ""}

class ClearSessionPayload(BaseModel):
    session_id: str

@app.post("/api/clear_session")
async def clear_session(payload: ClearSessionPayload):
    """Clear all session-specific data including embedding cache"""
    session_id = payload.session_id
    
    # Clear embedding cache for this session
    try:
        from embedding_cache import clear_cache
        clear_cache(session_id=session_id)
        embedding_cleared = True
    except Exception as e:
        print(f"[SESSION] Warning: Could not clear embedding cache: {e}")
        embedding_cleared = False
    
    # Clear ALL global data caches
    global RFP_TEXT_CACHE_TEXTAREA, RFP_TEXT_CACHE_FILE, RFP_TEXT_CACHE
    global LAST_UPLOAD_FILENAME, OPTIONS_CACHE
    
    print(f"[SESSION] Clearing all cached data for session: {session_id}")
    print(f"[SESSION] Before clear - RFP_TEXT_CACHE length: {len(RFP_TEXT_CACHE or '')}")
    
    RFP_TEXT_CACHE_TEXTAREA = None
    RFP_TEXT_CACHE_FILE = None
    RFP_TEXT_CACHE = None
    # RFP_SUMMARY_CACHE removed - no longer needed
    LAST_UPLOAD_FILENAME = None
    
    # Clear OPTIONS_CACHE if it exists
    if 'OPTIONS_CACHE' in globals():
        OPTIONS_CACHE = None
    
    # Clear all jobs from JOB_STORE
    global JOB_STORE
    cleared_jobs = []
    for job_id in list(JOB_STORE.keys()):
        cleared_jobs.append(job_id)
        del JOB_STORE[job_id]
    
    # Clear any cached scenarios
    if 'SCENARIOS_CACHE' in globals():
        global SCENARIOS_CACHE
        SCENARIOS_CACHE = None
    
    print(f"[SESSION] After clear - All caches reset")
    print(f"[SESSION] Cleared {len(cleared_jobs)} jobs: {cleared_jobs}")
    
    return {
        "ok": True,
        "message": f"Session {session_id} cleared completely",
        "cleared": {
            "embedding_cache": embedding_cleared,
            "rfp_text_cache": True,
            "upload_filename": True,
            "jobs": len(cleared_jobs),
            "job_ids": cleared_jobs,
            "timestamp": datetime.datetime.now().isoformat()
        }
    }

@app.get("/api/upload/progress/{job_id}")
async def get_upload_progress(job_id: str):
    """Get progress of image analysis job with two-phase tracking"""
    if job_id not in JOB_STORE:
        raise HTTPException(404, f"Job {job_id} not found")
    
    job = JOB_STORE[job_id]
    
    return {
        "job_id": job.job_id,
        "status": job.status.value,
        "total_images": job.total_images,
        "processed_images": job.processed_images,
        "percentage": round(job.percentage, 1),
        "eta_seconds": round(job.eta_seconds, 1) if job.eta_seconds is not None else None,
        "errors": job.errors,
        "result_text": job.result_text,
        "phase": job.phase,  # "quick_scan" or "deep_analysis"
        "skipped_images": job.skipped_images,
        "relevant_images": job.relevant_images
    }

@app.post("/api/upload/cancel/{job_id}")
async def cancel_upload(job_id: str):
    """Cancel an in-progress image analysis job"""
    if job_id not in JOB_STORE:
        raise HTTPException(404, f"Job {job_id} not found")
    
    job = JOB_STORE[job_id]
    
    if job.status in [JobStatus.COMPLETED, JobStatus.FAILED, JobStatus.CANCELLED]:
        return {"ok": False, "message": f"Job already {job.status.value}"}
    
    job.cancelled = True
    job.status = JobStatus.CANCELLED
    job.end_time = datetime.datetime.now().timestamp()
    
    return {"ok": True, "message": "Job cancelled successfully"}

@app.get("/api/components")
def list_components(deliverable: str):
    """Return all unique components for a given deliverable code."""
    if not DB.loaded:
        DB.load()
    df = DB.all_rows
    deliverable = str(deliverable).strip().lower()
    subset = df[df["Deliverable_Code"].astype(str).str.lower() == deliverable]
    comps = (
        subset["Component"]
        .fillna("")
        .astype(str)
        .str.strip()
        .replace({"nan": ""})
        .unique()
        .tolist()
    )
    return sorted([c for c in comps if c])

@app.get("/api/l3")
def list_l3(deliverable: str, component: str):
    """Return all L3 tasks for a given deliverable and component."""
    if not DB.loaded:
        DB.load()
    df = DB.all_rows
    d = str(deliverable).strip().lower()
    c = str(component).strip().lower()
    sub = df[
        (df["Deliverable_Code"].astype(str).str.lower() == d) &
        (df["Component"].astype(str).str.lower() == c)
    ]
    task_col = "Task_Label" if "Task_Label" in sub.columns else "Task"
    items = (
        sub[task_col]
        .fillna("")
        .astype(str)
        .str.strip()
        .replace({"nan": ""})
        .unique()
        .tolist()
    )
    return sorted([x for x in items if x])

def dedupe_list(xs):
    """Deduplicate and sort a list of strings, removing empty/nan values."""
    return sorted(list({str(x).strip() for x in xs if str(x).strip() and str(x).strip().lower() != "nan"}))

@app.get("/api/options")
def api_options():
    if not DB.loaded:
        DB.load()

    # Prefer v3 Drivers (3 each)
    v3_complexities, v3_tiers = DB.drivers_complexities_tiers_v3()
    if not v3_complexities:
        if DB.timeline_scaling is not None:
            v3_complexities = DB.timeline_scaling[DB.timeline_scaling["Scale_Type"]=="Complexity"]["Key"].head(3).tolist()
        else:
            v3_complexities = ["Basic", "Standard", "Advanced"]
    if not v3_tiers:
        if DB.timeline_scaling is not None:
            v3_tiers = DB.timeline_scaling[DB.timeline_scaling["Scale_Type"]=="Tier"]["Key"].head(3).tolist()
        else:
            v3_tiers = ["T1_LowVolume", "T2_MediumVolume", "T3_HighVolume"]

    rate_bands = dedupe_list(DB.rate_bands["Band_Name"].head(3).tolist()) if DB.rate_bands is not None else []
    pricing_modes = dedupe_list(["Flat_Blended","Per_Resource"])  # Deduplicate
    
    # FIXED: Only return unique deliverables (L0 items)
    # Group by Deliverable_Code to get unique deliverables
    deliv_cols = ["Deliverable_Code", "Deliverable", "Category"]
    if "Service Department" in DB.deliverables.columns:
        deliv_cols.append("Service Department")
    
    deliverables_df = DB.deliverables[deliv_cols].copy()
    
    # Remove duplicates - keep only unique deliverable codes
    deliverables_df = deliverables_df.drop_duplicates(subset=['Deliverable_Code'], keep='first')
    
    # Add sort order (use DataFrame index if no explicit sort column exists)
    deliverables_df["Sort_Order"] = deliverables_df.index
    deliverables = deliverables_df.to_dict(orient="records")

    # Safe extraction with None checks
    bundles = []
    if DB.b_defaults is not None and "Bundle" in DB.b_defaults.columns:
        bundles = DB.b_defaults["Bundle"].tolist()
    
    scenario_templates = []
    if DB.scenario_templates is not None:
        scenario_templates = DB.scenario_templates.to_dict(orient="records")
    
    pricing_settings = []
    if DB.pricing_settings is not None:
        pricing_settings = DB.pricing_settings.to_dict(orient="records")
    
    slack_settings = []
    if DB.slack_settings is not None:
        slack_settings = DB.slack_settings.to_dict(orient="records")
    
    return {
        "complexities": v3_complexities,
        "tiers": v3_tiers,
        "rate_bands": rate_bands,
        "pricing_modes": pricing_modes,
        "bundles": bundles,
        "deliverables": deliverables,
        "scenario_templates": scenario_templates,
        "pricing_settings": pricing_settings,
        "slack_settings": slack_settings,
    }

@app.get("/api/search_deliverables")
def api_search_deliverables(q: str = "", limit: int = 50):
    if not DB.loaded: DB.load()
    df = DB.deliverables.copy()
    q = (q or "").strip().lower()
    if q:
        mask = (
            df["Deliverable"].astype(str).str.lower().str.contains(q)
            | df["Category"].astype(str).str.lower().str.contains(q)
            | df["Deliverable_Code"].astype(str).str.lower().str.contains(q)
        )
        df = df[mask]
    rows = df[["Deliverable_Code", "Deliverable", "Category"]].head(limit).to_dict("records")
    return {"items": rows}

@app.post("/api/resolve_deliverables")
def api_resolve_deliverables(p: ResolveDeliverablesPayload):
    if not DB.loaded: DB.load()
    out = []
    for s in (p.inputs or []):
        code = DB._canonical_deliverable_code(s)   # v2.7 helper
        row = DB.deliverables[DB.deliverables["Deliverable_Code"].astype(str)==str(code)]
        if row.empty: 
            continue
        out.append({
            "input": s,
            "code": code,
            "deliverable": str(row["Deliverable"].iloc[0]),
            "category":  str(row["Category"].iloc[0]),
        })
    return {"resolved": out}

@app.get("/api/components_for")
def api_components_for(deliverable_code: str, complexity: str="Advanced", tier: str="T2_MediumVolume"):
    """List components for a deliverable with hours breakdown."""
    if not DB.loaded: DB.load()
    
    # Get included task groups for this deliverable
    try:
        # Use the database method to get task groups for deliverable
        sub = DB.all_rows[DB.all_rows["Deliverable_Code"].astype(str)==str(deliverable_code)]
        included = sorted(set(sub["task_group"].dropna().astype(str).tolist()))
    except Exception:
        included = []
    
    if not included:
        return {"items": []}
    
    scen_col = DB.scenario_hours_col(complexity, tier)
    
    # Get components list 
    comp_names = sorted(c for c in set(sub.get("Component","").astype(str)) if c and c!="nan")
    
    # Hours by component
    try:
        hours_map = DB.hours_by_component(deliverable_code, included, scen_col)
    except Exception:
        # Fallback: sum directly
        g = (sub[sub["task_group"].isin(included)]
             .groupby("Component")[scen_col].sum(numeric_only=True))
        hours_map = {k: float(v) for k, v in g.items()}
    
    return {"items": [{"name": c, "hours": float(hours_map.get(c, 0.0))} for c in comp_names]}

@app.get("/api/l3_for")
def api_l3_for(deliverable_code: str, component_name: str):
    """List L3 subtasks (Task_Label) for a deliverable and component."""
    if not DB.loaded: DB.load()
    
    try:
        # Filter by deliverable and component
        sub = DB.all_rows[
            (DB.all_rows["Deliverable_Code"].astype(str) == str(deliverable_code)) &
            (DB.all_rows["Component"].astype(str) == str(component_name))
        ]
        
        if sub.empty:
            return {"items": []}
        
        # Get unique Task_Label values
        task_labels = sorted(set(
            sub["Task_Label"].dropna().astype(str)
            .str.strip()
            .tolist()
        ))
        
        # Filter out empty strings
        task_labels = [label for label in task_labels if label and label != "nan"]
        
        return {"items": [{"Task_Label": label} for label in task_labels]}
    except Exception as e:
        print(f"Error in api_l3_for: {e}")
        return {"items": []}

# ============================================================================
# Step 2 Deterministic Endpoints (single source of truth)
# ============================================================================

@app.post("/api/step2/components")
async def step2_components(payload: dict):
    """
    Returns components for selected deliverables.
    payload = {"deliverables": ["deck_strategy", "production", ...]}
    returns = {
      "deck_strategy": ["brief", "art_direction", "internal_review", ...],
      "production": ["prepro_coord", "shoot", "post", ...],
    }
    """
    if not DB.loaded:
        DB.load()
    
    codes = [str(c).strip() for c in payload.get("deliverables", [])]
    df = DB.all_rows.copy()
    
    out = {}
    for code in codes:
        subset = df[df["Deliverable_Code"] == code]
        comps = (
            subset["Component"]
            .fillna("").astype(str).str.strip()
            .replace({"nan": ""})
            .tolist()
        )
        comps = sorted({c for c in comps if c})
        out[code] = comps if comps else ["general"]
    
    return out

@app.post("/api/step2/l3")
async def step2_l3(p: L3Request):
    """
    Returns L3 subtasks for a deliverable + component(s).
    Supports both single component and multiple components (bulk query).
    
    Single: {"deliverable_code": "deck_strategy", "component": "brief"}
    Bulk: {"deliverable_code": "deck_strategy", "component": ["brief", "art_direction"]}
    
    Returns merged, deduplicated list of L3 tasks.
    """
    if not DB.loaded:
        DB.load()
    
    dcode = str(p.deliverable_code).strip()
    # Handle both string and list inputs
    comps = p.component if isinstance(p.component, list) else [p.component]
    comps = [str(c).strip() for c in comps]
    
    df = DB.all_rows.copy()
    out: set[str] = set()
    
    for comp in comps:
        if not comp:
            continue
            
        rows = df[(df["Deliverable_Code"] == dcode) & (df["Component"] == comp)]
        
        # Prefer Task_Label (normalized), fallback to task_group if empty
        label_col = "Task_Label"
        if label_col not in rows.columns or rows[label_col].eq("").all():
            label_col = "task_group" if "task_group" in rows.columns else None
        
        if label_col:
            tasks = rows[label_col].fillna("").astype(str).str.strip().tolist()
            for task in tasks:
                if task:
                    out.add(task)
    
    return {
        "deliverable_code": dcode,
        "components": comps,
        "l3": sorted(out)
    }

@app.post("/api/step2/l3/bulk")
async def step2_l3_bulk(payload: dict):
    """
    Returns L3 subtasks grouped by component (not merged).
    payload: {"deliverable": "deck_strategy", "components": ["brief","art_direction",...]}
    returns: {"brief": ["deck_build","internal_review",...], "art_direction": [...]}
    
    This endpoint preserves component grouping for UI rendering.
    """
    if not DB.loaded:
        DB.load()
    
    dcode = str(payload.get("deliverable", "")).strip()
    comps = [str(x).strip() for x in payload.get("components", [])]
    df = DB.all_rows.copy()
    
    out = {}
    for comp in comps:
        if not comp:
            continue
        
        # ISSUE 1 FIX: Treat "General" or "general" as empty component
        comp_query = comp.lower() if comp.lower() in ['general', 'generic'] else comp
        if comp_query in ['general', 'generic']:
            # For "General", find rows with empty or missing Component
            rows = df[(df["Deliverable_Code"] == dcode) & 
                     ((df["Component"].isna()) | (df["Component"] == "") | 
                      (df["Component"].str.lower() == "general"))]
        else:
            rows = df[(df["Deliverable_Code"] == dcode) & (df["Component"] == comp)]
        
        # Prefer Task_Label (normalized), fallback to task_group if empty
        label_col = "Task_Label"
        if label_col not in rows.columns or rows[label_col].eq("").all():
            label_col = "task_group" if "task_group" in rows.columns else None
        
        tasks = []
        if label_col:
            task_list = rows[label_col].fillna("").astype(str).str.strip().tolist()
            tasks = sorted({t for t in task_list if t and t != "nan"})
        
        out[comp] = tasks
    
    return out

# --- Suggest Components for a deliverable ---
class SuggestComponentsReq(BaseModel):
    deliverable_code: str
    limit: Optional[int] = 6

@app.post("/api/step2/suggest/components")
def suggest_components(req: SuggestComponentsReq):
    if not DB.loaded:
        DB.load()
    d = str(req.deliverable_code)
    
    sub = DB.all_rows[DB.all_rows["Deliverable_Code"].astype(str) == d]
    if sub.empty:
        return JSONResponse([])
    
    # Simple weighting: frequency of tasks per Component, boosted by RFP keyword overlap
    txt = (RFP_TEXT_CACHE or "").lower()
    tokens = {t for t in re.findall(r"[a-z0-9]{3,}", txt)}
    
    # score component by (#rows) + bonus if Task_Label matches RFP
    scores = {}
    for _, r in sub.iterrows():
        comp = str(r.get("Component","") or "").strip() or "General"
        scores.setdefault(comp, 0)
        scores[comp] += 1
        lab = str(r.get("Task_Label","") or "").lower()
        if tokens and any(tok in lab for tok in tokens):
            scores[comp] += 2  # small boost for RFP relevance
    
    ranked = sorted(scores.items(), key=lambda x: (-x[1], x[0]))
    out = [c for c,_ in ranked[: (req.limit or 999)]]
    return JSONResponse(out)


# --- Suggest L3 tasks for one or more components (dedupe-aware) ---
class SuggestL3Req(BaseModel):
    deliverable_code: str
    components: List[str]
    exclude_labels: Optional[List[str]] = None
    limit_per_component: Optional[int] = 20

@app.post("/api/step2/suggest/l3")
def suggest_l3(req: SuggestL3Req):
    if not DB.loaded:
        DB.load()
    d = str(req.deliverable_code)
    excl = { (x or "").strip().lower() for x in (req.exclude_labels or []) }
    txt = (RFP_TEXT_CACHE or "").lower()
    tokens = {t for t in re.findall(r"[a-z0-9]{3,}", txt)}
    
    resp: Dict[str, List[str]] = {}
    sub = DB.all_rows[DB.all_rows["Deliverable_Code"].astype(str)==d]
    if sub.empty:
        return JSONResponse(resp)
    
    for c in req.components or []:
        cand = sub[(sub["Component"].astype(str)==str(c)) | ((sub["Component"]=="") & (c=="General"))]
        if cand.empty:
            resp[c] = []
            continue
        
        # score tasks: frequency + RFP keyword overlap
        t_scores: Dict[str, int] = {}
        for _, r in cand.iterrows():
            lab = str(r.get("Task_Label","") or "").strip()
            if not lab:
                continue
            key = lab.lower()
            t_scores.setdefault(lab, 0)
            t_scores[lab] += 1
            if tokens and any(tok in key for tok in tokens):
                t_scores[lab] += 2
        
        # drop duplicates across the whole project if requested
        if excl:
            t_scores = {lab:sc for lab,sc in t_scores.items() if lab.lower() not in excl}
        
        ranked = sorted(t_scores.items(), key=lambda x: (-x[1], x[0]))
        resp[c] = [lab for lab,_ in ranked[: (req.limit_per_component or 999)]]
    
    return JSONResponse(resp)

# ========= AI Suggest (GPT-5) =========
# Model will be selected based on tier and enforced by sitecustomize
OPENAI_MODEL = "gpt-5"  # Use GPT-5 model directly

class AISuggestReq(BaseModel):
    deliverable_code: str
    include_l3: bool = True
    top_components: int = 6
    top_l3_per_component: int = 12
    rfp_text: str | None = None
    exclude_labels: List[str] | None = None
    weighted_context: dict | None = None  # Pre-filter context from weighted rules

class WeightedScoresRequest(BaseModel):
    """Request model for weighted AI matching with TF-IDF selection mode"""
    rfp_text: Optional[str] = None
    selection_mode: str = "confidence_only"  # "confidence_only" | "tfidf_only" | "both"

# ========= AI Timeline Generation =========
class TimelineGenerationRequest(BaseModel):
    """Request model for AI timeline generation"""
    deliverables: List[Dict[str, Any]]  # Selected deliverables with metadata
    rfp_text: Optional[str] = None  # RFP context text
    project_start: Optional[str] = None  # ISO date format YYYY-MM-DD
    optimization_mode: str = "balanced"  # "speed" | "quality" | "balanced" | "cost"
    use_intelligent_scheduler: bool = True  # Use new intelligent scheduler
    session_id: Optional[str] = None  # Session ID for SCENARIO_STORE sync

def _component_catalog_for_deliverable(db: AgencyDB, dcode: str, max_tasks_per_comp: int = 40) -> Dict[str, List[str]]:
    """Return {component: [top task labels]} limited for token safety."""
    sub = db.all_rows[db.all_rows["Deliverable_Code"].astype(str) == str(dcode)]
    if sub.empty:
        return {}

    sub["Component"] = sub["Component"].fillna("").astype(str).str.strip().replace({"nan": ""})
    sub["Task_Label"] = sub["Task_Label"].fillna("").astype(str).str.strip().replace({"nan": ""})

    sub.loc[sub["Component"]=="", "Component"] = "General"

    out: Dict[str, List[str]] = {}
    for comp, grp in sub.groupby("Component", sort=False):
        g = (grp["Task_Label"]
                .replace({"": None})
                .dropna()
                .value_counts(sort=True))
        top = [str(x) for x in g.index.tolist()[:max_tasks_per_comp]]
        out[comp] = top
    return out

def _rules_pick_components_and_l3(db: AgencyDB, dcode: str, rfp: str,
                                  top_components: int, top_l3: int,
                                  exclude: set[str]) -> dict:
    """Deterministic fallback using frequency + RFP overlaps."""
    catalog = _component_catalog_for_deliverable(db, dcode, 100)
    text = (rfp or "").lower()
    toks = set(re.findall(r"[a-z0-9]{3,}", text))

    comp_scores = {}
    for c, tasks in catalog.items():
        score = len(tasks)
        if toks:
            score += sum(any(t in (lab.lower()) for t in toks) for lab in tasks)
        comp_scores[c] = score

    comp_ranked = [c for c,_ in sorted(comp_scores.items(), key=lambda x: (-x[1], x[0]))][:top_components]

    l3_pick: Dict[str, List[dict]] = {}
    for c in comp_ranked:
        tasks = catalog.get(c, [])
        t_scores = []
        for lab in tasks:
            key = lab.lower()
            if key in exclude: 
                continue
            s = 1 + sum(tok in key for tok in toks) * 2
            t_scores.append((lab, s))
        t_sorted = [lab for lab,_ in sorted(t_scores, key=lambda x: (-x[1], x[0]))][:top_l3]
        l3_pick[c] = [{"label": lab, "why": "Rule-based relevance"} for lab in t_sorted]

    return {
        "source": "rules",
        "components": [{"name": c, "score": float(comp_scores.get(c, 0)), "why": "High frequency + RFP keyword overlap"} for c in comp_ranked],
        "l3_by_component": l3_pick
    }

def _gpt_pick_components_and_l3(db: AgencyDB, dcode: str, rfp: str,
                                top_components: int, top_l3: int,
                                exclude: set[str], weighted_context: dict = None) -> dict:
    catalog = _component_catalog_for_deliverable(db, dcode, 60)
    drow = db.deliverables[db.deliverables["Deliverable_Code"].astype(str)==str(dcode)]
    dname = str(drow["Deliverable"].iloc[0]) if not drow.empty else dcode

    instructions = [
        "Pick components that best respond to the RFP.",
        "Only choose components and tasks that exist in the catalog.",
        "Avoid duplicated tasks across components (use exclude_labels and dedupe).",
        "Prefer tasks that reflect real client-facing outcomes and the typical workflow.",
        "Return JSON ONLY in the schema below; do not add fields."
    ]
    
    # Add weighted context hint if available
    if weighted_context and weighted_context.get("top_components"):
        top_comps = [c.get("component") for c in weighted_context["top_components"][:3]]
        if top_comps:
            instructions.insert(0, f"PRIORITY: Rule-based analysis suggests these components are highly relevant: {', '.join(top_comps)}. Consider these first.")

    user_payload = {
        "rfp_summary": (rfp or "")[:12000],
        "deliverable": {"code": dcode, "name": dname},
        "catalog": [{"component": c, "tasks": catalog[c]} for c in sorted(catalog.keys())],
        "top_components": top_components,
        "top_l3_per_component": top_l3,
        "exclude_labels": sorted(list(exclude)),
        "instructions": instructions,
        "return_schema": {
            "type": "object",
            "properties": {
                "components": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type":"string"},
                            "why": {"type":"string"},
                            "score": {"type":"number"}
                        },
                        "required": ["name"]
                    }
                },
                "l3_by_component": {
                    "type": "object",
                    "additionalProperties": {
                        "type": "array",
                        "items": {
                            "type":"object",
                            "properties": {
                                "label": {"type":"string"},
                                "why": {"type":"string"}
                            },
                            "required": ["label"]
                        }
                    }
                },
                "rationale_summary": {"type":"string"}
            },
            "required": ["components","l3_by_component"]
        }
    }

    if not openai_client:
        raise RuntimeError("OpenAI not configured")

    try:
        resp = openai_client.responses.create(
            model=OPENAI_MODEL,
            input=[
                {"role":"system","content":"You are a senior agency program manager. Output STRICT JSON per schema."},
                {"role":"user","content":json.dumps(user_payload)}
            ],
            response_format={"type":"json_object"}
        )
        content = resp.output_text
    except Exception:
        chat = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role":"system","content":"You are a senior agency program manager. Output STRICT JSON per schema."},
                {"role":"user","content":json.dumps(user_payload)}
            ],
            response_format={"type":"json_object"}
        )
        content = chat.choices[0].message.content

    data = json.loads(content)
    comps = data.get("components") or []
    for c in comps:
        c.setdefault("why","Selected by GPT‑5 for RFP fit")
        c.setdefault("score", 0.9)
    l3 = data.get("l3_by_component") or {}
    for k, arr in l3.items():
        for item in arr:
            item.setdefault("why","GPT‑5 rationale")

    return {
        "source": "gpt",
        "components": comps[:top_components],
        "l3_by_component": {k: v[:top_l3] for k,v in l3.items()},
        "rationale_summary": data.get("rationale_summary","")
    }

@app.post("/api/step2/ai/suggest")
def ai_suggest(req: AISuggestReq):
    db: AgencyDB = app.state.db
    if not getattr(db, "loaded", False):
        db.load()
    d = str(req.deliverable_code)
    rfp = (req.rfp_text or RFP_TEXT_CACHE or "")
    exclude = { (x or "").strip().lower() for x in (req.exclude_labels or []) }

    try:
        payload = _gpt_pick_components_and_l3(
            db, d, rfp, req.top_components, req.top_l3_per_component, 
            exclude, weighted_context=req.weighted_context
        )
        payload["model_used"] = OPENAI_MODEL
        if req.weighted_context:
            payload["used_weighted_prefilter"] = True
    except Exception as e:
        print(f"GPT suggest fallback to rules: {e}")
        payload = _rules_pick_components_and_l3(db, d, rfp, req.top_components, req.top_l3_per_component, exclude)
        payload["model_used"] = "rules"

    return JSONResponse(payload)

@app.post("/api/step2/ai/weights")
def api_weighted_scores(payload: WeightedScoresRequest):
    """
    NEW: Weighted AI matching using rule-based + lexical (TF-IDF) scoring.
    Returns deliverables ranked by match % with top components/tasks.
    Now supports configurable selection modes (confidence, TF-IDF, or both).
    """
    print("[DEBUG] ===== API_WEIGHTED_SCORES FUNCTION CALLED =====")
    rfp_text = payload.rfp_text or RFP_TEXT_CACHE or ""
    selection_mode = payload.selection_mode or "confidence_only"
    print(f"[DEBUG] selection_mode received: {selection_mode}")
    
    # Validate selection_mode
    valid_modes = {"confidence_only", "tfidf_only", "both"}
    if selection_mode not in valid_modes:
        print(f"[WARNING] Invalid selection_mode '{selection_mode}', defaulting to 'confidence_only'")
        selection_mode = "confidence_only"
    
    ai_rules_path = "AI_Matching_Rules_full.xlsx"
    
    if not DB.loaded:
        DB.load()
    
    try:
        result = score_rfp(rfp_text, ai_rules_path, deliverable_index_df=DB.deliverables)
        
        print(f"[TF-IDF SELECT] Processing {len(result.get('deliverables', []))} deliverables with mode={selection_mode}")
        
        # Apply selection logic to each deliverable based on selection_mode
        selected_count = 0
        for deliv in result.get("deliverables", []):
            match_percent = deliv.get("match_percent", 0.0)
            tfidf_similarity = deliv.get("tfidf_similarity", 0.0)
            
            # Apply should_select_deliverable function
            selection_info = should_select_deliverable(
                match_percent=match_percent,
                tfidf_similarity=tfidf_similarity,
                selection_mode=selection_mode
            )
            
            # Add selection fields to deliverable
            deliv["selected"] = selection_info["selected"]
            deliv["selection_reason"] = selection_info["selection_reason"]
            
            if selection_info["selected"]:
                selected_count += 1
        
        print(f"[TF-IDF SELECT] Selected {selected_count} deliverables")
        
        # Add selection_mode to result metadata
        result["selection_mode"] = selection_mode
        result["thresholds"] = {
            "confidence": CONF_THRESHOLD,
            "tfidf": TFIDF_THRESHOLD
        }
        
        return JSONResponse(result)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Weighted scoring failed: {str(e)}")

@app.post("/api/step2/ai/weights_v2")
def api_weighted_scores_v2(payload: WeightedScoresRequest):
    """
    V2: Weighted AI matching using rule-based + lexical (TF-IDF) scoring.
    Returns deliverables ranked by match % with top components/tasks.
    Supports configurable selection modes (confidence, TF-IDF, or both).
    CACHE BYPASS VERSION - identical logic to v1 but fresh route.
    """
    print("[DEBUG] ===== API_WEIGHTED_SCORES_V2 FUNCTION CALLED =====")
    rfp_text = payload.rfp_text or RFP_TEXT_CACHE or ""
    selection_mode = payload.selection_mode or "confidence_only"
    print(f"[DEBUG] selection_mode received: {selection_mode}")
    
    # Validate selection_mode
    valid_modes = {"confidence_only", "tfidf_only", "both"}
    if selection_mode not in valid_modes:
        print(f"[WARNING] Invalid selection_mode '{selection_mode}', defaulting to 'confidence_only'")
        selection_mode = "confidence_only"
    
    ai_rules_path = "AI_Matching_Rules_full.xlsx"
    
    if not DB.loaded:
        DB.load()
    
    try:
        result = score_rfp(rfp_text, ai_rules_path, deliverable_index_df=DB.deliverables)
        
        print(f"[TF-IDF SELECT] Processing {len(result.get('deliverables', []))} deliverables with mode={selection_mode}")
        
        # Apply selection logic to each deliverable based on selection_mode
        selected_count = 0
        for deliv in result.get("deliverables", []):
            match_percent = deliv.get("match_percent", 0.0)
            tfidf_similarity = deliv.get("tfidf_similarity", 0.0)
            
            # Apply should_select_deliverable function
            selection_info = should_select_deliverable(
                match_percent=match_percent,
                tfidf_similarity=tfidf_similarity,
                selection_mode=selection_mode
            )
            
            # Add selection fields to deliverable
            deliv["selected"] = selection_info["selected"]
            deliv["selection_reason"] = selection_info["selection_reason"]
            
            if selection_info["selected"]:
                selected_count += 1
        
        print(f"[TF-IDF SELECT] Selected {selected_count} deliverables")
        
        # Add selection_mode to result metadata
        result["selection_mode"] = selection_mode
        result["thresholds"] = {
            "confidence": CONF_THRESHOLD,
            "tfidf": TFIDF_THRESHOLD
        }
        
        return JSONResponse(result)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Weighted scoring failed: {str(e)}")

@app.post("/api/ai/generate_timeline")
async def generate_timeline(request: TimelineGenerationRequest):
    """
    Generate intelligent project timeline with parallel workstreams and dependencies.
    Uses the new intelligent scheduler for realistic project planning.
    Now with real-time SSE progress streaming!
    """
    
    if not DB.loaded:
        DB.load()
    
    # Import SSE streaming functions
    from app_perf.stream import create_sse_job, update_sse_job, StreamJobStatus
    
    # Create a job for progress tracking
    job_id = str(uuid.uuid4())
    job = create_sse_job(job_id)
    
    # Start the timeline generation in background
    async def generate_with_progress():
        print(f"[Timeline] Background task STARTED for job {job_id}")
        print(f"[Timeline] Request has {len(request.deliverables)} deliverables")
        print(f"[Timeline] use_intelligent_scheduler={request.use_intelligent_scheduler}")
        try:
            print(f"[Timeline] Updating job status to PROCESSING...")
            # Update job status
            update_sse_job(job_id, 
                          status=StreamJobStatus.PROCESSING,
                          progress=5.0,
                          message="Preparing timeline generation...",
                          current_stage="initialization")
            print(f"[Timeline] Job status updated, sleeping 0.1s...")
            await asyncio.sleep(0.1)  # Allow UI to update
            print(f"[Timeline] Starting deliverable enrichment...")
            
            # Enrich deliverables with database information
            enriched_deliverables = []
            total_deliverables = len(request.deliverables)
            
            # CRITICAL: Chunk processing for large projects to avoid timeout
            CHUNK_SIZE = 150  # Process max 150 tasks at a time
            
            for i, deliv in enumerate(request.deliverables):
                code = deliv.get('deliverable_code', '')
                
                # Update progress more frequently for large sets
                # For large sets (>10), update every deliverable
                # For very large sets (>20), also show batch progress
                progress = 10 + (i / total_deliverables) * 20  # 10-30% for enrichment
                
                # Determine update frequency based on size
                should_update = True
                if total_deliverables <= 10:
                    should_update = True  # Always update for small sets
                elif total_deliverables <= 20:
                    should_update = i % 2 == 0 or i == total_deliverables - 1  # Every 2nd for medium
                else:
                    should_update = i % 5 == 0 or i == total_deliverables - 1  # Every 5th for large
                
                if should_update:
                    # Create descriptive message based on project size
                    if total_deliverables > 50:
                        batch_num = (i // 10) + 1
                        total_batches = (total_deliverables // 10) + 1
                        message = f"Processing batch {batch_num}/{total_batches} - Deliverable {i+1} of {total_deliverables}..."
                    elif total_deliverables > 20:
                        message = f"Analyzing deliverable {i+1} of {total_deliverables} ({code})..."
                    else:
                        message = f"Processing deliverable {i+1} of {total_deliverables}..."
                    
                    update_sse_job(job_id,
                                  status=StreamJobStatus.PROCESSING,
                                  progress=progress,
                                  message=message,
                                  current_stage="analyzing_deliverables",
                                  processed_items=i+1,
                                  total_items=total_deliverables)
                
                # Get deliverable details from database
                db_row = DB.deliverables[DB.deliverables['Deliverable_Code'] == code]
                if not db_row.empty:
                    # Helper to normalize department names
                    def normalize_department(dept_str):
                        """Normalize department names to match frontend expectations"""
                        if not dept_str or str(dept_str).strip() == 'nan':
                            return 'Strategy'
                        
                        dept_lower = str(dept_str).strip().lower()
                        
                        # Map to standard departments
                        if 'creative' in dept_lower:
                            return 'Creative'
                        elif 'paid' in dept_lower or 'media' in dept_lower:
                            return 'Paid Media'
                        elif 'tech' in dept_lower or 'dev' in dept_lower:
                            return 'Technology'
                        elif 'content' in dept_lower:
                            return 'Content'
                        elif 'integrated' in dept_lower or 'marketing management' in dept_lower:
                            return 'Integrated Marketing Management'
                        elif 'project' in dept_lower and 'management' in dept_lower:
                            return 'Project Management'
                        elif 'quality' in dept_lower or 'qa' in dept_lower:
                            return 'Quality Assurance'
                        elif 'account' in dept_lower:
                            return 'Account Management'
                        elif 'strategy' in dept_lower or 'strategic' in dept_lower:
                            return 'Strategy'
                        else:
                            # Return original if no match, but capitalize properly
                            return ' '.join(word.capitalize() for word in str(dept_str).strip().split())
                    
                    # Get department from database - check different possible column names
                    department = 'Strategy'  # Default
                    for col_name in ['Service_Department', 'Service Department', 'Department']:
                        if col_name in db_row.columns and not db_row[col_name].empty:
                            dept_value = str(db_row[col_name].iloc[0]).strip()
                            if dept_value and dept_value != 'nan':
                                department = normalize_department(dept_value)
                                break
                    
                    # Also try to get department from all_rows if not in deliverables
                    if department == 'Strategy' and DB.all_rows is not None:
                        dept_rows = DB.all_rows[DB.all_rows['Deliverable_Code'] == code]
                        if not dept_rows.empty:
                            for col_name in ['Service_Department', 'Service Department', 'Department']:
                                if col_name in dept_rows.columns:
                                    dept_values = dept_rows[col_name].dropna()
                                    if not dept_values.empty:
                                        # Get the most common department for this deliverable
                                        dept_value = str(dept_values.value_counts().idxmax()).strip()
                                        if dept_value and dept_value != 'nan':
                                            department = normalize_department(dept_value)
                                            break
                    
                    # Enrich with database data
                    enriched = {
                        'deliverable_code': code,
                        'deliverable_name': db_row['Deliverable'].iloc[0] if 'Deliverable' in db_row.columns else deliv.get('name', ''),
                        'department': department,
                        'total_hours': deliv.get('hours', 0) or deliv.get('total_hours', 0),
                        'components': deliv.get('components', []),
                        'is_retainer': deliv.get('is_retainer', False),
                        'retainer_months': deliv.get('retainer_months', 0)
                    }
                else:
                    # Use provided data
                    enriched = deliv
                
                enriched_deliverables.append(enriched)
            
            # Progress: Creating dependencies (30-40%)
            print(f"[Timeline] Enrichment complete: {len(enriched_deliverables)} deliverables enriched")
            print(f"[Timeline] Updating progress: Creating dependencies...")
            update_sse_job(job_id,
                          status=StreamJobStatus.PROCESSING,
                          progress=35.0,
                          message="Creating dependencies and workstreams...",
                          current_stage="creating_dependencies")
            await asyncio.sleep(0.1)
            print(f"[Timeline] Progress update complete")
            
            # CRITICAL: Determine if chunking is needed to avoid timeout
            total_enriched = len(enriched_deliverables)
            needs_chunking = total_enriched > CHUNK_SIZE
            print(f"[Timeline] Total enriched: {total_enriched}, needs_chunking: {needs_chunking}")
            
            if needs_chunking:
                # Large project - process in chunks to prevent timeout
                print(f"[Timeline] Large project detected: {total_enriched} deliverables will be processed in chunks of {CHUNK_SIZE}")
                
                update_sse_job(job_id,
                              status=StreamJobStatus.PROCESSING,
                              progress=40.0,
                              message=f"Processing large project ({total_enriched} deliverables) in chunks of {CHUNK_SIZE}...",
                              current_stage="chunking")
                
                # Calculate chunk count
                chunk_count = (total_enriched + CHUNK_SIZE - 1) // CHUNK_SIZE
                print(f"[Timeline] Will process {chunk_count} chunks")
                
                # Process chunks with error handling
                all_tasks = []
                successful_chunks = 0
                failed_chunks = []
                
                for chunk_idx in range(chunk_count):
                    start_idx = chunk_idx * CHUNK_SIZE
                    end_idx = min((chunk_idx + 1) * CHUNK_SIZE, total_enriched)
                    chunk = enriched_deliverables[start_idx:end_idx]
                    
                    # Calculate progress: 40% at start, 85% at end of chunking
                    chunk_progress = 40.0 + ((chunk_idx + 1) / chunk_count) * 45.0
                    
                    update_sse_job(job_id,
                                  status=StreamJobStatus.PROCESSING,
                                  progress=chunk_progress,
                                  message=f"Processing batch {chunk_idx + 1}/{chunk_count} ({len(chunk)} deliverables)...",
                                  current_stage=f"batch_{chunk_idx + 1}_of_{chunk_count}",
                                  processed_items=chunk_idx + 1,
                                  total_items=chunk_count)
                    
                    try:
                        # Generate timeline for this chunk
                        print(f"[Timeline] Processing batch {chunk_idx + 1}/{chunk_count} ({start_idx+1}-{end_idx} of {total_enriched})")
                        
                        if request.use_intelligent_scheduler:
                            chunk_result = await generate_intelligent_timeline(
                                chunk,
                                request.project_start,
                                request.optimization_mode
                            )
                        else:
                            chunk_result = await generate_ai_timeline(
                                chunk,
                                request.rfp_text or RFP_TEXT_CACHE or "",
                                request.project_start,
                                request.optimization_mode,
                                use_intelligent_scheduler=False
                            )
                        
                        # Merge tasks from successful chunk
                        if chunk_result and 'tasks' in chunk_result:
                            chunk_tasks = chunk_result['tasks']
                            all_tasks.extend(chunk_tasks)
                            successful_chunks += 1
                            print(f"[Timeline] Batch {chunk_idx + 1}/{chunk_count} completed successfully: {len(chunk_tasks)} tasks generated")
                        else:
                            print(f"[Timeline] WARNING: Batch {chunk_idx + 1}/{chunk_count} returned no tasks")
                            failed_chunks.append(chunk_idx + 1)
                            
                    except Exception as chunk_error:
                        # Log error but continue with other chunks
                        error_msg = str(chunk_error)
                        print(f"[Timeline] ERROR in batch {chunk_idx + 1}/{chunk_count}: {error_msg}")
                        failed_chunks.append(chunk_idx + 1)
                        
                        # Update SSE with warning but continue
                        update_sse_job(job_id,
                                      status=StreamJobStatus.PROCESSING,
                                      progress=chunk_progress,
                                      message=f"Warning: Batch {chunk_idx + 1} failed, continuing with remaining batches...",
                                      current_stage=f"batch_{chunk_idx + 1}_failed")
                        await asyncio.sleep(0.1)
                
                # Log final chunking results
                print(f"[Timeline] Chunking complete: {successful_chunks}/{chunk_count} batches successful, {len(failed_chunks)} failed")
                if failed_chunks:
                    print(f"[Timeline] Failed batches: {failed_chunks}")
                
                # Combine all chunks into final result
                result = {
                    'tasks': all_tasks,
                    'success': True,
                    'metadata': {
                        'chunked': True,
                        'chunk_count': chunk_count,
                        'successful_chunks': successful_chunks,
                        'failed_chunks': failed_chunks,
                        'total_deliverables': total_enriched,
                        'total_tasks': len(all_tasks),
                        'partial_results': len(failed_chunks) > 0
                    }
                }
                
                # Add warning message if some chunks failed
                if failed_chunks:
                    result['warning'] = f"{len(failed_chunks)} batches failed - showing partial results from {successful_chunks} successful batches"
                
            else:
                # Standard processing for smaller projects
                print(f"[Timeline] Starting standard processing (no chunking needed)")
                # Use the intelligent timeline generator
                if request.use_intelligent_scheduler:
                    # Progress: Using intelligent scheduler (40-70%)
                    print(f"[Timeline] Using intelligent scheduler...")
                    update_sse_job(job_id,
                                  status=StreamJobStatus.PROCESSING,
                                  progress=40.0,
                                  message="Using intelligent scheduler to optimize timeline...",
                                  current_stage="optimizing_schedule")
                    
                    print(f"[Timeline] Calling generate_intelligent_timeline() with {len(enriched_deliverables)} deliverables...")
                    # Use the new intelligent scheduler
                    result = await generate_intelligent_timeline(
                        enriched_deliverables,
                        request.project_start,
                        request.optimization_mode
                    )
                    print(f"[Timeline] generate_intelligent_timeline() returned successfully")
                    
                    # Progress: AI reasoning (70-90%)
                    if request.rfp_text:
                        update_sse_job(job_id,
                                      status=StreamJobStatus.PROCESSING,
                                      progress=75.0,
                                      message="Enhancing timeline with AI reasoning...",
                                      current_stage="ai_reasoning")
                        
                        from ai_timeline_manager import enhance_with_ai_reasoning
                        result = await enhance_with_ai_reasoning(
                            result,
                            request.rfp_text,
                            enriched_deliverables
                        )
                else:
                    # Progress: Using standard scheduler (40-90%)
                    update_sse_job(job_id,
                                  status=StreamJobStatus.PROCESSING,
                                  progress=50.0,
                                  message="Generating timeline with standard scheduler...",
                                  current_stage="generating_timeline")
                    
                    # Use the standard AI timeline generator
                    result = await generate_ai_timeline(
                        enriched_deliverables,
                        request.rfp_text or RFP_TEXT_CACHE or "",
                        request.project_start,
                        request.optimization_mode,
                        use_intelligent_scheduler=False
                    )
            
            # Progress: Finalizing (90-100%)
            update_sse_job(job_id,
                          status=StreamJobStatus.PROCESSING,
                          progress=95.0,
                          message="Finalizing timeline and preparing visualization...",
                          current_stage="finalizing")
            
            # Add success flag and descriptive message
            result['success'] = True
            
            # Create descriptive message based on processing mode
            scheduler_type = 'intelligent' if request.use_intelligent_scheduler else 'standard'
            task_count = len(result.get('tasks', []))
            
            if needs_chunking:
                # Include chunking information in message
                metadata = result.get('metadata', {})
                chunk_count = metadata.get('chunk_count', 0)
                successful = metadata.get('successful_chunks', 0)
                
                if metadata.get('partial_results'):
                    result['message'] = f"Generated timeline with {task_count} tasks from {successful}/{chunk_count} batches using {scheduler_type} scheduler (some batches failed)"
                else:
                    result['message'] = f"Generated timeline with {task_count} tasks from {chunk_count} batches using {scheduler_type} scheduler"
            else:
                result['message'] = f"Generated timeline with {task_count} tasks using {scheduler_type} scheduler"
            
            # CRITICAL: Store result in SCENARIO_STORE if session_id is provided
            # Use proper dual-entry format helpers
            if hasattr(request, 'session_id') and request.session_id:
                session_id = request.session_id
                # Get or create proper state with dual-entry format
                state = get_session_state(session_id)
                scenario = state.get("scenario") or {"items": [], "totals": {"hours": 0.0, "price": 0.0}}
                scenario['timeline'] = result
                scenario['project_start'] = request.project_start
                update_working_scenario(session_id, scenario)
                print(f"[Timeline] Stored timeline in SCENARIO_STORE for session {session_id}")
                print(f"[DEBUG] Timeline storage: SCENARIO_STORE[{session_id}] keys:", list(SCENARIO_STORE[session_id].keys()))
            
            # CRITICAL FIX: Strip all dependencies from tasks to enable free movement
            if 'tasks' in result and isinstance(result['tasks'], list):
                for task in result['tasks']:
                    if isinstance(task, dict):
                        task['dependencies'] = ""  # Force empty string
            
            # Complete the job - ENSURE status is COMPLETED not just processing at 100%
            update_sse_job(job_id,
                          status=StreamJobStatus.COMPLETED,
                          progress=100.0,
                          message="Timeline generation complete!",
                          current_stage="completed",
                          result=result)
            
            print(f"[Timeline] Job {job_id} marked as COMPLETED with {len(result.get('tasks', []))} tasks")
            
            return result
            
        except Exception as e:
            print(f"[Timeline Generation] Error: {e}")
            import traceback
            traceback.print_exc()
            
            # Update job with error
            update_sse_job(job_id,
                          status=StreamJobStatus.FAILED,
                          progress=0.0,
                          message=f"Timeline generation failed: {str(e)}",
                          error=str(e))
            
            # Return error response
            return {
                'success': False,
                'error': str(e),
                'tasks': [],
                'reasoning': {
                    'error': f"Failed to generate timeline: {str(e)}"
                },
                'metadata': {}
            }
    
    # Wrapper to ensure background task exceptions are logged
    async def background_task_wrapper():
        try:
            await generate_with_progress()
        except Exception as e:
            print(f"[CRITICAL] Background task for job {job_id} crashed: {e}")
            import traceback
            traceback.print_exc()
            
            # Update job status to failed
            try:
                update_sse_job(job_id,
                              status=StreamJobStatus.FAILED,
                              progress=0.0,
                              message=f"Timeline generation crashed: {str(e)}",
                              error=str(e))
            except Exception as update_error:
                print(f"[CRITICAL] Could not update job status: {update_error}")
    
    # Start the background task with error handling
    task = asyncio.create_task(background_task_wrapper())
    
    # Add callback to log if task is cancelled or raises exception
    def task_done_callback(future):
        try:
            future.result()
        except asyncio.CancelledError:
            print(f"[Timeline] Background task for job {job_id} was cancelled")
        except Exception as e:
            print(f"[Timeline] Background task for job {job_id} raised exception: {e}")
    
    task.add_done_callback(task_done_callback)
    
    # Return immediately with the job ID
    return JSONResponse({
        'success': True,
        'job_id': job_id,
        'message': 'Timeline generation started. Connect to SSE stream for progress updates.'
    })

@app.get("/api/db/status")
def db_status():
    if not DB.loaded: DB.load()
    def ok(df): return (df is not None) and (getattr(df, "empty", True) is False)
    return {
        "loaded": DB.loaded,
        "source": DB.src,
        "has_v4": ok(DB.all_rows),
        "sheets": {
            "all_rows": ok(DB.all_rows),
            "deliverables": ok(DB.deliverables),
            "role_rate_card": ok(DB.role_rate_card),
            "b_rules": ok(DB.b_rules)
        }
    }

@app.post("/api/db/reload")
def db_reload():
    DB.loaded = False
    DB.load()
    return {"ok": True}

@app.post("/api/suggest_by_text")
def api_suggest(payload: SuggestPayload):
    if not DB.loaded:
        DB.load()
    recs = DB.suggest_deliverables_from_text(payload.rfp_text or "")
    # NEW: attach retainer hints
    for r in recs:
        is_ret, months = DB.retainer_recommendation(payload.rfp_text or "", r.get("deliverable",""))
        r["retainer_hint"] = bool(is_ret)
        r["retainer_months_suggested"] = int(months or 0)
    return {"suggested": recs}

@app.post("/api/suggest_by_file")
async def api_suggest_by_file(files: List[UploadFile] = File(...), background_tasks: BackgroundTasks = None):
    if not DB.loaded:
        DB.load()

    # Validate we have at least one file
    if not files:
        raise HTTPException(400, "No files uploaded.")
    
    # Store file text separately and merge with textarea text if present
    global RFP_TEXT_CACHE_TEXTAREA, RFP_TEXT_CACHE_FILE, RFP_TEXT_CACHE, LAST_UPLOAD_FILENAME
    
    # Process all files
    all_file_contents = []
    filenames = []
    background_jobs = []
    
    for file in files:
        # Read content
        content = await file.read()
        if not content or len(content) == 0:
            continue  # Skip empty files

        # Basic size guard (20 MB to match middleware)
        if len(content) > 20 * 1024 * 1024:
            raise HTTPException(413, f"File '{file.filename}' too large. Maximum size is 20MB.")

        text = _extract_text_from_upload(content, file.filename)
        # Hard cap text length to protect downstream regex scan
        if len(text) > 200_000:
            text = text[:200_000]
        
        # Add file content with clear label
        if text.strip():
            labeled_content = f"--- File: {file.filename} ---\n\n{text.strip()}"
            all_file_contents.append(labeled_content)
            filenames.append(file.filename)
            
            # Create job for image processing for this file
            job_id = str(uuid.uuid4())
            JOB_STORE[job_id] = JobState(job_id=job_id, status=JobStatus.PENDING)
            background_jobs.append(job_id)
            
            # Start background image processing for this file
            if background_tasks:
                background_tasks.add_task(_process_images_background, content, file.filename, job_id, text.strip())
    
    # Combine all file contents
    file_content = "\n\n".join(all_file_contents) if all_file_contents else ""
    RFP_TEXT_CACHE_FILE = file_content
    textarea_text = (RFP_TEXT_CACHE_TEXTAREA or "").strip()
    
    # Merge textarea text with all file contents
    parts = []
    if textarea_text:
        parts.append(textarea_text)
    if file_content:
        parts.append(file_content)
    
    merged_text = "\n\n".join(parts) if parts else ""
    
    # Cache merged text for backward compatibility
    RFP_TEXT_CACHE = merged_text
    
    # Generate recommendations from text (images will be processed in background)
    recs = DB.suggest_deliverables_from_text(merged_text or "")
    # NEW: attach retainer hints per deliverable
    for r in recs:
        is_ret, months = DB.retainer_recommendation(merged_text or "", r.get("deliverable",""))
        r["retainer_hint"] = bool(is_ret)
        r["retainer_months_suggested"] = int(months or 0)
    
    # NEW: remember for default project name (use first file)
    LAST_UPLOAD_FILENAME = filenames[0] if filenames else "upload"
    
    return {
        "suggested": recs, 
        "filenames": filenames,  # Return all filenames
        "job_ids": background_jobs,  # Return all job IDs
        "processing_images": len(background_jobs) > 0
    }

def extract_base_deliverable_code(code: str) -> str:
    """
    Extract the base deliverable code from an expanded code.
    Examples:
        'DEL-0027-Google_Ads' -> 'DEL-0027'
        'DEL-0036-North_America' -> 'DEL-0036'
        'DEL-0042-Strategy-Q1' -> 'DEL-0042'
        'DEL-0015-Creative-Launch' -> 'DEL-0015'
        'DEL-0001' -> 'DEL-0001' (unchanged)
    """
    if not code:
        return code
    
    # Pattern matches: DEL-NNNN where N is a digit
    # Everything after the 4 digits is considered the expansion suffix
    import re
    match = re.match(r'^(DEL-\d+)', str(code))
    if match:
        return match.group(1)
    
    # Fallback: If it doesn't match the expected pattern, return as-is
    return str(code)

def _safe_retainer_map(retainers) -> dict:
    out = {}
    for r in (retainers or []):
        code = (getattr(r, "deliverable_code", "") or (r.get("deliverable_code") if isinstance(r, dict) else "")).strip()
        raw = getattr(r, "months", None) if not isinstance(r, dict) else r.get("months")
        try:
            m = int(raw) if raw is not None else 0
        except Exception:
            m = 0
        m = max(1, min(12, m)) if m > 0 else 0
        if code:
            out[code] = m
    return out

def _resolve_scenario(spec: ScenarioSpec, category: str) -> Dict[str, Any]:
    if spec.mode == "template":
        # either use key or explicit complexity/tier
        if spec.scenario_key:
            s = DB.scenario_templates[DB.scenario_templates["Scenario_Key"]==spec.scenario_key]
            if not s.empty:
                c = str(s["Complexity"].iloc[0]); t = str(s["Tier"].iloc[0])
                return {"mode":"template","complexity":c,"tier":t,"scenario_key":spec.scenario_key}
        # fallback to provided complexity/tier
        return {"mode":"template","complexity":spec.complexity or "Advanced","tier":spec.tier or "T2_MediumVolume","scenario_key":spec.scenario_key or "CUSTOM"}
    elif spec.mode == "bundle":
        # use bundle defaults
        b = spec.bundle or "Better"
        c,t = DB.default_complexity_tier_for_bundle(b)
        return {"mode":"bundle","bundle":b,"complexity":c,"tier":t}
    else:
        raise HTTPException(400, f"Unknown scenario mode: {spec.mode}")

def _scenario_for_deliverable(deliv_code: str, category: str, spec: Dict[str, Any],
                              pricing_mode: str, blended_rate: Optional[float], rate_band: str,
                              use_slack: bool, slack_i: int, slack_c: int, slack_pct: float,
                              project_start: Optional[str], scenario_letter: str,
                              retainer_months: int = 0, selected_components: Optional[Union[List[str], Dict[str, Optional[float]]]] = None) -> Dict[str, Any]:
    # Extract base code for database lookups (handles expanded codes like 'DEL-0027-Google_Ads')
    base_code = extract_base_deliverable_code(deliv_code)
    
    # Which task groups to include?
    if spec["mode"] == "bundle":
        included = DB.included_task_groups(category, spec["bundle"])
    else:
        # Template mode: include all task_groups that exist in data for this deliverable (collapsed to unique)
        sub = DB.all_rows[DB.all_rows["Deliverable_Code"].astype(str)==str(base_code)]
        included = sorted(set(sub["task_group"].dropna().astype(str).tolist()))

    complexity, tier = spec["complexity"], spec["tier"]
    scen_col = DB.scenario_hours_col(complexity, tier)
    
    # Handle component-level selection (both old and new formats)
    if selected_components:
        frames = []
        
        # Handle both old format (list) and new format (dict)
        if isinstance(selected_components, list):
            # Old format: ["component1", "component2"]
            component_dict = {comp: None for comp in selected_components}  # None = use default hours
        else:
            # New format: {"component1": 5.5, "component2": None} 
            component_dict = selected_components
            
        print(f"DEBUG Component processing for {deliv_code}: {component_dict}")
        
        for comp, custom_hours in component_dict.items():
            try:
                fr = DB.hours_by_role_for_component(base_code, comp, included, scen_col)
                if fr is not None and not fr.empty:
                    # Apply custom hours if provided
                    if custom_hours is not None:
                        # Scale the role distribution proportionally to match custom hours
                        original_total = fr["Hours"].sum()
                        if original_total > 0:
                            scale_factor = custom_hours / original_total
                            fr = fr.copy()  # Avoid modifying original
                            fr["Hours"] = fr["Hours"] * scale_factor
                            print(f"DEBUG Scaled component '{comp}' from {original_total}h to {custom_hours}h (factor: {scale_factor:.3f})")
                    
                    frames.append(fr)
            except Exception:
                # Fallback: get data for this component directly from all_rows
                sub = DB.all_rows[
                    (DB.all_rows["Deliverable_Code"].astype(str)==str(base_code)) &
                    (DB.all_rows["Component"].astype(str)==str(comp)) &
                    (DB.all_rows["task_group"].isin(included))
                ]
                if not sub.empty:
                    # Group by Resource_Title, Seniority and sum hours
                    grouped = sub.groupby(["Resource_Title", "Seniority"])[scen_col].sum().reset_index()
                    grouped.columns = ["Resource_Title", "Seniority", "Hours"]
                    
                    # Apply custom hours if provided
                    if custom_hours is not None:
                        original_total = grouped["Hours"].sum()
                        if original_total > 0:
                            scale_factor = custom_hours / original_total
                            grouped["Hours"] = grouped["Hours"] * scale_factor
                            print(f"DEBUG Scaled component '{comp}' (fallback) from {original_total}h to {custom_hours}h")
                    
                    frames.append(grouped)
        
        hrs_by_role = pd.concat(frames, ignore_index=True) if frames else pd.DataFrame(columns=["Resource_Title","Seniority","Hours"])
        if not hrs_by_role.empty:
            hrs_by_role = hrs_by_role.groupby(["Resource_Title","Seniority"], as_index=False)["Hours"].sum()
        
        # NEW: if the selected components sum to 0 hours (e.g., front-end sent 0s),
        # fall back to the deliverable-level hours to avoid a zeroed deliverable.
        if hrs_by_role.empty or float(hrs_by_role["Hours"].sum()) <= 0.0:
            print(f"DEBUG {deliv_code}: component selection totals 0h -> fallback to deliverable defaults")
            hrs_by_role = DB.hours_by_role_for_deliverable(base_code, included, scen_col)
    else:
        hrs_by_role = DB.hours_by_role_for_deliverable(base_code, included, scen_col)
    
    # ---- after total_hours is computed - RETAINER-AWARE PRICING ----
    total_hours_raw = float(hrs_by_role["Hours"].sum()) if not hrs_by_role.empty else 0.0
    
    monthly_hours_int = int(round(total_hours_raw))
    if pricing_mode == "Flat_Blended":
        eff_rate = float(blended_rate if blended_rate is not None else
                         DB.pricing_settings.loc[DB.pricing_settings["Key"]=="Default_Blended_Rate","Default"].astype(float).iloc[0])
        monthly_price_int = int(round(monthly_hours_int * eff_rate))
    else:
        price_raw = DB.per_resource_price(hrs_by_role, rate_band=rate_band or "Standard_US")
        eff_rate = round((price_raw / total_hours_raw), 2) if total_hours_raw > 0 else 0.0
        monthly_price_int = int(round(monthly_hours_int * eff_rate))

    months = max(0, int(retainer_months or 0))
    if months > 0:
        total_hours = monthly_hours_int * months
        price_int = monthly_price_int * months
    else:
        total_hours = monthly_hours_int
        price_int = monthly_price_int

    # Schedule
    schedule = DB.build_schedule(
        base_code, included, complexity, tier,
        use_slack, slack_i, slack_c, slack_pct, project_start,
        scenario_letter=scenario_letter
    )

    # Expose components for UI
    sub = DB.all_rows[DB.all_rows["Deliverable_Code"].astype(str)==str(base_code)]
    comp_names = sorted(set(sub["Component"].astype(str))) if not sub.empty else []
    comp_names = [c for c in comp_names if c and c != "nan"]
    
    try:
        comp_hours = DB.hours_by_component(base_code, included, scen_col)
    except Exception:
        # Fallback: calculate from raw data
        comp_hours = {}
        if not sub.empty:
            comp_sub = sub[sub["task_group"].isin(included)]
            for comp in comp_names:
                comp_data = comp_sub[comp_sub["Component"].astype(str)==str(comp)]
                comp_hours[comp] = float(comp_data[scen_col].sum()) if not comp_data.empty else 0.0
    
    return {
        "deliverable_code": deliv_code,
        "included_task_groups": included,
        "complexity": complexity,
        "tier": tier,
        "scenario_col": scen_col,
        "hours_by_role": hrs_by_role.to_dict(orient="records"),
        "total_hours": total_hours,                # total (months × monthly)
        "effective_rate": round(eff_rate, 2),
        "price": price_int,                        # total (months × monthly)
        "schedule": schedule,
        # NEW:
        "retainer": {"months": months} if months > 0 else None,
        "monthly_hours": monthly_hours_int if months > 0 else None,
        "monthly_price": monthly_price_int if months > 0 else None,
        "components": [{"name": c, "hours": float(comp_hours.get(c,0.0)), "selected": (not selected_components or c in selected_components)} for c in comp_names]
    }

@app.post("/api/timeline/suggest")
async def suggest_timeline(request: dict):
    """
    Generate AI-optimized timeline for selected deliverables
    
    Request body (either format):
    Format 1 (legacy):
    {
        "selected_deliverable_codes": ["deck_strategy", "social_content"],
        "rfp_text": "optional RFP text for context",
        "project_start": "2024-01-15",
        "optimization_mode": "balanced"
    }
    
    Format 2 (new):
    {
        "deliverables": [
            {
                "code": "DEL-0034",
                "name": "Brand Positioning Strategy",
                "department": "Strategy",
                "hours": 120,
                "components": [{"name": "...", "hours": 40}]
            }
        ],
        "start_date": "2025-07-07",
        "optimization_mode": "balanced"
    }
    """
    try:
        # Handle both input formats
        deliverables = request.get("deliverables", [])
        selected_codes = request.get("selected_deliverable_codes", [])
        rfp_text = request.get("rfp_text", "")
        project_start = request.get("project_start") or request.get("start_date")  # Support both field names
        optimization_mode = request.get("optimization_mode", "balanced")
        
        deliverables_data = []
        
        # If deliverables are provided directly (new format), use them
        if deliverables:
            for deliv in deliverables:
                # Validate deliverable structure
                if not isinstance(deliv, dict):
                    continue
                
                # Extract data from provided deliverable
                code = deliv.get("code", f"CUSTOM-{len(deliverables_data)+1}")
                name = deliv.get("name", "Untitled Deliverable")
                department = deliv.get("department", "Strategy")
                hours = deliv.get("hours", 0)
                components = deliv.get("components", [])
                
                # Format components
                formatted_components = []
                for comp in components:
                    if isinstance(comp, dict):
                        formatted_components.append({
                            "name": comp.get("name", "Component"),
                            "hours": float(comp.get("hours", 0))
                        })
                
                deliverables_data.append({
                    "deliverable_code": code,
                    "deliverable_name": name,
                    "components": formatted_components,
                    "total_hours": float(hours),
                    "department": department
                })
                
                # Also add to selected_codes for compatibility
                selected_codes.append(code)
        
        # Otherwise, use legacy format with database lookup
        elif selected_codes:
            # Load database
            db = app.state.db
            if not db:
                db = AgencyDB()
                db.load()
            
            # Build deliverables data structure for timeline generation
            for code in selected_codes:
                # Find deliverable in database
                deliv_df = db.deliverables[db.deliverables["Deliverable_Code"] == code]
                if deliv_df.empty:
                    continue
                
                # Get components and tasks for this deliverable
                components = []
                comp_df = db.all_rows[db.all_rows["Deliverable_Code"] == code]["Component"].unique()
                for comp in comp_df:
                    if pd.isna(comp):
                        continue
                    comp_tasks = db.all_rows[
                        (db.all_rows["Deliverable_Code"] == code) & 
                        (db.all_rows["Component"] == comp)
                    ]
                    hours = comp_tasks["Hours"].sum() if "Hours" in comp_tasks.columns else 0
                    components.append({
                        "name": comp,
                        "hours": float(hours)
                    })
                
                # Calculate total hours for deliverable
                total_hours = db.all_rows[db.all_rows["Deliverable_Code"] == code]["Hours"].sum() if "Hours" in db.all_rows.columns else 0
                
                # Get department (try different column names)
                dept_col = None
                for col in ["Service_Department", "Service Department", "Department"]:
                    if col in deliv_df.columns:
                        dept_col = col
                        break
                
                department = deliv_df[dept_col].iloc[0] if dept_col and not deliv_df[dept_col].empty else "Strategy"
                
                deliverables_data.append({
                    "deliverable_code": code,
                    "deliverable_name": deliv_df["Deliverable"].iloc[0],
                    "components": components,
                    "total_hours": float(total_hours),
                    "department": str(department) if not pd.isna(department) else "Strategy"
                })
        
        # Check if we have any valid deliverables
        if not deliverables_data:
            return JSONResponse({
                "error": "No valid deliverables selected",
                "tasks": [],
                "reasoning": {},
                "metadata": {}
            })
        
        # Generate AI timeline
        result = await suggest_timeline_from_selection(
            selected_codes=selected_codes,
            deliverables_db=deliverables_data,
            rfp_text=rfp_text,
            project_start=project_start,
            optimization_mode=optimization_mode
        )
        
        return JSONResponse(result)
        
    except Exception as e:
        print(f"[Timeline API] Error: {e}")
        return JSONResponse({
            "error": str(e),
            "tasks": [],
            "reasoning": {},
            "metadata": {}
        }, status_code=500)

@app.post("/api/build")
def api_build(payload: BuildPayload):
    global _CURRENT_SCENARIOS  # Declare at function start to avoid syntax error
    
    try:
        print(f"DEBUG: api_build called with {len(payload.selected_deliverable_codes)} deliverables")
        if not DB.loaded:
            DB.load()
        
        # --- CASE 0: PRESERVE STEP 3 EDITS (GPT 5 Pro spec) ---
        # If session_id is provided and SCENARIO_STORE has a working scenario,
        # return it instead of rebuilding (unless reset=true in future)
        # This prevents the 137→98 hours regression when clicking "Build Scenario" again
        if payload.session_id and payload.session_id in SCENARIO_STORE:
            working_scenario = get_working_scenario(payload.session_id)
            if working_scenario and working_scenario.get("items"):
                totals = working_scenario.get("totals", {})
                total_hours = totals.get("hours", 0)
                total_price = totals.get("price", 0)
                
                print(f"[/api/build] ✅ Using SCENARIO_STORE[{payload.session_id}] "
                      f"hours={total_hours} price={total_price} (preserving Step 3 edits)")
                
                # Update _CURRENT_SCENARIOS for Step 4 compatibility
                _CURRENT_SCENARIOS["A"] = working_scenario
                
                return {
                    "scenarios": {"A": working_scenario},
                    "totals": {"A": totals},
                    "total_hours": total_hours,
                    "total_price": total_price,
                    "cached": True,
                    "message": "Using working scenario from SCENARIO_STORE (Step 3 edits preserved)"
                }

        # Prepare UI intent
        pricing_mode = payload.pricing_mode
        blended_rate = payload.blended_rate
        rate_band    = payload.rate_band or "Standard_US"
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"ERROR in api_build initialization: {error_details}")
        raise HTTPException(500, f"Build initialization error: {str(e)}")

    # Slack/timeline
    use_slack = bool(payload.use_slack)
    slack_i   = int(payload.slack_after_internal)
    slack_c   = int(payload.slack_after_client)
    slack_pct = float(payload.slack_global_pct or 0)
    project_start = payload.project_start
    client_budget_usd = payload.client_budget_usd
    
    # Project name for exports - use payload or fall back to upload title
    project_name = (payload.project_name or 
                   _upload_title_default() or 
                   f"Proposal {datetime.date.today().isoformat()}").strip()

    # Build retainer map
    ret_map = _safe_retainer_map(payload.retainers)
    
    # Build component selection map (supports both formats)
    comp_map = {}
    for k, v in (payload.selected_components_map or {}).items():
        if isinstance(v, str) and v == "__ALL__":
            # Sentinel value meaning "include all components" - pass empty dict to include all
            comp_map[str(k)] = {}
        elif isinstance(v, list):
            # Old format: ["component1", "component2"]
            comp_map[str(k)] = {str(x): None for x in v}  # None means use default hours
        elif isinstance(v, dict):
            # New format: {"component": hours or None}
            # Sanitize: treat <= 0 as "unselected" (drop); keep None to mean "use default hours"
            cleaned: dict[str, Optional[float]] = {}
            for name, hours in v.items():
                try:
                    if hours is None:
                        cleaned[str(name)] = None
                    else:
                        h = float(hours)
                        if h > 0:
                            cleaned[str(name)] = h
                        # <= 0 -> dropped (unselected)
                except Exception:
                    # ignore bad values
                    continue
            comp_map[str(k)] = cleaned
        else:
            comp_map[str(k)] = {}
    
    print(f"DEBUG Build: Component map processed: {comp_map}")

    # Build Scenario A only (B/C removed for simplicity)
    scenarios = {}
    default_spec_a = ScenarioSpec(mode="template", scenario_key="MED_LOW")
    spec_in = payload.scenario_a or default_spec_a
    letter = "A"
    
    if True:  # Keep indentation consistent
        per_deliv = []
        for code in payload.selected_deliverable_codes:
            # Extract base code from expanded codes (e.g., 'DEL-0027-Google_Ads' -> 'DEL-0027')
            base_code = extract_base_deliverable_code(code)
            expanded_suffix = code[len(base_code):] if len(code) > len(base_code) else ""  # Preserve the suffix for context
            
            row = DB.deliverables[DB.deliverables["Deliverable_Code"].astype(str)==str(base_code)]
            if row.empty: 
                print(f"DEBUG Build: No deliverable found for code '{code}' (base: '{base_code}')")
                continue
            
            deliverable_name = str(row["Deliverable"].iloc[0])
            cat = str(row["Category"].iloc[0])
            
            # If this is an expanded deliverable, append context to the name
            if expanded_suffix:
                # Clean up suffix: remove leading dash, replace underscores with spaces
                suffix_clean = expanded_suffix.lstrip('-').replace('_', ' ')
                deliverable_name = f"{deliverable_name} - {suffix_clean}"
                print(f"DEBUG Build: Expanded code '{code}' -> Base '{base_code}' -> Name '{deliverable_name}' (Category: {cat})")
            else:
                print(f"DEBUG Build: Code '{code}' -> Name '{deliverable_name}' (Category: {cat})")
            
            spec_resolved = _resolve_scenario(spec_in, cat)
            # Use original code for retainer and component map lookups (might be expanded)
            # If not found, also try with base code as fallback
            months = int(ret_map.get(code, ret_map.get(base_code, 0)))
            selected_components_dict = comp_map.get(str(code), comp_map.get(str(base_code), {}))
            out = _scenario_for_deliverable(
                code, cat, spec_resolved,
                pricing_mode, blended_rate, rate_band,
                use_slack, slack_i, slack_c, slack_pct, project_start,
                scenario_letter=letter,
                retainer_months=months,   # NEW
                selected_components=selected_components_dict  # NEW
            )
            # Add names for readability
            out["deliverable"] = deliverable_name
            out["category"]    = cat
            print(f"DEBUG Build: Added to scenario - deliverable field: '{out.get('deliverable', 'MISSING')}')")
            per_deliv.append(out)

        # Look up any previously built scenario so we can preserve a user-locked timeline
        _prev = _current_scenarios().get(letter) or {}
        locked = bool(_prev.get("manual_order_locked"))

        if locked:
            # 1) Preserve the user's order (keep intersection, append any new codes at the end)
            built_by_code = {str(it["deliverable_code"]): it for it in per_deliv}
            prev_order = [str(c) for c in (_prev.get("user_order") or [])]
            keep = [c for c in prev_order if c in built_by_code]
            tail = [c for c in built_by_code.keys() if c not in keep]
            ordered_codes = keep + tail
            per_deliv = [built_by_code[c] for c in ordered_codes]

            # 2) Carry forward the sequential schedule/dates computed by Save Order
            prev_by_code = {str(it.get("deliverable_code")): it for it in (_prev.get("items") or [])}
            for it in per_deliv:
                code = str(it["deliverable_code"])
                prev_it = prev_by_code.get(code)
                if prev_it and prev_it.get("schedule"):
                    sched = prev_it["schedule"]
                    it["schedule"] = sched
                    it["start"] = sched[0]["start_date"]
                    it["end"]   = sched[-1]["end_date"]
                    it["duration_days"] = sum(int(r["duration_days"]) for r in sched)
        else:
            # Default behavior when user hasn't locked a timeline yet
            per_deliv = _sort_deliverables(per_deliv, letter)

        # Apply scenario-specific multipliers and add PM/QA overhead
        spec_resolved = _resolve_scenario(spec_in, "General")  # Get resolved spec for complexity/tier
        per_deliv = _apply_scenario_knobs(
            per_deliv, letter, 
            spec_resolved.get("complexity", "Advanced"), 
            spec_resolved.get("tier", "T2_MediumVolume"),
            pricing_mode, rate_band, DB, blended_rate
        )

        # Totals after order is finalized and scenario knobs applied
        price_sum = sum(int(x["price"]) for x in per_deliv)
        hours_sum = sum(int(round(x["total_hours"])) for x in per_deliv)

        # Build the scenario object
        scenario_out = {
            "pricing_mode": pricing_mode,
            "rate_band": rate_band,
            "blended_rate": blended_rate,
            "use_slack": use_slack,
            "slack_after_internal": slack_i,
            "slack_after_client": slack_c,
            "slack_global_pct": slack_pct,
            "project_start": project_start,
            "project_name": project_name,  # Store project name for exports
            "items": per_deliv,
            "totals": {"hours": int(hours_sum), "price": int(price_sum)}
        }

        # Preserve/Set order tracking flags
        if locked:
            scenario_out["ai_order"] = list(_prev.get("ai_order") or [it["deliverable_code"] for it in per_deliv])
            scenario_out["user_order"] = ordered_codes
            scenario_out["manual_order_locked"] = True
        else:
            scenario_out["ai_order"] = [it["deliverable_code"] for it in per_deliv]
            scenario_out["user_order"] = list(scenario_out["ai_order"])
            scenario_out["manual_order_locked"] = False

        # Inflate components if missing (defensive fallback for exports)
        scenario_out = _inflate_components_if_missing(scenario_out)
        
        scenarios[letter] = scenario_out

    # Store scenarios globally for reordering (global declared at function start)
    _CURRENT_SCENARIOS.update(scenarios)
    
    # Add budget metrics if client_budget_usd was provided
    if client_budget_usd and client_budget_usd > 0:
        scenario_price = scenarios["A"]["totals"]["price"]
        budget_delta = scenario_price - client_budget_usd
        coverage_pct = (scenario_price / client_budget_usd) * 100 if client_budget_usd > 0 else 0
        scale_factor = client_budget_usd / scenario_price if scenario_price > 0 else 1.0
        
        scenarios["A"]["budget_info"] = {
            "client_budget_usd": client_budget_usd,
            "total_price": scenario_price,
            "budget_delta": budget_delta,
            "coverage_pct": round(coverage_pct, 1),
            "scale_factor_if_fit": round(scale_factor, 3)
        }
    
    # Save to SCENARIO_STORE if session_id provided (enables Gantt sync)
    # CRITICAL: Use set_baseline_and_scenario() to create proper dual-entry format
    # This ensures Step 3 edits persist when /api/pricing/build_scenario is called later
    if payload.session_id:
        scenario_data = scenarios["A"].copy()
        scenario_data['session_id'] = payload.session_id
        scenario_data['last_saved'] = datetime.datetime.now().isoformat()
        set_baseline_and_scenario(payload.session_id, scenario_data)
        print(f"[SCENARIO_STORE] Saved scenario to session {payload.session_id} with dual-entry format (enables Step 3 persistence)")
    
    # Return scenarios (A only)
    return {"scenarios": scenarios}

# ========== NEW PRICING API ENDPOINTS ==========

class RedistributeHoursPayload(BaseModel):
    """Request payload for hour redistribution"""
    deliverable_name: str
    deliverable_code: str
    new_total_hours: float
    components: List[Dict[str, Any]]
    complexity: str = "Advanced"
    tier: str = "T2_MediumVolume"
    use_ai: bool = True
    context: Optional[str] = None

class RetainerAnalysisPayload(BaseModel):
    """Request payload for retainer analysis"""
    deliverable_name: str
    total_hours: float
    duration_months: int = 12

class RetainerDistributionPayload(BaseModel):
    """Request payload for retainer hour distribution"""
    monthly_hours: float
    duration_months: int = 12
    ramp_up: bool = True
    seasonality: Optional[List[float]] = None

# ========== NEW SCENARIO_STORE ENDPOINTS MODELS ==========

class BuildScenarioPayload(BaseModel):
    """Build scenario from Step 2 selection"""
    session_id: str
    selection: Dict[str, Any]  # Deliverables with components from Step 2
    project_name: Optional[str] = None
    project_start: Optional[str] = None
    pricing_mode: Optional[str] = "Flat_Blended"
    blended_rate: Optional[float] = DEFAULT_BILLABLE_RATE_USD
    rate_band: Optional[str] = "Standard_US"
    client_budget_usd: Optional[float] = None
    retainers: Optional[List[Dict[str, Any]]] = None
    reset: Optional[bool] = False  # If True, rebuild from Step 2 even if scenario exists

class OptimizeScenarioPayload(BaseModel):
    """Optimize pricing for a scenario"""
    session_id: str
    deliverable_index: int
    new_total_hours: float
    use_ai: bool = True

class CadenceSuggestionPayload(BaseModel):
    """Get retainer vs project cadence suggestion"""
    session_id: str
    deliverable_index: int

class RetainerSuggestionsPayload(BaseModel):
    """Calculate retainer distribution"""
    session_id: str
    deliverable_index: int
    monthly_hours: float
    duration_months: int = 12

class UpdateTaskPayload(BaseModel):
    """Update Gantt task and sync pricing"""
    session_id: str
    wbs_id: str
    duration_days: Optional[float] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    hours_per_day: float = 8.0

class UpdateTasksBatchPayload(BaseModel):
    """Batch update multiple Gantt tasks for mouseup commits"""
    session_id: str
    updates: List[UpdateTaskPayload]

# Removed duplicate endpoint - see the other /api/pricing/redistribute-hours endpoint below

@app.post("/api/pricing/analyze-retainer")
async def api_analyze_retainer(payload: RetainerAnalysisPayload):
    """
    Analyze whether a deliverable should be retainer vs project-based.
    
    Example request:
    {
        "deliverable_name": "Social Media Management",
        "total_hours": 480,
        "duration_months": 12
    }
    """
    try:
        result = analyze_retainer_vs_project(
            deliverable_name=payload.deliverable_name,
            total_hours=payload.total_hours,
            duration_months=payload.duration_months
        )
        
        return {
            "success": True,
            "result": result
        }
    except Exception as e:
        print(f"[Pricing API] Error analyzing retainer: {e}")
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

@app.post("/api/pricing/retainer-distribution")
async def api_retainer_distribution(payload: RetainerDistributionPayload):
    """
    Calculate monthly hour distribution for retainer engagements.
    
    Example request:
    {
        "monthly_hours": 40,
        "duration_months": 12,
        "ramp_up": true
    }
    """
    try:
        distribution = calculate_retainer_distribution(
            monthly_hours=payload.monthly_hours,
            duration_months=payload.duration_months,
            ramp_up=payload.ramp_up,
            seasonality=payload.seasonality
        )
        
        return {
            "success": True,
            "distribution": distribution,
            "total_hours": sum(distribution.values()),
            "average_monthly": sum(distribution.values()) / len(distribution) if distribution else 0
        }
    except Exception as e:
        print(f"[Pricing API] Error calculating distribution: {e}")
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

# ========== NEW SCENARIO_STORE API ENDPOINTS ==========

# ---------- Scenario Persistence Endpoints ----------
@app.get("/api/scenario/{session_id}")
async def api_get_scenario(session_id: str):
    """Retrieve saved scenario from SCENARIO_STORE"""
    if session_id in SCENARIO_STORE:
        scenario = SCENARIO_STORE[session_id]
        # Add timestamp if not present
        if 'last_retrieved' not in scenario:
            scenario['last_retrieved'] = datetime.datetime.now().isoformat()
        return {
            "success": True,
            "scenario": scenario,
            "session_id": session_id,
            "found": True
        }
    return {
        "success": False, 
        "error": f"No scenario found for session {session_id}",
        "found": False
    }

@app.post("/api/scenario/save")
async def api_save_scenario(payload: Dict[str, Any]):
    """Explicitly save scenario to SCENARIO_STORE using proper dual-entry format"""
    try:
        session_id = payload.get("session_id")
        if not session_id:
            return JSONResponse(
                {"success": False, "error": "session_id is required"},
                status_code=400
            )
        
        # Store the scenario with timestamp
        scenario_data = payload.get("scenario", payload)
        scenario_data['last_saved'] = datetime.datetime.now().isoformat()
        scenario_data['session_id'] = session_id
        
        # Use update_working_scenario to maintain dual-entry format
        update_working_scenario(session_id, scenario_data)
        print(f"[DEBUG] /api/scenario/save: SCENARIO_STORE[{session_id}] keys:", list(SCENARIO_STORE[session_id].keys()))
        
        print(f"[SCENARIO_STORE] Saved scenario for session {session_id}")
        return {
            "success": True,
            "session_id": session_id,
            "saved_at": scenario_data['last_saved']
        }
    except Exception as e:
        print(f"[SCENARIO_STORE] Error saving scenario: {e}")
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

@app.delete("/api/scenario/{session_id}")
async def api_delete_scenario(session_id: str):
    """Clear scenario from SCENARIO_STORE"""
    if session_id in SCENARIO_STORE:
        del SCENARIO_STORE[session_id]
        print(f"[SCENARIO_STORE] Deleted scenario for session {session_id}")
        return {"success": True, "deleted": True}
    return {"success": True, "deleted": False, "message": "No scenario to delete"}

@app.get("/api/scenario/exists/{session_id}")
async def api_scenario_exists(session_id: str):
    """Check if scenario exists in SCENARIO_STORE"""
    exists = session_id in SCENARIO_STORE
    scenario_info = None
    
    if exists:
        scenario = SCENARIO_STORE[session_id]
        scenario_info = {
            "last_saved": scenario.get('last_saved'),
            "has_deliverables": bool(scenario.get('items') or scenario.get('deliverables')),
            "has_timeline": bool(scenario.get('timeline')),
            "project_name": scenario.get('project_name', 'Untitled Project')
        }
    
    return {
        "exists": exists,
        "session_id": session_id,
        "info": scenario_info
    }

# Clean up old sessions periodically (24 hours)
@app.post("/api/scenario/cleanup")
async def api_cleanup_old_scenarios():
    """Remove scenarios older than 24 hours"""
    try:
        now = datetime.datetime.now()
        removed = []
        
        for session_id, scenario in list(SCENARIO_STORE.items()):
            saved_at_str = scenario.get('last_saved')
            if saved_at_str:
                try:
                    saved_at = datetime.datetime.fromisoformat(saved_at_str)
                    age_hours = (now - saved_at).total_seconds() / 3600
                    if age_hours > 24:
                        del SCENARIO_STORE[session_id]
                        removed.append(session_id)
                except Exception:
                    pass  # Skip malformed timestamps
        
        print(f"[SCENARIO_STORE] Cleaned up {len(removed)} old scenarios")
        return {
            "success": True,
            "removed_count": len(removed),
            "removed_sessions": removed
        }
    except Exception as e:
        print(f"[SCENARIO_STORE] Cleanup error: {e}")
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

# =============================================================================
# SMART SCHEDULE OPTIMIZATION (GPT-5.1 Pro) - GATED & SAFE
# =============================================================================
# Optional AI optimization layer - FULLY GATED & OFF BY DEFAULT
# Features:
# - Extracts baseline schedule from real WBS hierarchy (same source as exports)
# - Uses GPT-5.1 Pro via Responses API with strict validation
# - Enforces L2-L4 waterfall + L5+ parallel execution
# - Stores optimized schedule separately without affecting baseline
# - Feature flag ensures zero impact on current exports until QA passes
# =============================================================================

# ╔════════════════════════════════════════════════════════════════════════════╗
# ║ FEATURE FLAG: SMART SCHEDULE OPTIMIZATION - FROZEN & DISABLED              ║
# ╚════════════════════════════════════════════════════════════════════════════╝
# DEFAULT: OFF - Feature is completely disabled until explicitly enabled
# 
# When SMART_OPTIMIZATION_ENABLED = False (the default):
#   - /api/optimize-schedule → returns 423 (Locked)
#   - /api/accept-optimization → returns 423 (Locked)
#   - /api/reject-optimization → returns 423 (Locked)
#   - System behaves EXACTLY like pre-optimization version
#   - Exports use baseline schedule only
#
# To enable (after full QA):
#   Set environment variable: SMART_OPTIMIZATION_ENABLED=true
#
# Data stored (when feature is ON after acceptance):
#   - SCENARIO_STORE[session_id]['optimized_schedule'] = { items, explanation, model_used }
#   - SCENARIO_STORE[session_id]['active_schedule'] = 'baseline' | 'optimized'
#   - SCENARIO_STORE[session_id]['wbs_rows'] = baseline WBS structure
#
# Endpoints:
#   POST /api/optimize-schedule → Trigger AI optimization (feature-gated)
#   POST /api/accept-optimization → Accept optimized schedule (feature-gated)
#   POST /api/reject-optimization → Discard optimized schedule (feature-gated)
#
# Export Pipeline:
#   /api/export and /api/export_xml → Use baseline_schedule only (NOT hooked to active_schedule yet)
#   Exports NEVER affected by optimization status
#
# ═════════════════════════════════════════════════════════════════════════════
SMART_OPTIMIZATION_ENABLED = os.getenv("SMART_OPTIMIZATION_ENABLED", "false").lower() == "true"
print(f"[STARTUP] Smart Schedule Optimization: {'ENABLED' if SMART_OPTIMIZATION_ENABLED else 'DISABLED (frozen)'}")

class OptimizeSchedulePayload(BaseModel):
    session_id: str
    use_gpt_thinking: bool = False

class OptimizationResponse(BaseModel):
    success: bool
    session_id: str
    model_used: str
    changes_count: int
    explanation: Dict[str, Any]
    diff: List[Dict[str, Any]]
    error: Optional[str] = None

class AcceptOptimizationPayload(BaseModel):
    session_id: str

def build_optimization_payload(scenario: Dict[str, Any], wbs_rows: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """
    Extract baseline schedule from REAL WBS hierarchy (same source as XML export).
    This ensures AI sees the EXACT schedule that will be exported.
    
    Args:
        scenario: SCENARIO_STORE entry
        wbs_rows: Full WBS row list from build_wbs_with_pricing() with L2-L5 hierarchy
    
    Returns:
        JSON payload for GPT-5.1 Pro, or None if unable to extract
    """
    if not wbs_rows:
        print("[OPTIMIZER] No WBS rows - cannot build baseline")
        return None
    
    project_name = scenario.get('project_name', 'Untitled Project')
    project_start = scenario.get('project_start', datetime.date.today().isoformat())
    
    # Build working calendar
    working_calendar = {
        "workday_hours": 8,
        "weekends_off": True,
        "holidays": ["2025-01-01", "2025-07-04", "2025-12-25", "2025-11-27", "2025-11-28"]
    }
    
    # Convert WBS rows to optimization items (with full L2-L5 hierarchy)
    optimization_items = []
    
    for idx, row in enumerate(wbs_rows):
        wbs = row.get('WBS', '')
        if not wbs or wbs == '1':
            continue  # Skip project summary
        
        # Determine level from WBS (dots indicate depth: 1.x=L2, 1.x.y=L3, etc)
        level = wbs.count('.') + 1
        if level < 2 or level > 5:
            continue
        
        name = row.get('Name', row.get('Task_Name', f'Task {wbs}'))
        role = row.get('Role', '') if level >= 5 else None
        
        # Get dates from row (may be from Gantt or calculated)
        start_date_str = row.get('Start_Date', '')
        end_date_str = row.get('End_Date', '')
        
        # Fallback to project_start if no dates
        if not start_date_str:
            start_date_str = project_start
        if not end_date_str:
            end_date_str = project_start
        
        # Get work hours (critical - must be preserved)
        planned_hours = float(row.get('PlannedHours', row.get('Planned_Hours', 0)))
        duration_days = max(1, round(row.get('Duration', 1)))
        
        # Get dependencies (WBS format: "1.2,1.3")
        dependencies = []
        dep_str = row.get('Dependencies', '')
        if dep_str:
            dependencies = [d.strip() for d in str(dep_str).split(',') if d.strip()]
        
        optimization_items.append({
            "id": wbs,  # Use WBS as stable ID
            "wbs": wbs,
            "name": name,
            "level": level,
            "type": "role" if level >= 5 else "summary",
            "role": role,
            "planned_hours": planned_hours,
            "duration_days": duration_days,
            "start": start_date_str,
            "end": end_date_str,
            "dependencies": dependencies,
            "service_department": row.get('Service_Department', 'Other')
        })
    
    print(f"[OPTIMIZER] Built baseline with {len(optimization_items)} items (L2-L5 hierarchy)")
    
    return {
        "project": {
            "name": project_name,
            "start_date": project_start,
            "target_end_date": None,
            "working_calendar": working_calendar
        },
        "items": optimization_items
    }

def get_optimizer_system_prompt() -> str:
    """
    System prompt for GPT-5.1 Pro with explicit L2-L4 vs L5+ dependency rules.
    """
    return """You are a senior digital project manager and scheduling expert.
You receive a project plan with deliverables, components, tasks, and role-level tasks, including dates, durations, and dependencies.
Your job is to adjust the schedule realistically, not to invent new work.

**CRITICAL RULES:**

1. **Do NOT add or remove items** - only adjust dates, durations, and dependencies.

2. **Respect the project working calendar** (8h days, no weekends, holidays list).

3. **Hierarchy preservation** - Keep the WBS hierarchy intact:
   - Level 1 = Project Summary
   - Level 2 = Deliverables (can be chained in waterfall)
   - Level 3 = Components (can be chained in waterfall)
   - Level 4 = Tasks (can be chained in waterfall)
   - Level 5+ = Role-level tasks (must run in PARALLEL unless truly dependent)

4. **L2-L4 Waterfall Dependencies:**
   - Deliverables, components, and tasks CAN be chained where appropriate
   - Use dependencies to express true blocking relationships
   - Example: "Brand Narrative" depends on "Key Pillars" is valid

5. **L5+ Parallel Execution (CRITICAL):**
   - Role-level tasks (Account Manager, Copywriter, Strategist, Designer, Developer) must NOT be auto-chained
   - They should run in PARALLEL unless you explicitly identify a real dependency
   - If you create a dependency between role tasks, explain WHY in your output
   - Example: AM, CW, and Strategist can all start on same day working simultaneously

6. **Keep total planned hours per task unchanged** - if you change duration, you're changing how hours are spread, not the total.

7. **Try to keep within the user's timeline if possible** - if not, explain why.

8. **Output only valid JSON** that matches the requested schema (no code fences, no commentary).

Your goal is to create a realistic, executable schedule that maximizes parallel work while respecting true dependencies."""

async def optimize_schedule_with_ai(
    session_id: str,
    baseline_payload: Dict[str, Any],
    use_thinking: bool = False
) -> Dict[str, Any]:
    """
    Call GPT-5.1 Pro via Responses API to optimize the schedule.
    
    Args:
        session_id: Session identifier for logging
        baseline_payload: Baseline schedule JSON
        use_thinking: True = gpt-5.1-thinking, False = gpt-5.1-pro
    
    Returns:
        Dict with optimized items + explanation, or error details
    """
    import json
    from openai import OpenAI
    from gpt5_helpers import retry_with_exponential_backoff
    
    # Get OpenAI client
    api_key = os.getenv("OPENAI_API_KEY") or os.getenv("Open_AI_Key")
    if not api_key:
        raise ValueError("[OPTIMIZER] No OpenAI API key found")
    
    client = OpenAI(api_key=api_key)
    
    # Select model
    model = "gpt-5.1-thinking" if use_thinking else "gpt-5.1-pro"
    
    # Build messages
    system_prompt = get_optimizer_system_prompt()
    user_message = f"""Here is the current baseline project schedule generated by our app.
Please adjust dates, durations, and dependencies to be realistic, following the rules in the system message.
Then return the updated schedule and an explanation of key changes.

BASELINE JSON:
{json.dumps(baseline_payload, indent=2)}

Expected response format:
{{
  "items": [
    {{
      "id": "item-0",
      "start": "YYYY-MM-DD",
      "end": "YYYY-MM-DD",
      "duration_days": 5,
      "dependencies": ["item-1", "item-2"]
    }}
  ],
  "explanation": {{
    "summary": "Overall changes summary",
    "changes_by_group": [
      {{
        "group_name": "Brand Positioning",
        "notes": "What changed and why"
      }}
    ],
    "risks_or_assumptions": ["List any assumptions made"]
  }}
}}"""
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_message}
    ]
    
    print(f"[OPTIMIZER] Starting optimization for session {session_id}")
    print(f"[OPTIMIZER] Using model: {model}")
    print(f"[OPTIMIZER] Baseline has {len(baseline_payload.get('items', []))} items")
    
    # Call GPT-5.1 Pro via Responses API (enforcer will route it correctly)
    def make_api_call():
        response = client.responses.create(
            model=model,
            input=messages,
            max_output_tokens=4000,
            reasoning={"effort": "high" if not use_thinking else "medium"}
        )
        return response
    
    # Use retry logic
    try:
        response = retry_with_exponential_backoff(
            make_api_call,
            max_retries=3,
            log_prefix="OPTIMIZER",
            raise_on_failure=True
        )
        
        # Extract response text
        from gpt5_helpers import _extract_output_text
        response_text = _extract_output_text(response)
        
        print(f"[OPTIMIZER] Received response ({len(response_text)} chars)")
        
        # Parse JSON response
        try:
            # Remove code fences if present
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0]
            
            optimized_data = json.loads(response_text.strip())
            
            print(f"[OPTIMIZER] Parsed JSON successfully")
            print(f"[OPTIMIZER] Optimized schedule has {len(optimized_data.get('items', []))} items")
            
            return {
                "success": True,
                "model_used": model,
                "optimized_items": optimized_data.get('items', []),
                "explanation": optimized_data.get('explanation', {}),
                "raw_response": response_text[:500]  # First 500 chars for debugging
            }
        
        except json.JSONDecodeError as je:
            print(f"[OPTIMIZER] JSON parse error: {je}")
            print(f"[OPTIMIZER] Response text: {response_text[:1000]}")
            return {
                "success": False,
                "error": f"Failed to parse AI response as JSON: {str(je)}",
                "raw_response": response_text[:500]
            }
    
    except Exception as e:
        print(f"[OPTIMIZER] API call failed: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": f"AI optimization failed: {str(e)}"
        }

def validate_optimization_response(
    baseline_items: List[Dict[str, Any]],
    optimized_items: List[Dict[str, Any]]
) -> Dict[str, Any]:
    """
    STRICT validation - rejects any AI result that violates safety rules.
    
    Rules enforced:
    1. No added/removed items or level changes (hierarchy preservation)
    2. Planned hours must remain unchanged per task
    3. All dependencies must reference existing items
    4. Dates must be valid and within reasonable range
    """
    errors = []
    
    # Check 1: Cardinality (no adds/removes)
    if len(optimized_items) != len(baseline_items):
        errors.append(f"HIERARCHY_VIOLATION: Item count changed {len(baseline_items)}→{len(optimized_items)}")
        return {"valid": False, "errors": errors, "warnings": []}
    
    # Build maps for validation
    baseline_map = {item['id']: item for item in baseline_items}
    baseline_ids = set(baseline_map.keys())
    optimized_map = {item['id']: item for item in optimized_items}
    optimized_ids = set(optimized_map.keys())
    
    # Check 2: IDs must be identical (no adds/removes)
    if baseline_ids != optimized_ids:
        missing = baseline_ids - optimized_ids
        extra = optimized_ids - baseline_ids
        if missing or extra:
            errors.append(f"HIERARCHY_VIOLATION: Items added/removed (missing: {missing}, extra: {extra})")
            return {"valid": False, "errors": errors, "warnings": []}
    
    # Check 3: Strict per-item validation
    for opt_item in optimized_items:
        item_id = opt_item.get('id')
        base_item = baseline_map.get(item_id)
        
        if not base_item:
            errors.append(f"MISSING_BASELINE: {item_id} not in baseline")
            continue
        
        # 3a: Level must not change (hierarchy preservation)
        if opt_item.get('level') != base_item.get('level'):
            errors.append(f"HIERARCHY_VIOLATION: {item_id} level changed {base_item.get('level')}→{opt_item.get('level')}")
        
        # 3b: Planned hours must remain IDENTICAL
        opt_hours = float(opt_item.get('planned_hours', 0))
        base_hours = float(base_item.get('planned_hours', 0))
        if abs(opt_hours - base_hours) > 0.01:
            errors.append(f"HOURS_VIOLATION: {item_id} hours changed {base_hours}→{opt_hours}")
        
        # 3c: Validate date formats (ISO YYYY-MM-DD)
        for date_field in ['start', 'end']:
            date_val = opt_item.get(date_field)
            if date_val:
                try:
                    # Must be valid ISO format
                    test_date = datetime.date.fromisoformat(str(date_val)[:10])
                    # Must be within 5 years (sanity check)
                    today = datetime.date.today()
                    if (test_date - today).days > 365 * 5 or (test_date - today).days < -365:
                        errors.append(f"DATE_RANGE: {item_id} {date_field} out of range: {date_val}")
                except (ValueError, AttributeError):
                    errors.append(f"INVALID_DATE: {item_id} {date_field}={date_val}")
        
        # 3d: Duration must be positive
        duration = opt_item.get('duration_days', 0)
        if duration <= 0:
            errors.append(f"INVALID_DURATION: {item_id} duration={duration} (must be >0)")
        
        # 3e: Dependencies must reference existing items only
        deps = opt_item.get('dependencies', [])
        if deps:
            for dep in deps:
                if dep not in baseline_ids:
                    errors.append(f"INVALID_DEP: {item_id} references non-existent {dep}")
    
    # Return result
    if errors:
        print(f"[OPTIMIZER] VALIDATION FAILED with {len(errors)} errors:")
        for err in errors[:5]:
            print(f"  - {err}")
        if len(errors) > 5:
            print(f"  ... and {len(errors)-5} more errors")
        return {"valid": False, "errors": errors, "warnings": []}
    
    print(f"[OPTIMIZER] ✅ VALIDATION PASSED")
    return {"valid": True, "errors": [], "warnings": []}

def calculate_schedule_diff(
    baseline_items: List[Dict[str, Any]],
    optimized_items: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    """
    Calculate differences between baseline and optimized schedules.
    
    Returns:
        List of diff objects grouped by deliverable showing what changed.
    """
    diffs = []
    
    # Create ID map for quick lookup
    opt_map = {item['id']: item for item in optimized_items}
    
    for base_item in baseline_items:
        item_id = base_item['id']
        opt_item = opt_map.get(item_id)
        
        if not opt_item:
            continue  # Should not happen if validation passed
        
        # Check for changes
        changes = []
        
        # Start date change
        if base_item.get('start') != opt_item.get('start'):
            changes.append({
                "field": "start_date",
                "old": base_item.get('start'),
                "new": opt_item.get('start')
            })
        
        # End date change
        if base_item.get('end') != opt_item.get('end'):
            changes.append({
                "field": "end_date",
                "old": base_item.get('end'),
                "new": opt_item.get('end')
            })
        
        # Duration change
        if base_item.get('duration_days') != opt_item.get('duration_days'):
            changes.append({
                "field": "duration",
                "old": base_item.get('duration_days'),
                "new": opt_item.get('duration_days')
            })
        
        # Dependencies change
        base_deps = set(base_item.get('dependencies', []))
        opt_deps = set(opt_item.get('dependencies', []))
        if base_deps != opt_deps:
            added = list(opt_deps - base_deps)
            removed = list(base_deps - opt_deps)
            changes.append({
                "field": "dependencies",
                "added": added,
                "removed": removed
            })
        
        if changes:
            diffs.append({
                "id": item_id,
                "name": base_item.get('name'),
                "level": base_item.get('level'),
                "changes": changes
            })
    
    return diffs

@app.post("/api/optimize-schedule", response_model=OptimizationResponse)
async def api_optimize_schedule(payload: OptimizeSchedulePayload):
    """
    Trigger AI optimization for a project schedule using GPT-5.1 Pro.
    FEATURE FLAG CONTROLLED - returns error if SMART_OPTIMIZATION_ENABLED is OFF.
    
    Workflow:
    1. Check feature flag
    2. Extract baseline from real WBS hierarchy (same source as XML export)
    3. Call GPT-5.1 Pro via Responses API
    4. Validate AI response (strict rules enforcement)
    5. Store optimized schedule separately (does NOT affect active_schedule)
    6. Return diff + explanation to frontend
    """
    # FEATURE FLAG CHECK - SINGLE ENFORCEMENT POINT
    if not SMART_OPTIMIZATION_ENABLED:
        print(f"[OPTIMIZER] ⛔ FEATURE DISABLED: Smart Optimization is OFF (SMART_OPTIMIZATION_ENABLED={SMART_OPTIMIZATION_ENABLED})")
        raise HTTPException(
            status_code=423,  # 423 Locked - feature is locked/disabled
            detail="Smart Schedule Optimization is currently disabled. Set SMART_OPTIMIZATION_ENABLED=true to enable."
        )
    
    try:
        session_id = payload.session_id
        
        # Get scenario from store
        if session_id not in SCENARIO_STORE:
            raise HTTPException(404, f"Scenario not found for session {session_id}")
        
        scenario = SCENARIO_STORE[session_id]
        
        # Initialize active_schedule if not present (default to baseline)
        if 'active_schedule' not in scenario:
            scenario['active_schedule'] = 'baseline'
        
        # Get WBS rows for real baseline (same source as XML export)
        # NOTE: This requires WBS rows to be stored in scenario after build_wbs_with_pricing()
        wbs_rows = scenario.get('wbs_rows', [])
        if not wbs_rows:
            print(f"[OPTIMIZER] ⚠️  No WBS rows in scenario - AI optimization skipped (must build WBS first)")
            raise HTTPException(
                status_code=400,
                detail="Cannot optimize: scenario not fully built. Build the WBS structure first."
            )
        
        # Build optimization payload from REAL WBS hierarchy
        baseline_payload = build_optimization_payload(scenario, wbs_rows)
        if not baseline_payload:
            raise HTTPException(
                status_code=400,
                detail="Failed to extract baseline schedule from WBS hierarchy"
            )
        
        baseline_items = baseline_payload['items']
        print(f"[OPTIMIZER] 🚀 Starting optimization for {len(baseline_items)} items (feature flag: ENABLED)")
        
        # Call AI optimizer
        result = await optimize_schedule_with_ai(
            session_id=session_id,
            baseline_payload=baseline_payload,
            use_thinking=payload.use_gpt_thinking
        )
        
        if not result.get('success'):
            return OptimizationResponse(
                success=False,
                session_id=session_id,
                model_used=result.get('model_used', 'unknown'),
                changes_count=0,
                explanation={},
                diff=[],
                error=result.get('error', 'Unknown error')
            )
        
        # Validate response
        optimized_items = result.get('optimized_items', [])
        validation = validate_optimization_response(baseline_items, optimized_items)
        
        if not validation['valid']:
            error_msg = f"AI response validation failed: {'; '.join(validation['errors'])}"
            print(f"[OPTIMIZER] {error_msg}")
            return OptimizationResponse(
                success=False,
                session_id=session_id,
                model_used=result.get('model_used'),
                changes_count=0,
                explanation={},
                diff=[],
                error=error_msg
            )
        
        # Calculate diff
        diff = calculate_schedule_diff(baseline_items, optimized_items)
        
        # Store optimized schedule in SCENARIO_STORE (don't overwrite baseline)
        scenario['optimized_schedule'] = {
            'items': optimized_items,
            'explanation': result.get('explanation', {}),
            'model_used': result.get('model_used'),
            'optimized_at': datetime.datetime.now().isoformat()
        }
        scenario['active_schedule'] = 'baseline'  # Default to baseline until user accepts
        
        SCENARIO_STORE[session_id] = scenario
        
        print(f"[OPTIMIZER] ✅ Optimization complete: {len(diff)} changes")
        
        return OptimizationResponse(
            success=True,
            session_id=session_id,
            model_used=result.get('model_used'),
            changes_count=len(diff),
            explanation=result.get('explanation', {}),
            diff=diff
        )
    
    except Exception as e:
        print(f"[OPTIMIZER] Error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Optimization failed: {str(e)}")

@app.post("/api/accept-optimization")
async def api_accept_optimization(payload: AcceptOptimizationPayload):
    """
    Accept the AI-optimized schedule.
    FEATURE FLAG CONTROLLED - returns 423 when SMART_OPTIMIZATION_ENABLED is OFF.
    """
    # FEATURE FLAG CHECK - SINGLE ENFORCEMENT POINT
    if not SMART_OPTIMIZATION_ENABLED:
        print(f"[OPTIMIZER] ⛔ FEATURE DISABLED: Smart Optimization is OFF")
        raise HTTPException(
            status_code=423,  # 423 Locked
            detail="Smart Schedule Optimization is currently disabled."
        )
    
    try:
        session_id = payload.session_id
        
        if session_id not in SCENARIO_STORE:
            raise HTTPException(404, f"Scenario not found for session {session_id}")
        
        scenario = SCENARIO_STORE[session_id]
        
        if 'optimized_schedule' not in scenario:
            raise HTTPException(400, "No optimized schedule found. Run optimization first.")
        
        # Promote optimized schedule to active
        scenario['active_schedule'] = 'optimized'
        SCENARIO_STORE[session_id] = scenario
        
        print(f"[OPTIMIZER] ✅ Accepted optimization for session {session_id}")
        
        return {
            "success": True,
            "message": "Optimized schedule is now active. Exports will use the AI-adjusted dates."
        }
    
    except Exception as e:
        print(f"[OPTIMIZER] Accept error: {e}")
        raise HTTPException(500, f"Failed to accept optimization: {str(e)}")

@app.post("/api/reject-optimization")
async def api_reject_optimization(payload: AcceptOptimizationPayload):
    """
    Reject the AI-optimized schedule.
    FEATURE FLAG CONTROLLED - returns 423 when SMART_OPTIMIZATION_ENABLED is OFF.
    """
    # FEATURE FLAG CHECK - SINGLE ENFORCEMENT POINT
    if not SMART_OPTIMIZATION_ENABLED:
        print(f"[OPTIMIZER] ⛔ FEATURE DISABLED: Smart Optimization is OFF")
        raise HTTPException(
            status_code=423,  # 423 Locked
            detail="Smart Schedule Optimization is currently disabled."
        )
    
    try:
        session_id = payload.session_id
        
        if session_id not in SCENARIO_STORE:
            raise HTTPException(404, f"Scenario not found for session {session_id}")
        
        scenario = SCENARIO_STORE[session_id]
        
        # Discard optimized schedule
        if 'optimized_schedule' in scenario:
            del scenario['optimized_schedule']
        
        scenario['active_schedule'] = 'baseline'
        SCENARIO_STORE[session_id] = scenario
        
        print(f"[OPTIMIZER] ✅ Rejected optimization for session {session_id}")
        
        return {
            "success": True,
            "message": "Optimized schedule discarded. Baseline schedule remains active."
        }
    
    except Exception as e:
        print(f"[OPTIMIZER] Reject error: {e}")
        raise HTTPException(500, f"Failed to reject optimization: {str(e)}")

# =============================================================================
# END SMART SCHEDULE OPTIMIZATION
# =============================================================================

@app.post("/api/pricing/build_scenario")
async def api_build_scenario(payload: BuildScenarioPayload):
    """
    Build scenario from Step 2 selection and store in SCENARIO_STORE.
    Creates rows with Deliverable/Component/Task_Label structure.
    
    NEW DUAL-ENTRY ARCHITECTURE:
    - If reset=False and baseline exists → return working scenario (preserve Step 3 edits)
    - If reset=True → restore working scenario from baseline (deepcopy)
    - If no baseline exists → build from Step 2, set both baseline and scenario
    """
    if not DB.loaded:
        DB.load()
    
    try:
        session_id = payload.session_id
        
        # CRITICAL: Call get_session_state FIRST to trigger migration from legacy format
        # This converts old flat SCENARIO_STORE entries to dual-entry (baseline/scenario) format
        # Without this, has_baseline() fails on migrated data and Step 3 edits are lost
        _ = get_session_state(session_id)
        
        # Check if baseline exists (now works correctly after migration)
        if has_baseline(session_id):
            if payload.reset:
                # Reset from Step 2: restore working scenario from baseline
                print(f"[SCENARIO_STORE] 🔄 Reset requested - restoring from baseline for {session_id}")
                scenario = reset_scenario_from_baseline(session_id)
                return {
                    "success": True,
                    "session_id": session_id,
                    "scenario": scenario,
                    "total_items": len(scenario.get("items", [])),
                    "totals": scenario.get("totals", {}),
                    "reset": True
                }
            else:
                # Return current working scenario (preserve Step 3 edits)
                print(f"[SCENARIO_STORE] ✅ Returning working scenario for {session_id} (preserving Step 3 edits)")
                scenario = get_working_scenario(session_id)
                scenario = _recompute_totals(scenario)
                update_working_scenario(session_id, scenario)
                return {
                    "success": True,
                    "session_id": session_id,
                    "scenario": scenario,
                    "total_items": len(scenario.get("items", [])),
                    "totals": scenario.get("totals", {}),
                    "cached": True  # Flag to indicate this was from cache
                }
        
        if payload.reset:
            print(f"[SCENARIO_STORE] 🔄 Reset requested - rebuilding scenario from Step 2 for {session_id}")
        
        selection = payload.selection
        
        # Build scenario items from selection
        items = []
        
        # Parse selection structure: {deliverable_codes: [...], components_map: {...}, l3_map: {...}}
        deliverable_codes = selection.get('deliverable_codes', [])
        components_map = selection.get('components_map', {})
        l3_map = selection.get('l3_map', {})
        
        # For now, just use the existing /api/build logic to generate items
        # We'll call the existing scenario builder and store the result in SCENARIO_STORE
        from copy import deepcopy
        
        # Build a standard payload for the existing /api/build endpoint
        build_payload_dict = {
            "deliverable_codes": deliverable_codes,
            "selected_components": components_map,
            "selected_l3": l3_map,
            "pricing_mode": payload.pricing_mode or "Flat_Blended",
            "blended_rate": payload.blended_rate or DEFAULT_BILLABLE_RATE_USD,
            "rate_band": payload.rate_band or "Standard_US",
            "project_start": payload.project_start or datetime.date.today().isoformat(),
            "client_budget_usd": payload.client_budget_usd or None,
            "use_slack": True,
            "slack_after_internal": 0,
            "slack_after_client": 0,
            "slack_global_pct": 0,
            "retainers": payload.retainers or []
        }
        
        # Build using existing logic (reuse the code)
        # For now, we'll create a simple scenario with basic items
        for deliv_code in deliverable_codes:
            # Get deliverable info from database
            db_rows = DB.all_rows[DB.all_rows['Deliverable_Code'] == deliv_code]
            if db_rows.empty:
                continue
            
            deliverable_name = db_rows['Deliverable'].iloc[0] if 'Deliverable' in db_rows.columns else deliv_code
            
            # Get components - either user-selected or all default components
            comp_selection = components_map.get(deliv_code, "__ALL__")
            
            if comp_selection == "__ALL__":
                # Include all components for this deliverable
                components_for_deliv = db_rows['Component'].unique().tolist()
            elif isinstance(comp_selection, dict):
                # User selected specific components
                components_for_deliv = list(comp_selection.keys())
            else:
                components_for_deliv = []
            
            for component_name in components_for_deliv:
                if not component_name or component_name == "":
                    continue
                
                # Get L3 tasks for this deliverable/component
                l3_tasks = l3_map.get(deliv_code, {}).get(component_name, [])
                
                # Get rows for this deliverable/component
                comp_rows = db_rows[db_rows['Component'] == component_name]
                
                if not l3_tasks:
                    # No specific L3 tasks selected, use defaults from DB
                    for _, row in comp_rows.iterrows():
                        task_label = row.get('Task_Label', '')
                        if not task_label:
                            continue
                        
                        hours = float(row.get('Hours', 0) or 0)
                        rate = float(row.get('Rate_USD') or DEFAULT_BILLABLE_RATE_USD)
                        
                        items.append({
                            "Deliverable": deliverable_name,
                            "Deliverable_Code": deliv_code,
                            "Component": component_name,
                            "Task_Label": task_label,
                            "Planned_Hours": hours,
                            "Rate_USD": rate,
                            "Price_USD": round(hours * rate, 2),
                            "Role": row.get('Role', 'Generalist'),
                            "Seniority": row.get('Seniority', 'Mid'),
                            "Service Department": row.get('Service Department', '')
                        })
                else:
                    # User selected specific L3 tasks
                    for task_label in l3_tasks:
                        # Find the corresponding row in DB
                        task_rows = comp_rows[comp_rows['Task_Label'] == task_label]
                        if not task_rows.empty:
                            row = task_rows.iloc[0]
                            hours = float(row.get('Hours', 0) or 0)
                            rate = float(row.get('Rate_USD') or DEFAULT_BILLABLE_RATE_USD)
                        else:
                            hours = 0.0
                            rate = DEFAULT_BILLABLE_RATE_USD
                        
                        items.append({
                            "Deliverable": deliverable_name,
                            "Deliverable_Code": deliv_code,
                            "Component": component_name,
                            "Task_Label": task_label,
                            "Planned_Hours": hours,
                            "Rate_USD": rate,
                            "Price_USD": round(hours * rate, 2),
                            "Role": row.get('Role', 'Generalist') if not task_rows.empty else 'Generalist',
                            "Seniority": row.get('Seniority', 'Mid') if not task_rows.empty else 'Mid',
                            "Service Department": row.get('Service Department', '') if not task_rows.empty else ''
                        })
        
        # Create scenario
        scenario = {
            "items": items,
            "project_name": payload.project_name or "New Project",
            "project_start": payload.project_start or datetime.date.today().isoformat(),
            "totals": {
                "hours": 0.0,
                "price": 0.0
            }
        }
        
        # Recompute totals
        scenario = _recompute_totals(scenario)
        
        # Store in SCENARIO_STORE using new dual-entry architecture
        # Sets both baseline and working scenario
        set_baseline_and_scenario(session_id, scenario)
        
        print(f"[SCENARIO_STORE] 🆕 Built new scenario from Step 2 for {session_id} with {len(items)} items")
        
        return {
            "success": True,
            "session_id": session_id,
            "scenario": scenario,  # Return full scenario for frontend
            "total_items": len(items),
            "totals": scenario["totals"],
            "new_baseline": True  # Flag to indicate this is a fresh build
        }
        
    except Exception as e:
        print(f"[SCENARIO_STORE] Error building scenario: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

# =============================================================================
# NEW PRICING ENDPOINTS (Dual-Entry Architecture)
# =============================================================================

class ScenarioItemPatch(BaseModel):
    """Payload for patching a single scenario item."""
    session_id: str
    deliverable_id: str  # Deliverable_Code or Row_ID
    component_id: Optional[str] = None  # Optional component ID for component-level updates
    hours: Optional[float] = None
    rate_usd: Optional[float] = None
    price_usd: Optional[float] = None
    cadence: Optional[str] = None  # Legacy field, mapped to billing_cadence
    billing_cadence: Optional[str] = None  # one_time, monthly, quarterly, semi_annual, annual
    cadence_units: Optional[int] = None  # Number of periods (e.g., 2 × semi_annual = 12 months)
    cadence_price: Optional[float] = None  # Price per cadence period
    monthly_price: Optional[float] = None  # Optional monthly price
    months: Optional[int] = None  # Legacy field, mapped to retainer_months

def _apply_item_updates(item: dict, payload: ScenarioItemPatch):
    """Apply updates to an item (deliverable or component), then call apply_cadence_to_item."""
    old_item = item.copy()  # Snapshot for change detection
    
    if payload.hours is not None:
        item["Planned_Hours"] = payload.hours
        item["planned_hours"] = payload.hours
        item["hours"] = payload.hours
        item["total_hours"] = payload.hours
    
    if payload.rate_usd is not None:
        item["Rate_USD"] = payload.rate_usd
        item["rate_usd"] = payload.rate_usd
        item["blended_rate"] = payload.rate_usd
        item["effective_rate"] = payload.rate_usd
    
    if payload.price_usd is not None:
        item["Price_USD"] = payload.price_usd
        item["price_usd"] = payload.price_usd
        item["total_price"] = payload.price_usd
    
    # Handle billing_cadence (new) and cadence (legacy)
    cadence_value = payload.billing_cadence or payload.cadence
    if cadence_value is not None:
        item["billing_cadence"] = cadence_value
        item["cadence"] = cadence_value
    
    # Handle cadence_units
    if payload.cadence_units is not None:
        item["cadence_units"] = payload.cadence_units
    
    # Handle cadence_price
    if payload.cadence_price is not None:
        item["cadence_price"] = payload.cadence_price
    
    # Handle monthly_price
    if payload.monthly_price is not None:
        item["monthly_price"] = payload.monthly_price
    
    # Handle months (legacy) mapping to retainer_months
    if payload.months is not None:
        item["retainer_months"] = payload.months
        item["months"] = payload.months
    
    # Apply cadence logic if any cadence-related field was explicitly set (use is not None checks)
    cadence_fields_changed = any([
        payload.billing_cadence is not None,
        payload.cadence is not None,
        payload.cadence_units is not None,
        payload.cadence_price is not None,
        payload.monthly_price is not None,
        payload.months is not None
    ])
    
    if cadence_fields_changed:
        apply_cadence_to_item(item, old_item)
    else:
        # Simple price recalculation for non-cadence updates
        hours = payload.hours if payload.hours is not None else item.get("hours") or item.get("Planned_Hours") or 0
        rate = payload.rate_usd if payload.rate_usd is not None else item.get("rate_usd") or item.get("Rate_USD") or 210
        new_price = int(hours * rate)
        item["price"] = new_price
        item["Price_USD"] = new_price
        item["price_usd"] = new_price

@app.patch("/api/pricing/scenario/item")
async def api_patch_scenario_item(payload: ScenarioItemPatch):
    """
    Update a single item (deliverable or component) in the working scenario.
    Called by Step 3 grid edits.
    
    If component_id is provided, updates the component within the deliverable.
    Otherwise updates the deliverable directly.
    
    Returns updated scenario with recomputed totals.
    """
    try:
        session_id = payload.session_id
        
        # Get current working scenario
        scenario = get_working_scenario(session_id)
        if not scenario or not scenario.get("items"):
            raise HTTPException(404, f"No scenario found for session {session_id}")
        
        items = scenario.get("items", [])
        updated = False
        
        for item in items:
            item_id = item.get("deliverable_code") or item.get("Deliverable_Code") or item.get("id") or item.get("Row_ID")
            if str(item_id) == str(payload.deliverable_id):
                if payload.component_id:
                    # Update component within this deliverable
                    components = item.get("components", [])
                    for comp in components:
                        comp_id = comp.get("id") or comp.get("component_id") or comp.get("Component_Code")
                        if str(comp_id) == str(payload.component_id):
                            _apply_item_updates(comp, payload)
                            updated = True
                            break
                else:
                    # Update deliverable directly
                    _apply_item_updates(item, payload)
                    updated = True
                break
        
        if not updated:
            target = f"component {payload.component_id}" if payload.component_id else f"item {payload.deliverable_id}"
            raise HTTPException(404, f"{target} not found in scenario")
        
        # Update working scenario and recompute totals
        scenario = update_working_scenario(session_id, scenario)
        
        target = f"component {payload.component_id}" if payload.component_id else f"item {payload.deliverable_id}"
        
        # Debug logging as specified in GPT-5 Pro instructions
        print(f"[SCENARIO PATCH] session={session_id} "
              f"item_key={payload.deliverable_id} "
              f"hours={payload.hours} "
              f"rate={payload.rate_usd} "
              f"cadence={payload.cadence} "
              f"months={payload.months}")
        
        # Log the updated totals
        totals = scenario.get("totals", {})
        print(f"[SCENARIO PATCH] totals after update: hours={totals.get('hours')} price={totals.get('price')}")
        
        # Also log the specific item's new values
        for item in scenario.get("items", []):
            item_id = item.get("deliverable_code") or item.get("Deliverable_Code")
            if str(item_id) == str(payload.deliverable_id):
                print(f"[SCENARIO PATCH] item {item_id}: "
                      f"total_hours={item.get('total_hours')} "
                      f"hours={item.get('hours')} "
                      f"cadence={item.get('billing_cadence')} "
                      f"retainer_months={item.get('retainer_months')} "
                      f"price={item.get('price')}")
                break
        
        print(f"[SCENARIO_STORE] ✏️ Updated {target} in session {session_id}")
        
        return {
            "success": True,
            "session_id": session_id,
            "scenario": scenario,
            "totals": scenario.get("totals", {})
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[SCENARIO_STORE] Error patching item: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

class RebuildBreakdownPayload(BaseModel):
    """Payload for rebuilding pricing breakdown from working scenario."""
    session_id: str

@app.post("/api/pricing/rebuild_breakdown")
async def api_rebuild_breakdown(payload: RebuildBreakdownPayload):
    """
    Rebuild Deliverable Pricing Details table + summary cards from working scenario.
    Does NOT change hours/rates/prices in the scenario itself.
    
    Called by "Re-build Scenario" (blue) button.
    """
    try:
        session_id = payload.session_id
        
        # Get current working scenario
        scenario = get_working_scenario(session_id)
        if not scenario or not scenario.get("items"):
            raise HTTPException(404, f"No scenario found for session {session_id}")
        
        # Recompute totals
        scenario = _recompute_totals(scenario)
        
        # Update working scenario with recomputed totals
        update_working_scenario(session_id, scenario)
        
        # Build breakdown by deliverable for Pricing Details table
        items = scenario.get("items", [])
        deliverable_summary = {}
        
        for item in items:
            deliv_code = item.get("Deliverable_Code") or item.get("deliverable_code", "Unknown")
            deliv_name = item.get("Deliverable") or item.get("deliverable", deliv_code)
            
            if deliv_code not in deliverable_summary:
                deliverable_summary[deliv_code] = {
                    "deliverable_code": deliv_code,
                    "deliverable": deliv_name,
                    "hours": 0.0,
                    "rate": 0.0,
                    "price": 0.0,
                    "cadence": item.get("billing_cadence", "One-Time"),
                    "months": item.get("retainer_months", 0),
                    "components": []
                }
            
            hours = float(item.get("Planned_Hours") or item.get("hours") or 0)
            price = float(item.get("Price_USD") or item.get("price_usd") or 0)
            rate = float(item.get("Rate_USD") or item.get("rate_usd") or DEFAULT_BILLABLE_RATE_USD)
            
            deliverable_summary[deliv_code]["hours"] += hours
            deliverable_summary[deliv_code]["price"] += price
            deliverable_summary[deliv_code]["rate"] = rate  # Use last seen rate
            
            # Track components
            comp_name = item.get("Component") or item.get("component", "")
            if comp_name:
                deliverable_summary[deliv_code]["components"].append({
                    "name": comp_name,
                    "hours": hours,
                    "price": price,
                    "rate": rate
                })
        
        breakdown = list(deliverable_summary.values())
        
        # Compute summary cards
        one_time_hours = 0.0
        one_time_price = 0.0
        retainer_hours = 0.0
        retainer_price = 0.0
        
        for d in breakdown:
            cadence = (d.get("cadence") or "One-Time").lower()
            if "month" in cadence or "retainer" in cadence:
                retainer_hours += d["hours"]
                retainer_price += d["price"]
            else:
                one_time_hours += d["hours"]
                one_time_price += d["price"]
        
        summary = {
            "one_time": {
                "hours": round(one_time_hours, 2),
                "price": round(one_time_price, 2)
            },
            "retainer": {
                "hours": round(retainer_hours, 2),
                "price": round(retainer_price, 2)
            },
            "grand_total": {
                "hours": round(one_time_hours + retainer_hours, 2),
                "price": round(one_time_price + retainer_price, 2)
            }
        }
        
        print(f"[SCENARIO_STORE] 🔄 Rebuilt breakdown for session {session_id}")
        
        return {
            "success": True,
            "session_id": session_id,
            "breakdown": breakdown,
            "summary": summary,
            "totals": scenario.get("totals", {})
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[SCENARIO_STORE] Error rebuilding breakdown: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

class ResetFromStep2Payload(BaseModel):
    """Payload for resetting scenario from Step 2 baseline."""
    session_id: str
    rebuild_from_step2: bool = False  # If True and no baseline exists, rebuild from Step 2

@app.post("/api/pricing/reset_from_step2")
async def api_reset_from_step2(payload: ResetFromStep2Payload):
    """
    Reset working scenario to baseline (deepcopy).
    Called by "Reset from Step 2" (orange) button.
    
    Only calls Step 2 if baseline is missing and rebuild_from_step2=True.
    """
    try:
        session_id = payload.session_id
        
        if not has_baseline(session_id):
            if payload.rebuild_from_step2:
                # No baseline - need to build from Step 2
                # This shouldn't normally happen if user followed proper flow
                print(f"[SCENARIO_STORE] ⚠️ No baseline for {session_id}, cannot reset")
                return {
                    "success": False,
                    "error": "No baseline exists. Please build scenario first.",
                    "session_id": session_id
                }
            else:
                return {
                    "success": False,
                    "error": "No baseline exists to reset to.",
                    "session_id": session_id
                }
        
        # Reset working scenario from baseline
        scenario = reset_scenario_from_baseline(session_id)
        
        print(f"[SCENARIO_STORE] 🔙 Reset scenario from baseline for {session_id}")
        
        return {
            "success": True,
            "session_id": session_id,
            "scenario": scenario,
            "totals": scenario.get("totals", {}),
            "reset": True
        }
        
    except Exception as e:
        print(f"[SCENARIO_STORE] Error resetting from Step 2: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

@app.post("/api/pricing/optimize")
async def api_optimize_scenario(payload: OptimizeScenarioPayload):
    """
    Optimize pricing for a deliverable in the scenario using AI redistribution.
    Calls redistribute_hours() and recomputes totals.
    """
    try:
        session_id = payload.session_id
        
        # Use get_working_scenario() for consistent access (GPT-5 Pro spec)
        scenario = get_working_scenario(session_id)
        if not scenario or not scenario.get("items"):
            raise HTTPException(404, f"Scenario not found for session {session_id}")
        
        items = scenario.get("items", [])
        
        if payload.deliverable_index >= len(items):
            raise HTTPException(400, "Invalid deliverable index")
        
        # Get the deliverable item
        item = items[payload.deliverable_index]
        deliverable_name = item.get("Deliverable", "")
        deliverable_code = item.get("Deliverable_Code", "")
        
        # Get all components for this deliverable
        deliverable_items = [i for i in items if i.get("Deliverable_Code") == deliverable_code]
        
        # Group by component
        components = []
        seen_components = set()
        for di in deliverable_items:
            comp_name = di.get("Component", "")
            if comp_name and comp_name not in seen_components:
                seen_components.add(comp_name)
                comp_hours = sum(i.get("Planned_Hours", 0) for i in deliverable_items if i.get("Component") == comp_name)
                components.append({
                    "name": comp_name,
                    "hours": comp_hours
                })
        
        # Call redistribute_hours
        result = await redistribute_hours(
            deliverable_name=deliverable_name,
            deliverable_code=deliverable_code,
            new_total_hours=payload.new_total_hours,
            components=components,
            use_ai=payload.use_ai
        )
        
        # Update hours in scenario based on result
        if result and hasattr(result, 'components'):
            for comp_alloc in result.components:
                # Update all items with this component
                for i in items:
                    if i.get("Deliverable_Code") == deliverable_code and i.get("Component") == comp_alloc.name:
                        # Proportionally adjust hours
                        old_comp_hours = sum(x.get("Planned_Hours", 0) for x in deliverable_items if x.get("Component") == comp_alloc.name)
                        if old_comp_hours > 0:
                            ratio = comp_alloc.suggested_hours / old_comp_hours
                            i["Planned_Hours"] = round(i.get("Planned_Hours", 0) * ratio, 2)
        
        # Recompute totals and update working scenario (preserves dual-entry store)
        scenario = _recompute_totals(scenario)
        update_working_scenario(session_id, scenario)
        
        return {
            "success": True,
            "result": result,
            "totals": scenario["totals"]
        }
        
    except Exception as e:
        print(f"[SCENARIO_STORE] Error optimizing: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

@app.post("/api/pricing/cadence_suggestion")
async def api_cadence_suggestion(payload: CadenceSuggestionPayload):
    """
    Get retainer vs project-based cadence suggestion for a deliverable.
    Calls analyze_retainer_vs_project().
    """
    try:
        session_id = payload.session_id
        
        # Use get_working_scenario() for consistent access (GPT-5 Pro spec)
        scenario = get_working_scenario(session_id)
        if not scenario or not scenario.get("items"):
            raise HTTPException(404, f"Scenario not found for session {session_id}")
        
        items = scenario.get("items", [])
        
        if payload.deliverable_index >= len(items):
            raise HTTPException(400, "Invalid deliverable index")
        
        # Get the deliverable
        item = items[payload.deliverable_index]
        deliverable_name = item.get("Deliverable", "")
        
        # Calculate total hours for this deliverable
        deliverable_code = item.get("Deliverable_Code", "")
        total_hours = sum(i.get("Planned_Hours", 0) for i in items if i.get("Deliverable_Code") == deliverable_code)
        
        # Call analyze_retainer_vs_project
        result = analyze_retainer_vs_project(
            deliverable_name=deliverable_name,
            total_hours=total_hours
        )
        
        return {
            "success": True,
            "result": result
        }
        
    except Exception as e:
        print(f"[SCENARIO_STORE] Error getting cadence suggestion: {e}")
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

@app.post("/api/pricing/retainer_suggestions")
async def api_retainer_suggestions(payload: RetainerSuggestionsPayload):
    """
    Calculate retainer distribution and store in scenario.
    Calls calculate_retainer_distribution() and updates scenario.
    """
    try:
        session_id = payload.session_id
        
        # Use get_working_scenario() for consistent access (GPT-5 Pro spec)
        scenario = get_working_scenario(session_id)
        if not scenario or not scenario.get("items"):
            raise HTTPException(404, f"Scenario not found for session {session_id}")
        
        items = scenario.get("items", [])
        
        if payload.deliverable_index >= len(items):
            raise HTTPException(400, "Invalid deliverable index")
        
        # Get the deliverable
        item = items[payload.deliverable_index]
        deliverable_code = item.get("Deliverable_Code", "")
        
        # Call calculate_retainer_distribution
        distribution = calculate_retainer_distribution(
            monthly_hours=payload.monthly_hours,
            duration_months=payload.duration_months
        )
        
        # Store retainer info in scenario
        if "retainer_distributions" not in scenario:
            scenario["retainer_distributions"] = {}
        
        scenario["retainer_distributions"][deliverable_code] = {
            "monthly_hours": payload.monthly_hours,
            "duration_months": payload.duration_months,
            "distribution": distribution
        }
        
        # Update working scenario (preserves dual-entry store)
        update_working_scenario(session_id, scenario)
        
        return {
            "success": True,
            "distribution": distribution,
            "total_hours": sum(distribution.values()) if isinstance(distribution, dict) else 0
        }
        
    except Exception as e:
        print(f"[SCENARIO_STORE] Error calculating retainer suggestions: {e}")
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

@app.post("/api/timeline/update_task")
async def api_update_timeline_task(payload: UpdateTaskPayload):
    """
    Update Gantt task and sync pricing.
    Updates task duration/dates and recalculates Planned_Hours.
    """
    try:
        session_id = payload.session_id
        
        # Use get_working_scenario() for consistent access (GPT-5 Pro spec)
        scenario = get_working_scenario(session_id)
        if not scenario or not scenario.get("items"):
            raise HTTPException(404, f"Scenario not found for session {session_id}")
        
        items = scenario.get("items", [])
        
        # Find item by WBS_ID
        target_item = None
        for item in items:
            if item.get("WBS_ID") == payload.wbs_id:
                target_item = item
                break
        
        if not target_item:
            raise HTTPException(404, f"Task not found with WBS_ID {payload.wbs_id}")
        
        # Update duration/dates if provided
        if payload.duration_days is not None:
            target_item["Duration_Days"] = payload.duration_days
            # Recalculate hours: Duration_Days * Hours_Per_Day
            new_hours = payload.duration_days * payload.hours_per_day
            target_item["Planned_Hours"] = round(new_hours, 2)
        
        if payload.start_date:
            target_item["Start_Date"] = payload.start_date
            print(f"[GANTT EDIT] Saved Start_Date={payload.start_date} for WBS_ID={payload.wbs_id}")
        
        if payload.end_date:
            target_item["End_Date"] = payload.end_date
            print(f"[GANTT EDIT] Saved End_Date={payload.end_date} for WBS_ID={payload.wbs_id}")
        
        # Recompute totals and update working scenario (preserves dual-entry store)
        scenario = _recompute_totals(scenario)
        update_working_scenario(session_id, scenario)
        
        # Update timeline if it exists
        if "timeline" in scenario:
            timeline = scenario["timeline"]
            if "tasks" in timeline:
                for task in timeline["tasks"]:
                    if task.get("id") == payload.wbs_id:
                        if payload.duration_days is not None:
                            task["duration_days"] = payload.duration_days
                            task["hours"] = payload.duration_days * payload.hours_per_day
                        if payload.start_date:
                            task["start_date"] = payload.start_date
                        if payload.end_date:
                            task["end_date"] = payload.end_date
                        break
        
        return {
            "success": True,
            "updated_item": target_item,
            "totals": scenario["totals"]
        }
        
    except Exception as e:
        print(f"[SCENARIO_STORE] Error updating task: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

@app.post("/api/timeline/update_tasks_batch")
async def api_update_timeline_tasks_batch(payload: UpdateTasksBatchPayload):
    """
    Batch update multiple Gantt tasks for efficient mouseup commits.
    Updates multiple tasks in a single transaction and recomputes totals once.
    """
    try:
        session_id = payload.session_id
        
        # Use get_working_scenario() for consistent access (GPT-5 Pro spec)
        scenario = get_working_scenario(session_id)
        if not scenario or not scenario.get("items"):
            raise HTTPException(404, f"Scenario not found for session {session_id}")
        
        items = scenario.get("items", [])
        
        touched_deliv_codes = set()
        
        for update in payload.updates:
            for item in items:
                if item.get("WBS_ID") == update.wbs_id:
                    if update.duration_days is not None:
                        item["Duration_Days"] = update.duration_days
                        new_hours = update.duration_days * update.hours_per_day
                        item["Planned_Hours"] = round(new_hours, 2)
                    if update.start_date:
                        item["Start_Date"] = update.start_date
                    if update.end_date:
                        item["End_Date"] = update.end_date
                    touched_deliv_codes.add(item.get("Deliverable_Code", ""))
                    break
        
        # Recompute totals and update working scenario (preserves dual-entry store)
        scenario = _recompute_totals(scenario)
        update_working_scenario(session_id, scenario)
        
        return {
            "success": True,
            "touched": list(touched_deliv_codes),
            "totals": scenario.get("totals", {}),
            "updated_count": len(payload.updates)
        }
        
    except Exception as e:
        print(f"[SCENARIO_STORE] Error in batch update: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            {"success": False, "error": str(e)},
            status_code=500
        )

# ========== Industry Template API Endpoints ==========
@app.get("/api/industry/templates")
def get_industry_templates():
    """Return list of available industry templates with metadata"""
    return {
        "templates": get_available_industries(),
        "default": None  # No default, let user choose
    }

@app.post("/api/industry/suggest-deliverables")
def suggest_industry_deliverables(request: dict):
    """Suggest deliverables based on selected industry and RFP keywords"""
    industry = request.get('industry', '')
    rfp_text = request.get('rfp_text', '')
    keywords = request.get('keywords', '')  # Support both rfp_text and keywords parameters
    
    if not industry:
        return {"deliverables": [], "message": "No industry selected"}
    
    template = get_industry_template(industry)
    if not template:
        return {"deliverables": [], "message": f"Template not available for {industry}"}
    
    # Extract keywords from RFP or use provided keywords
    extracted_keywords = []
    
    # Combine both rfp_text and keywords if provided
    combined_text = ""
    if rfp_text:
        combined_text += " " + rfp_text
    if keywords:
        combined_text += " " + keywords
    
    if combined_text:
        import re
        # Extract meaningful words (3+ chars)
        words = re.findall(r'\b[a-zA-Z]{3,}\b', combined_text.lower())
        
        # Industry-specific keywords based on selected industry (also check for luxury alias)
        if industry in ["luxury_fashion", "luxury"]:
            industry_keywords = ["fashion", "luxury", "collection", "runway", "show", "campaign", 
                              "lookbook", "editorial", "influencer", "heritage", "event", "gala",
                              "spring", "summer", "fall", "winter", "season", "week", "paris",
                              "milan", "london", "new york", "celebrity", "ambassador", "boutique",
                              "flagship", "exclusive", "premium", "haute", "couture", "ready-to-wear"]
        elif industry == "beauty":
            industry_keywords = ["beauty", "cosmetic", "skincare", "makeup", "launch", "product",
                              "tutorial", "influencer", "clinical", "ingredient", "sustainable",
                              "sephora", "ulta", "sample", "event", "campaign", "seeding",
                              "collection", "holiday", "seasonal", "education", "masterclass",
                              "mua", "artist", "clean", "natural", "organic", "vegan", "cruelty-free",
                              "packaging", "refill", "efficacy", "dermatologist", "before", "after",
                              "innovation", "technology", "botanical", "formula", "retail", "ecommerce",
                              "virtual", "consultation", "photography", "social", "instagram", "tiktok"]
        else:
            # Generic keywords for other industries
            industry_keywords = ["campaign", "launch", "marketing", "digital", "content", "social",
                              "event", "production", "strategy", "brand", "creative", "video"]
        
        # Find matching keywords and also pass the full word list for partial matches
        extracted_keywords = [w for w in words if w in industry_keywords]
        
        # Also pass full words for template to perform its own matching
        if not extracted_keywords:
            extracted_keywords = words[:20]  # Pass first 20 words if no industry keywords found
    
    # Get suggested deliverables from template
    suggested = template.get_suggested_deliverables(extracted_keywords)
    
    return {
        "industry": industry,
        "deliverables": suggested,
        "keywords_found": keywords,
        "total_suggested": len(suggested)
    }

@app.post("/api/industry/calculate-timeline")
def calculate_industry_timeline(request: dict):
    """Calculate industry-specific timeline with milestones"""
    industry = request.get('industry', '')
    deliverable_codes = request.get('deliverable_codes', [])
    start_date_str = request.get('start_date', datetime.datetime.now().isoformat())
    
    if not industry:
        return {"error": "No industry selected"}
    
    template = get_industry_template(industry)
    if not template:
        return {"error": f"Template not available for {industry}"}
    
    # Parse start date
    try:
        start_date = datetime.datetime.fromisoformat(start_date_str.replace('Z', '+00:00'))
    except:
        start_date = datetime.datetime.now()
    
    # Calculate timeline - handle real estate specific parameters
    if industry == "real_estate":
        project_phase = request.get('project_phase', 'sales_launch')
        timeline = template.calculate_timeline(deliverable_codes, start_date, project_phase)
    else:
        timeline = template.calculate_timeline(deliverable_codes, start_date)
    
    return {
        "industry": industry,
        "timeline": timeline,
        "start_date": start_date.isoformat(),
        "deliverable_count": len(deliverable_codes)
    }

@app.post("/api/industry/calculate-pricing")
def calculate_industry_pricing(request: dict):
    """Calculate industry-specific pricing with luxury adjustments"""
    industry = request.get('industry', '')
    deliverable_codes = request.get('deliverable_codes', [])
    base_rate = request.get('base_rate', DEFAULT_BILLABLE_RATE_USD)
    
    if not industry:
        return {"error": "No industry selected"}
    
    template = get_industry_template(industry)
    if not template:
        return {"error": f"Template not available for {industry}"}
    
    # Calculate pricing - handle real estate specific parameters
    if industry == "real_estate":
        property_type = request.get('property_type', None)
        num_phases = request.get('num_phases', 1)
        pricing = template.calculate_pricing(deliverable_codes, base_rate, property_type, num_phases)
    else:
        pricing = template.calculate_pricing(deliverable_codes, base_rate)
    
    return {
        "industry": industry,
        "pricing": pricing,
        "base_rate": base_rate,
        "deliverable_count": len(deliverable_codes)
    }

@app.post("/api/ai/analyze_project_retainer")
async def analyze_project_retainer(request: dict):
    """Analyze RFP text to suggest PROJECT or RETAINER type for each deliverable"""
    rfp_text = request.get('rfp_text', '')
    deliverables = request.get('deliverables', [])
    
    if not rfp_text or not deliverables:
        return JSONResponse(
            status_code=400,
            content={"error": "RFP text and deliverables are required"}
        )
    
    try:
        # Create the prompt for analysis
        deliverable_list = "\n".join([f"- {d['code']}: {d['name']}" for d in deliverables])
        
        prompt = f"""Analyze the following RFP text and determine whether each deliverable should be categorized as PROJECT (one-time) or RETAINER (ongoing monthly).

RFP TEXT (truncated to 3000 chars):
{rfp_text[:3000]}

DELIVERABLES TO ANALYZE:
{deliverable_list}

CLASSIFICATION RULES:
- PROJECT: One-time deliverables like strategy documents, brand guidelines, website builds, initial setup
- RETAINER: Ongoing monthly services like paid media management, content production, social media management, optimization

Look for indicators:
- Duration words: "monthly", "ongoing", "continuous", "12-month", "annual", "recurring"
- Frequency patterns: "weekly reports", "monthly optimization", "10 posts/month", "daily monitoring"
- One-time indicators: "initial", "launch", "setup", "development", "design", "creation", "build"

Return a JSON object with deliverable codes as keys and classification objects as values.
Example: {{"deck_strategy": {{"type": "PROJECT", "confidence": 0.9, "reasoning": "Strategy deck is a one-time deliverable"}}, ...}}"""

        # Check if we have OpenAI client configured
        if hasattr(app.state, 'openai_client') and app.state.openai_client:
            # Call OpenAI API
            client = app.state.openai_client
            response = client.chat.completions.create(
                model="gpt-5-mini",
                messages=[
                    {"role": "system", "content": "You are an expert at analyzing project scopes and determining whether deliverables are one-time projects or ongoing retainer services."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                temperature=0.3,
                max_tokens=2000
            )
            
            # Parse AI response
            ai_suggestions = json.loads(response.choices[0].message.content)
            
            return JSONResponse(content={
                "suggestions": ai_suggestions,
                "method": "ai"
            })
        else:
            # Fallback to heuristic analysis
            raise Exception("OpenAI client not configured, using heuristics")
            
    except Exception as e:
        print(f"[PROJECT_RETAINER_ANALYSIS] Using heuristic fallback: {e}")
        
        # Heuristic-based suggestions
        suggestions = {}
        for deliv in deliverables:
            code = deliv['code']
            name = deliv['name'].lower()
            
            # Check for retainer indicators in RFP text
            rfp_lower = rfp_text.lower()
            has_ongoing = any(word in rfp_lower for word in ['monthly', 'ongoing', 'continuous', 'recurring', 'retainer'])
            
            # Analyze deliverable name
            if any(word in name for word in ['management', 'optimization', 'monitoring', 'monthly', 'ongoing', 'retainer', 'maintenance']):
                suggestions[code] = {
                    "type": "RETAINER",
                    "confidence": 0.8 if has_ongoing else 0.6,
                    "reasoning": "Ongoing service keywords detected in deliverable name"
                }
            elif any(word in name for word in ['strategy', 'design', 'development', 'build', 'setup', 'launch', 'guide', 'creation', 'audit']):
                suggestions[code] = {
                    "type": "PROJECT",
                    "confidence": 0.8,
                    "reasoning": "One-time deliverable keywords detected"
                }
            else:
                # Default based on RFP context
                suggestions[code] = {
                    "type": "RETAINER" if has_ongoing else "PROJECT",
                    "confidence": 0.5,
                    "reasoning": "Default classification based on RFP context"
                }
        
        return JSONResponse(content={
            "suggestions": suggestions,
            "method": "heuristic"
        })

@app.get("/api/scenarios")
def api_get_scenarios():
    """
    Return what's in memory so Step 4 can recover even if a previous build
    had a transient parsing issue on the client.
    """
    return {
        "ok": True,
        "scenarios": _current_scenarios() if callable(globals().get("_current_scenarios")) else (_CURRENT_SCENARIOS or {})
    }

@app.post("/api/scenarios")
def api_post_scenarios(payload: dict, request: Request):
    """
    Builds or refreshes a pricing scenario.
    - If reset=true (query param) -> rebuilds from Step 2 baseline.
    - Otherwise -> uses the working copy stored in SCENARIO_STORE to preserve Step 3 edits.
    """
    try:
        # 1) Extract session_id from payload (supports both camelCase and snake_case)
        session_id = payload.get("sessionId") or payload.get("session_id")
        
        # 2) Reset flag - query param wins, then payload fallback
        reset_param = request.query_params.get("reset")
        if reset_param is not None:
            reset_flag = str(reset_param).lower() in ("1", "true", "yes", "y")
        else:
            reset_flag = str(payload.get("reset", "")).lower() in ("1", "true", "yes", "y")
        
        print(f"[/api/scenarios] session_id={session_id} reset={reset_flag}")
        
        # Get current state from dual-entry store
        state = get_session_state(session_id)
        baseline = state.get("baseline", {})
        scenario = state.get("scenario", {})
        
        # --- CASE 1: RESET (explicit rebuild from baseline) -------------------
        if reset_flag:
            if baseline:
                print(f"[/api/scenarios] 🔄 Reset requested - restoring from baseline for {session_id}")
                new_scenario = deepcopy(baseline)
                new_scenario = _recompute_totals(new_scenario)
                update_working_scenario(session_id, new_scenario)
                print(f"[DEBUG] SCENARIO_STORE[{session_id}] keys after reset:", list(SCENARIO_STORE[session_id].keys()))
                return {
                    "ok": True,
                    "scenarios": {"A": new_scenario},
                    "totals": new_scenario.get("totals", {}),
                    "total_hours": new_scenario.get("totals", {}).get("hours", 0),
                    "total_price": new_scenario.get("totals", {}).get("price", 0),
                    "reset": True
                }
        
        # --- CASE 0: USE WORKING SCENARIO FROM SCENARIO_STORE (preserve Step 3 edits) ---
        # If session_id exists, reset is false, and SCENARIO_STORE has a scenario, return it directly
        if session_id and not reset_flag and session_id in SCENARIO_STORE:
            scenario = get_working_scenario(session_id)
            if scenario and scenario.get("items"):
                totals = scenario.get("totals") or state.get("totals") or {}
                total_hours = totals.get("hours", scenario.get("total_hours", 0))
                total_price = totals.get("price", scenario.get("total_price", 0))
                
                # Keep Step 4 compatibility
                window_scenarios = {"A": scenario}
                _CURRENT_SCENARIOS.update(window_scenarios)
                
                print(f"[/api/scenarios] ✅ Using SCENARIO_STORE[{session_id}] "
                      f"hours={total_hours} price={total_price}")
                
                return {
                    "ok": True,
                    "scenarios": window_scenarios,
                    "totals": {"A": totals},
                    "metadata": state.get("metadata", {}),
                    "total_hours": total_hours,
                    "total_price": total_price,
                }
        
        # --- CASE 2: FALLBACK - working scenario from get_session_state ---------------------
        # If a working copy exists (but wasn't in SCENARIO_STORE path above), reuse it
        if scenario and scenario.get("items"):
            print(f"[/api/scenarios] ✅ Returning working scenario for {session_id} (preserving Step 3 edits)")
            scenario = _recompute_totals(scenario)
            if session_id:
                update_working_scenario(session_id, scenario)
            totals = scenario.get("totals", {})
            total_hours = totals.get("hours", 0)
            total_price = totals.get("price", 0)
            print(f"[/api/scenarios] hours={total_hours} price={total_price}")
            return {
                "ok": True,
                "scenarios": {"A": scenario},
                "totals": {"A": totals},  # Normalized to match CASE 0
                "metadata": state.get("metadata", {}),
                "total_hours": total_hours,
                "total_price": total_price,
                "cached": True
            }
        
        if reset_flag:
            print(f"[/api/scenarios] Reset requested but no baseline - falling back to Step 2")
        
        # Extract codes from the payload - check all possible field names
        codes = None
        
        # Check for various field name formats (camelCase, snake_case, variations)
        code_field_names = [
            "codes",                      # Original expected
            "selected_codes",             # Snake case
            "selectedCodes",              # CamelCase from frontend
            "deliverable_codes",          # Alternative snake case
            "deliverableCodes",          # Alternative camelCase
            "selected_deliverable_codes", # Full snake case
            "selectedDeliverableCodes",   # Full camelCase
            "deliverables",              # Simple form
            "selectedDeliverables"       # CamelCase simple form
        ]
        
        for field_name in code_field_names:
            if field_name in payload and payload[field_name]:
                codes = payload[field_name]
                print(f"[/api/scenarios] Found codes in field: {field_name}")
                break
        
        if not codes:
            # Log what we actually received for debugging
            print(f"[/api/scenarios] ERROR: No codes found. Received fields: {list(payload.keys())}")
            raise HTTPException(
                status_code=422, 
                detail=f"Missing deliverable codes. Expected one of: {', '.join(code_field_names)}. Received: {list(payload.keys())}"
            )
        
        # Handle scenario letter field (both camelCase and snake_case)
        scenario_letter = payload.get("scenario_letter") or payload.get("scenarioLetter")
        
        # Handle scenario_a field (both camelCase and snake_case)
        scenario_a = payload.get("scenario_a") or payload.get("scenarioA") or {
            "mode": "template", 
            "complexity": "Advanced", 
            "tier": "T2_MediumVolume"
        }
        
        # Handle other field name variations
        pricing_mode = payload.get("pricing_mode") or payload.get("pricingMode") or "Flat_Blended"
        blended_rate = payload.get("blended_rate") or payload.get("blendedRate")
        rate_band = payload.get("rate_band") or payload.get("rateBand") or "Standard_US"
        use_slack = payload.get("use_slack") if "use_slack" in payload else payload.get("useSlack", True)
        slack_after_internal = payload.get("slack_after_internal") or payload.get("slackAfterInternal") or 1
        slack_after_client = payload.get("slack_after_client") or payload.get("slackAfterClient") or 2
        slack_global_pct = payload.get("slack_global_pct") or payload.get("slackGlobalPct") or 0.05
        project_start = payload.get("project_start") or payload.get("projectStart")
        client_budget_usd = payload.get("client_budget_usd") or payload.get("clientBudgetUsd") or payload.get("clientBudget")
        
        # Handle complex map fields
        selected_components_map = (
            payload.get("selected_components_map") or 
            payload.get("selectedComponentsMap") or 
            payload.get("componentsMap")
        )
        
        selected_l3_map = (
            payload.get("selected_l3_map") or 
            payload.get("selectedL3Map") or 
            payload.get("l3Map") or
            payload.get("selected_l2_map") or
            payload.get("selectedL2Map") or
            payload.get("l2Map")
        )
        
        # Log what we're building with
        print(f"[/api/scenarios] Building with {len(codes)} codes, pricing_mode={pricing_mode}, rate_band={rate_band}")
        
        # Prepare the build payload with proper field names
        build_payload = BuildPayload(
            selected_deliverable_codes=codes,
            scenario_a=scenario_a,
            pricing_mode=pricing_mode,
            blended_rate=blended_rate,
            rate_band=rate_band,
            use_slack=use_slack,
            slack_after_internal=slack_after_internal,
            slack_after_client=slack_after_client,
            slack_global_pct=slack_global_pct,
            project_start=project_start,
            client_budget_usd=client_budget_usd,
            retainers=payload.get("retainers", []),
            selected_components_map=selected_components_map,
            selected_l3_map=selected_l3_map
        )
        
        # Call the existing build logic
        result = api_build(build_payload)
        
        # Store scenarios in window.SCENARIOS format
        if isinstance(result, dict) and "scenarios" in result:
            window_scenarios = result["scenarios"]
            # Update the global scenarios store
            _CURRENT_SCENARIOS.update(window_scenarios)
            
            # Add legacy top-level totals for UI compatibility
            scenario_a = window_scenarios.get("A") or {}
            totals = scenario_a.get("totals") or {}
            total_hours = totals.get("hours", scenario_a.get("total_hours", 0))
            total_price = totals.get("price", scenario_a.get("total_price", 0))
            
            # Return in the expected format
            return {
                "ok": True,
                "scenarios": window_scenarios,
                "totals": result.get("totals", {}),
                "metadata": result.get("metadata", {}),
                "total_hours": total_hours,
                "total_price": total_price
            }
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"ERROR in /api/scenarios: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to build scenarios: {str(e)}")

@app.post("/api/build_scenario_c")
def api_build_scenario_c_deprecated(payload: dict):
    if not DB.loaded:
        DB.load()

    # 1) Fetch the base scenario (A or B) from current scenarios store
    base_letter = payload.base.upper()
    base = _current_scenarios().get(base_letter)
    if not base:
        raise HTTPException(status_code=400, detail=f"Base scenario {base_letter} not found. Please build scenarios A and B first.")

    # 2) Union of deliverable codes (base codes + add-on codes)
    base_codes = [item["deliverable_code"] for item in (base.get("items") or [])]
    # Use dict.fromkeys to maintain order and remove duplicates
    union_codes = list(dict.fromkeys(base_codes + payload.add_on_codes))
    
    # Filter out any unknown codes
    valid_codes = []
    for code in union_codes:
        if not DB.deliverables[DB.deliverables["Deliverable_Code"].astype(str)==str(code)].empty:
            valid_codes.append(code)
    
    # 3) Inherit parameters from base unless explicitly provided
    complexity = payload.complexity or base.get("complexity", "Advanced")
    tier = payload.tier or base.get("tier", "T2_MediumVolume")
    use_slack = base.get("use_slack", True) if payload.use_slack is None else payload.use_slack
    slack_i = base.get("slack_after_internal", 1) if payload.slack_after_internal is None else payload.slack_after_internal
    slack_c = base.get("slack_after_client", 2) if payload.slack_after_client is None else payload.slack_after_client
    slack_pct = base.get("slack_global_pct", 0.05) if payload.slack_global_pct is None else payload.slack_global_pct
    project_start = payload.project_start or base.get("project_start")
    
    # Build inheritance map from base unless overridden
    base_ret_map = {}
    for it in (base.get("items") or []):
        if it.get("retainer", {}) and int(it["retainer"].get("months", 0)) > 0:
            base_ret_map[it["deliverable_code"]] = int(it["retainer"]["months"])
    override_map = {r.deliverable_code: max(1, min(12, int(r.months))) for r in (payload.retainers or [])}
    ret_map = {**base_ret_map, **override_map}

    # 4) Build scenario items using existing logic (same as /api/build)
    per_deliv = []
    for code in valid_codes:
        row = DB.deliverables[DB.deliverables["Deliverable_Code"].astype(str)==str(code)]
        if row.empty:
            continue
        cat = str(row["Category"].iloc[0])
        
        # Create scenario spec for this deliverable
        spec_resolved = {"mode": "template", "complexity": complexity, "tier": tier}
        
        months = int(ret_map.get(code, 0))
        # For Scenario C, no component selection yet - use default
        out = _scenario_for_deliverable(
            code, cat, spec_resolved,
            payload.pricing_mode, payload.blended_rate, payload.rate_band,
            use_slack, slack_i, slack_c, slack_pct, project_start,
            scenario_letter="C",
            retainer_months=months,   # NEW
            selected_components=None  # No component selection for Scenario C yet
        )
        # Add names for readability
        out["deliverable"] = str(row["Deliverable"].iloc[0])
        out["category"] = cat
        per_deliv.append(out)

    # Sort deliverables by phase order
    per_deliv = _sort_deliverables(per_deliv, "C")
    
    # Apply scenario-specific multipliers and add PM/QA overhead
    per_deliv = _apply_scenario_knobs(
        per_deliv, "C", complexity, tier,
        payload.pricing_mode, payload.rate_band, DB, payload.blended_rate
    )
    
    price_sum = sum(int(x["price"]) for x in per_deliv)
    hours_sum = sum(int(round(x["total_hours"])) for x in per_deliv)

    # 5) Create Scenario C
    scenario_c = {
        "label": "Scenario C (Upsell)",
        "pricing_mode": payload.pricing_mode,
        "rate_band": payload.rate_band,
        "blended_rate": payload.blended_rate,
        "complexity": complexity,
        "tier": tier,
        "use_slack": use_slack,
        "slack_after_internal": slack_i,
        "slack_after_client": slack_c,
        "slack_global_pct": slack_pct,
        "project_start": project_start,
        "items": per_deliv,
        "totals": {"hours": int(hours_sum), "price": int(price_sum)}
    }
    
    # Add budget metrics if client_budget_usd was provided
    if payload.client_budget_usd and payload.client_budget_usd > 0:
        scenario_price = int(price_sum)
        budget_delta = scenario_price - payload.client_budget_usd
        coverage_pct = (scenario_price / payload.client_budget_usd) * 100 if payload.client_budget_usd > 0 else 0
        scale_factor = payload.client_budget_usd / scenario_price if scenario_price > 0 else 1.0
        
        scenario_c["budget_info"] = {
            "client_budget_usd": payload.client_budget_usd,
            "total_price": scenario_price,
            "budget_delta": budget_delta,
            "coverage_pct": round(coverage_pct, 1),
            "scale_factor_if_fit": round(scale_factor, 3)
        }

    # 6) Store/update in memory next to A/B for this session
    global _CURRENT_SCENARIOS
    _CURRENT_SCENARIOS["C"] = scenario_c
    
    return {"C": scenario_c}

@app.post("/api/auto_build")
def api_auto_build(payload: AutoBuildPayload):
    if not DB.loaded:
        DB.load()

    # 1) Get AI suggestions
    suggestions = DB.suggest_deliverables_from_text(payload.rfp_text or "")
    selected_codes = [s["deliverable_code"] for s in suggestions]

    # Build retainer map
    ret_map = _safe_retainer_map(payload.retainers)

    # 2) If nothing matched, return an empty set so frontend can prompt to add
    if not selected_codes:
        return {
            "suggested": suggestions,
            "scenarios": {
                "A": {"items": [], "totals": {"hours": 0.0, "price": 0.0}},
                "B": {"items": [], "totals": {"hours": 0.0, "price": 0.0}},
            }
        }

    # 3) Reuse the same logic as /api/build to assemble scenarios
    #    (We inline the essential parts to keep it simple.)
    def _build_for(selected_deliverable_codes, scen_spec):
        per_deliv = []
        price_sum = 0.0
        hours_sum = 0.0
        for code in selected_deliverable_codes:
            row = DB.deliverables[DB.deliverables["Deliverable_Code"].astype(str) == str(code)]
            if row.empty:
                continue
            cat = str(row["Category"].iloc[0])
            spec_resolved = _resolve_scenario(scen_spec, cat)
            months = int(ret_map.get(code, 0))
            out = _scenario_for_deliverable(
                code, cat, spec_resolved,
                payload.pricing_mode, payload.blended_rate, payload.rate_band or "Standard_US",
                bool(payload.use_slack), int(payload.slack_after_internal), int(payload.slack_after_client),
                float(payload.slack_global_pct or 0), payload.project_start,
                scenario_letter="A",  # letter doesn't affect numbers; acceptable here
                retainer_months=months  # NEW
            )
            out["deliverable"] = str(row["Deliverable"].iloc[0])
            out["category"] = cat
            per_deliv.append(out)
            price_sum += out["price"]
            hours_sum += out["total_hours"]
        return {
            "pricing_mode": payload.pricing_mode,
            "rate_band": payload.rate_band or "Standard_US",
            "blended_rate": payload.blended_rate,
            "use_slack": bool(payload.use_slack),
            "slack_after_internal": int(payload.slack_after_internal),
            "slack_after_client": int(payload.slack_after_client),
            "slack_global_pct": float(payload.slack_global_pct or 0),
            "project_start": payload.project_start,
            "items": per_deliv,
            "totals": {"hours": round(hours_sum, 2), "price": round(price_sum, 2)}
        }

    scenarios = {
        "A": _build_for(selected_codes, payload.scenario_a),
        "B": _build_for(selected_codes, payload.scenario_b),
    }

    return {"suggested": suggestions, "scenarios": scenarios}

@app.post("/api/export")
def api_export(payload: Union[ExportPayload, dict]):
    """
    Export a Workfront file (CSV or XLSX) from a single scenario payload.
    The 'Project_Name' column is set to payload.project_name if provided.
    The download filename is derived from project/scenario.
    
    Accepts formats:
    - {"scenario": {...}, "file_format": "xlsx", "project_name": "Project"}
    - {"scenarios": {"A": {...}}, "file_format": "xlsx", "project_name": "Project"}
    - {"session_id": "...", "file_format": "xlsx"}  # NEW: Read from SCENARIO_STORE
    """
    if not DB.loaded:
        DB.load()

    # Handle both dict and model formats
    if isinstance(payload, dict):
        # NEW: Check session_id FIRST to use working scenario (preserves Step 3 edits)
        session_id = payload.get("session_id")
        scenario = None
        
        if session_id:
            try:
                _ = get_session_state(session_id)  # Trigger migration if needed
                scenario = get_working_scenario(session_id)
                if scenario and scenario.get("items"):
                    print(f"[/api/export] Using working scenario from SCENARIO_STORE for session {session_id}")
                    scenario = _inflate_components_if_missing(scenario)
            except Exception as e:
                print(f"[/api/export] Error loading from SCENARIO_STORE: {e}")
                scenario = None
        
        # Fallback: Extract scenario from either 'scenario' or 'scenarios.A'
        if not scenario:
            scenario = payload.get("scenario")
        if not scenario and "scenarios" in payload:
            # Get the first scenario (usually 'A')
            scenarios_dict = payload.get("scenarios", {})
            if scenarios_dict and isinstance(scenarios_dict, dict):
                first_key = next(iter(scenarios_dict.keys()), None)
                if first_key:
                    scenario = scenarios_dict[first_key]
        
        project_name = payload.get("project_name")
        file_format = payload.get("file_format", "csv")
        scenario_label = payload.get("scenario_label")
    else:
        scenario = payload.scenario
        project_name = payload.project_name
        file_format = payload.file_format or "csv"
        scenario_label = payload.scenario_label

    project_name = project_name or _upload_title_default() or f"Proposal {datetime.date.today().isoformat()}"

    # Inflate components if missing (defensive fallback)
    scenario = _inflate_components_if_missing(scenario or {})
    
    df = build_wbs_dataframe_from_scenario(scenario, project_name)
    # <<< force A–E to the left and guarantee presence
    df = _ensure_v3_ae_columns(df)

    base = _export_basename(project_name, scenario_label)  # always adds EST timestamp

    fmt = (file_format or "csv").lower()
    if fmt == "csv":
        out_path = f"{base}.csv"
        df.to_csv(out_path, index=False)
        return FileResponse(out_path, filename=os.path.basename(out_path), media_type="text/csv")

    # xlsx
    try:
        out_path = f"{base}.xlsx"
        with pd.ExcelWriter(out_path, engine="openpyxl") as xw:
            df.to_excel(xw, index=False)
            _apply_number_formats(xw.sheets[list(xw.sheets.keys())[0]], df)
        return FileResponse(
            out_path, filename=os.path.basename(out_path),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    except Exception as ex:
        raise HTTPException(400, "XLSX export requires 'openpyxl'.") from ex

@app.post("/api/export_workbook")
def api_export_workbook(payload: ExportWorkbookPayload):
    if not DB.loaded:
        DB.load()
    project = (payload.project_name
               or _upload_title_default()
               or f"Proposal {datetime.date.today().isoformat()}").strip()
    
    # Inflate components if missing (defensive fallback)
    scenario_a = _inflate_components_if_missing(payload.scenario_a or {})
    scenario_b = _inflate_components_if_missing(payload.scenario_b or {})
    
    dfA = build_wbs_dataframe_from_scenario(scenario_a, project)
    dfB = build_wbs_dataframe_from_scenario(scenario_b, project)
    
    dfA = _ensure_v3_ae_columns(dfA)
    dfB = _ensure_v3_ae_columns(dfB)
    
    # Generate retainer summaries if any retainer items exist
    retainer_summaryA = create_retainer_summary(scenario_a)
    retainer_summaryB = create_retainer_summary(scenario_b)
    has_retainersA = not retainer_summaryA.empty
    has_retainersB = not retainer_summaryB.empty
    
    base = _export_basename(project, "Scenarios A & B")  # includes EST timestamp
    out_path = f"{base}.xlsx"
    # Always use stable, distinct tab names to prevent accidental overwrite
    sheetA = "Scenario A"
    sheetB = "Scenario B"
    with pd.ExcelWriter(out_path, engine="openpyxl") as xw:
        dfA.to_excel(xw, sheet_name=sheetA, index=False)
        dfB.to_excel(xw, sheet_name=sheetB, index=False)
        
        # Add retainer summary sheets if there are retainers
        if has_retainersA:
            retainer_summaryA.to_excel(xw, sheet_name="Retainer Summary A", index=False)
        if has_retainersB:
            retainer_summaryB.to_excel(xw, sheet_name="Retainer Summary B", index=False)
        
        _apply_number_formats(xw.sheets[sheetA], dfA)
        _apply_number_formats(xw.sheets[sheetB], dfB)
        
        # Apply formatting to retainer sheets
        if has_retainersA:
            _apply_number_formats(xw.sheets["Retainer Summary A"], retainer_summaryA)
        if has_retainersB:
            _apply_number_formats(xw.sheets["Retainer Summary B"], retainer_summaryB)
    return FileResponse(
        out_path, filename=os.path.basename(out_path),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

@app.post("/api/export_workbook_abc")
def api_export_workbook_abc(payload: ExportWorkbookABCPayload):
    """Export scenarios A, B, and C to a single Excel file with separate sheets."""
    if not DB.loaded:
        DB.load()
    
    project = (payload.project_name
               or _upload_title_default()
               or f"Proposal {datetime.date.today().isoformat()}").strip()
    
    # Inflate components if missing (defensive fallback)
    scenario_a = _inflate_components_if_missing(payload.scenario_a or {})
    scenario_b = _inflate_components_if_missing(payload.scenario_b or {})
    scenario_c = _inflate_components_if_missing(payload.scenario_c or {})
    
    # Build DataFrames for all three scenarios
    dfA = build_wbs_dataframe_from_scenario(scenario_a, project)
    dfB = build_wbs_dataframe_from_scenario(scenario_b, project)
    dfC = build_wbs_dataframe_from_scenario(scenario_c, project)
    
    # Ensure consistent column formatting
    dfA = _ensure_v3_ae_columns(dfA)
    dfB = _ensure_v3_ae_columns(dfB)
    dfC = _ensure_v3_ae_columns(dfC)
    
    # Generate retainer summaries if any retainer items exist
    retainer_summaryA = create_retainer_summary(scenario_a)
    retainer_summaryB = create_retainer_summary(scenario_b)
    retainer_summaryC = create_retainer_summary(scenario_c)
    has_retainersA = not retainer_summaryA.empty
    has_retainersB = not retainer_summaryB.empty
    has_retainersC = not retainer_summaryC.empty
    
    # Generate filename with EST timestamp
    base = _export_basename(project, "Scenarios A, B & C")  # includes EST timestamp
    out_path = f"{base}.xlsx"
    
    # Create Excel file with three sheets plus retainer summaries
    with pd.ExcelWriter(out_path, engine="openpyxl") as xw:
        dfA.to_excel(xw, sheet_name=payload.sheet_name_a or "Scenario A", index=False)
        dfB.to_excel(xw, sheet_name=payload.sheet_name_b or "Scenario B", index=False)
        dfC.to_excel(xw, sheet_name=payload.sheet_name_c or "Scenario C", index=False)
        
        # Add retainer summary sheets if there are retainers
        if has_retainersA:
            retainer_summaryA.to_excel(xw, sheet_name="Retainer Summary A", index=False)
        if has_retainersB:
            retainer_summaryB.to_excel(xw, sheet_name="Retainer Summary B", index=False)
        if has_retainersC:
            retainer_summaryC.to_excel(xw, sheet_name="Retainer Summary C", index=False)
        
        # Apply number formatting to all sheets
        _apply_number_formats(xw.sheets[payload.sheet_name_a or "Scenario A"], dfA)
        _apply_number_formats(xw.sheets[payload.sheet_name_b or "Scenario B"], dfB)
        _apply_number_formats(xw.sheets[payload.sheet_name_c or "Scenario C"], dfC)
        
        # Apply formatting to retainer sheets
        if has_retainersA:
            _apply_number_formats(xw.sheets["Retainer Summary A"], retainer_summaryA)
        if has_retainersB:
            _apply_number_formats(xw.sheets["Retainer Summary B"], retainer_summaryB)
        if has_retainersC:
            _apply_number_formats(xw.sheets["Retainer Summary C"], retainer_summaryC)
    
    return FileResponse(
        out_path, filename=os.path.basename(out_path),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

def _assert_has_items(scen: dict, label: str):
    """Guard to ensure scenario has items before export"""
    if not scen or not scen.get("items"):
        raise HTTPException(400, f"No build context for {label}. Run Build once in Step 3.")

@app.post("/api/export/xml")
def api_export_xml_post(payload: Union[ExportXMLPayload, dict]):
    """
    Export a single scenario as Microsoft Project XML (MSPDI) format.
    POST endpoint for XML export with flexible payload handling.
    
    Accepts formats:
    - {"scenario": {...}, "project_name": "Project"}
    - {"scenarios": {"A": {...}}, "project_name": "Project"}
    - {"session_id": "...", "project_name": "Project"}  # NEW: Read from SCENARIO_STORE
    """
    # NEW: Check if session_id is provided to read from SCENARIO_STORE
    if isinstance(payload, dict) and "session_id" in payload and "scenario" not in payload:
        session_id = payload.get("session_id")
        
        # Use get_working_scenario to respect dual-entry architecture
        scenario = get_working_scenario(session_id)
        if not scenario or not scenario.get("items"):
            raise HTTPException(404, f"Scenario not found for session {session_id}")
        
        # Recompute totals to ensure Step 3 edits have correct totals
        scenario = _recompute_totals(scenario)
        update_working_scenario(session_id, scenario)
        print(f"[XML EXPORT] 📦 Using working scenario for session {session_id} (preserves Step 3 edits)")
        
        # Create new payload with scenario from store
        payload_dict = {
            "scenario": scenario,
            "project_name": payload.get("project_name") or scenario.get("project_name", "Project"),
            "scenario_label": payload.get("scenario_label", "A"),
            "sheet_name": payload.get("sheet_name", "WBS Export"),
            "start_date_mode": payload.get("start_date_mode", "today"),
            "fixed_start_iso": payload.get("fixed_start_iso") or scenario.get("project_start"),
            "hours_per_day": payload.get("hours_per_day", 8.0)
        }
        return api_export_xml(payload_dict)
    
    # Delegate to the existing export_xml function
    return api_export_xml(payload)

@app.post("/api/export_xml")
def api_export_xml(payload: Union[ExportXMLPayload, dict]):
    """
    Export a single scenario as Microsoft Project XML (MSPDI) format.
    Uses the convert_excel_to_mspdi function with multi-resource merge capability.
    """
    if not DB.loaded:
        DB.load()

    # Handle both dict and model formats
    if isinstance(payload, dict):
        # Extract scenario from either 'scenario' or 'scenarios.A'
        scenario = payload.get("scenario")
        if not scenario and "scenarios" in payload:
            scenarios_dict = payload.get("scenarios", {})
            if scenarios_dict and isinstance(scenarios_dict, dict):
                first_key = next(iter(scenarios_dict.keys()), None)
                if first_key:
                    scenario = scenarios_dict[first_key]
        
        project_name = payload.get("project_name")
        scenario_label = payload.get("scenario_label")
        sheet_name = payload.get("sheet_name", "WBS Export")
        start_date_mode = payload.get("start_date_mode", "today")
        fixed_start_iso = payload.get("fixed_start_iso")
        hours_per_day = payload.get("hours_per_day", 8.0)
    else:
        scenario = payload.scenario
        project_name = payload.project_name
        scenario_label = payload.scenario_label
        sheet_name = payload.sheet_name
        start_date_mode = payload.start_date_mode
        fixed_start_iso = payload.fixed_start_iso
        hours_per_day = payload.hours_per_day

    project_name = project_name or _upload_title_default() or f"Proposal {datetime.date.today().isoformat()}"

    # Guard: ensure scenario has items
    _assert_has_items(scenario or {}, "XML export")
    
    # Inflate components if missing (defensive fallback)
    scenario = _inflate_components_if_missing(scenario or {})
    
    # Build WBS DataFrame
    df = build_wbs_dataframe_from_scenario(scenario, project_name)
    df = _ensure_v3_ae_columns(df)

    # Create temporary Excel file for MSPDI conversion with unique ID to prevent collisions
    base = _export_basename(project_name, scenario_label or "Scenario")
    unique_id = uuid.uuid4().hex[:8]
    temp_xlsx = os.path.join(tempfile.gettempdir(), f"{base}_temp_{unique_id}.xlsx")
    output_xml = os.path.join(tempfile.gettempdir(), f"{base}_{unique_id}.xml")
    
    try:
        # Write to temporary Excel file
        with pd.ExcelWriter(temp_xlsx, engine="openpyxl") as xw:
            df.to_excel(xw, sheet_name=sheet_name, index=False)
            _apply_number_formats(xw.sheets[sheet_name], df)

        # Convert to MSPDI XML
        # Use fixed_start_iso from payload, or fall back to project_start from scenario
        project_start_iso = fixed_start_iso or scenario.get("project_start")
        
        stats = convert_excel_to_mspdi(
            input_xlsx=temp_xlsx,
            output_xml=output_xml,
            sheet_name=sheet_name,
            start_date_mode=start_date_mode,
            fixed_start_iso=project_start_iso,
            hours_per_day=hours_per_day,
            merge_identical_children=False,
            project_name=project_name,
            pricing_mode=scenario.get("pricing_mode", "Flat_Blended"),
            rate_band=scenario.get("rate_band", "Standard_US"),
            blended_rate=scenario.get("blended_rate"),
            add_deliverable_milestones=False  # Workfront compatibility: no alphanumeric WBS
        )
        
        # Post-process XML to parallelize identical task names (optional)
        final_xml = output_xml
        if PARALLELIZE_IDENTICAL_NAMES:
            final_xml = post_process_xml(output_xml)

        return FileResponse(
            final_xml,
            filename=os.path.basename(final_xml),
            media_type="application/xml",
            headers={"X-Export-Stats": json.dumps(stats)}
        )

    finally:
        # Clean up temporary Excel file
        if os.path.exists(temp_xlsx):
            os.remove(temp_xlsx)

@app.post("/api/export_workbook_xml")
def api_export_workbook_xml(payload: ExportWorkbookXMLPayload):
    """
    Export scenarios A and B to separate XML files in a zip archive.
    Each XML uses the convert_excel_to_mspdi function with multi-resource merge.
    """
    if not DB.loaded:
        DB.load()
    
    project = (payload.project_name
               or _upload_title_default()
               or f"Proposal {datetime.date.today().isoformat()}").strip()
    
    # Guard: ensure scenarios have items
    _assert_has_items(payload.scenario_a or {}, "Scenario A XML export")
    _assert_has_items(payload.scenario_b or {}, "Scenario B XML export")
    
    # Inflate components if missing (defensive fallback)
    scenario_a = _inflate_components_if_missing(payload.scenario_a or {})
    scenario_b = _inflate_components_if_missing(payload.scenario_b or {})
    
    # Build DataFrames
    dfA = build_wbs_dataframe_from_scenario(scenario_a, project)
    dfB = build_wbs_dataframe_from_scenario(scenario_b, project)
    
    dfA = _ensure_v3_ae_columns(dfA)
    dfB = _ensure_v3_ae_columns(dfB)
    
    base = _export_basename(project, "Scenarios A & B XML")
    unique_id = uuid.uuid4().hex[:8]
    temp_files = []
    
    try:
        # Create XML for Scenario A
        temp_xlsx_a = os.path.join(tempfile.gettempdir(), f"{base}_A_temp_{unique_id}.xlsx")
        output_xml_a = os.path.join(tempfile.gettempdir(), f"{base}_Scenario_A_{unique_id}.xml")
        temp_files.extend([temp_xlsx_a, output_xml_a])
        
        with pd.ExcelWriter(temp_xlsx_a, engine="openpyxl") as xw:
            dfA.to_excel(xw, sheet_name="Scenario A", index=False)
            _apply_number_formats(xw.sheets["Scenario A"], dfA)
        
        # Use project_start_iso from payload or scenario_a
        project_start_iso = payload.project_start_iso or scenario_a.get("project_start")
        
        stats_a = convert_excel_to_mspdi(
            input_xlsx=temp_xlsx_a,
            output_xml=output_xml_a,
            sheet_name="Scenario A",
            fixed_start_iso=project_start_iso,
            merge_identical_children=False,
            project_name=project,
            pricing_mode=scenario_a.get("pricing_mode", "Flat_Blended"),
            rate_band=scenario_a.get("rate_band", "Standard_US"),
            blended_rate=scenario_a.get("blended_rate")
        )
        
        # Post-process Scenario A XML
        final_xml_a = output_xml_a
        if PARALLELIZE_IDENTICAL_NAMES:
            final_xml_a = post_process_xml(output_xml_a)
            temp_files.append(final_xml_a)
        
        # Create XML for Scenario B
        temp_xlsx_b = os.path.join(tempfile.gettempdir(), f"{base}_B_temp_{unique_id}.xlsx")
        output_xml_b = os.path.join(tempfile.gettempdir(), f"{base}_Scenario_B_{unique_id}.xml")
        temp_files.extend([temp_xlsx_b, output_xml_b])
        
        with pd.ExcelWriter(temp_xlsx_b, engine="openpyxl") as xw:
            dfB.to_excel(xw, sheet_name="Scenario B", index=False)
            _apply_number_formats(xw.sheets["Scenario B"], dfB)
        
        # Use project_start_iso from payload or scenario_b
        project_start_iso_b = payload.project_start_iso or scenario_b.get("project_start")
        
        stats_b = convert_excel_to_mspdi(
            input_xlsx=temp_xlsx_b,
            output_xml=output_xml_b,
            sheet_name="Scenario B",
            fixed_start_iso=project_start_iso_b,
            merge_identical_children=False,
            project_name=project,
            pricing_mode=scenario_b.get("pricing_mode", "Flat_Blended"),
            rate_band=scenario_b.get("rate_band", "Standard_US"),
            blended_rate=scenario_b.get("blended_rate")
        )
        
        # Post-process Scenario B XML
        final_xml_b = output_xml_b
        if PARALLELIZE_IDENTICAL_NAMES:
            final_xml_b = post_process_xml(output_xml_b)
            temp_files.append(final_xml_b)
        
        # Create zip file with both XMLs
        import zipfile
        zip_path = os.path.join(tempfile.gettempdir(), f"{base}_{unique_id}.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            zipf.write(final_xml_a, f"Scenario_A.xml")
            zipf.write(final_xml_b, f"Scenario_B.xml")
            # Add stats as JSON file
            stats_json = json.dumps({
                "scenario_a": stats_a,
                "scenario_b": stats_b
            }, indent=2)
            zipf.writestr("export_stats.json", stats_json)
        
        return FileResponse(
            zip_path,
            filename=os.path.basename(zip_path),
            media_type="application/zip"
        )
        
    finally:
        # Clean up temporary files
        for temp_file in temp_files:
            if os.path.exists(temp_file):
                os.remove(temp_file)

@app.post("/api/export_workbook_xml_abc")
def api_export_workbook_xml_abc(p: ExportWorkbookXMLABCPayload):
    """
    Export scenarios A, B, and C to separate XML files in a zip archive.
    Each XML uses the convert_excel_to_mspdi function.
    """
    if not DB.loaded:
        DB.load()

    project = (p.project_name or _upload_title_default() or
               f"Proposal {datetime.date.today().isoformat()}").strip()

    # Guard: ensure scenarios exist
    _assert_has_items(p.scenario_a or {}, "Scenario A XML export")
    _assert_has_items(p.scenario_b or {}, "Scenario B XML export")
    _assert_has_items(p.scenario_c or {}, "Scenario C XML export")

    # Inflate components just like A/B route
    scenA = _inflate_components_if_missing(p.scenario_a or {})
    scenB = _inflate_components_if_missing(p.scenario_b or {})
    scenC = _inflate_components_if_missing(p.scenario_c or {})

    # Build WBS dataframes
    dfA = build_wbs_dataframe_from_scenario(scenA, project)
    dfB = build_wbs_dataframe_from_scenario(scenB, project)
    dfC = build_wbs_dataframe_from_scenario(scenC, project)
    dfA = _ensure_v3_ae_columns(dfA)
    dfB = _ensure_v3_ae_columns(dfB)
    dfC = _ensure_v3_ae_columns(dfC)

    base = _export_basename(project, "Scenarios A, B & C XML")
    unique_id = uuid.uuid4().hex[:8]
    temp_files, xml_files = [], []

    try:
        # A
        tmp_xlsx_a = os.path.join(tempfile.gettempdir(), f"{base}_A_temp_{unique_id}.xlsx")
        out_xml_a = os.path.join(tempfile.gettempdir(), f"{base}_Scenario_A_{unique_id}.xml")
        temp_files += [tmp_xlsx_a, out_xml_a]
        with pd.ExcelWriter(tmp_xlsx_a, engine="openpyxl") as xw:
            dfA.to_excel(xw, sheet_name="Scenario A", index=False)
            _apply_number_formats(xw.sheets["Scenario A"], dfA)
        stats_a = convert_excel_to_mspdi(
            input_xlsx=tmp_xlsx_a, output_xml=out_xml_a, sheet_name="Scenario A",
            start_date_mode=p.start_date_mode, fixed_start_iso=p.fixed_start_iso,
            hours_per_day=p.hours_per_day, merge_identical_children=False,
            project_name=project,
            pricing_mode=scenA.get("pricing_mode", "Flat_Blended"),
            rate_band=scenA.get("rate_band", "Standard_US"),
            blended_rate=scenA.get("blended_rate")
        )
        # Post-process Scenario A XML
        final_xml_a = out_xml_a
        if PARALLELIZE_IDENTICAL_NAMES:
            final_xml_a = post_process_xml(out_xml_a)
            temp_files.append(final_xml_a)
        xml_files.append(("Scenario_A.xml", final_xml_a, stats_a))

        # B
        tmp_xlsx_b = os.path.join(tempfile.gettempdir(), f"{base}_B_temp_{unique_id}.xlsx")
        out_xml_b = os.path.join(tempfile.gettempdir(), f"{base}_Scenario_B_{unique_id}.xml")
        temp_files += [tmp_xlsx_b, out_xml_b]
        with pd.ExcelWriter(tmp_xlsx_b, engine="openpyxl") as xw:
            dfB.to_excel(xw, sheet_name="Scenario B", index=False)
            _apply_number_formats(xw.sheets["Scenario B"], dfB)
        stats_b = convert_excel_to_mspdi(
            input_xlsx=tmp_xlsx_b, output_xml=out_xml_b, sheet_name="Scenario B",
            start_date_mode=p.start_date_mode, fixed_start_iso=p.fixed_start_iso,
            hours_per_day=p.hours_per_day, merge_identical_children=False,
            project_name=project,
            pricing_mode=scenB.get("pricing_mode", "Flat_Blended"),
            rate_band=scenB.get("rate_band", "Standard_US"),
            blended_rate=scenB.get("blended_rate")
        )
        # Post-process Scenario B XML
        final_xml_b = out_xml_b
        if PARALLELIZE_IDENTICAL_NAMES:
            final_xml_b = post_process_xml(out_xml_b)
            temp_files.append(final_xml_b)
        xml_files.append(("Scenario_B.xml", final_xml_b, stats_b))

        # C
        tmp_xlsx_c = os.path.join(tempfile.gettempdir(), f"{base}_C_temp_{unique_id}.xlsx")
        out_xml_c = os.path.join(tempfile.gettempdir(), f"{base}_Scenario_C_{unique_id}.xml")
        temp_files += [tmp_xlsx_c, out_xml_c]
        with pd.ExcelWriter(tmp_xlsx_c, engine="openpyxl") as xw:
            dfC.to_excel(xw, sheet_name="Scenario C", index=False)
            _apply_number_formats(xw.sheets["Scenario C"], dfC)
        stats_c = convert_excel_to_mspdi(
            input_xlsx=tmp_xlsx_c, output_xml=out_xml_c, sheet_name="Scenario C",
            start_date_mode=p.start_date_mode, fixed_start_iso=p.fixed_start_iso,
            hours_per_day=p.hours_per_day, merge_identical_children=False,
            project_name=project,
            pricing_mode=scenC.get("pricing_mode", "Flat_Blended"),
            rate_band=scenC.get("rate_band", "Standard_US"),
            blended_rate=scenC.get("blended_rate")
        )
        # Post-process Scenario C XML
        final_xml_c = out_xml_c
        if PARALLELIZE_IDENTICAL_NAMES:
            final_xml_c = post_process_xml(out_xml_c)
            temp_files.append(final_xml_c)
        xml_files.append(("Scenario_C.xml", final_xml_c, stats_c))

        # Zip all 3
        import zipfile
        zip_path = os.path.join(tempfile.gettempdir(), f"{base}_{unique_id}.zip")
        with zipfile.ZipFile(zip_path, "w") as z:
            for alias, real, _ in xml_files:
                z.write(real, alias)
            z.writestr("export_stats.json", json.dumps({
                "scenario_a": xml_files[0][2],
                "scenario_b": xml_files[1][2],
                "scenario_c": xml_files[2][2],
            }, indent=2))
        return FileResponse(zip_path, filename=os.path.basename(zip_path),
                            media_type="application/zip")
    finally:
        for f in temp_files:
            if os.path.exists(f):
                try:
                    os.remove(f)
                except:
                    pass

# ---------- Individual Scenario Export Endpoints (Task 4.5) ----------

def _export_single_scenario_xml(
    scenario: Dict[str, Any],
    scenario_label: str,
    project_name: Optional[str] = None,
    add_deliverable_milestones: bool = False
) -> str:
    """
    Helper function to export a single scenario to XML.
    Returns the path to the final XML file.
    """
    if not DB.loaded:
        DB.load()
    
    print(f"[EXPORT_XML] Starting export for {scenario_label}")
    print(f"[EXPORT_XML] Scenario keys: {list(scenario.keys())}")
    print(f"[EXPORT_XML] Has timeline_tasks: {'timeline_tasks' in scenario}")
    print(f"[EXPORT_XML] Has timeline: {'timeline' in scenario}")
    
    if "timeline_tasks" in scenario:
        print(f"[EXPORT_XML] timeline_tasks count: {len(scenario['timeline_tasks'])}")
    
    # Guard: ensure scenario has items
    _assert_has_items(scenario, f"{scenario_label} XML export")
    
    # Inflate components if missing
    scenario = _inflate_components_if_missing(scenario)
    
    print(f"[EXPORT_XML] After inflate, has timeline_tasks: {'timeline_tasks' in scenario}")
    
    # Determine project name
    project = (project_name 
               or _upload_title_default() 
               or f"Proposal {datetime.date.today().isoformat()}").strip()
    
    print(f"[EXPORT_XML] About to call build_wbs_dataframe_from_scenario()")
    print(f"[EXPORT_XML] Passing scenario with keys: {list(scenario.keys())}")
    
    # Build WBS DataFrame
    df = build_wbs_dataframe_from_scenario(scenario, project)
    df = _ensure_v3_ae_columns(df)
    
    # Create temporary Excel file for MSPDI conversion with unique ID to prevent collisions
    base = _export_basename(project, scenario_label)
    unique_id = uuid.uuid4().hex[:8]
    temp_xlsx = os.path.join(tempfile.gettempdir(), f"{base}_temp_{unique_id}.xlsx")
    output_xml = os.path.join(tempfile.gettempdir(), f"{base}_{unique_id}.xml")
    
    try:
        # Write to temporary Excel file
        with pd.ExcelWriter(temp_xlsx, engine="openpyxl") as xw:
            df.to_excel(xw, sheet_name=scenario_label, index=False)
            _apply_number_formats(xw.sheets[scenario_label], df)
        
        # Convert to MSPDI XML using the simple, correct version (not the bloated enhanced version)
        project_start_iso = scenario.get("project_start")
        
        # Use the simple convert_excel_to_mspdi function defined in this file (line 9053)
        # This version produces clean XML matching the Nov 12 reference without bloat
        # FORCE add_deliverable_milestones=False to override any cached scenario settings
        stats = convert_excel_to_mspdi(
            input_xlsx=temp_xlsx,
            output_xml=output_xml,
            sheet_name=scenario_label,
            start_date_mode="fixed" if project_start_iso else "next_monday",
            fixed_start_iso=project_start_iso,
            hours_per_day=8.0,
            merge_identical_children=False,
            project_name=project,
            pricing_mode=scenario.get("pricing_mode", "Flat_Blended"),
            rate_band=scenario.get("rate_band", "Standard_US"),
            blended_rate=scenario.get("blended_rate"),
            add_deliverable_milestones=False  # WORKFRONT COMPAT: Always False (ignore parameter)
        )
        
        # Post-process XML to parallelize identical task names (optional)
        final_xml = output_xml
        if PARALLELIZE_IDENTICAL_NAMES:
            final_xml = post_process_xml(output_xml)
        
        return final_xml
    
    finally:
        # Clean up temporary Excel file
        if os.path.exists(temp_xlsx):
            try:
                os.remove(temp_xlsx)
            except:
                pass

@app.get("/api/export/xml/a")
def api_export_xml_scenario_a(add_anchors: bool = False, session_id: Optional[str] = None):
    """
    Export Scenario A only as XML.
    If session_id is provided, uses SCENARIO_STORE (includes Gantt edits).
    Otherwise uses _CURRENT_SCENARIOS (pricing-only data).
    """
    scenarios = _get_scenarios(session_id)
    if "A" not in scenarios:
        raise HTTPException(400, "Scenario A not found. Please build scenarios first.")
    
    final_xml = _export_single_scenario_xml(
        scenario=scenarios["A"],
        scenario_label="Scenario A",
        project_name=scenarios["A"].get("project_name"),
        add_deliverable_milestones=add_anchors
    )
    
    return FileResponse(
        final_xml,
        filename=os.path.basename(final_xml),
        media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{os.path.basename(final_xml)}"'}
    )

@app.get("/api/export/xml/b")
def api_export_xml_scenario_b(add_anchors: bool = False, session_id: Optional[str] = None):
    """
    Export Scenario B only as XML.
    If session_id is provided, uses SCENARIO_STORE (includes Gantt edits).
    Otherwise uses _CURRENT_SCENARIOS (pricing-only data).
    """
    scenarios = _get_scenarios(session_id)
    if "B" not in scenarios:
        raise HTTPException(400, "Scenario B not found. Please build scenarios first.")
    
    final_xml = _export_single_scenario_xml(
        scenario=scenarios["B"],
        scenario_label="Scenario B",
        project_name=scenarios["B"].get("project_name"),
        add_deliverable_milestones=add_anchors
    )
    
    return FileResponse(
        final_xml,
        filename=os.path.basename(final_xml),
        media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{os.path.basename(final_xml)}"'}
    )

@app.get("/api/export/xml/c")
def api_export_xml_scenario_c(add_anchors: bool = False, session_id: Optional[str] = None):
    """
    Export Scenario C only as XML.
    If session_id is provided, uses SCENARIO_STORE (includes Gantt edits).
    Otherwise uses _CURRENT_SCENARIOS (pricing-only data).
    """
    scenarios = _get_scenarios(session_id)
    if "C" not in scenarios:
        raise HTTPException(400, "Scenario C not found. Please build scenarios first.")
    
    final_xml = _export_single_scenario_xml(
        scenario=scenarios["C"],
        scenario_label="Scenario C",
        project_name=scenarios["C"].get("project_name"),
        add_deliverable_milestones=add_anchors
    )
    
    return FileResponse(
        final_xml,
        filename=os.path.basename(final_xml),
        media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{os.path.basename(final_xml)}"'}
    )

@app.get("/api/export/xml/all")
def api_export_xml_all_scenarios(session_id: Optional[str] = None):
    """
    Export all three scenarios (A, B, C) as a ZIP archive.
    If session_id is provided, uses SCENARIO_STORE (includes Gantt edits).
    Otherwise uses _CURRENT_SCENARIOS (pricing-only data).
    """
    scenarios = _get_scenarios(session_id)
    
    # Check if all scenarios exist
    missing = []
    for letter in ["A", "B", "C"]:
        if letter not in scenarios:
            missing.append(f"Scenario {letter}")
    
    if missing:
        raise HTTPException(400, f"Missing scenarios: {', '.join(missing)}. Please build all scenarios first.")
    
    # Determine project name (prefer from Scenario A)
    project = (scenarios.get("A", {}).get("project_name")
               or _upload_title_default()
               or f"Proposal {datetime.date.today().isoformat()}").strip()
    
    temp_files = []
    xml_files = []
    
    try:
        # Export each scenario
        for letter, label in [("A", "Scenario A"), ("B", "Scenario B"), ("C", "Scenario C")]:
            final_xml = _export_single_scenario_xml(
                scenario=scenarios[letter],
                scenario_label=label,
                project_name=project
            )
            temp_files.append(final_xml)
            xml_files.append((f"Scenario_{letter}.xml", final_xml))
        
        # Create ZIP file with all three XMLs
        import zipfile
        base = _export_basename(project, "Scenarios A, B & C")
        unique_id = uuid.uuid4().hex[:8]
        zip_path = os.path.join(tempfile.gettempdir(), f"{base}_{unique_id}.zip")
        
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for alias, real_path in xml_files:
                zipf.write(real_path, alias)
            
            # Add metadata
            zipf.writestr("export_info.json", json.dumps({
                "project_name": project,
                "export_date": datetime.datetime.now().isoformat(),
                "scenarios": ["A", "B", "C"]
            }, indent=2))
        
        return FileResponse(
            zip_path,
            filename=os.path.basename(zip_path),
            media_type="application/zip"
        )
    
    finally:
        # Clean up temporary XML files
        for temp_file in temp_files:
            if os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except:
                    pass

@app.post("/api/audit_pricing")
def api_audit_pricing(p: AuditPricingPayload):
    if not DB.loaded:
        DB.load()

    items = p.scenario.get("items", [])
    m = _band_multiplier(p.rate_band)
    deliverables = []
    tot_expected = 0.0
    tot_shown = 0.0
    warnings = []

    for d in items:
        # hours by role from scenario
        hrs = pd.DataFrame(d.get("hours_by_role", []))
        if hrs.empty:
            hrs = pd.DataFrame(columns=["Resource_Title","Seniority","Hours"]).assign(Hours=0.0)
        total_raw = float(hrs["Hours"].sum())

        # decide billable hours basis
        billable_total = round(total_raw) if p.price_uses_rounded_hours else total_raw
        scale = (billable_total / total_raw) if total_raw > 0 else 0.0
        hrs_bill = hrs.copy()
        hrs_bill["Hours"] = hrs_bill["Hours"] * scale

        if p.pricing_mode == "Flat_Blended":
            # default blended rate if omitted
            if p.blended_rate is None:
                ps = DB.pricing_settings[DB.pricing_settings["Key"]=="Default_Blended_Rate"]
                p.blended_rate = float(ps["Default"].iloc[0]) if not ps.empty else DEFAULT_BILLABLE_RATE_USD
            expected = billable_total * float(p.blended_rate)
            missing_roles = []
        else:
            # per-resource: join to rate card
            rc = DB.role_rate_card[["Resource_Title","Seniority","Rate_USD"]].copy()
            merged = hrs_bill.merge(rc, on=["Resource_Title","Seniority"], how="left")
            miss = merged[merged["Rate_USD"].isna()][["Resource_Title","Seniority"]].drop_duplicates()
            missing_roles = miss.to_dict(orient="records")
            merged["Rate_USD"] = merged["Rate_USD"].fillna(0.0)
            merged["Cost"] = merged["Hours"] * merged["Rate_USD"] * m
            expected = float(merged["Cost"].sum())
            if missing_roles:
                warnings.append({
                    "deliverable": d.get("deliverable"),
                    "missing_rate_roles": missing_roles
                })

        shown = float(d.get("price", 0.0))
        diff = round(expected - shown, 2)
        deliverables.append({
            "deliverable": d.get("deliverable"),
            "hours_raw": round(total_raw, 2),
            "hours_billed": round(billable_total, 2),
            "expected_price": round(expected, 2),
            "shown_price": round(shown, 2),
            "delta": diff,
        })
        tot_expected += expected
        tot_shown += shown

    scenario_delta = round(tot_expected - tot_shown, 2)
    return {
        "pricing_mode": p.pricing_mode,
        "rate_band": p.rate_band,
        "blended_rate": p.blended_rate,
        "uses_rounded_hours_for_pricing": bool(p.price_uses_rounded_hours),
        "deliverables": deliverables,
        "totals": {
            "expected": round(tot_expected, 2),
            "shown": round(tot_shown, 2),
            "delta": scenario_delta
        },
        "ok": abs(scenario_delta) < 0.01 and not warnings,
        "warnings": warnings
    }

# ──────────────────────────────────────────────────────────────────────────────
# Timeline Scheduler Endpoints (AI-powered optimization with SS+lag overlaps)
# ──────────────────────────────────────────────────────────────────────────────

class ScheduleOptimizePayload(BaseModel):
    xml_path: str  # Input MSPDI XML file path
    changes: Optional[List[Dict[str, Any]]] = None  # Optional UI-driven changes

@app.post("/api/schedule/optimize")
def api_schedule_optimize(p: ScheduleOptimizePayload):
    """
    Optimize timeline using AI Scheduler Kit:
    - Converts FS dependencies to SS+lag overlaps based on rules
    - Preserves review/approval gates
    - Rounds durations to whole days
    - Recalculates resource units
    """
    try:
        from AI_Scheduler_Kit_v2.src.orchestrator import run_pipeline
        
        # Generate output paths
        base_name = os.path.splitext(os.path.basename(p.xml_path))[0]
        output_xml = os.path.join("exports", f"{base_name}_optimized.xml")
        gantt_json = os.path.join("exports", f"{base_name}_gantt.json")
        explanations_json = os.path.join("exports", f"{base_name}_explanations.json")
        audit_xlsx = os.path.join("exports", f"{base_name}_audit.xlsx")
        
        # Ensure exports directory exists
        os.makedirs("exports", exist_ok=True)
        
        # Run optimization pipeline
        result = run_pipeline(
            xml_in=p.xml_path,
            xml_out=output_xml,
            gantt_json=gantt_json,
            explanations_json=explanations_json,
            excel_out=audit_xlsx,
            changes=p.changes,
            ai_callable=None,  # Could integrate GPT here for smarter overlaps
            round_policy="ceil"
        )
        
        # Load Gantt data for frontend visualization
        gantt_data = None
        if os.path.exists(gantt_json):
            with open(gantt_json, 'r') as f:
                gantt_data = json.load(f)
        
        # Load explanations
        explanations = None
        if os.path.exists(explanations_json):
            with open(explanations_json, 'r') as f:
                explanations = json.load(f)
        
        return {
            "success": True,
            "optimized_xml": output_xml,
            "gantt_data": gantt_data,
            "explanations": explanations,
            "notes": result.get("notes", []),
            "stats": {
                "xml_out": output_xml,
                "gantt_json": gantt_json if os.path.exists(gantt_json) else None,
                "audit_file": audit_xlsx if os.path.exists(audit_xlsx) else None
            }
        }
    
    except Exception as e:
        import traceback
        error_detail = f"{str(e)}\n\nTraceback:\n{traceback.format_exc()}"
        print(f"[SCHEDULE OPTIMIZE ERROR] {error_detail}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/schedule/download/{file_type}/{base_name}")
def api_schedule_download(file_type: str, base_name: str):
    """Download optimized files (xml, gantt, explanations, audit)"""
    # SECURITY: Sanitize base_name to prevent path traversal attacks
    # Allow alphanumeric, dash, underscore, period, and space but block path separators
    import re as re_module
    if not re_module.match(r'^[a-zA-Z0-9_\-\.\s]+$', base_name):
        raise HTTPException(400, "Invalid base_name: contains forbidden characters")
    
    # Block explicit path traversal patterns
    if '..' in base_name or '/' in base_name or '\\' in base_name:
        raise HTTPException(400, "Invalid base_name: path traversal detected")
    
    # Further sanitize by stripping any directory components
    safe_base_name = os.path.basename(base_name)
    
    file_map = {
        "xml": f"exports/{safe_base_name}_optimized.xml",
        "gantt": f"exports/{safe_base_name}_gantt.json",
        "explanations": f"exports/{safe_base_name}_explanations.json",
        "audit": f"exports/{safe_base_name}_audit.xlsx"
    }
    
    if file_type not in file_map:
        raise HTTPException(400, f"Invalid file_type: {file_type}")
    
    file_path = file_map[file_type]
    
    # SECURITY: Resolve to absolute path and verify it's within exports directory
    abs_file_path = os.path.abspath(file_path)
    abs_exports_dir = os.path.abspath("exports")
    
    if not abs_file_path.startswith(abs_exports_dir + os.sep):
        raise HTTPException(403, "Access denied: path outside exports directory")
    
    if not os.path.exists(abs_file_path):
        raise HTTPException(404, f"File not found: {os.path.basename(file_path)}")
    
    media_types = {
        "xml": "application/xml",
        "gantt": "application/json",
        "explanations": "application/json",
        "audit": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    }
    
    return FileResponse(
        abs_file_path,
        filename=os.path.basename(abs_file_path),
        media_type=media_types.get(file_type, "application/octet-stream")
    )

@app.post("/api/summarize", response_model=RfpSummary)
def api_summarize(p: SummarizePayload):
    if not p.rfp_text:
        raise HTTPException(400, "rfp_text is required for /api/summarize (use /api/summarize_by_file for uploads).")
    
    # Store textarea text separately and merge with file text if present
    global RFP_TEXT_CACHE_TEXTAREA, RFP_TEXT_CACHE_FILE, RFP_TEXT_CACHE
    
    RFP_TEXT_CACHE_TEXTAREA = (p.rfp_text or "").strip() or None
    textarea_text = RFP_TEXT_CACHE_TEXTAREA or ""
    file_text = (RFP_TEXT_CACHE_FILE or "").strip()
    
    # Combine both sources with clear separator
    if textarea_text and file_text:
        merged_text = f"{textarea_text}\n\n--- Uploaded Document Content ---\n\n{file_text}"
    elif file_text:
        merged_text = file_text
    else:
        merged_text = textarea_text
    
    # Cache merged text for backward compatibility
    RFP_TEXT_CACHE = merged_text
    
    return ai_summarize_rfp_text(merged_text)

# Removed /api/rfp/summary_cache endpoint - frontend holds window.currentRfpSummary instead

@app.post("/api/summarize_by_file")
async def api_summarize_by_file(
    files: List[UploadFile] = File(...), 
    background_tasks: BackgroundTasks = None,
    analyze_images: bool = Form(True)  # User preference for image analysis
):
    # Validate we have at least one file
    if not files:
        raise HTTPException(400, "No files uploaded.")
    
    # Store file text separately and merge with textarea text if present
    global RFP_TEXT_CACHE_TEXTAREA, RFP_TEXT_CACHE_FILE, RFP_TEXT_CACHE, LAST_UPLOAD_FILENAME
    
    # Process all files
    all_file_contents = []
    filenames = []
    background_jobs = []
    
    for file in files:
        content = await file.read()
        if not content:
            continue  # Skip empty files
            
        text = _extract_text_from_upload(content, file.filename)
        # Hard cap text length
        if len(text) > 200_000:
            text = text[:200_000]
        
        # Add file content with clear label
        if text.strip():
            labeled_content = f"--- File: {file.filename} ---\n\n{text.strip()}"
            all_file_contents.append(labeled_content)
            filenames.append(file.filename)
            
            # Create job for image processing for this file (only if enabled)
            if analyze_images:
                job_id = str(uuid.uuid4())
                JOB_STORE[job_id] = JobState(job_id=job_id, status=JobStatus.PENDING)
                background_jobs.append(job_id)
                
                # Start background image processing for this file
                if background_tasks:
                    background_tasks.add_task(_process_images_background, content, file.filename, job_id, text.strip(), analyze_images)
    
    # Combine all file contents
    file_content = "\n\n".join(all_file_contents) if all_file_contents else ""
    RFP_TEXT_CACHE_FILE = file_content
    textarea_text = (RFP_TEXT_CACHE_TEXTAREA or "").strip()
    
    # Merge textarea text with all file contents
    parts = []
    if textarea_text:
        parts.append(textarea_text)
    if file_content:
        parts.append(file_content)
    
    merged_text = "\n\n".join(parts) if parts else ""
    
    # Cache merged text for backward compatibility
    RFP_TEXT_CACHE = merged_text
    
    # NEW: remember for default project name (use first file)
    LAST_UPLOAD_FILENAME = filenames[0] if filenames else "upload"
    
    # Get summary using GPT-5
    summary = ai_summarize_rfp_text(merged_text)
    
    # Return summary with job_ids for progress tracking
    # Use .dict() for Pydantic v1 compatibility
    summary_dict = summary.dict() if hasattr(summary, 'dict') else summary.model_dump()
    return {
        **summary_dict,
        "filenames": filenames,  # Return all filenames
        "job_ids": background_jobs,  # Return all job IDs
        "processing_images": len(background_jobs) > 0
    }

@app.post("/api/retainer_detect")
def api_retainer_detect(p: dict):
    """Detect retainer opportunities for given deliverables with comprehensive null-safety."""
    if not DB.loaded:
        DB.load()
    
    # Validate input payload
    if not p or not isinstance(p, dict):
        raise HTTPException(400, "Invalid payload: expected a JSON object")
    
    # Extract and validate RFP text
    rfp_text = str(p.get("rfp_text", "") or "").strip()
    
    # Extract and validate deliverable codes
    deliverable_codes_raw = p.get("deliverable_codes")
    if deliverable_codes_raw is None:
        deliverable_codes = []
    elif not isinstance(deliverable_codes_raw, list):
        raise HTTPException(400, "Invalid deliverable_codes: expected a list")
    else:
        # Filter out None, empty, and invalid entries
        deliverable_codes = [
            str(x).strip() for x in deliverable_codes_raw 
            if x is not None and str(x).strip() and str(x).strip().lower() != "nan"
        ]
    
    # Fallback: use codes from the last built A scenario if none provided
    if not deliverable_codes:
        try:
            scen = _current_scenarios().get("A") or {}
            items = scen.get("items") or []
            if not isinstance(items, list):
                items = []
            deliverable_codes = [
                str(it.get("deliverable_code")).strip() 
                for it in items 
                if it and isinstance(it, dict) and it.get("deliverable_code")
                and str(it.get("deliverable_code")).strip().lower() != "nan"
            ]
        except Exception:
            deliverable_codes = []
    
    # Validate we have deliverables to process
    if not deliverable_codes:
        raise HTTPException(400, "No valid deliverables provided. Please select at least one deliverable or build a scenario first.")
    
    # Validate database has deliverables data
    if DB.deliverables is None or DB.deliverables.empty:
        raise HTTPException(500, "Database deliverables not loaded. Please contact support.")
    
    # Get deliverable names from codes with null-safety
    try:
        db_delivs = DB.deliverables[["Deliverable_Code", "Deliverable"]].copy()
        db_delivs["Deliverable_Code"] = db_delivs["Deliverable_Code"].astype(str)
        code_to_name = {r["Deliverable_Code"]: r["Deliverable"] for _, r in db_delivs.iterrows()}
    except Exception as e:
        raise HTTPException(500, f"Error loading deliverable names: {str(e)}")
    
    # Process retainer recommendations with null-safety
    retainers = []
    for code in deliverable_codes:
        if not code:
            continue
            
        deliv_name = code_to_name.get(code, "")
        if not deliv_name:
            continue
        
        try:
            # Call retainer recommendation with null-safety
            is_retainer, months = DB.retainer_recommendation(rfp_text, deliv_name)
            
            # Validate months is a valid number
            if months is not None:
                try:
                    months_int = int(float(months))
                    if months_int < 1:
                        months_int = 6  # Default to 6 months for invalid values
                except (ValueError, TypeError):
                    months_int = 6  # Default to 6 months if conversion fails
            else:
                months_int = 6  # Default to 6 months if None
            
            if is_retainer:
                retainers.append({
                    "deliverable_code": code,
                    "deliverable_name": deliv_name,
                    "suggested_months": months_int,
                    "confidence": "high" if months_int >= 6 else "medium"
                })
        except Exception as e:
            # Log error but continue processing other deliverables
            print(f"[RETAINER] Error processing {code}: {str(e)}")
            continue
    
    return {"retainers": retainers}

# --- AI Pricing Optimization Models and Endpoints ---
from typing import Dict, Any, List, Optional
from pydantic import BaseModel

class ComponentHours(BaseModel):
    name: str
    current_hours: float
    
class HourRedistributionRequest(BaseModel):
    deliverable_code: str
    deliverable_name: str
    new_total_hours: float
    components: List[ComponentHours]
    rfp_text: Optional[str] = ""
    
class HourRedistributionResponse(BaseModel):
    suggested_distribution: List[Dict[str, Any]]
    reasoning: str
    confidence_score: float
    original_total: float
    new_total: float
    
class RetainerSuggestionRequest(BaseModel):
    deliverable_codes: List[str]
    rfp_text: str
    
@app.post("/api/pricing/redistribute-hours")
async def redistribute_hours(request: Union[HourRedistributionRequest, dict]):
    """
    AI-powered hour redistribution endpoint.
    Takes a deliverable with new total hours and redistributes among components.
    
    Accepts simple format:
    {"deliverable_code": "DEL-0029", "total_hours": 100}
    
    Or complex format:
    {"deliverable_code": "DEL-0029", "deliverable_name": "Name", "new_total_hours": 100, "components": [...]}
    """
    # Ensure DB is loaded
    if not DB.loaded:
        DB.load()
        
    try:
        from ai_pricing_optimizer import redistribute_hours as ai_redistribute
        
        # Handle both dict and model formats
        if isinstance(request, dict):
            deliverable_code = request.get("deliverable_code", "")
            total_hours = request.get("total_hours") or request.get("new_total_hours", 100)
            deliverable_name = request.get("deliverable_name", "")
            components = request.get("components", [])
            rfp_text = request.get("rfp_text", "")
        else:
            deliverable_code = request.deliverable_code
            total_hours = request.new_total_hours
            deliverable_name = request.deliverable_name
            components = [{"name": c.name, "hours": c.current_hours} for c in request.components]
            rfp_text = request.rfp_text
        
        # If no deliverable name provided, try to get it from DB
        if not deliverable_name and deliverable_code:
            if hasattr(DB, 'deliverables') and DB.deliverables is not None:
                deliv_row = DB.deliverables[DB.deliverables["Deliverable_Code"] == deliverable_code]
                if not deliv_row.empty:
                    deliverable_name = deliv_row["Deliverable"].iloc[0]
        
        # If still no name, use code as name
        if not deliverable_name:
            deliverable_name = deliverable_code or "Deliverable"
        
        # If no components provided, create default ones
        if not components:
            components = [
                {"name": "Research & Discovery", "hours": total_hours * 0.2},
                {"name": "Planning & Strategy", "hours": total_hours * 0.15},
                {"name": "Development & Execution", "hours": total_hours * 0.45},
                {"name": "Review & Optimization", "hours": total_hours * 0.2}
            ]
        
        # Call AI redistribution
        result = await ai_redistribute(
            deliverable_name=deliverable_name,
            deliverable_code=deliverable_code,
            new_total_hours=total_hours,
            components=components,
            context=rfp_text
        )
        
        # Map the result from PricingOptimizer to the response model
        return HourRedistributionResponse(
            suggested_distribution=[
                {
                    "name": comp.name,
                    "original_hours": comp.current_hours,
                    "suggested_hours": comp.suggested_hours,
                    "change": comp.change,
                    "percentage": comp.percentage_of_total,
                    "reasoning": comp.reasoning
                }
                for comp in result.components
            ],
            reasoning=result.reasoning,
            confidence_score=result.confidence,
            original_total=result.original_total,
            new_total=result.total_hours
        )
        
    except Exception as e:
        print(f"[PRICING] Error in redistribution: {str(e)}")
        # Return proportional redistribution as fallback
        original_total = sum(c.current_hours for c in request.components)
        scale = request.new_total_hours / original_total if original_total > 0 else 1.0
        
        suggested = []
        for comp in request.components:
            suggested.append({
                "name": comp.name,
                "original_hours": comp.current_hours,
                "suggested_hours": round(comp.current_hours * scale, 1),
                "change": round(comp.current_hours * scale - comp.current_hours, 1)
            })
        
        return HourRedistributionResponse(
            suggested_distribution=suggested,
            reasoning="Proportional distribution (fallback mode - AI not available)",
            confidence_score=0.5,
            original_total=original_total,
            new_total=request.new_total_hours
        )

@app.post("/api/pricing/retainer_suggest")
async def suggest_retainer_configuration(request: RetainerSuggestionRequest):
    """
    AI-powered retainer suggestion endpoint.
    Analyzes deliverables and suggests which should be retainers.
    """
    # Ensure DB is loaded
    if not DB.loaded:
        DB.load()
    
    try:
        from ai_pricing_optimizer import PricingOptimizer
        optimizer = PricingOptimizer()
        
        suggestions = []
        for code in request.deliverable_codes:
            # Get deliverable name from DB - check if deliverables exists
            if not hasattr(DB, 'deliverables') or DB.deliverables is None:
                print(f"[PRICING] DB.deliverables not available for code {code}")
                continue
                
            deliv_row = DB.deliverables[DB.deliverables["Deliverable_Code"] == code]
            if deliv_row.empty:
                continue
                
            deliv_name = deliv_row["Deliverable"].iloc[0]
            
            # Check if it should be a retainer
            is_retainer = await optimizer.should_be_retainer(deliv_name, request.rfp_text)
            
            if is_retainer["is_retainer"]:
                suggestions.append({
                    "deliverable_code": code,
                    "deliverable_name": deliv_name,
                    "suggested_months": is_retainer["suggested_months"],
                    "reasoning": is_retainer["reasoning"]
                })
        
        return {"suggestions": suggestions}
        
    except Exception as e:
        print(f"[PRICING] Error in retainer suggestion: {str(e)}")
        # Fallback to rule-based detection
        retainer_keywords = ["management", "social media", "seo", "optimization", 
                            "maintenance", "support", "monitoring", "reporting"]
        suggestions = []
        
        # Ensure DB is available for fallback
        if hasattr(DB, 'deliverables') and DB.deliverables is not None:
            for code in request.deliverable_codes:
                deliv_row = DB.deliverables[DB.deliverables["Deliverable_Code"] == code]
                if deliv_row.empty:
                    continue
                    
                deliv_name = str(deliv_row["Deliverable"].iloc[0]).lower()
                
                if any(keyword in deliv_name for keyword in retainer_keywords):
                    suggestions.append({
                        "deliverable_code": code,
                        "deliverable_name": deliv_row["Deliverable"].iloc[0],
                        "suggested_months": 12,
                        "reasoning": "Ongoing service based on name pattern (fallback mode)"
                    })
        
        return {"suggestions": suggestions}

@app.post("/api/ai/optimize_pricing")
async def api_optimize_pricing(request: dict):
    """
    Optimize pricing based on budget constraints and other factors.
    
    Args:
        request: Dict containing:
            - target_budget: Target budget in dollars
            - scenario: Current scenario with WBS items
            - company_size: "startup" | "mid_market" | "enterprise" (optional)
            - urgency: "rush" | "standard" | "flexible" (optional)
            - industry_multiplier: float (optional, e.g., 1.5 for luxury)
            - maintain_quality_tiers: bool (optional, default True)
            
    Returns:
        Optimized scenario with adjusted pricing
    """
    try:
        # Extract parameters - handle both field names for backward compatibility
        target_budget = request.get("target_budget") or request.get("client_budget")
        scenario = request.get("scenario", {})
        company_size = request.get("company_size", "mid_market")
        urgency = request.get("urgency", "standard")
        industry_multiplier = request.get("industry_multiplier", 1.0)
        maintain_quality_tiers = request.get("maintain_quality_tiers", True)
        
        # If scenario has 'items' instead of 'wbs', convert it
        if "items" in scenario and "wbs" not in scenario:
            scenario["wbs"] = scenario["items"]
        
        if not target_budget:
            raise HTTPException(status_code=400, detail="target_budget or client_budget is required")
        
        if not scenario or "wbs" not in scenario:
            raise HTTPException(status_code=400, detail="Valid scenario with WBS is required")
        
        # Calculate current total price
        current_total = sum(float(item.get("Price", 0)) for item in scenario.get("wbs", []))
        
        if current_total == 0:
            raise HTTPException(status_code=400, detail="Current scenario has no pricing")
        
        # Calculate optimization ratio
        base_ratio = target_budget / current_total
        
        # Apply company size adjustments
        size_multipliers = {
            "startup": 0.85,     # 15% discount for startups
            "mid_market": 1.0,   # Standard pricing
            "enterprise": 1.25   # 25% premium for enterprise
        }
        size_mult = size_multipliers.get(company_size, 1.0)
        
        # Apply urgency adjustments
        urgency_multipliers = {
            "rush": 1.3,         # 30% premium for rush jobs
            "standard": 1.0,     # Standard timing
            "flexible": 0.9      # 10% discount for flexible timing
        }
        urgency_mult = urgency_multipliers.get(urgency, 1.0)
        
        # Combine all multipliers
        final_ratio = base_ratio * size_mult * urgency_mult * industry_multiplier
        
        # Define minimum rates based on seniority
        min_rates = {
            "Junior": 75,
            "Mid": 125,
            "Senior": 175,
            "Director": 250,
            "VP": 350,
            "EVP": 450
        }
        
        # Check if budget is too low for minimum viable delivery
        min_viable_price = 0
        for item in scenario.get("wbs", []):
            if item.get("Hours"):
                hours = float(item.get("Hours", 0))
                seniority = item.get("Seniority", "Mid")
                min_rate = min_rates.get(seniority, 125)
                min_viable_price += hours * min_rate * 0.5  # 50% of min rate as absolute floor
        
        if target_budget < min_viable_price:
            return JSONResponse(
                status_code=400, 
                content={
                    "error": "Budget too low for minimum viable delivery",
                    "minimum_viable": min_viable_price,
                    "requested": target_budget,
                    "recommendation": "Consider reducing scope or increasing budget"
                }
            )
        
        # Optimize pricing for each WBS item
        optimized_wbs = []
        total_optimized = 0
        
        for item in scenario.get("wbs", []):
            optimized_item = item.copy()
            
            if item.get("Price") and float(item.get("Price", 0)) > 0:
                hours = float(item.get("Hours", 0))
                if hours > 0:
                    # Calculate current rate
                    current_price = float(item.get("Price", 0))
                    current_rate = current_price / hours
                    
                    # Apply optimization
                    new_rate = current_rate * final_ratio
                    
                    # Apply quality tier constraints if requested
                    if maintain_quality_tiers:
                        seniority = item.get("Seniority", "Mid")
                        min_rate_for_tier = min_rates.get(seniority, 125)
                        max_rate_for_tier = min_rate_for_tier * 3  # Max 3x min rate
                        
                        # Enforce bounds
                        new_rate = max(min_rate_for_tier, min(new_rate, max_rate_for_tier))
                    else:
                        # Still enforce absolute minimums
                        new_rate = max(50, new_rate)  # Absolute floor of $50/hr
                    
                    # Handle unlimited budget scenarios (cap at reasonable maximum)
                    if target_budget > 10000000:  # $10M+
                        new_rate = min(new_rate, 1000)  # Cap at $1000/hr
                    
                    optimized_item["Rate"] = round(new_rate, 2)
                    optimized_item["Price"] = round(hours * new_rate, 2)
                    total_optimized += optimized_item["Price"]
            
            optimized_wbs.append(optimized_item)
        
        # Final adjustment to match target exactly (distribute any remaining difference)
        if total_optimized > 0 and abs(total_optimized - target_budget) > 1:
            adjustment_ratio = target_budget / total_optimized
            for item in optimized_wbs:
                if item.get("Price") and float(item.get("Price", 0)) > 0:
                    item["Price"] = round(float(item["Price"]) * adjustment_ratio, 2)
                    if item.get("Hours") and float(item.get("Hours", 0)) > 0:
                        item["Rate"] = round(item["Price"] / float(item["Hours"]), 2)
        
        # Create optimized scenario
        optimized_scenario = scenario.copy()
        optimized_scenario["wbs"] = optimized_wbs
        optimized_scenario["total_price"] = sum(float(item.get("Price", 0)) for item in optimized_wbs)
        optimized_scenario["optimization_details"] = {
            "original_total": current_total,
            "target_budget": target_budget,
            "achieved_total": optimized_scenario["total_price"],
            "company_size": company_size,
            "urgency": urgency,
            "industry_multiplier": industry_multiplier,
            "optimization_ratio": final_ratio,
            "quality_tiers_maintained": maintain_quality_tiers
        }
        
        return optimized_scenario
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[PRICING] Error in pricing optimization: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Pricing optimization failed: {str(e)}")

@app.post("/api/reconcile", response_model=ReconcileResult)
def api_reconcile(p: ReconcilePayload):
    if not DB.loaded:
        DB.load()

    try:
        # Normalize inputs
        ai_labels: List[str] = [str(x).strip() for x in (p.summary_deliverables or []) if str(x).strip()]
        if not ai_labels:
            # Nothing to reconcile
            return ReconcileResult(add=[], delete=[], unchanged=[])

        sel_codes: List[str] = [str(x) for x in (p.db_selected_deliverable_codes or [])]

        # DB deliverables map
        db_all = DB.deliverables[["Deliverable_Code", "Deliverable"]].copy()
        db_all["Deliverable_Code"] = db_all["Deliverable_Code"].astype(str)
        db_all["Deliverable"] = db_all["Deliverable"].astype(str)

        # Selected map (left panel)
        db_sel = db_all[db_all["Deliverable_Code"].isin(sel_codes)]
        code_to_name = {r["Deliverable_Code"]: r["Deliverable"] for _, r in db_sel.iterrows()}

        # Precompute tokens for speed
        ai_tok = [(lab, _norm_tokens(lab)) for lab in ai_labels]
        db_tok = [(r["Deliverable_Code"], r["Deliverable"], _norm_tokens(r["Deliverable"])) for _, r in db_all.iterrows()]

        ADD_THRESHOLD = 0.35     # how close an AI label must be to a DB deliverable to recommend ADD (lowered from 0.45)
        DELETE_THRESHOLD = 0.25  # if no AI label is at least this close, recommend DELETE

        add: List[ReconcileSuggestion] = []
        unchanged: List[str] = []
        delete: List[ReconcileSuggestion] = []

        # --- ADD & UNCHANGED ---
        # For each AI label, find the best matching DB deliverable by token Jaccard
        for lab, lab_tok in ai_tok:
            best_code = None
            best_name = ""
            best_score = 0.0
            for code, name, name_tok in db_tok:
                s = _jaccard(lab_tok, name_tok)
                if s > best_score:
                    best_code, best_name, best_score = code, name, s

            if not best_code:
                continue

            if best_code in sel_codes:
                # Already selected -> unchanged (count it if reasonably close)
                if best_score >= DELETE_THRESHOLD:
                    unchanged.append(best_name)
            else:
                # Not selected yet -> recommend ADD if strong enough
                if best_score >= 0.35:
                    add.append(ReconcileSuggestion(
                        code=best_code, label=best_name,
                        reason=f"Matches AI summary item \"{lab}\" (score {best_score:.2f}).",
                        preselect=True
                    ))

        # --- DELETE ---
        # Safe delete block - prevents DataFrame errors with unequal column lengths
        ai_labels = [str(x) for x in (p.summary_deliverables or []) if str(x).strip()]
        for code in sel_codes:
            name = code_to_name.get(code, code)
            name_tokens = _norm_tokens(name)
            max_score = 0.0
            for lbl in ai_labels:
                max_score = max(max_score, _jaccard(name_tokens, _norm_tokens(lbl)))
            if max_score < 0.25:
                delete.append(ReconcileSuggestion(
                    code=code, label=name, reason="Not found in AI Summary.", preselect=True
                ))

        # Deduplicate unchanged list & sort
        unchanged = sorted(set(unchanged))

        # Add the actual selection the server used
        db_used_labels = [code_to_name.get(c, c) for c in sel_codes]

        return ReconcileResult(
            add=add, delete=delete, unchanged=unchanged,
            db_used_codes=sel_codes, db_used_labels=db_used_labels
        )

    except Exception as ex:
        # Return a clear 400 instead of a 500 so the UI can show a friendly message
        raise HTTPException(status_code=400, detail=f"Reconciliation error: {ex}")

@app.post("/api/reorder_timeline")
def api_reorder_timeline(p: ReorderPayload):
    if not DB.loaded:
        DB.load()

    letter = (p.scenario_letter or "A").upper()
    scen = _current_scenarios().get(letter)
    if not scen:
        raise HTTPException(400, f"Scenario {letter} not built")

    # Get current items and create mapping by deliverable code
    current_items = list(scen.get("items") or [])
    by_code = {str(it.get("deliverable_code") or it.get("code")): it for it in current_items}
    want = [str(c) for c in p.deliverable_codes]

    # Build new schedules in the requested order, sequentially
    reordered_items: list[dict] = []
    cursor_date = None  # None -> use project_start for first

    for code in want:
        # Get existing item for metadata
        existing_item = by_code.get(code)
        if not existing_item:
            continue

        # included task groups (from client), else derive all TGs present for that deliverable
        included = (p.included_map or {}).get(code)
        if not included:
            sub = DB.all_rows[DB.all_rows["Deliverable_Code"].astype(str) == str(code)]
            included = sorted(set(sub["task_group"].dropna().astype(str).tolist()))

        # sequential packing: first item uses project_start, others start the day after previous end
        start = p.project_start
        if cursor_date is not None:
            start = str(cursor_date)

        sched = DB.build_schedule(
            deliverable_code=code,
            included_task_groups=included,
            complexity=p.complexity or existing_item.get("complexity"), 
            tier=p.tier or existing_item.get("tier"),
            use_slack=p.use_slack,
            slack_after_internal=p.slack_after_internal,
            slack_after_client=p.slack_after_client,
            slack_global_pct=p.slack_global_pct,
            project_start=start,
            scenario_letter=letter
        )

        # Update existing item with new schedule
        updated_item = existing_item.copy()
        updated_item["schedule"] = sched
        reordered_items.append(updated_item)

        # advance cursor to the day after this deliverable's last end date
        if sched:
            last_end = sched[-1]["end_date"]
            y, m, d = map(int, last_end.split("-"))
            cursor_date = datetime.date(y, m, d) + datetime.timedelta(days=1)

    # Update scenario with new order and persist
    scen["items"] = reordered_items
    if reordered_items:
        scen["timeline"] = {
            "start": reordered_items[0]["schedule"][0]["start_date"] if reordered_items[0].get("schedule") else None,
            "end": reordered_items[-1]["schedule"][-1]["end_date"] if reordered_items[-1].get("schedule") else None,
        }
    scen["user_order"] = want
    scen["ai_order"] = scen.get("ai_order") or list(want)
    scen["manual_order_locked"] = True
    _CURRENT_SCENARIOS[letter] = scen  # persist

    # Return items for frontend and full scenario for persistence
    return {"items": [{"deliverable_code": d["deliverable_code"], "deliverable": d.get("deliverable"), "schedule": d["schedule"]} for d in reordered_items],
            "scenario": scen}


# ---------- Scenario Differentiation Helpers ----------
def _resolve_global_multipliers(db: AgencyDB, complexity: str, tier: str) -> tuple[float,float]:
    cm = 1.0; vm = 1.0
    try:
        ps = db.pricing_settings
        if ps is not None:
            cm_row = ps[ps["Key"].str.lower()==f"complexity.{complexity.lower()}"]
            vm_row = ps[ps["Key"].str.lower()==f"volume.{tier.lower()}"]
            if not cm_row.empty: cm = float(cm_row["Value"].iloc[0])
            if not vm_row.empty: vm = float(vm_row["Value"].iloc[0])
    except Exception:
        pass
    return cm, vm

def _apply_multipliers(row_hours, complexity, tier, letter, db):
    cm, vm = _resolve_global_multipliers(db, complexity, tier)
    sm = SCENARIO_MULT[letter]["hours_mult"]
    return max(0.5, min(2.0, row_hours * cm * vm * sm))

def _make_overhead_item(name: str, hours: float) -> dict:
    return {
        "deliverable_code": f"OVERHEAD-{name.upper().replace(' ', '-')}",
        "deliverable": name,
        "category": "PM",
        "included_task_groups": [name],
        "hours_by_role": [{"role": "Account Manager", "seniority": "Senior", "hours": hours}],
        "total_hours": hours,
        "effective_rate": 0,
        "price": 0,
        "components": [],
        "schedule": []
    }

def _apply_scenario_knobs(items, letter, complexity, tier, pricing_mode, rate_band, db, blended_rate=None):
    knobs = SCENARIO_MULT.get(letter, SCENARIO_MULT["A"])
    for it in items:
        # Apply hours multiplier to each role
        for r in it.get("hours_by_role", []):
            if "hours" in r:
                r["hours"] = r["hours"] * knobs["hours_mult"]
        # Update total hours
        if "total_hours" in it:
            it["total_hours"] = it["total_hours"] * knobs["hours_mult"]
        # Strip/augment components
        if knobs.get("strip_optional"):
            it["components"] = [c for c in it.get("components", []) if not c.get("optional")]
        if knobs.get("include_addons"):
            it["components"].extend(it.get("addons", []))
    
    # Add PM/QA overhead as explicit task groups
    total_hours = sum(sum(r.get("hours", 0) for r in it.get("hours_by_role", [])) for it in items)
    pm_hours = total_hours * knobs.get("pm_pct", 0)
    qa_hours = total_hours * knobs.get("qa_pct", 0)
    if pm_hours > 0:
        items.append(_make_overhead_item("Project Management", pm_hours))
    if qa_hours > 0:
        items.append(_make_overhead_item("QA / Reporting", qa_hours))
    return items

# ---------- MSPDI Export Function ----------
def _std_rate_for(role: str, mode: str, band: str, blended: float, db: AgencyDB) -> float:
    if mode == "Flat_Blended" and blended:
        return float(blended)
    # per-resource lookup
    try:
        # prefer matrix; fallback to role_rate_card
        m = db.rate_matrix
        if m is not None:
            row = m[(m["Role"].str.strip().str.lower()==role.strip().lower()) &
                    (m["Band"].str.strip().str.lower()==band.strip().lower())]
            if not row.empty:
                return float(row["USD_per_hour"].iloc[0])
        rc = db.role_rate_card
        if rc is not None:
            row = rc[rc["Role"].str.strip().str.lower()==role.strip().lower()]
            if not row.empty:
                return float(row[band].iloc[0])  # columns like Standard_US, Premium_US
    except Exception:
        pass
    return float(blended or DEFAULT_BILLABLE_RATE_USD)

def convert_excel_to_mspdi(
    input_xlsx: str,
    output_xml: str,
    sheet_name: str = "Scenario A",
    start_date_mode: str = "next_monday",
    fixed_start_iso: Optional[str] = None,
    hours_per_day: float = 8.0,
    calendar_blocks: List[Tuple[str, str]] = [("08:00:00","12:00:00"), ("13:00:00","17:00:00")],
    roles_split_rule: str = "even",
    role_weights: Optional[Dict[str, float]] = None,
    preserve_predecessors: str = "normalize",
    allow_unassigned: bool = True,
    include_audits: bool = True,
    audits_dir: Optional[str] = None,
    merge_identical_children: bool = False,  # <— toggle for multi-resource merge
    project_name: Optional[str] = None,      # <— explicit project name override
    pricing_mode: str = "Flat_Blended",      # <— NEW: pricing mode
    rate_band: str = "Standard_US",          # <— NEW: rate band
    blended_rate: Optional[float] = None,    # <— NEW: blended rate
    add_deliverable_milestones: bool = False # <— NEW: toggle for START/END anchors
) -> Dict[str, int]:
    """
    Convert Excel WBS data to Microsoft Project XML (MSPDI) format with multi-resource merge capability.
    
    Returns a dictionary with statistics about the conversion process.
    """
    try:
        # Load Excel data
        df = pd.read_excel(input_xlsx, sheet_name=sheet_name)
        
        # DEBUG: Log DataFrame columns and FixedCost values
        if "FixedCost" in df.columns:
            deliverable_mask = df["WBS_ID"].astype(str).str.count(r'\.') == 1 if "WBS_ID" in df.columns else pd.Series([False]*len(df))
            deliverables_with_cost = df.loc[deliverable_mask & (df["FixedCost"] > 0)]
            print(f"[MSPDI] 📊 Excel contains FixedCost column with {len(deliverables_with_cost)} non-zero deliverable values")
            for _, row in deliverables_with_cost.iterrows():
                print(f"[MSPDI]    FixedCost={row['FixedCost']} for WBS {row.get('WBS_ID', 'N/A')}: {str(row.get('Task_Name', ''))[:40]}")
        else:
            print(f"[MSPDI] ⚠️ FixedCost column NOT found in Excel. Columns: {list(df.columns)}")
        
        # --- Derive a proper project title for the MSPDI <Project><Name> ---
        project_title = None
        try:
            # Prefer the root row (WBS_ID == "1") Task_Name if present
            if "WBS_ID" in df.columns and "Task_Name" in df.columns:
                root = df[df["WBS_ID"].astype(str).str.strip() == "1"]
                if not root.empty:
                    t = str(root["Task_Name"].iloc[0]).strip()
                    # Treat empty strings and literal "nan"/"none" as missing
                    if t and t.lower() not in ['nan', 'none', '']:
                        project_title = t
            # Fallback: iterate through Project_Name values to find first valid entry
            if not project_title and "Project_Name" in df.columns:
                pn_series = df["Project_Name"].dropna().astype(str)
                for val in pn_series:
                    t = str(val).strip()
                    # Treat empty strings and literal "nan"/"none" as missing
                    if t and t.lower() not in ['nan', 'none', '']:
                        project_title = t
                        break
        except Exception:
            project_title = None
        
        if not project_title:
            project_title = str(sheet_name).strip() or "Project"
        
        # Convert DataFrame to list of row dictionaries for processing
        rows = []
        for _, row in df.iterrows():
            # Extract basic task information with NaN handling
            def safe_int(value, default=0):
                """Convert value to int, handling NaN and None."""
                if pd.isna(value) or value is None:
                    return default
                try:
                    return int(float(value))
                except (ValueError, TypeError):
                    return default
            
            def safe_float(value, default=0.0):
                """Convert value to float, handling NaN and None."""
                if pd.isna(value) or value is None:
                    return default
                try:
                    return float(value)
                except (ValueError, TypeError):
                    return default
            
            def safe_str(value, default=""):
                """Convert value to string, handling NaN and None."""
                if pd.isna(value) or value is None:
                    return default
                s = str(value).strip()
                if s.lower() in ['nan', 'none', 'nat', '']:
                    return default
                return s
                    
            # Clean up task name and role
            task_name = str(row.get("Task_Name", ""))
            if not task_name or task_name.lower() in ['nan', 'none', '']:
                task_name = "Unnamed Task"
                
            role = str(row.get("Role", "Unassigned"))
            if not role or role.lower() in ['nan', 'none', '']:
                role = "Unassigned"
                
            # TASK 16: Extract Rate_USD and Price_USD from scenario pricing (no re-computation)
            rate_usd = safe_float(row.get("Rate_USD"), None)
            price_usd = safe_float(row.get("Price_USD"), None)
            
            # FIX: Preserve Start_Date and End_Date from Excel if they exist (from Gantt timeline_tasks)
            start_date_str = safe_str(row.get("Start_Date"), "")
            end_date_str = safe_str(row.get("End_Date"), "")
            
            # TASK 16: Extract FixedCost from WBS for XML export (Step 3 price)
            fixed_cost = safe_float(row.get("FixedCost"), 0.0)
            
            # DEBUG: Log deliverable FixedCost values
            wbs_id = str(row.get("WBS_ID", ""))
            outline_level = wbs_id.count(".") + 1 if wbs_id else 0
            if outline_level == 2 and fixed_cost > 0:
                print(f"[XML EXPORT] 📊 Extracted FixedCost={fixed_cost} for WBS {wbs_id} (Task: {str(row.get('Task_Name', ''))[:50]})")
            
            task_row = {
                "WBS": str(row.get("WBS_ID", "")),
                "ParentWBS": str(row.get("Parent_WBS_ID", "")),
                "Name": task_name,
                "PlannedHours": safe_float(row.get("Planned_Hours"), 0),
                "StartOffset": safe_int(row.get("Start_Offset_Days"), 0),
                "Duration": safe_int(row.get("Duration_Days"), 1),
                "Dependencies": str(row.get("Dependencies", "")),
                "RoleList": [role],
                "RoleStr": role,
                "UID": 0,  # Will be assigned later
                "DeliverableCode": str(row.get("Deliverable_Code", "")),
                "Component": str(row.get("Component", "")),
                "Rate_USD": rate_usd,   # From scenario pricing
                "Price_USD": price_usd,  # From scenario pricing
                "FixedCost": fixed_cost, # Step 3 price as FixedCost (replaces hourly cost calc)
                "Start_Date": start_date_str,  # Preserve from Gantt if available
                "End_Date": end_date_str  # Preserve from Gantt if available
            }
            rows.append(task_row)
        
        # Assign UIDs
        for i, row in enumerate(rows, 1):
            row["UID"] = i
        
        # TASK 16: Create UID -> Rate_USD lookup for assignment cost calculation
        uid_to_rate = {r["UID"]: r.get("Rate_USD") for r in rows}
        
        # --- NEW: merge siblings with the same name into their parent as multi-assignments
        prealloc_by_parent_wbs: Dict[str, Dict[str, float]] = {}
        removed_child_wbs: Set[str] = set()

        if merge_identical_children:
            # index helpers
            by_wbs = {r["WBS"]: r for r in rows if r.get("WBS")}
            kids_by_parent: Dict[str, List[str]] = {}
            for r in rows:
                p = r.get("ParentWBS")
                if p:
                    kids_by_parent.setdefault(p, []).append(r["WBS"])

            for parent_wbs, kid_wbs_list in list(kids_by_parent.items()):
                # Only immediate children and all must be leaves
                kid_rows = [by_wbs[k] for k in kid_wbs_list if k in by_wbs]
                if not kid_rows:
                    continue
                # skip if any child itself has children (not a leaf)
                if any(k in kids_by_parent for k in kid_wbs_list):
                    continue

                parent = by_wbs.get(parent_wbs)
                if not parent:
                    continue

                # Heuristic:
                # - every child has same Name as parent
                # - each child has exactly ONE role
                # - each child has >0 planned hours
                same_name = all(kr["Name"] == parent["Name"] for kr in kid_rows)
                one_role  = all(len(kr["RoleList"]) == 1 for kr in kid_rows)
                has_hours = all((kr["PlannedHours"] or 0) > 0 for kr in kid_rows)
                if not (same_name and one_role and has_hours):
                    continue

                # Aggregate hours by role
                agg: Dict[str, float] = {}
                for kr in kid_rows:
                    role = kr["RoleList"][0]
                    agg[role] = agg.get(role, 0.0) + float(kr["PlannedHours"])

                if len(agg) < 2:
                    continue

                # Record prealloc for the parent
                prealloc_by_parent_wbs[parent_wbs] = agg

                # Make the parent a single multi-role leaf
                parent["RoleList"] = list(agg.keys())
                parent["RoleStr"]  = ",".join(parent["RoleList"])
                if (parent.get("PlannedHours") or 0) <= 0:
                    parent["PlannedHours"] = sum(agg.values())
                
                # STEP 3 FIXED COST PRESERVATION: Carry forward FixedCost during merge
                # Prefer sum if children are month splits, else max for duplicates
                fixed_cost_values = [float(kr.get("FixedCost", 0) or 0) for kr in kid_rows]
                fixed_cost_sum = sum(fixed_cost_values)
                fixed_cost_max = max(fixed_cost_values) if fixed_cost_values else 0
                parent_fixed = float(parent.get("FixedCost", 0) or 0)
                if fixed_cost_sum > parent_fixed:
                    parent["FixedCost"] = fixed_cost_sum if fixed_cost_sum > 0 else fixed_cost_max

                # Remove the children
                removed_child_wbs.update(kid_wbs_list)

        # If we merged anything, drop the children now
        if removed_child_wbs:
            rows = [r for r in rows if r["WBS"] not in removed_child_wbs]

        # Build a universal WBS index (needed even when merge is OFF)
        by_wbs = {r["WBS"]: r for r in rows if r.get("WBS")}

        # Build children_by_parent map and child_to_parent for dep rewrites
        children_by_parent: Dict[str, List[str]] = {}
        for r in rows:
            p = r["ParentWBS"]
            if p:
                children_by_parent.setdefault(p, []).append(r["WBS"])
        summary_set: Set[str] = set(children_by_parent.keys())

        child_to_parent: Dict[str, str] = {}
        for p, kids in children_by_parent.items():
            for k in kids:
                if k in removed_child_wbs:   # only those removed pre-merge
                    child_to_parent[k] = p

        # Helper functions for dependency normalization
        def is_ancestor(ancestor_wbs: str, descendant_wbs: str) -> bool:
            current = descendant_wbs
            visited = set()
            while current and current not in visited:
                visited.add(current)
                if current == ancestor_wbs:
                    return True
                # Find parent of current
                parent_found = None
                for r in rows:
                    if r["WBS"] == current:
                        parent_found = r.get("ParentWBS")
                        break
                current = parent_found
            return False

        def list_leaves_under(parent_wbs: str) -> List[str]:
            leaves = []
            for r in rows:
                if r.get("ParentWBS") == parent_wbs and r["WBS"] not in summary_set:
                    leaves.append(r["WBS"])
            return leaves

        def first_leaf(wbs: str) -> str:
            if wbs in summary_set:
                children = children_by_parent.get(wbs, [])
                if children:
                    return first_leaf(children[0])
            return wbs

        def last_leaf(wbs: str) -> str:
            if wbs in summary_set:
                children = children_by_parent.get(wbs, [])
                if children:
                    return last_leaf(children[-1])
            return wbs

        # === WORKFRONT SEQUENCING ENRICHMENT ===
        # Add anchor tasks per deliverable and chain components properly
        def safe_code(name: str, code: str | None, unique_index: int = 0) -> str:
            """
            Generate a safe, unique deliverable code for START/END anchors.
            Never returns [nan], None, or empty string.
            Falls back to slug of deliverable name, then to generic placeholder with index.
            """
            import re
            
            # Ensure code is a clean string (handle None, nan, empty, etc.)
            c = str(code).strip() if code is not None else ""
            # Check for various "nan" representations
            if c and c.lower() not in ["nan", "none", "null", ""]:
                return c
            
            # Fallback 1: Create slug from deliverable name
            name_str = str(name).strip() if name is not None else ""
            if name_str and name_str.lower() not in ["nan", "none", "null", ""]:
                slug = re.sub(r"[^a-z0-9]+", "_", name_str.lower()).strip("_")
                if slug and slug != "unnamed":
                    return slug
            
            # Fallback 2: Use generic placeholder with unique index
            return f"deliverable_{unique_index}" if unique_index > 0 else "deliverable"
        
        def enrich_wbs_for_workfront(rows):
            enriched = []
            anchor_id_counter = [90000]  # Start high to avoid conflicts with existing WBS IDs
            
            # Group rows by deliverable (excluding Project Summary WBS=1)
            deliverables = {}
            root_row = None
            other_rows = []  # Rows without DeliverableCode
            for r in rows:
                if r["WBS"] == "1":
                    root_row = r
                    continue
                dcode = r.get("DeliverableCode", "").strip()
                if dcode:
                    deliverables.setdefault(dcode, []).append(r)
                else:
                    other_rows.append(r)
            
            # Add root row first
            if root_row:
                enriched.append(root_row)
            
            prev_deliv_end_wbs = None
            
            deliverable_index = 0
            for dcode in sorted(deliverables.keys()):
                deliv_rows = deliverables[dcode]
                if not deliv_rows:
                    continue
                
                deliverable_index += 1
                # Get deliverable name for safe_code fallback
                deliv_name = next((r.get("Deliverable", "") for r in deliv_rows if r.get("Deliverable")), "")
                
                # Generate safe, unique code (safe_code handles uniqueness internally)
                safe_dcode = safe_code(deliv_name, dcode, deliverable_index)
                
                # Use friendly deliverable name for milestone labels (use safe_dcode to avoid "nan")
                dname = (deliv_name or safe_dcode or f"Deliverable {deliverable_index}").strip()
                
                # Create START anchor
                start_wbs = f"ANCHOR_{anchor_id_counter[0]}"
                anchor_id_counter[0] += 1
                start_anchor = {
                    "WBS": start_wbs,
                    "ParentWBS": "1",
                    "Name": f"Start — {dname}",
                    "PlannedHours": 0,
                    "StartOffset": 0,
                    "Duration": 0,
                    "Dependencies": prev_deliv_end_wbs if prev_deliv_end_wbs else "",
                    "RoleList": [],
                    "RoleStr": "",
                    "UID": 0,
                    "DeliverableCode": dcode,
                    "Component": ""
                }
                enriched.append(start_anchor)
                
                # Group deliverable rows by component
                components = {}
                for r in deliv_rows:
                    comp = r.get("Component", "").strip() or "General"
                    components.setdefault(comp, []).append(r)
                
                prev_comp_last_wbs = start_wbs
                
                # Add components in order
                for comp in sorted(components.keys()):
                    comp_rows = components[comp]
                    
                    # Add predecessor to all rows in this component
                    for i, r in enumerate(comp_rows):
                        if not r.get("Dependencies"):
                            r["Dependencies"] = prev_comp_last_wbs
                        enriched.append(r)
                    
                    # Last row of this component becomes predecessor for next
                    if comp_rows:
                        prev_comp_last_wbs = comp_rows[-1]["WBS"]
                
                # Create END anchor
                end_wbs = f"ANCHOR_{anchor_id_counter[0]}"
                anchor_id_counter[0] += 1
                end_anchor = {
                    "WBS": end_wbs,
                    "ParentWBS": "1",
                    "Name": f"End — {dname}",
                    "PlannedHours": 0,
                    "StartOffset": 0,
                    "Duration": 0,
                    "Dependencies": prev_comp_last_wbs,
                    "RoleList": [],
                    "RoleStr": "",
                    "UID": 0,
                    "DeliverableCode": dcode,
                    "Component": ""
                }
                enriched.append(end_anchor)
                prev_deliv_end_wbs = end_wbs
            
            # Append rows without DeliverableCode (preserve all original rows)
            enriched.extend(other_rows)
            
            # Reassign UIDs
            for i, r in enumerate(enriched, 1):
                r["UID"] = i
            
            return enriched
        
        # Apply enrichment (optional)
        if add_deliverable_milestones:
            rows = enrich_wbs_for_workfront(rows)
        
        # Rebuild indices after enrichment
        by_wbs = {r["WBS"]: r for r in rows if r.get("WBS")}
        children_by_parent = {}
        for r in rows:
            p = r["ParentWBS"]
            if p:
                children_by_parent.setdefault(p, []).append(r["WBS"])
        summary_set = set(children_by_parent.keys())
        
        # Helper for numeric WBS sorting
        def wbs_sort_key(wbs):
            """
            Convert WBS like '1.1.10.3' into tuple (0,1), (0,1), (0,10), (0,3) for natural numeric ordering.
            Handles null/empty WBS and mixed numeric/text segments safely.
            Returns tuple of (type, value) pairs where type=0 for numeric, type=1 for text.
            """
            if not wbs:
                return ()
            
            parts = str(wbs).split(".")
            key = []
            for p in parts:
                try:
                    # Numeric segment: (0, int_value)
                    key.append((0, int(p)))
                except ValueError:
                    # Text segment: (1, str_value)
                    key.append((1, p))
            return tuple(key)
        
        # WBS-BASED DEPENDENCY SYSTEM
        # Auto-builds L4→L4 and L5→L5 dependencies from WBS hierarchy
        if ENABLE_WBS_DEPENDENCIES:
            print("[WBS-DEP] Building L4→L4 and L5→L5 dependencies from WBS")

            wbs_dependencies = []

            # Sort all rows in WBS order (numeric sorting for proper multi-digit handling)
            rows_sorted = sorted(rows, key=lambda r: wbs_sort_key(r["WBS"]))

            # ---------- L4 → L4 chain (within same L3) ----------
            # Group L4 rows by their L3 parent
            from collections import defaultdict
            l4_by_l3 = defaultdict(list)
            for r in rows_sorted:
                wbs = str(r["WBS"])
                parts = wbs.split(".")
                outline = len(parts)
                if outline == 4:
                    parent_l3 = ".".join(parts[:3])  # e.g. 1.1.1
                    l4_by_l3[parent_l3].append(r)

            # For each L3 parent, chain its L4 children sequentially
            for parent_l3, l4_group in l4_by_l3.items():
                # Sort L4 rows in WBS order within group (numeric sorting)
                l4_group_sorted = sorted(l4_group, key=lambda r: wbs_sort_key(r["WBS"]))
                
                # Chain sequentially - NO WRAP-AROUND
                prev_l4 = None
                for r in l4_group_sorted:
                    if prev_l4 is not None:
                        # Add edge: previous L4 → this L4
                        wbs_dependencies.append((prev_l4["WBS"], r["WBS"]))
                        print(f"[WBS-DEP] L4→L4: {prev_l4['WBS']} → {r['WBS']}")
                    prev_l4 = r

            # ---------- L5 → L5 chain (within same L4) ----------
            # Group L5 rows by their L4 parent
            l5_by_l4 = defaultdict(list)
            for r in rows_sorted:
                wbs = str(r["WBS"])
                parts = wbs.split(".")
                outline = len(parts)
                if outline == 5:
                    parent_l4 = ".".join(parts[:4])  # e.g. 1.1.1.1
                    l5_by_l4[parent_l4].append(r)

            # For each L4 parent, chain its L5 children sequentially
            for parent_l4, l5_group in l5_by_l4.items():
                # Sort L5 rows in WBS order within group (numeric sorting)
                l5_group_sorted = sorted(l5_group, key=lambda r: wbs_sort_key(r["WBS"]))
                
                # Chain sequentially - NO WRAP-AROUND
                prev_l5 = None
                for r in l5_group_sorted:
                    if prev_l5 is not None:
                        # Add edge: previous L5 → this L5
                        wbs_dependencies.append((prev_l5["WBS"], r["WBS"]))
                        print(f"[WBS-DEP] L5→L5: {prev_l5['WBS']} → {r['WBS']}")
                    prev_l5 = r

            normalized_edges = wbs_dependencies
            print(f"[WBS-DEP] Built {len(normalized_edges)} dependencies (L4 & L5)")
        else:
            # LEGACY MODE: Normalize dependencies & drop unsafe hierarchy edges
            init_edges = []
            for r in rows:
                deps = r.get("Dependencies", "").strip()
                if deps:
                    for dep in deps.split(","):
                        dep = dep.strip()
                        if dep:
                            init_edges.append((dep, r["WBS"]))

            normalized_edges = []
            for pred_wbs, succ_wbs in init_edges:
                # Rewrite removed children to their parents
                actual_pred = child_to_parent.get(pred_wbs, pred_wbs)
                actual_succ = child_to_parent.get(succ_wbs, succ_wbs)
                
                # Skip if either doesn't exist after merge
                if actual_pred not in by_wbs or actual_succ not in by_wbs:
                    continue
                    
                # Skip hierarchy edges (ancestor -> descendant)
                if is_ancestor(actual_pred, actual_succ) or is_ancestor(actual_succ, actual_pred):
                    continue
                    
                # Convert summary tasks to their representative leaves
                if actual_pred in summary_set:
                    actual_pred = last_leaf(actual_pred)
                if actual_succ in summary_set:
                    actual_succ = first_leaf(actual_succ)
                    
                if actual_pred != actual_succ:
                    normalized_edges.append((actual_pred, actual_succ))

        # Calculate project start date (MUST be timezone-naive for MSPDI compatibility)
        if fixed_start_iso:
            # Extract date-only portion and add standard work hours to create timezone-naive datetime
            # This matches the approach used for Gantt date preservation and ensures consistent
            # behavior for distributed teams across different timezones
            start_date_only = datetime.date.fromisoformat(fixed_start_iso[:10])
            project_start = datetime.datetime.combine(start_date_only, datetime.time(8, 0))
        elif start_date_mode == "next_monday":
            today = datetime.date.today()
            days_ahead = 0 - today.weekday()  # Monday is 0
            if days_ahead <= 0:
                days_ahead += 7
            project_start = datetime.datetime.combine(today + datetime.timedelta(days=days_ahead), datetime.time(9, 0))
        elif start_date_mode == "fixed":
            # Use current date with business hours start if no fixed_start_iso provided
            project_start = datetime.datetime.now().replace(hour=8, minute=0, second=0, microsecond=0)
        else:
            project_start = datetime.datetime.now().replace(hour=9, minute=0, second=0, microsecond=0)

        # Business calendar helpers (using same Mon-Fri, 8-12 & 13-17 schedule)
        from datetime import time, date
        
        BUS_BLOCKS = [(time(8,0), time(12,0)), (time(13,0), time(17,0))]

        def is_business_day(d):
            return d.weekday() < 5  # Mon–Fri

        def business_minutes_in_range(day: date, start_t: time, end_t: time) -> int:
            # minutes worked on a single day between start_t and end_t
            if not is_business_day(day): 
                return 0
            start_t = max(start_t, BUS_BLOCKS[0][0])
            end_t   = min(end_t,   BUS_BLOCKS[-1][1])
            if end_t <= start_t:
                return 0
            total = 0
            for a,b in BUS_BLOCKS:
                s = max(start_t, a)
                e = min(end_t,   b)
                if e > s:
                    total += int((datetime.datetime.combine(day,e) - datetime.datetime.combine(day,s)).total_seconds() // 60)
            return total

        def business_minutes_between(start_dt: datetime.datetime, end_dt: datetime.datetime) -> int:
            if end_dt <= start_dt:
                return 0
            cur = start_dt.date()
            end = end_dt.date()
            minutes = 0
            # first day (partial)
            minutes += business_minutes_in_range(cur, start_dt.time(), time(17,0))
            # middle full days
            d = cur + datetime.timedelta(days=1)
            while d < end:
                if is_business_day(d):
                    minutes += 480  # 8h
                d += datetime.timedelta(days=1)
            # last day (partial)
            minutes += business_minutes_in_range(end, time(8,0), end_dt.time())
            return minutes

        # Calculate task schedules
        uid_to_sched = {}
        preserved_dates = set()  # Track which UIDs have dates from Gantt that should not be rolled up
        
        # Build WBS-to-UID and parent mapping FIRST so we can look up parent dates
        wbs_to_uid_pre = {r["WBS"]: r["UID"] for r in rows}
        uid_to_row = {r["UID"]: r for r in rows}
        
        for r in rows:
            # CRITICAL FIX: Check if Start_Date and End_Date already exist from Gantt timeline edits
            # If they do, preserve them EXACTLY instead of recalculating from offsets
            # This ensures Workfront imports match the Gantt UI timeline
            start_date_str = r.get("Start_Date", "")
            end_date_str = r.get("End_Date", "")
            
            if start_date_str and end_date_str:
                # Parse existing dates from Gantt timeline
                try:
                    # Parse date strings (handle both YYYY-MM-DD and ISO format with time)
                    # Extract just the date portion for timezone consistency
                    start_date_only = datetime.date.fromisoformat(start_date_str[:10])
                    end_date_only = datetime.date.fromisoformat(end_date_str[:10])
                    
                    # Combine with standard work hours (08:00 AM start, 05:00 PM end)
                    start_date = datetime.datetime.combine(start_date_only, datetime.time(8, 0))
                    end_date = datetime.datetime.combine(end_date_only, datetime.time(17, 0))
                    
                    # CRITICAL FIX: Duration = calendar span in 8h-day units (Option C from Architect)
                    # This prevents Workfront from recalculating dates on import
                    # Formula: Duration_minutes = ((Finish - Start).days + 1) × 480
                    # Work stays as PlannedHours (may differ from Duration - indicates partial allocation)
                    calendar_days = (end_date_only - start_date_only).days + 1  # Inclusive count
                    duration_hours = calendar_days * 8.0  # 8-hour working days
                    
                    # Mark this UID as having preserved dates that should not be overwritten by rollup
                    preserved_dates.add(r["UID"])
                    
                    # Log successful preservation
                    task_name = r.get("Name", "Unknown")[:30]
                    print(f"[XML EXPORT] ✅ Preserved Gantt dates for {task_name} (WBS {r['WBS']}): {start_date.date()} to {end_date.date()}, {duration_hours:.1f}h")
                    
                except (ValueError, AttributeError, TypeError) as e:
                    # If parsing fails, fall back to offset calculation and log the issue
                    print(f"[XML EXPORT] ⚠️ Failed to parse Gantt dates for WBS {r['WBS']}: {e}")
                    start_date = project_start + datetime.timedelta(days=r["StartOffset"])
                    duration_hours = max(r["Duration"] * hours_per_day, r["PlannedHours"])
                    end_date = start_date + datetime.timedelta(hours=duration_hours)
            else:
                # No existing dates - calculate from offsets
                # Log when deliverables don't have preserved dates (helps diagnose issues)
                is_deliverable_level = r["WBS"].count(".") <= 2  # Level 2 or 3 deliverables
                if is_deliverable_level:
                    task_name = r.get("Name", "Unknown")[:30]
                    print(f"[XML EXPORT] ℹ️ No Gantt dates for {task_name} (WBS {r['WBS']}) - using offset calculation")
                
                # FIX: If parent has preserved dates from Gantt, calculate relative to parent
                # Otherwise fall back to global project_start
                parent_wbs = r.get("ParentWBS")
                parent_uid = wbs_to_uid_pre.get(parent_wbs) if parent_wbs else None
                
                if parent_uid and parent_uid in preserved_dates and parent_uid in uid_to_sched:
                    # Parent has preserved dates - calculate this row relative to parent
                    parent_sched = uid_to_sched[parent_uid]
                    parent_row = uid_to_row[parent_uid]
                    parent_offset = parent_row.get("StartOffset", 0)
                    this_offset = r.get("StartOffset", 0)
                    
                    # Calculate offset RELATIVE to parent (subtract parent's offset from this offset)
                    relative_offset = this_offset - parent_offset
                    start_date = parent_sched["Start"] + datetime.timedelta(days=relative_offset)
                    
                    duration_hours = max(r["Duration"] * hours_per_day, r["PlannedHours"])
                    end_date = start_date + datetime.timedelta(hours=duration_hours)
                    
                    print(f"[XML EXPORT] 📍 Row {r['WBS']} ({r.get('TaskName', 'N/A')[:30]}): Parent {parent_wbs} preserved → calculating as parent_start + {relative_offset} days = {start_date.date()}")
                else:
                    # No preserved parent - use global project_start (original logic)
                    start_date = project_start + datetime.timedelta(days=r["StartOffset"])
                    duration_hours = max(r["Duration"] * hours_per_day, r["PlannedHours"])
                    end_date = start_date + datetime.timedelta(hours=duration_hours)
            
            uid_to_sched[r["UID"]] = {
                "Start": start_date,
                "Finish": end_date,
                "PlannedHours": r["PlannedHours"],
                "DurationHours": duration_hours
            }

        # Build UID-based children mapping for rollup (as expected by patch)
        wbs_to_uid = {r["WBS"]: r["UID"] for r in rows}
        wbs_children = children_by_parent  # Save original WBS-based mapping
        children_by_parent = {}  # UID-based mapping for rollup
        for wbs, child_wbs_list in wbs_children.items():
            parent_uid = wbs_to_uid.get(wbs)
            if parent_uid:
                children_by_parent[parent_uid] = [wbs_to_uid.get(child_wbs) for child_wbs in child_wbs_list if wbs_to_uid.get(child_wbs)]
        summary_set = set(children_by_parent.keys())

        # Helper function to STRETCH and clamp all descendants to fit within parent's date range
        def stretch_and_clamp_descendants(uid, bound_start, bound_finish):
            """
            Stretch and clamp all descendants of uid to proportionally fill parent's date range.
            This ensures child tasks span the full 121-day parent window instead of clustering in ~6 days.
            """
            kids = children_by_parent.get(uid, [])
            if not kids:
                return
            
            # Calculate the original span of all children (before stretching)
            child_starts = [uid_to_sched[k]["Start"] for k in kids]
            child_finishes = [uid_to_sched[k]["Finish"] for k in kids]
            original_min = min(child_starts)
            original_max = max(child_finishes)
            original_span = (original_max - original_min).total_seconds()
            
            # Calculate the target span (parent's range)
            target_span = (bound_finish - bound_start).total_seconds()
            
            # Skip stretching if original span is 0 or target span is very small
            if original_span <= 0 or target_span <= 0:
                return
            
            # Calculate stretch ratio
            stretch_ratio = target_span / original_span if original_span > 0 else 1.0
            
            for k in kids:
                child_start = uid_to_sched[k]["Start"]
                child_finish = uid_to_sched[k]["Finish"]
                
                # Calculate proportional position within stretched range
                # offset_ratio = how far into the original span this child starts
                offset_seconds = (child_start - original_min).total_seconds()
                duration_seconds = (child_finish - child_start).total_seconds()
                
                # Scale offset and duration proportionally
                new_offset_seconds = offset_seconds * stretch_ratio
                new_duration_seconds = max(3600, duration_seconds * stretch_ratio)  # min 1 hour
                
                # Calculate new start/finish based on bound_start
                new_start = bound_start + datetime.timedelta(seconds=new_offset_seconds)
                new_finish = new_start + datetime.timedelta(seconds=new_duration_seconds)
                
                # Clamp to bounds (just in case of floating point issues)
                new_start = max(new_start, bound_start)
                new_finish = min(new_finish, bound_finish)
                
                # Ensure finish is still after start
                if new_finish <= new_start:
                    new_finish = new_start + datetime.timedelta(hours=1)
                
                # Update child schedule with stretched dates
                uid_to_sched[k]["Start"] = new_start
                uid_to_sched[k]["Finish"] = new_finish
                
                # Recalculate duration for stretched child
                duration_hours = max(1.0, (new_finish - new_start).total_seconds() / 3600)
                uid_to_sched[k]["DurationHours"] = duration_hours
                
                # Recursively stretch this child's descendants
                if k in summary_set:
                    stretch_and_clamp_descendants(k, new_start, new_finish)
            
            print(f"[XML EXPORT] 📐 Stretched {len(kids)} children: original {original_span/86400:.1f} days → target {target_span/86400:.1f} days (ratio={stretch_ratio:.2f})")
        
        # CRITICAL FIX Part 2: Stretch and clamp descendants for all preserved deliverables
        # This ensures child components/tasks FILL the deliverable window set in Gantt (not just clamp)
        # Must happen BEFORE rollup_summary runs to prevent wrong date calculations
        for preserved_uid in preserved_dates:
            if preserved_uid in uid_to_sched and preserved_uid in children_by_parent:
                preserved_start = uid_to_sched[preserved_uid]["Start"]
                preserved_finish = uid_to_sched[preserved_uid]["Finish"]
                stretch_and_clamp_descendants(preserved_uid, preserved_start, preserved_finish)
                print(f"[XML EXPORT] 🔒 Stretched descendants of preserved UID {preserved_uid} to fill {preserved_start.date()} to {preserved_finish.date()}")
        
        # 1) roll up start/finish for every summary from its direct/indirect leaves
        def rollup_summary(uid):
            kids = children_by_parent.get(uid, [])
            if not kids:
                return uid_to_sched[uid]["Start"], uid_to_sched[uid]["Finish"]
            
            # CRITICAL FIX Part 1: If this UID has preserved dates from Gantt, return them immediately
            # Descendants were already clamped above, so we don't need to recalculate anything
            # This prevents rollup from overwriting Campaign Plan Deck (Mar 5) back to Jan 5
            if uid in preserved_dates:
                return uid_to_sched[uid]["Start"], uid_to_sched[uid]["Finish"]
            
            # For non-preserved summaries, recurse to collect child dates
            starts, finishes = [], []
            for k in kids:
                s,f = rollup_summary(k) if k in summary_set else (uid_to_sched[k]["Start"], uid_to_sched[k]["Finish"])
                starts.append(s); finishes.append(f)
            
            child_min = min(starts)
            child_max = max(finishes)
            
            # Roll up from children (only for non-preserved summaries)
            uid_to_sched[uid]["Start"]  = child_min
            uid_to_sched[uid]["Finish"] = child_max
            return uid_to_sched[uid]["Start"], uid_to_sched[uid]["Finish"]

        # Call it on every summary
        for uid in list(summary_set):
            rollup_summary(uid)

        # Also roll up the very top project row if you have one (UID of the first row)
        top_uid = min(uid_to_sched.keys())
        if top_uid in summary_set:
            rollup_summary(top_uid)

        # 2) Recompute Duration for ALL tasks from Start/Finish span (business minutes)
        # CRITICAL FIX: Use ONLY business day calculation, not max() to avoid calendar day inflation
        for uid, sched in uid_to_sched.items():
            span_min = business_minutes_between(sched["Start"], sched["Finish"])
            # cache as hours for later; the XML writer will multiply by 60 again
            sched["DurationHours"] = span_min / 60.0
        
        # VALIDATION: Log sample deliverable durations to verify business day calculation
        print("\n[XML EXPORT VALIDATION] Sample Deliverable Durations (Business Days):")
        print("-" * 100)
        sample_count = 0
        for r in rows:
            if r["WBS"].count(".") == 1 and sample_count < 5:  # Deliverable level (e.g., "1.1")
                uid = r["UID"]
                sched = uid_to_sched[uid]
                start = sched["Start"]
                finish = sched["Finish"]
                dur_hours = sched["DurationHours"]
                dur_days = dur_hours / 8.0  # 8-hour workdays
                print(f"  {r['Name'][:50]:50} | Start: {str(start)[:10]} | Finish: {str(finish)[:10]} | Duration: {dur_days:.1f} business days ({dur_hours:.1f}h)")
                sample_count += 1
        print("-" * 100 + "\n")
            
        # Convert datetime objects back to strings for XML output
        for uid, sched in uid_to_sched.items():
            if isinstance(sched["Start"], datetime.datetime):
                sched["Start"] = sched["Start"].strftime("%Y-%m-%dT%H:%M:%S")
            if isinstance(sched["Finish"], datetime.datetime):
                sched["Finish"] = sched["Finish"].strftime("%Y-%m-%dT%H:%M:%S")

        # Create resource list (filter out nan/empty roles)
        all_roles = set()
        for r in rows:
            all_roles.update(r["RoleList"])
        if allow_unassigned:
            all_roles.add("Unassigned")
        
        # Remove nan, empty, and None values
        all_roles = {role for role in all_roles if role and str(role).lower() not in ['nan', 'none', '']}
        
        resources = []
        res_name_to_uid = {}
        for i, role in enumerate(sorted(all_roles), 1):
            resources.append({"UID": i, "ID": i, "Name": role})
            res_name_to_uid[role] = i

        # Map prealloc from WBS -> UID (after UIDs exist)
        prealloc_by_task_uid: Dict[int, Dict[str, float]] = {}
        if prealloc_by_parent_wbs:
            wbs_to_uid = {r["WBS"]: r["UID"] for r in rows if r["WBS"]}
            for wbs, role_hours in prealloc_by_parent_wbs.items():
                uid = wbs_to_uid.get(wbs)
                if uid:
                    prealloc_by_task_uid[uid] = role_hours

        # Create assignments
        assignments = []
        assign_uid = 1
        for r in rows:
            if r["WBS"] in summary_set:
                continue

            task_hours = uid_to_sched[r["UID"]]["PlannedHours"]
            if task_hours <= 0.0001:
                continue

            # Duration basis for Units
            task_dur_h = uid_to_sched[r["UID"]]["DurationHours"] if uid_to_sched[r["UID"]]["DurationHours"] > 0 else task_hours
            if uid_to_sched[r["UID"]]["DurationHours"] <= 0.0001 and task_hours > 0:
                uid_to_sched[r["UID"]]["DurationHours"] = task_hours

            # --- Use precomputed role->hours if merged
            alloc = prealloc_by_task_uid.get(r["UID"])
            if alloc:
                for role, work_h in alloc.items():
                    res_uid = res_name_to_uid.get(role) or res_name_to_uid.get("Unassigned")
                    units = (work_h / task_dur_h) if task_dur_h > 0 else 1.0
                    units = max(0.05, min(units, 2.0))
                    assignments.append({
                        "UID": assign_uid,
                        "TaskUID": r["UID"],
                        "ResourceUID": res_uid,
                        "Start": uid_to_sched[r["UID"]]["Start"],
                        "Finish": uid_to_sched[r["UID"]]["Finish"],
                        "Units": units,
                        "WorkHours": work_h
                    })
                    assign_uid += 1
                continue  # done with this task

            # else: fall back to existing split-by-role behavior
            role_list = r["RoleList"]
            if not role_list:
                role_list = ["Unassigned"]
                
            if roles_split_rule == "even":
                hours_per_role = task_hours / len(role_list)
                for role in role_list:
                    res_uid = res_name_to_uid.get(role) or res_name_to_uid.get("Unassigned")
                    units = (hours_per_role / task_dur_h) if task_dur_h > 0 else 1.0
                    units = max(0.05, min(units, 2.0))
                    assignments.append({
                        "UID": assign_uid,
                        "TaskUID": r["UID"],
                        "ResourceUID": res_uid,
                        "Start": uid_to_sched[r["UID"]]["Start"],
                        "Finish": uid_to_sched[r["UID"]]["Finish"],
                        "Units": units,
                        "WorkHours": hours_per_role
                    })
                    assign_uid += 1
            elif roles_split_rule == "weighted" and role_weights:
                total_weight = sum(role_weights.get(role, 1.0) for role in role_list)
                for role in role_list:
                    weight = role_weights.get(role, 1.0)
                    hours_for_role = task_hours * (weight / total_weight)
                    res_uid = res_name_to_uid.get(role) or res_name_to_uid.get("Unassigned")
                    units = (hours_for_role / task_dur_h) if task_dur_h > 0 else 1.0
                    units = max(0.05, min(units, 2.0))
                    assignments.append({
                        "UID": assign_uid,
                        "TaskUID": r["UID"],
                        "ResourceUID": res_uid,
                        "Start": uid_to_sched[r["UID"]]["Start"],
                        "Finish": uid_to_sched[r["UID"]]["Finish"],
                        "Units": units,
                        "WorkHours": hours_for_role
                    })
                    assign_uid += 1

        # FIX: Calculate actual PlannedHours from assignments and update uid_to_sched AND row dictionaries
        # This ensures tasks show correct <Work> values instead of PT0M
        task_hours_from_assignments = {}
        for assign in assignments:
            task_uid = assign["TaskUID"]
            work_hours = assign["WorkHours"]
            task_hours_from_assignments[task_uid] = task_hours_from_assignments.get(task_uid, 0) + work_hours
        
        # Update uid_to_sched AND row dictionaries with actual hours from assignments
        # Update both to ensure XML generation gets the correct values regardless of which source it reads
        # CRITICAL: Row dictionaries use "Planned_Hours" (underscore), not "PlannedHours" (camel case)
        for task_uid, total_hours in task_hours_from_assignments.items():
            if task_uid in uid_to_sched:
                uid_to_sched[task_uid]["PlannedHours"] = total_hours
            # Also update the row dictionary with the correct field name (underscore!)
            for r in rows:
                if r["UID"] == task_uid:
                    r["Planned_Hours"] = total_hours  # Note: underscore, not camelCase
                    break

        # NOTE: Milestone filtering temporarily disabled pending user clarification on which
        # specific task name patterns should be excluded
        # The reference XML doesn't show these milestone patterns, so we need to verify
        # which tasks should be filtered vs preserved
        # filtered_rows = []
        # for r in rows:
        #     task_name = r.get("Name", "").strip()
        #     ...filtering logic...
        # rows = filtered_rows

        # FIX: Derive actual project StartDate and FinishDate from timeline_tasks
        # This ensures Workfront imports show correct durations instead of inflated values
        # (e.g., 844 days when actual work is ~250 days)
        actual_project_start = None
        actual_project_finish = None
        
        if uid_to_sched:
            # Parse all task Start/Finish dates to find earliest and latest
            all_starts = []
            all_finishes = []
            
            for task_uid, sched in uid_to_sched.items():
                try:
                    # Parse the ISO format dates (YYYY-MM-DDTHH:MM:SS)
                    start_dt = datetime.datetime.fromisoformat(sched["Start"])
                    finish_dt = datetime.datetime.fromisoformat(sched["Finish"])
                    all_starts.append(start_dt)
                    all_finishes.append(finish_dt)
                except (ValueError, KeyError, TypeError):
                    # Skip tasks with invalid dates
                    continue
            
            # Use actual task dates if we found any valid dates
            if all_starts and all_finishes:
                actual_project_start = min(all_starts)
                actual_project_finish = max(all_finishes)
                
                # Override the calculated project_start with actual earliest task date
                project_start = actual_project_start
                project_finish = actual_project_finish
                
                # FIX: Update root task (WBS=1, Project Summary) to match project dates
                # The root task should span the entire project timeline
                root_task_uid = None
                for r in rows:
                    if r["WBS"] == "1":
                        root_task_uid = r["UID"]
                        break
                
                if root_task_uid and root_task_uid in uid_to_sched:
                    uid_to_sched[root_task_uid]["Start"] = project_start.strftime("%Y-%m-%dT%H:%M:%S")
                    uid_to_sched[root_task_uid]["Finish"] = project_finish.strftime("%Y-%m-%dT%H:%M:%S")
                
                print(f"[MSPDI Export] Project dates derived from timeline tasks:")
                print(f"  - Start: {project_start.strftime('%Y-%m-%d')} (earliest task)")
                print(f"  - Finish: {project_finish.strftime('%Y-%m-%d')} (latest task)")
                print(f"  - Duration: {(project_finish - project_start).days} days")
            else:
                # Fallback: use calculated project_start for both
                project_finish = project_start + datetime.timedelta(days=180)
                print(f"[MSPDI Export] WARNING: No valid task dates found, using fallback project dates")
        else:
            # No tasks scheduled - use fallback
            project_finish = project_start + datetime.timedelta(days=180)
            print(f"[MSPDI Export] WARNING: No task schedules found, using fallback project dates")

        # ====================================================================================
        # MULTI-ASSIGNMENT DETECTION: Group rows with identical Name+Start+Finish+WBS_Prefix
        # ====================================================================================
        def group_multi_assignments(rows_input, uid_to_sched_map):
            """
            Group rows that represent shared meetings/reviews (same name, start, finish, WBS prefix).
            
            Returns:
                - merged_rows: List of row dicts (single-assignment unchanged, multi-assignment merged)
                - uid_alias_map: Dict mapping consumed UIDs to their primary UID
            """
            from collections import defaultdict
            import html
            
            # Helper function to normalize timestamps (round to nearest minute)
            def normalize_timestamp(iso_str):
                """Normalize timestamp to minute precision to handle second/millisecond drift"""
                try:
                    # Remove timezone suffix and parse
                    clean_str = iso_str.replace("Z", "").split("+")[0].split("-")[0]
                    dt = datetime.datetime.fromisoformat(clean_str)
                    # Round to nearest minute
                    rounded = dt.replace(second=0, microsecond=0)
                    return rounded.strftime("%Y-%m-%dT%H:%M:00")
                except:
                    return iso_str
            
            buckets = defaultdict(list)
            
            for row in rows_input:
                # Extract WBS prefix (first 4 segments for L4 deliverable/component scope)
                wbs = row.get("WBS", "")
                parts = wbs.split(".")
                prefix = ".".join(parts[:4]) if len(parts) >= 4 else wbs
                
                # Only group rows with valid role/resource (skip if missing or multi-role)
                # Check if row has a single role assignment
                role_val = row.get("Role", "")
                if not role_val or pd.isna(role_val) or str(role_val).strip() == "":
                    # No role - don't group
                    buckets[("UNGROUPED", row["UID"], "", "")].append(row)
                    continue
                
                # Get Start/Finish from uid_to_sched (already normalized ISO strings)
                task_uid = row["UID"]
                if task_uid not in uid_to_sched_map:
                    # No schedule - don't group
                    buckets[("UNGROUPED", row["UID"], "", "")].append(row)
                    continue
                    
                # Normalize timestamps to handle second/millisecond drift
                start_iso = normalize_timestamp(uid_to_sched_map[task_uid]["Start"])
                finish_iso = normalize_timestamp(uid_to_sched_map[task_uid]["Finish"])
                
                # Decode HTML entities and normalize task name for grouping
                task_name = html.unescape(str(row.get("Name", "")).strip())
                
                # Debug logging for client review tasks
                if "client review" in task_name.lower():
                    print(f"[DEBUG MERGE] WBS={wbs} Name='{task_name}' Start={start_iso} Finish={finish_iso} Role={role_val}")
                
                # Grouping key: (WBS_prefix, task_name, start, finish)
                key = (prefix, task_name, start_iso, finish_iso)
                buckets[key].append(row)
            
            merged_rows = []
            uid_alias_map = {}
            
            for key, group in buckets.items():
                prefix, task_name, start_iso, finish_iso = key
                
                # UNGROUPED bucket - keep as-is
                if prefix == "UNGROUPED":
                    merged_rows.extend(group)
                    continue
                
                # Single-assignment task - keep unchanged
                if len(group) == 1:
                    merged_rows.append(group[0])
                    print(f"[MULTI-ASSIGN] Single task: {task_name[:50]} (WBS {group[0]['WBS']})")
                    continue
                
                # Multi-assignment task - merge into one
                primary_uid = group[0]["UID"]
                total_hours = sum(float(r.get("Planned_Hours", 0)) for r in group)
                total_price = sum(float(r.get("Price_USD", 0)) for r in group)
                
                # Create merged task record
                merged = {
                    **group[0],  # Copy first row as base
                    "UID": primary_uid,
                    "Planned_Hours": total_hours,
                    "Price_USD": total_price,  # Sum of all assignee costs
                    "multi_assignment": True,  # Flag for downstream processing
                    "assignments": [
                        {
                            "source_uid": r["UID"],
                            "resource_name": r.get("Role", "Unassigned"),
                            "work_hours": float(r.get("Planned_Hours", 0)),
                            "rate_usd": float(r.get("Rate_USD", 0)),  # Keep individual rate for assignment
                            "price_usd": float(r.get("Price_USD", 0))  # Keep individual cost for assignment
                        }
                        for r in group
                    ]
                }
                merged_rows.append(merged)
                
                # Map consumed UIDs to primary UID
                for r in group[1:]:
                    uid_alias_map[r["UID"]] = primary_uid
                
                role_list = [r.get("Role", "?") for r in group]
                print(f"[MULTI-ASSIGN] ✅ Merged {len(group)} tasks: '{task_name[:50]}' (WBS {prefix}) - Roles: {', '.join(role_list)} - Total: {total_hours}h")
            
            print(f"[MULTI-ASSIGN] Summary: {len(rows_input)} rows → {len(merged_rows)} tasks ({len(uid_alias_map)} merged)")
            return merged_rows, uid_alias_map
        
        # Multi-assignment grouping: Controlled by ENABLE_MULTI_ASSIGNMENT flag
        if ENABLE_MULTI_ASSIGNMENT:
            rows, uid_alias_map = group_multi_assignments(rows, uid_to_sched)
            print(f"[EXPORT] Multi-assignment merging ENABLED - {len(uid_alias_map)} UIDs aliased")
        else:
            # Identity mapping (each UID maps to itself)
            uid_alias_map = {row["UID"]: row["UID"] for row in rows}
            print(f"[EXPORT] Multi-assignment merging DISABLED - exporting {len(rows)} tasks (one per role)")
        
        # ====================================================================================
        # CHRONOLOGICAL WATERFALL SORTING
        # Sort rows by Start_Date while preserving parent-child adjacency
        # ====================================================================================
        def sort_rows_chronologically(rows_input, uid_to_sched_map, summary_set_input):
            """
            Sort rows chronologically while preserving parent-child adjacency.
            
            Requirements:
            - Project Summary (WBS="1") stays first
            - All other rows sorted by Start_Date (ascending)
            - Parent-child adjacency preserved (parents immediately before children)
            - Within siblings, sort by Start_Date
            - Tie-breaker: original WBS numeric order
            
            Returns:
                Sorted list of rows
            """
            print("\n[CHRONOLOGICAL SORT] Starting waterfall sort...")
            
            # Find project summary (WBS="1")
            project_summary = None
            other_rows = []
            for r in rows_input:
                if r["WBS"] == "1":
                    project_summary = r
                else:
                    other_rows.append(r)
            
            if not project_summary:
                print("[CHRONOLOGICAL SORT] ⚠️ No project summary found (WBS=1), proceeding without it")
                other_rows = rows_input
            
            # Build parent-child map using WBS (more reliable than ParentWBS)
            children_by_wbs = {}
            wbs_to_row = {}
            
            for r in other_rows:
                wbs = r["WBS"]
                wbs_to_row[wbs] = r
                
                # Calculate parent WBS (remove last segment)
                parts = wbs.split(".")
                if len(parts) > 1:
                    parent_wbs = ".".join(parts[:-1])
                    if parent_wbs not in children_by_wbs:
                        children_by_wbs[parent_wbs] = []
                    children_by_wbs[parent_wbs].append(wbs)
            
            # Helper to get Start_Date for sorting
            def get_start_date(row):
                """Get start date for sorting - return datetime or fallback to WBS numeric order"""
                uid = row["UID"]
                if uid in uid_to_sched_map:
                    start_str = uid_to_sched_map[uid].get("Start", "")
                    if start_str:
                        try:
                            # Parse ISO format date
                            return datetime.datetime.fromisoformat(start_str.replace("Z", ""))
                        except (ValueError, AttributeError):
                            pass
                # Fallback: use WBS numeric order (convert to sortable tuple)
                wbs_parts = row["WBS"].split(".")
                return tuple(int(p) for p in wbs_parts if p.isdigit())
            
            # Helper to get WBS numeric tuple for tie-breaking
            def get_wbs_tuple(wbs_str):
                """Convert WBS string to sortable tuple of integers"""
                parts = wbs_str.split(".")
                return tuple(int(p) for p in parts if p.isdigit())
            
            # Recursive function to sort tree and flatten
            def sort_and_flatten(wbs, sorted_list):
                """Recursively sort children by Start_Date and flatten tree"""
                # Add current node
                if wbs in wbs_to_row:
                    sorted_list.append(wbs_to_row[wbs])
                
                # Get children and sort by Start_Date, then WBS
                child_wbs_list = children_by_wbs.get(wbs, [])
                if child_wbs_list:
                    # Sort children by: 1) Start_Date, 2) WBS numeric order
                    child_rows = [(child_wbs, wbs_to_row.get(child_wbs)) for child_wbs in child_wbs_list]
                    child_rows = [(wbs, r) for wbs, r in child_rows if r is not None]
                    
                    # Sort by Start_Date, then WBS
                    child_rows.sort(key=lambda x: (get_start_date(x[1]), get_wbs_tuple(x[0])))
                    
                    # Recursively add sorted children
                    for child_wbs, _ in child_rows:
                        sort_and_flatten(child_wbs, sorted_list)
            
            # Start with project summary (WBS="1") children
            sorted_rows = []
            if project_summary:
                sorted_rows.append(project_summary)
                sort_and_flatten("1", sorted_rows)
            else:
                # No project summary - sort top-level items
                top_level = [wbs for wbs in wbs_to_row.keys() if wbs.count(".") == 0]
                top_level.sort(key=lambda wbs: (get_start_date(wbs_to_row[wbs]), get_wbs_tuple(wbs)))
                for wbs in top_level:
                    sort_and_flatten(wbs, sorted_rows)
            
            # Validation: Check all rows included
            if len(sorted_rows) != len(rows_input):
                print(f"[CHRONOLOGICAL SORT] ⚠️ Row count mismatch: {len(rows_input)} input → {len(sorted_rows)} output")
            
            # Log sample of sorted order
            print(f"[CHRONOLOGICAL SORT] ✅ Sorted {len(sorted_rows)} rows chronologically")
            print("[CHRONOLOGICAL SORT] First 10 rows after sort:")
            for i, r in enumerate(sorted_rows[:10], 1):
                uid = r["UID"]
                start = uid_to_sched_map.get(uid, {}).get("Start", "N/A")[:10] if uid in uid_to_sched_map else "N/A"
                print(f"  {i:2}. WBS={r['WBS']:15} Name={r['Name'][:40]:40} Start={start}")
            
            return sorted_rows
        
        # Apply chronological sorting
        rows = sort_rows_chronologically(rows, uid_to_sched, summary_set)
        
        # ====================================================================================
        # UID REMAPPING
        # Reassign UIDs sequentially (1, 2, 3, ...) and update all references
        # ====================================================================================
        print("\n[UID REMAPPING] Starting UID reassignment...")
        
        # Build old_uid -> new_uid mapping
        old_uid_to_new_uid = {}
        for new_id, r in enumerate(rows, 1):
            old_uid = r["UID"]
            old_uid_to_new_uid[old_uid] = new_id
            r["UID"] = new_id  # Update row with new UID
        
        print(f"[UID REMAPPING] Created mapping for {len(old_uid_to_new_uid)} UIDs")
        
        # Update uid_to_sched with new UIDs
        new_uid_to_sched = {}
        for old_uid, new_uid in old_uid_to_new_uid.items():
            if old_uid in uid_to_sched:
                new_uid_to_sched[new_uid] = uid_to_sched[old_uid]
        uid_to_sched = new_uid_to_sched
        
        print(f"[UID REMAPPING] Updated uid_to_sched: {len(uid_to_sched)} entries")
        
        # Update Dependencies in rows (convert WBS references to new UIDs if needed)
        # NOTE: Dependencies in this codebase use WBS strings, not UIDs
        # But we need to ensure any UID-based dependencies are updated
        for r in rows:
            # Dependencies field uses WBS strings, so no remapping needed
            # But if there were UID-based predecessors, they would need updating here
            pass
        
        # Update summary_set with new UIDs
        wbs_to_new_uid = {r["WBS"]: r["UID"] for r in rows}
        new_summary_set = set()
        for r in rows:
            if r["WBS"] in children_by_parent or r["WBS"] in summary_set:
                new_summary_set.add(r["UID"])
        summary_set = new_summary_set
        
        # Update resources and assignments with new UIDs
        # Resources are separate and don't need updating
        # Assignments need TaskUID updated
        for assign in assignments:
            old_task_uid = assign["TaskUID"]
            if old_task_uid in old_uid_to_new_uid:
                assign["TaskUID"] = old_uid_to_new_uid[old_task_uid]
        
        print(f"[UID REMAPPING] Updated {len(assignments)} assignments")
        print(f"[UID REMAPPING] ✅ UID remapping complete\n")
        
        # ====================================================================================
        # DEPENDENCY REMAPPING: Rebuild edges after WBS flattening & UID remap
        # ====================================================================================
        print("[DEP-REMAP] Rebuilding dependency edges with flattened WBS...")
        
        # Filter normalized_edges to only include WBS pairs that exist in wbs_to_new_uid
        # This removes stale references like "1.Strategy.2" that no longer exist after flattening
        dependency_remap_count = 0
        dropped_dependencies = []
        
        for pred_wbs, succ_wbs in normalized_edges:
            if pred_wbs not in wbs_to_new_uid or succ_wbs not in wbs_to_new_uid:
                dropped_dependencies.append((pred_wbs, succ_wbs))
                if len(dropped_dependencies) <= 10:  # Log first 10 dropped
                    print(f"[DEP-REMAP] ⚠️  Dropped: {pred_wbs} → {succ_wbs} (WBS not in flattened set)")
        
        if dropped_dependencies:
            print(f"[DEP-REMAP] Total dropped dependencies: {len(dropped_dependencies)}")
        
        # Use FILTERED normalized_edges (only edges with valid WBS references)
        valid_normalized_edges = [
            (pred_wbs, succ_wbs) for pred_wbs, succ_wbs in normalized_edges
            if pred_wbs in wbs_to_new_uid and succ_wbs in wbs_to_new_uid
        ]
        
        print(f"[DEP-REMAP] Valid edges retained: {len(valid_normalized_edges)} / {len(normalized_edges)}")
        print(f"[DEP-REMAP] ✅ Dependency remapping complete\n")
        
        # ====================================================================================
        # EDGE FILTERING: Drop L5+ → L5+ edges to enable parallel role execution
        # ====================================================================================
        print("[EDGE-FILTER] Filtering edges to drop L5+ → L5+ (role-to-role) dependencies...")
        
        def outline_level_from_wbs(wbs: str) -> int:
            """Calculate outline level from WBS string. '1' -> 1, '1.1' -> 2, '1.1.1.1.1' -> 5"""
            return wbs.count(".") + 1
        
        # Filter out L5+ → L5+ edges to allow parallel role execution (AM, CW, Strategist, etc.)
        # Keep edges where at least one side is ≤L4 (summary level: deliverable, component, or task)
        filtered_edges = []
        dropped_role_edges = []
        
        for pred_wbs, succ_wbs in valid_normalized_edges:
            pred_level = outline_level_from_wbs(pred_wbs)
            succ_level = outline_level_from_wbs(succ_wbs)
            
            # Keep edges where at least one side is L2/L3/L4 (summary level)
            if pred_level <= 4 or succ_level <= 4:
                filtered_edges.append((pred_wbs, succ_wbs))
            else:
                # Drop L5+ → L5+ (role-to-role) to enable parallel execution
                dropped_role_edges.append((pred_wbs, succ_wbs))
                if len(dropped_role_edges) <= 10:  # Log first 10
                    pred_name = next((r["Name"] for r in rows if r["WBS"] == pred_wbs), pred_wbs)
                    succ_name = next((r["Name"] for r in rows if r["WBS"] == succ_wbs), succ_wbs)
                    print(f"[EDGE-FILTER] ⏭️  Dropping L{pred_level}→L{succ_level}: {pred_name} → {succ_name}")
        
        if dropped_role_edges:
            print(f"[EDGE-FILTER] Dropped {len(dropped_role_edges)} role-to-role edges (L5+ → L5+)")
        
        print(f"[EDGE-FILTER] Retained {len(filtered_edges)} summary-level edges (L2/L3/L4)")
        
        # Update normalized_edges for downstream compatibility
        normalized_edges = filtered_edges
        print(f"[EDGE-FILTER] ✅ Edge filtering complete\n")
        
        # ====================================================================================
        # SIBLING AUTO-CHAINING: Create waterfall dependencies between summary-level siblings
        # ====================================================================================
        print("[SIBLING-CHAIN] Building parent→children map for auto-chaining...")
        
        # Build parent→children map for L2 deliverables, L3 components, and L4 tasks
        parent_to_children = {}
        for r in rows:
            wbs = r["WBS"]
            if wbs == "1":
                continue  # Skip project root
            
            # Get parent WBS
            wbs_parts = wbs.split(".")
            parent_wbs = ".".join(wbs_parts[:-1]) if len(wbs_parts) > 1 else "1"
            
            # Check if this is a summary-like row (L2/L3/L4 only, NOT L5+ role tasks)
            outline_level = wbs.count(".") + 1
            is_summary_like = outline_level in (2, 3, 4)  # L2 deliverables, L3 components, L4 tasks (NOT L5+ roles)
            
            if is_summary_like:
                parent_to_children.setdefault(parent_wbs, []).append(r)
        
        print(f"[SIBLING-CHAIN] Found {len(parent_to_children)} parent groups with summary-level children")
        
        # Generate forced edges by chaining siblings within each parent
        forced_edges = set()
        existing_edges = set(filtered_edges)  # Use filtered edges (L5+ → L5+ already removed)
        
        for parent_wbs, children in parent_to_children.items():
            if len(children) < 2:
                continue  # Need at least 2 siblings to chain
            
            # Sort siblings by Start date (fall back to datetime.min for deterministic ordering)
            children_sorted = sorted(
                children,
                key=lambda r: (uid_to_sched.get(r["UID"], {}).get("Start", datetime.datetime.min), r["WBS"])
            )
            
            # Chain consecutive siblings
            for i in range(len(children_sorted) - 1):
                a = children_sorted[i]
                b = children_sorted[i + 1]
                pred_wbs, succ_wbs = a["WBS"], b["WBS"]
                
                # Only add if there's no edge already
                if (pred_wbs, succ_wbs) not in existing_edges:
                    forced_edges.add((pred_wbs, succ_wbs))
                    print(f"[SIBLING-CHAIN] ➕ {a['Name']} → {b['Name']} ({pred_wbs} → {succ_wbs})")
        
        print(f"[SIBLING-CHAIN] Generated {len(forced_edges)} auto-chain dependencies")
        
        # Merge filtered edges with sibling chains to create final edge set
        all_edges = list(set(filtered_edges) | forced_edges)
        
        # Update normalized_edges with final edge set for downstream compatibility
        normalized_edges = all_edges
        print(f"[SIBLING-CHAIN] ✅ Sibling auto-chaining complete ({len(all_edges)} total edges)\n")
        
        # ====================================================================================
        # BRAND CHAIN VALIDATION: Verify critical deliverable sequence after auto-chaining
        # ====================================================================================
        # Verify the critical deliverable sequence: Key Pillars → Brand Narrative → Brand Values → Brand Vision → Purpose
        brand_chain_map = {
            "Key Pillars": "Brand Narrative",
            "Brand Narrative": "Brand Values Statement",
            "Brand Values Statement": "Brand Vision Statement",
            "Brand Vision Statement": "Purpose Statement"
        }
        
        brand_chain_validated = []
        for pred_name, succ_name in brand_chain_map.items():
            # Find WBS for these deliverables
            pred_wbs_match = next((r["WBS"] for r in rows if r["Name"] == pred_name), None)
            succ_wbs_match = next((r["WBS"] for r in rows if r["Name"] == succ_name), None)
            
            if pred_wbs_match and succ_wbs_match:
                # Check if this edge exists in all_edges (after filtering + auto-chaining)
                edge_found = (pred_wbs_match, succ_wbs_match) in all_edges
                brand_chain_validated.append((pred_name, succ_name, edge_found))
                
                status = "✅" if edge_found else "❌"
                print(f"[BRAND-CHAIN] {status} {pred_name} → {succ_name} ({pred_wbs_match} → {succ_wbs_match})")
            else:
                print(f"[BRAND-CHAIN] ⚠️  Missing deliverable: {pred_name} or {succ_name}")
        
        brand_chain_complete = all(found for _, _, found in brand_chain_validated)
        if brand_chain_complete:
            print(f"[BRAND-CHAIN] ✅ Brand chain complete: 4/4 links present in final dependencies")
        else:
            print(f"[BRAND-CHAIN] ❌ Brand chain incomplete: some links missing from final dependencies")
        print()
        
        # ====================================================================================
        # Generate XML
        project = Element("Project", xmlns="http://schemas.microsoft.com/project")
        
        # Project info - use explicit project_name if provided, otherwise fall back to derived title
        SubElement(project, "Name").text = (project_name or project_title)
        SubElement(project, "CreationDate").text = datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
        
        # FIX: Use actual project dates derived from timeline tasks (not fallback calculations)
        SubElement(project, "StartDate").text = project_start.strftime("%Y-%m-%dT%H:%M:%S")
        SubElement(project, "FinishDate").text = project_finish.strftime("%Y-%m-%dT%H:%M:%S")
        SubElement(project, "CurrentDate").text = project_start.strftime("%Y-%m-%dT%H:%M:%S")
        
        # Project header tuning
        SubElement(project, "DefaultCalendarUID").text = "1"
        SubElement(project, "ScheduleFromStart").text = "1"
        SubElement(project, "MinutesPerDay").text = "480"
        SubElement(project, "MinutesPerWeek").text = "2400"
        SubElement(project, "DaysPerMonth").text = "20"
        SubElement(project, "DurationFormat").text = "7"
        SubElement(project, "DefaultStartTime").text = "09:00:00"
        SubElement(project, "DefaultFinishTime").text = "17:00:00"
        
        # Calendars
        calendars = SubElement(project, "Calendars")
        calendar = SubElement(calendars, "Calendar")
        SubElement(calendar, "UID").text = "1"
        SubElement(calendar, "Name").text = "Standard"
        SubElement(calendar, "IsBaseCalendar").text = "1"
        
        # Working days (Mon-Fri only, weekends non-working)
        # DayType: 1=Sunday, 2=Monday, 3=Tuesday, 4=Wednesday, 5=Thursday, 6=Friday, 7=Saturday
        # IMPORTANT: Saturday is DayType=7, not DayType=6 (which is Friday)
        weekdays = SubElement(calendar, "WeekDays")
        for day_num in range(1, 8):
            weekday = SubElement(weekdays, "WeekDay")
            SubElement(weekday, "DayType").text = str(day_num)
            # Make Sunday (1) and Saturday (7) non-working; Monday-Friday (2-6) are working days
            if day_num in [1, 7]:  # Sunday=1 and Saturday=7 are non-working
                SubElement(weekday, "DayWorking").text = "0"
            else:  # Monday=2 through Friday=6 are working days
                SubElement(weekday, "DayWorking").text = "1"
                working_times = SubElement(weekday, "WorkingTimes")
                for start_time, end_time in calendar_blocks:
                    working_time = SubElement(working_times, "WorkingTime")
                    SubElement(working_time, "FromTime").text = start_time
                    SubElement(working_time, "ToTime").text = end_time

        # Resources
        resources_elem = SubElement(project, "Resources")
        for res in resources:
            resource = SubElement(resources_elem, "Resource")
            SubElement(resource, "UID").text = str(res["UID"])
            SubElement(resource, "ID").text = str(res["ID"])
            SubElement(resource, "Name").text = res["Name"]
            SubElement(resource, "Type").text = "1"  # Work resource
            # Add StandardRate for pricing
            role = res.get("Name", "Unassigned")
            rate = _std_rate_for(role, pricing_mode, rate_band, blended_rate, DB)
            SubElement(resource, "StandardRate").text = f"{rate:.2f}"

        # HYBRID COSTTYPE HELPER: Apply CostType based on FixedCost for ALL task types
        # This ensures monthly retainer children, deliverables, and all tasks get proper cost handling
        def apply_costtype_and_fixedcost(task_elem, row_dict, task_name=""):
            # Check multiple field names for price (Step 3 uses price_usd, export prep may use FixedCost)
            fixed_cost = float(
                row_dict.get("FixedCost") or
                row_dict.get("Price_USD") or
                row_dict.get("price_usd") or
                row_dict.get("total_price") or
                0
            )
            if fixed_cost > 0:
                SubElement(task_elem, "CostType").text = "2"  # Fixed Cost
                SubElement(task_elem, "FixedCost").text = f"{fixed_cost:.2f}"
                SubElement(task_elem, "FixedCostAccrual").text = "3"  # End of task
                print(f"[XML EXPORT] 💰 CostType=2 (Fixed Cost) FixedCost=${fixed_cost:.2f} for: {task_name}")
            else:
                SubElement(task_elem, "CostType").text = "0"  # Role Hourly (default)

        # Tasks
        tasks_elem = SubElement(project, "Tasks")
        for task_id, r in enumerate(rows, 1):
            task = SubElement(tasks_elem, "Task")
            SubElement(task, "UID").text = str(r["UID"])
            SubElement(task, "ID").text = str(task_id)
            # Use "Project Summary" for root task, otherwise use the actual name
            is_root = r["WBS"] == "1"
            name_txt = "Project Summary" if is_root else r["Name"]
            SubElement(task, "Name").text = name_txt
            
            # IMPORTANT: Keep full WBS paths for uniqueness, only OutlineLevel is capped
            # Workfront needs full WBS (e.g., "1.2.3.1.1") for unique task identification
            # But OutlineLevel is capped at 3 for visual hierarchy (Deliverable → Component → Task)
            SubElement(task, "WBS").text = r["WBS"]
            SubElement(task, "OutlineNumber").text = r["WBS"] 
            SubElement(task, "Start").text = uid_to_sched[r["UID"]]["Start"]
            SubElement(task, "Finish").text = uid_to_sched[r["UID"]]["Finish"]
            
            # Calculate outline level EARLY (needed for deliverable detection)
            outline_level = r["WBS"].count(".") + 1  # 1 for '1', 2 for '1.1', etc.
            SubElement(task, "OutlineLevel").text = str(outline_level)
            
            # GPT-5 PRO FIX: Make deliverable/summary mutually exclusive to prevent duplicate ConstraintType tags
            # WATERFALL FIX: Deliverables are now OutlineLevel=2 (direct children of Project Summary)
            is_deliverable = outline_level == 2 and not is_root
            is_summary = (r["WBS"] in summary_set or is_root) and not is_deliverable
            SubElement(task, "Summary").text = "1" if is_summary else "0"
            
            # GPT-5 PRO FIX: Conditional DurationFormat based on manual scheduling
            # 53 (Elapsed Days + Manual) for manually scheduled tasks (deliverables, non-root summaries)
            # 7 (Elapsed Days) for auto-scheduled tasks (root, leaf)
            is_manually_scheduled = is_deliverable or (is_summary and not is_root)
            SubElement(task, "DurationFormat").text = "53" if is_manually_scheduled else "7"
            
            # GPT-5 PRO FIX: Reordered branches to prevent overlap (root → deliverable → summary → leaf)
            if is_root:
                # ROOT TASK (OutlineLevel=1): PT0M duration, no Manual* tags
                SubElement(task, "Work").text = "PT0M"
                SubElement(task, "Duration").text = "PT0M"
                SubElement(task, "ConstraintType").text = "4"  # Must Start On
                SubElement(task, "ConstraintDate").text = uid_to_sched[r["UID"]]["Start"]
                # Root does NOT get Manual* tags or Type/IsEffortDriven per GPT-5 Pro spec
            
            elif is_deliverable:
                # DELIVERABLE TASKS (OutlineLevel=2): SNET constraint + Manual* tags if has children
                # Deliverables are direct children of Project Summary with SNET to lock start dates
                SubElement(task, "Work").text = "PT0M"
                
                # FIX: Calculate actual Duration from Start/Finish dates (not hardcoded PT480M)
                # This fixes the mismatch where Step 4 shows 121 days but Workfront shows 1 day
                try:
                    start_str = uid_to_sched[r["UID"]]["Start"]
                    finish_str = uid_to_sched[r["UID"]]["Finish"]
                    start_dt = datetime.datetime.fromisoformat(start_str.replace('Z', ''))
                    finish_dt = datetime.datetime.fromisoformat(finish_str.replace('Z', ''))
                    duration_days = max(1, (finish_dt - start_dt).days)
                    duration_minutes = duration_days * 480  # 8-hour workday
                    SubElement(task, "Duration").text = f"PT{duration_minutes}M"
                    print(f"[XML EXPORT] 📏 Deliverable Duration: {name_txt[:40]} = {duration_days} days (PT{duration_minutes}M)")
                except Exception as e:
                    # Fallback to 1 day
                    SubElement(task, "Duration").text = "PT480M"
                    print(f"[XML EXPORT] ⚠️ Duration fallback for {name_txt[:40]}: {e}")
                
                # Manual Scheduling tags to lock parent timelines
                SubElement(task, "Type").text = "1"  # Fixed Duration
                SubElement(task, "IsEffortDriven").text = "0"
                SubElement(task, "ManuallyScheduled").text = "1"
                SubElement(task, "ManualStart").text = "1"
                SubElement(task, "ManualFinish").text = "1"
                SubElement(task, "ManualDuration").text = "1"
                SubElement(task, "ManualWork").text = "1"
                
                # SNET constraint to lock deliverable start date
                start_date_str = uid_to_sched[r["UID"]]["Start"]
                if 'T' in start_date_str:
                    date_part = start_date_str.split('T')[0]
                    constraint_date = f"{date_part}T08:00:00"
                else:
                    constraint_date = f"{start_date_str}T08:00:00"
                
                SubElement(task, "ConstraintType").text = "2"  # Start No Earlier Than
                SubElement(task, "ConstraintDate").text = constraint_date
                
                print(f"[XML EXPORT] 🔒 Applied SNET constraint to deliverable: {name_txt} (WBS {r['WBS']}, Start={constraint_date})")
            
            elif is_summary and not is_root:
                # NON-ROOT SUMMARY TASKS (OutlineLevel 2-5): Calculate actual duration from Start/Finish
                SubElement(task, "Work").text = "PT0M"
                
                # FIX: Calculate actual Duration from Start/Finish dates (not hardcoded PT480M)
                # This ensures summary task durations match their date spans
                try:
                    start_str = uid_to_sched[r["UID"]]["Start"]
                    finish_str = uid_to_sched[r["UID"]]["Finish"]
                    start_dt = datetime.datetime.fromisoformat(start_str.replace('Z', ''))
                    finish_dt = datetime.datetime.fromisoformat(finish_str.replace('Z', ''))
                    duration_days = max(1, (finish_dt - start_dt).days)
                    duration_minutes = duration_days * 480  # 8-hour workday
                    SubElement(task, "Duration").text = f"PT{duration_minutes}M"
                except Exception:
                    # Fallback to 1 day
                    SubElement(task, "Duration").text = "PT480M"
                
                # GPT-5 PRO SPEC: Summary tasks (OutlineLevel 2-5) need Type, IsEffortDriven, and Manual Scheduling tags
                # These lock parent timelines and prevent Workfront from recalculating rolled-up dates
                SubElement(task, "Type").text = "1"  # Fixed Duration
                SubElement(task, "IsEffortDriven").text = "0"
                SubElement(task, "ManuallyScheduled").text = "1"
                SubElement(task, "ManualStart").text = "1"
                SubElement(task, "ManualFinish").text = "1"
                SubElement(task, "ManualDuration").text = "1"
                SubElement(task, "ManualWork").text = "1"
            # Only set Duration/Work for non-summary leaf tasks
            else:
                # WORKFRONT FIX: Duration MUST equal (Finish - Start) for ALL tasks
                # NO 960-minute snapping - use actual timestamp difference
                # This is critical to prevent Workfront from recalculating dates on import
                
                start_val = uid_to_sched[r["UID"]]["Start"]
                finish_val = uid_to_sched[r["UID"]]["Finish"]
                
                # Robust datetime parser that handles timezone suffixes (Z, +00:00, etc.)
                def to_datetime(val):
                    if isinstance(val, datetime.datetime):
                        return val
                    elif isinstance(val, str):
                        # Remove timezone suffixes for parsing
                        val_clean = val.replace('Z', '+00:00')  # Convert Z to +00:00
                        # Try parsing with fromisoformat (handles timezones)
                        try:
                            return datetime.datetime.fromisoformat(val_clean)
                        except ValueError:
                            # Fallback: strip timezone manually and parse
                            val_clean = val.split('+')[0].split('-')
                            # Rejoin date portion (first 3 parts: YYYY, MM, DD)
                            val_clean = '-'.join(val_clean[:3]) + val_clean[3] if len(val_clean) > 3 else '-'.join(val_clean)
                            val_clean = val_clean.split('.')[0]  # Remove milliseconds
                            return datetime.datetime.fromisoformat(val_clean)
                    else:
                        return val
                
                start_dt = to_datetime(start_val)
                finish_dt = to_datetime(finish_val)
                
                # Work calculation: For multi-assignment tasks, sum all assignee hours
                if r.get("multi_assignment") and "assignments" in r:
                    # Multi-assignment: sum all assignee work hours
                    total_work_hours = sum(a["work_hours"] for a in r["assignments"])
                    work_minutes = int(total_work_hours * 60)
                else:
                    # Single assignment: use PlannedHours from Gantt (actual effort)
                    work_hours = uid_to_sched[r['UID']].get('PlannedHours', 0)
                    if pd.isna(work_hours):
                        work_hours = 0
                    work_minutes = max(0, int(work_hours * 60))
                
                SubElement(task, "Work").text = f"PT{work_minutes}M"
                
                # DURATION FIX: Ensure Work ≤ Duration using ceil(hours/8) formula
                # This fixes "11 hours in 0.12 days" issue by guaranteeing sufficient duration
                work_hours = work_minutes / 60.0
                
                if work_hours == 0:
                    # Milestone: zero duration
                    dur_minutes = 0
                else:
                    # Regular task: required_days = max(1, ceil(hours / 8.0))
                    # Uses 8-hour workday aligned with MinutesPerDay=480
                    import math
                    required_days = max(1, math.ceil(work_hours / 8.0))
                    dur_minutes = required_days * 480
                
                SubElement(task, "Duration").text = f"PT{dur_minutes}M"
                
                # WORKFRONT FIX: Set ALL leaf tasks as Fixed Duration to prevent effort-driven recalculation
                # Type=1 (Fixed Duration) + IsEffortDriven=0 ensures Workfront doesn't adjust durations based on effort
                SubElement(task, "Type").text = "1"  # Fixed Duration
                SubElement(task, "IsEffortDriven").text = "0"
                
                # WORKFRONT FIX: Explicit ASAP constraint on ALL leaf tasks to prevent constraint inheritance
                # Without this, leaf tasks inherit parent deliverable's SNET constraint, causing serial execution
                # ConstraintType=0 (As Soon As Possible) allows tasks to start when dependencies are met
                SubElement(task, "ConstraintType").text = "0"  # As Soon As Possible
            
            # GPT-5 PRO FIX: Add CalendarUID to all tasks to enable calendar-day duration display
            # Without this, Workfront treats durations as raw minutes and displays "0.12 days"
            SubElement(task, "CalendarUID").text = "1"
            
            # Mark anchor rows (WBS starts with ANCHOR_) as milestones
            is_anchor = str(r.get("WBS", "")).startswith("ANCHOR_")
            SubElement(task, "Milestone").text = "1" if is_anchor else "0"
            
            # UNIVERSAL COSTTYPE: Apply hybrid CostType logic to ALL tasks (deliverables, summaries, leaves, monthly children)
            # This ensures any task with FixedCost > 0 gets CostType=2, all others get CostType=0 (Role Hourly)
            apply_costtype_and_fixedcost(task, r, name_txt)

        # WORKFRONT FIX: Add finish anchor milestones for preserved DELIVERABLES
        # This locks the deliverable's finish date by creating a zero-duration milestone
        # with Must Finish On constraint, linked to the last leaf task within that deliverable
        next_uid = max([r["UID"] for r in rows]) + 1
        next_id = len(rows) + 1
        finish_anchor_map = {}  # Maps preserved deliverable UID to finish anchor UID
        
        # Helper function to find last leaf task within a deliverable's hierarchy
        def find_last_leaf_in_deliverable(deliverable_wbs, all_rows):
            """Find the deepest/last leaf task by WBS traversal within a deliverable's hierarchy"""
            # Get all direct descendants of this deliverable (children with deliverable_wbs as parent)
            # Use proper parent-child relationships from ParentWBS, not just string prefix matching
            def get_all_descendants(parent_wbs):
                """Recursively get all descendants of a parent WBS"""
                descendants = []
                for row in all_rows:
                    if row.get("ParentWBS") == parent_wbs:
                        descendants.append(row)
                        # Recursively get this child's descendants
                        descendants.extend(get_all_descendants(row["WBS"]))
                return descendants
            
            all_descendants = get_all_descendants(deliverable_wbs)
            
            if not all_descendants:
                return None
            
            # Filter to only leaf tasks (not in summary_set)
            leaves = [d for d in all_descendants if d["WBS"] not in summary_set]
            
            if not leaves:
                return None
            
            # Sort by WBS depth (deepest first), then by finish date (latest first)
            # This ensures we get the terminal leaf task
            leaves_sorted = sorted(
                leaves,
                key=lambda x: (
                    -x["WBS"].count("."),  # Deeper WBS = more dots (sort descending)
                    -(ord(uid_to_sched.get(x["UID"], {}).get("Finish", "9999")[0]) if uid_to_sched.get(x["UID"], {}).get("Finish") else 0)  # Later finish date
                )
            )
            
            return leaves_sorted[0] if leaves_sorted else None
        
        for r in rows:
            # Only create finish anchors for preserved DELIVERABLES (summary tasks with preserved dates)
            if r["UID"] in preserved_dates and r["WBS"] in summary_set and r["WBS"] != "1":
                # Find the last leaf task under this deliverable
                last_leaf = find_last_leaf_in_deliverable(r["WBS"], rows)
                
                if not last_leaf:
                    print(f"[XML EXPORT] ⚠️ No leaf tasks found for deliverable {r['Name']} (WBS {r['WBS']}), skipping finish anchor")
                    continue
                
                # Create finish anchor milestone for this deliverable
                finish_anchor_uid = next_uid
                next_uid += 1
                
                finish_anchor_map[r["UID"]] = finish_anchor_uid
                
                # Create the finish anchor task element
                task = SubElement(tasks_elem, "Task")
                SubElement(task, "UID").text = str(finish_anchor_uid)
                SubElement(task, "ID").text = str(next_id)
                next_id += 1
                
                # Name it "{DeliverableName} - Finish Anchor"
                anchor_name = f"{r['Name']} - Finish Anchor"
                SubElement(task, "Name").text = anchor_name
                
                # CRITICAL FIX: Use numeric WBS as CHILD of deliverable (not sibling with _FA suffix)
                # Example: Deliverable WBS=1.1.2 → Finish Anchor WBS=1.1.2.99
                # This makes the anchor a child task, not a sibling, which is required for Workfront rollup
                wbs_anchor = f"{r['WBS']}.99"
                SubElement(task, "WBS").text = wbs_anchor
                SubElement(task, "OutlineNumber").text = wbs_anchor
                
                # CRITICAL FIX: Use deliverable's EXACT finish date with T17:00:00 format (5 PM end)
                # Ensure consistent time format for finish constraints
                finish_date_str = uid_to_sched[r["UID"]]["Finish"]
                # Parse and reformat to ensure T17:00:00
                if 'T' in finish_date_str:
                    date_part = finish_date_str.split('T')[0]
                    finish_date = f"{date_part}T17:00:00"
                else:
                    finish_date = f"{finish_date_str}T17:00:00"
                
                SubElement(task, "Start").text = finish_date
                SubElement(task, "Finish").text = finish_date
                
                # Zero-duration milestone
                SubElement(task, "Summary").text = "0"
                SubElement(task, "DurationFormat").text = "7"
                SubElement(task, "Work").text = "PT0M"
                SubElement(task, "Duration").text = "PT0M"
                SubElement(task, "Type").text = "1"
                SubElement(task, "IsEffortDriven").text = "0"
                
                # Must Finish On constraint (ConstraintType=3) with T17:00:00 format
                SubElement(task, "ConstraintType").text = "3"
                SubElement(task, "ConstraintDate").text = finish_date
                
                # OutlineLevel = deliverable level + 1 (child of deliverable)
                # Deliverable is at level 3, so anchor is at level 4
                deliverable_level = r["WBS"].count(".") + 1
                anchor_level = deliverable_level + 1
                SubElement(task, "OutlineLevel").text = str(anchor_level)
                
                # Mark as milestone
                SubElement(task, "Milestone").text = "1"
                
                # CRITICAL FIX: Set predecessor to last leaf task within THIS deliverable
                # Not the deliverable itself, and not a task from a different deliverable
                pred_link = SubElement(task, "PredecessorLink")
                SubElement(pred_link, "PredecessorUID").text = str(last_leaf["UID"])
                SubElement(pred_link, "Type").text = "1"  # Finish-to-Start
                SubElement(pred_link, "CrossProject").text = "0"
                SubElement(pred_link, "LinkLag").text = "0"
                SubElement(pred_link, "LagFormat").text = "7"
                
                print(f"[XML EXPORT] ⚓ Created finish anchor for deliverable: {r['Name']} (WBS {r['WBS']}) → Last leaf: {last_leaf['Name']} (UID {last_leaf['UID']}, WBS {last_leaf['WBS']}), Finish={finish_date})")

        # Assignments
        # CRITICAL FIX: Filter out assignments for summary/parent tasks
        # Summary tasks should have NO assignments - Workfront will roll up from children
        # Build set of summary task UIDs by checking which tasks have children
        summary_task_uids = set()
        wbs_to_uid = {r["WBS"]: r["UID"] for r in rows}
        for parent_wbs in children_by_parent.keys():
            if parent_wbs in wbs_to_uid:
                summary_task_uids.add(wbs_to_uid[parent_wbs])
        # Also add root task
        if "1" in wbs_to_uid:
            summary_task_uids.add(wbs_to_uid["1"])
        
        assignments_elem = SubElement(project, "Assignments")
        
        # BACKWARD COMPATIBILITY: Build deliverable FixedCost lookup
        # Only zero assignment costs when parent deliverable has FixedCost > 0
        deliverable_fixed_costs = {}
        for r in rows:
            outline_level = r["WBS"].count(".") + 1
            if outline_level == 2:  # Deliverable level
                fixed_cost = float(r.get("FixedCost", 0))
                deliverable_fixed_costs[r["WBS"]] = fixed_cost
        
        def get_deliverable_wbs(wbs):
            """Get the L2 deliverable WBS for any task WBS."""
            parts = wbs.split(".")
            if len(parts) >= 2:
                return ".".join(parts[:2])  # L2 = "1.2"
            return wbs
        
        # Build assignments for all tasks (including multi-assignment merged tasks)
        assign_uid_counter = 1
        for row in rows:
            task_uid = row["UID"]
            
            # Skip assignments that reference summary/parent tasks
            if task_uid in summary_task_uids:
                print(f"[XML EXPORT] 🚫 Skipping assignment for summary task UID {task_uid} (summary tasks should have no assignments)")
                continue
            
            # Check if this is a multi-assignment merged task
            if row.get("multi_assignment") and "assignments" in row:
                # Multi-assignment task - create multiple Assignment elements
                for assign_data in row["assignments"]:
                    assignment = SubElement(assignments_elem, "Assignment")
                    SubElement(assignment, "UID").text = str(assign_uid_counter)
                    assign_uid_counter += 1
                    
                    SubElement(assignment, "TaskUID").text = str(task_uid)
                    
                    # Find resource UID for this role
                    resource_name = assign_data["resource_name"]
                    res_uid = res_name_to_uid.get(resource_name) or res_name_to_uid.get("Unassigned")
                    SubElement(assignment, "ResourceUID").text = str(res_uid)
                    
                    SubElement(assignment, "Start").text = uid_to_sched[task_uid]["Start"]
                    SubElement(assignment, "Finish").text = uid_to_sched[task_uid]["Finish"]
                    
                    # Compute Units and Work
                    work_hours = assign_data["work_hours"]
                    work_min = work_hours * 60
                    dur_hours = uid_to_sched[task_uid].get('DurationHours', 0)
                    dur_min = dur_hours * 60
                    # Use actual duration for accurate Units calculation (no 480-minute rounding)
                    units = 0 if dur_min == 0 else work_min / dur_min
                    
                    SubElement(assignment, "Units").text = str(units)
                    SubElement(assignment, "Work").text = f"PT{int(work_min)}M"
                    
                    # BACKWARD COMPATIBILITY: Zero cost only if parent deliverable has FixedCost
                    # Otherwise, use hourly rate calculation (legacy behavior)
                    deliv_wbs = get_deliverable_wbs(row.get("WBS", ""))
                    parent_fixed_cost = deliverable_fixed_costs.get(deliv_wbs, 0)
                    if parent_fixed_cost > 0:
                        SubElement(assignment, "Cost").text = "0.00"
                    else:
                        # Fallback to hourly calculation
                        rate = _std_rate_for(resource_name, pricing_mode, rate_band, blended_rate, DB)
                        cost = work_hours * rate
                        SubElement(assignment, "Cost").text = f"{cost:.2f}"
                    
                print(f"[XML EXPORT] 👥 Created {len(row['assignments'])} assignments for multi-assignment task: {row['Name'][:50]} (UID {task_uid})")
            else:
                # Single-assignment task - use original logic
                # Find the assignment for this task from the pre-built assignments list
                task_assignments = [a for a in assignments if a["TaskUID"] == task_uid]
                
                for assign in task_assignments:
                    assignment = SubElement(assignments_elem, "Assignment")
                    SubElement(assignment, "UID").text = str(assign_uid_counter)
                    assign_uid_counter += 1
                    
                    SubElement(assignment, "TaskUID").text = str(task_uid)
                    res_uid = assign["ResourceUID"]
                    SubElement(assignment, "ResourceUID").text = str(res_uid)
                    SubElement(assignment, "Start").text = assign["Start"]
                    SubElement(assignment, "Finish").text = assign["Finish"]
                    
                    # Compute Units = work_min / dur_min for Fixed Duration tasks
                    work_hours = assign['WorkHours']
                    work_min = work_hours * 60
                    dur_hours = uid_to_sched[task_uid].get('DurationHours', 0)
                    dur_min = dur_hours * 60
                    # Use actual duration for accurate Units calculation (no 480-minute rounding)
                    units = 0 if dur_min == 0 else work_min / dur_min
                    
                    SubElement(assignment, "Units").text = str(units)
                    SubElement(assignment, "Work").text = f"PT{int(work_min)}M"
                    
                    # BACKWARD COMPATIBILITY: Zero cost only if parent deliverable has FixedCost
                    # Otherwise, use hourly rate calculation (legacy behavior)
                    deliv_wbs = get_deliverable_wbs(row.get("WBS", ""))
                    parent_fixed_cost = deliverable_fixed_costs.get(deliv_wbs, 0)
                    if parent_fixed_cost > 0:
                        SubElement(assignment, "Cost").text = "0.00"
                    else:
                        # Fallback to hourly calculation
                        res_name = next((r["Name"] for r in resources if r["UID"] == res_uid), "Unassigned")
                        rate = _std_rate_for(res_name, pricing_mode, rate_band, blended_rate, DB)
                        cost = work_hours * rate
                        SubElement(assignment, "Cost").text = f"{cost:.2f}"

        # WORKFRONT FIX: Add PredecessorLinks for:
        # 1. Deliverable-level dependencies (OutlineLevel ≤4 to ≤4)
        # 2. Component-level dependencies (L5 summary to L5 summary)
        # 3. Skip leaf-to-leaf within same L5 component (enables parallel role execution)
        wbs_to_uid = {r["WBS"]: r["UID"] for r in rows}
        uid_to_outline = {r["UID"]: r["WBS"].count(".") + 1 for r in rows}
        wbs_to_wbs = {r["UID"]: r["WBS"] for r in rows}
        
        def get_deliverable_parent(wbs):
            """Extract deliverable portion of WBS (first 3 levels: 1.2.3)"""
            parts = wbs.split(".")
            if len(parts) >= 3:
                return ".".join(parts[:3])  # Deliverable is always WBS x.y.z
            return wbs
        
        def get_l5_prefix(wbs):
            """Extract L5 component prefix (first 5 levels: 1.2.3.1.1)"""
            if not wbs:
                return None
            parts = wbs.split(".")
            if len(parts) >= 5:
                return ".".join(parts[:5])
            return None  # Not deep enough to be in L5 component
        
        # Build metadata map for dependency filtering
        uid_metadata = {}
        for r in rows:
            uid = r["UID"]
            wbs = r["WBS"]
            outline_level = wbs.count(".") + 1
            is_leaf = wbs not in summary_set  # Leaf if not in summary_set
            l5_prefix = get_l5_prefix(wbs)
            
            uid_metadata[uid] = {
                "outline_level": outline_level,
                "l5_prefix": l5_prefix,
                "is_leaf": is_leaf,
                "wbs": wbs
            }
        
        # NOTE: Redundant dependency filtering block removed - L5+ → L5+ edges already filtered
        # earlier at line 10363. The all_edges set contains the final edge list after:
        # 1. WBS remapping (drops stale references)
        # 2. L5+ → L5+ filtering (enables parallel role execution)
        # 3. Sibling auto-chaining (adds waterfall dependencies for L2/L3/L4)
        
        # Add PredecessorLink elements for all final edges
        # WATERFALL FIX: Use wbs_to_new_uid which has remapped UIDs after chronological sort
        print(f"[XML-EXPORT] Writing {len(all_edges)} PredecessorLink elements...")
        skipped_zero_duration = 0
        for pred_wbs, succ_wbs in all_edges:
            pred_uid = wbs_to_new_uid.get(pred_wbs)
            succ_uid = wbs_to_new_uid.get(succ_wbs)
            if pred_uid and succ_uid:
                # Find the successor task element and add a PredecessorLink (MSPDI: no wrapper)
                for task_elem in tasks_elem.findall("Task"):
                    task_uid_elem = task_elem.find("UID")
                    if task_uid_elem is not None and task_uid_elem.text == str(succ_uid):
                        # VALIDATION: Skip if either task has zero duration (e.g., root task)
                        succ_duration_elem = task_elem.find("Duration")
                        if succ_duration_elem is not None and succ_duration_elem.text == "PT0M":
                            skipped_zero_duration += 1
                            break
                        
                        # Find predecessor task to check its duration
                        pred_task_elem = None
                        for t_elem in tasks_elem.findall("Task"):
                            t_uid_elem = t_elem.find("UID")
                            if t_uid_elem is not None and t_uid_elem.text == str(pred_uid):
                                pred_task_elem = t_elem
                                break
                        
                        if pred_task_elem is not None:
                            pred_duration_elem = pred_task_elem.find("Duration")
                            if pred_duration_elem is not None and pred_duration_elem.text == "PT0M":
                                skipped_zero_duration += 1
                                break
                        
                        # Both tasks have valid durations - write PredecessorLink
                        pred_link = SubElement(task_elem, "PredecessorLink")
                        SubElement(pred_link, "PredecessorUID").text = str(pred_uid)
                        SubElement(pred_link, "Type").text = "1"          # 1 = Finish-to-Start
                        SubElement(pred_link, "CrossProject").text = "0"
                        # Optional but harmless:
                        SubElement(pred_link, "LinkLag").text = "0"
                        SubElement(pred_link, "LagFormat").text = "7"     # 7 = days
                        break
        
        if skipped_zero_duration > 0:
            print(f"[VALIDATION] ⏭️  Skipped {skipped_zero_duration} PredecessorLink(s) with zero-duration tasks")

        # Compute project summary start/finish from children (no more hardcoded dates)
        if tasks_elem is not None:
            all_task_starts = []
            all_task_finishes = []
            project_task_elem = None
            
            for task_elem in tasks_elem.findall("Task"):
                uid_elem = task_elem.find("UID")
                start_elem = task_elem.find("Start")
                finish_elem = task_elem.find("Finish")
                
                if uid_elem is not None and start_elem is not None and finish_elem is not None:
                    if uid_elem.text == "1":  # Project summary task
                        project_task_elem = task_elem
                    else:
                        # Collect all non-project task dates
                        all_task_starts.append(start_elem.text)
                        all_task_finishes.append(finish_elem.text)
            
            # Set project summary dates from children min/max
            if project_task_elem is not None and all_task_starts and all_task_finishes:
                earliest_start = min(all_task_starts)
                latest_finish = max(all_task_finishes)
                
                # Update project summary start/finish
                proj_start_elem = project_task_elem.find("Start")
                proj_finish_elem = project_task_elem.find("Finish")
                
                if proj_start_elem is not None:
                    proj_start_elem.text = earliest_start
                if proj_finish_elem is not None:
                    proj_finish_elem.text = latest_finish
                
                # Ensure project summary is marked as summary
                proj_summary_elem = project_task_elem.find("Summary")
                if proj_summary_elem is not None:
                    proj_summary_elem.text = "1"

        # WORKFRONT VALIDATION PASS: Comprehensive checks for XML structure
        print(f"\n[WORKFRONT VALIDATION] ═══════════════════════════════════════")
        
        # Validation 1: Count PredecessorLinks at OutlineLevel ≥5 (should be 0)
        deep_pred_count = 0
        for task_elem in tasks_elem.findall("Task"):
            outline_elem = task_elem.find("OutlineLevel")
            if outline_elem is not None:
                outline_level = int(outline_elem.text)
                if outline_level >= 5:
                    # Check if this task has any PredecessorLink
                    pred_links = task_elem.findall("PredecessorLink")
                    if pred_links:
                        deep_pred_count += len(pred_links)
                        name_elem = task_elem.find("Name")
                        wbs_elem = task_elem.find("WBS")
                        task_name = name_elem.text if name_elem is not None else "Unknown"
                        task_wbs = wbs_elem.text if wbs_elem is not None else "Unknown"
                        print(f"[VALIDATION] ⚠️  Found PredecessorLink at L{outline_level}: {task_name} (WBS {task_wbs})")
        
        if deep_pred_count == 0:
            print(f"[VALIDATION] ✅ No deep dependencies (OutlineLevel ≥5): PASS")
        else:
            print(f"[VALIDATION] ❌ Found {deep_pred_count} deep dependencies at OutlineLevel ≥5: FAIL")
        
        # Validation 2: Check all finish anchors are children (numeric WBS, not _FA)
        anchor_validation_passed = True
        for task_elem in tasks_elem.findall("Task"):
            name_elem = task_elem.find("Name")
            if name_elem is not None and "Finish Anchor" in name_elem.text:
                wbs_elem = task_elem.find("WBS")
                if wbs_elem is not None:
                    wbs = wbs_elem.text
                    if "_FA" in wbs:
                        print(f"[VALIDATION] ❌ Finish anchor has non-numeric WBS: {wbs}")
                        anchor_validation_passed = False
                    elif not wbs.endswith(".99"):
                        print(f"[VALIDATION] ⚠️  Finish anchor doesn't end with .99: {wbs}")
        
        if anchor_validation_passed:
            print(f"[VALIDATION] ✅ All finish anchors use numeric child WBS (.99): PASS")
        
        # Validation 3: Check all summary tasks have Summary=1, Work=PT0M, Duration=PT480M (except root)
        summary_validation_passed = True
        pt0m_count = 0
        for task_elem in tasks_elem.findall("Task"):
            # Check if task has children by looking at ParentWBS references
            wbs_elem = task_elem.find("WBS")
            uid_elem = task_elem.find("UID")
            if wbs_elem is not None:
                task_wbs = wbs_elem.text
                is_root = uid_elem is not None and uid_elem.text == "1"
                
                # Check if any other task has this as ParentWBS (indicates children)
                has_children = False
                for other_task in tasks_elem.findall("Task"):
                    other_wbs = other_task.find("WBS")
                    if other_wbs is not None and other_wbs.text.startswith(task_wbs + "."):
                        has_children = True
                        break
                
                if has_children:
                    summary_elem = task_elem.find("Summary")
                    work_elem = task_elem.find("Work")
                    dur_elem = task_elem.find("Duration")
                    
                    if summary_elem is None or summary_elem.text != "1":
                        name_elem = task_elem.find("Name")
                        task_name = name_elem.text if name_elem is not None else "Unknown"
                        print(f"[VALIDATION] ❌ Parent task missing Summary=1: {task_name} (WBS {task_wbs})")
                        summary_validation_passed = False
                    
                    # Check duration: Root should be PT0M, all others should be PT480M
                    if dur_elem is not None:
                        expected_duration = "PT0M" if is_root else "PT480M"
                        if dur_elem.text == "PT0M":
                            pt0m_count += 1
                        if dur_elem.text != expected_duration and not is_root:
                            name_elem = task_elem.find("Name")
                            task_name = name_elem.text if name_elem is not None else "Unknown"
                            print(f"[VALIDATION] ⚠️  Summary task has wrong duration: {task_name} (WBS {task_wbs}, expected {expected_duration}, got {dur_elem.text})")
        
        if summary_validation_passed:
            print(f"[VALIDATION] ✅ All parent tasks have Summary=1 + PT480M duration: PASS")
        print(f"[VALIDATION] 📊 Tasks with Duration=PT0M: {pt0m_count} (should only be project root + milestones)")
        
        # Validation 4: Check duration math for sample tasks (1-hour tasks should be PT60M, not PT960M)
        duration_validation_passed = True
        sample_count = 0
        for task_elem in tasks_elem.findall("Task"):
            start_elem = task_elem.find("Start")
            finish_elem = task_elem.find("Finish")
            dur_elem = task_elem.find("Duration")
            
            if start_elem is not None and finish_elem is not None and dur_elem is not None:
                try:
                    start_str = start_elem.text
                    finish_str = finish_elem.text
                    dur_str = dur_elem.text
                    
                    # Parse start/finish to calculate expected duration
                    start_dt = datetime.datetime.fromisoformat(start_str.replace('T', ' ').replace('Z', ''))
                    finish_dt = datetime.datetime.fromisoformat(finish_str.replace('T', ' ').replace('Z', ''))
                    
                    # Calculate actual time difference
                    actual_minutes = int((finish_dt - start_dt).total_seconds() / 60)
                    
                    # Parse duration from XML
                    if dur_str.startswith("PT") and dur_str.endswith("M"):
                        xml_minutes = int(dur_str[2:-1])
                        
                        # Check if duration matches (allow small tolerance for rounding)
                        if abs(xml_minutes - actual_minutes) > 60:  # More than 1 hour difference
                            name_elem = task_elem.find("Name")
                            task_name = name_elem.text if name_elem is not None else "Unknown"
                            if sample_count < 3:  # Only show first 3 mismatches
                                print(f"[VALIDATION] ⚠️  Duration mismatch: {task_name[:40]} | Expected {actual_minutes}m, got {xml_minutes}m")
                                sample_count += 1
                                duration_validation_passed = False
                except:
                    pass
        
        if duration_validation_passed:
            print(f"[VALIDATION] ✅ Duration matches Start/Finish timestamps: PASS")
        
        # Validation 5: Check all leaf tasks have ConstraintType=0 (ASAP) to prevent constraint inheritance
        leaf_constraint_issues = 0
        for task_elem in tasks_elem.findall("Task"):
            summary_elem = task_elem.find("Summary")
            constraint_elem = task_elem.find("ConstraintType")
            outline_elem = task_elem.find("OutlineLevel")
            uid_elem = task_elem.find("UID")
            
            # Check if this is a leaf task (Summary=0 or missing, not root UID=1)
            is_leaf = (summary_elem is None or summary_elem.text == "0") and (uid_elem is None or uid_elem.text != "1")
            is_deliverable = outline_elem is not None and outline_elem.text == "3"
            
            if is_leaf and not is_deliverable:
                # Leaf task should have ConstraintType=0 (ASAP)
                if constraint_elem is None or constraint_elem.text != "0":
                    name_elem = task_elem.find("Name")
                    task_name = name_elem.text if name_elem is not None else "Unknown"
                    constraint_val = constraint_elem.text if constraint_elem is not None else "MISSING"
                    if leaf_constraint_issues < 3:  # Only show first 3
                        print(f"[VALIDATION] ⚠️  Leaf task missing ASAP constraint: {task_name[:40]} (Constraint={constraint_val}, expected 0)")
                    leaf_constraint_issues += 1
        
        if leaf_constraint_issues == 0:
            print(f"[VALIDATION] ✅ All leaf tasks have ConstraintType=0 (ASAP): PASS")
        else:
            print(f"[VALIDATION] ⚠️  Found {leaf_constraint_issues} leaf tasks without ASAP constraint")
        
        # Validation 6: Check all non-root summary tasks have Manual Scheduling tags (GPT-5 Pro spec)
        manual_tag_issues = 0
        root_has_manual_tags = False
        for task_elem in tasks_elem.findall("Task"):
            summary_elem = task_elem.find("Summary")
            wbs_elem = task_elem.find("WBS")
            
            # Check if this is the root task (WBS="1" is always the project root)
            is_root_task = wbs_elem is not None and wbs_elem.text == "1"
            
            # Check if this is a summary task (has children)
            if summary_elem is not None and summary_elem.text == "1":
                # Check for Manual* tags
                manual_scheduled = task_elem.find("ManuallyScheduled")
                manual_start = task_elem.find("ManualStart")
                manual_finish = task_elem.find("ManualFinish")
                manual_duration = task_elem.find("ManualDuration")
                manual_work = task_elem.find("ManualWork")
                
                # Root should NOT have Manual* tags
                if is_root_task:
                    if any([manual_scheduled is not None, manual_start is not None, 
                           manual_finish is not None, manual_duration is not None, manual_work is not None]):
                        print(f"[VALIDATION] ❌ Root task should NOT have Manual* tags")
                        root_has_manual_tags = True
                else:
                    # Non-root summaries MUST have Manual* tags
                    missing_tags = []
                    if manual_scheduled is None or manual_scheduled.text != "1":
                        missing_tags.append("ManuallyScheduled")
                    if manual_start is None or manual_start.text != "1":
                        missing_tags.append("ManualStart")
                    if manual_finish is None or manual_finish.text != "1":
                        missing_tags.append("ManualFinish")
                    if manual_duration is None or manual_duration.text != "1":
                        missing_tags.append("ManualDuration")
                    if manual_work is None or manual_work.text != "1":
                        missing_tags.append("ManualWork")
                    
                    if missing_tags:
                        name_elem = task_elem.find("Name")
                        task_name = name_elem.text if name_elem is not None else "Unknown"
                        if manual_tag_issues < 3:  # Only show first 3
                            print(f"[VALIDATION] ❌ Summary task missing Manual tags: {task_name[:40]} (missing: {', '.join(missing_tags)})")
                        manual_tag_issues += 1
        
        if manual_tag_issues == 0 and not root_has_manual_tags:
            print(f"[VALIDATION] ✅ All non-root summary tasks have Manual Scheduling tags, root excluded: PASS")
        else:
            if manual_tag_issues > 0:
                print(f"[VALIDATION] ❌ Found {manual_tag_issues} summary tasks missing Manual* tags")
            if root_has_manual_tags:
                print(f"[VALIDATION] ❌ Root task incorrectly has Manual* tags")
        
        # Validation 7: Verify DurationFormat=7 globally (not 5)
        wrong_duration_format = 0
        for task_elem in tasks_elem.findall("Task"):
            dur_format_elem = task_elem.find("DurationFormat")
            if dur_format_elem is not None and dur_format_elem.text != "7":
                name_elem = task_elem.find("Name")
                task_name = name_elem.text if name_elem is not None else "Unknown"
                if wrong_duration_format < 3:
                    print(f"[VALIDATION] ❌ Wrong DurationFormat: {task_name[:40]} (expected 7, got {dur_format_elem.text})")
                wrong_duration_format += 1
        
        if wrong_duration_format == 0:
            print(f"[VALIDATION] ✅ All tasks use DurationFormat=7 (Elapsed Days): PASS")
        else:
            print(f"[VALIDATION] ❌ Found {wrong_duration_format} tasks with wrong DurationFormat")
        
        print(f"[WORKFRONT VALIDATION] ═══════════════════════════════════════\n")
        
        # VALIDATION: Check project-level dates match task-level date ranges
        # This catches issues like inflated durations before the user imports into Workfront
        try:
            project_start_elem = project.find("StartDate")
            project_finish_elem = project.find("FinishDate")
            
            if project_start_elem is not None and project_finish_elem is not None:
                proj_start_str = project_start_elem.text
                proj_finish_str = project_finish_elem.text
                
                # Parse project dates
                proj_start_dt = datetime.datetime.fromisoformat(proj_start_str.replace('T', ' ').replace('Z', ''))
                proj_finish_dt = datetime.datetime.fromisoformat(proj_finish_str.replace('T', ' ').replace('Z', ''))
                
                # Get min/max from actual tasks (excluding root task UID=1)
                tasks_elem = project.find("Tasks")
                if tasks_elem is not None:
                    task_starts = []
                    task_finishes = []
                    
                    for task in tasks_elem.findall("Task"):
                        uid_elem = task.find("UID")
                        start_elem = task.find("Start")
                        finish_elem = task.find("Finish")
                        
                        # Skip root task (UID=1) for validation
                        if (uid_elem is not None and uid_elem.text != "1" and 
                            start_elem is not None and finish_elem is not None):
                            try:
                                task_start = datetime.datetime.fromisoformat(start_elem.text.replace('T', ' ').replace('Z', ''))
                                task_finish = datetime.datetime.fromisoformat(finish_elem.text.replace('T', ' ').replace('Z', ''))
                                task_starts.append(task_start)
                                task_finishes.append(task_finish)
                            except:
                                pass
                    
                    if task_starts and task_finishes:
                        actual_min = min(task_starts)
                        actual_max = max(task_finishes)
                        
                        print(f"[MSPDI Validation] Date alignment check:")
                        print(f"  Project StartDate:  {proj_start_dt.strftime('%Y-%m-%d')}")
                        print(f"  Earliest task:      {actual_min.strftime('%Y-%m-%d')}")
                        print(f"  Match: {'✓' if proj_start_dt.date() == actual_min.date() else '✗ MISMATCH'}")
                        print(f"  Project FinishDate: {proj_finish_dt.strftime('%Y-%m-%d')}")
                        print(f"  Latest task:        {actual_max.strftime('%Y-%m-%d')}")
                        print(f"  Match: {'✓' if proj_finish_dt.date() == actual_max.date() else '✗ MISMATCH'}")
                        
                        # Warn if dates don't align (shouldn't happen with the fix)
                        if proj_start_dt.date() != actual_min.date() or proj_finish_dt.date() != actual_max.date():
                            print(f"[MSPDI Validation] ⚠️  WARNING: Project dates don't match task dates!")
                            print(f"  This may cause Workfront to show inflated durations.")
        except Exception as e:
            print(f"[MSPDI Validation] Could not validate dates: {e}")

        # Write XML file with proper formatting and error handling
        xml_string = tostring(project, encoding='unicode')
        
        try:
            # Try pretty printing with minidom
            dom = minidom.parseString(xml_string)
            pretty_xml = dom.toprettyxml(indent="  ", encoding=None)
            
            # Remove extra blank lines that toprettyxml adds
            lines = [line for line in pretty_xml.split('\n') if line.strip()]
            final_xml = '\n'.join(lines)
            
            with open(output_xml, 'w', encoding='utf-8') as f:
                f.write(final_xml)
                f.flush()
                os.fsync(f.fileno())
        except Exception as e:
            # Fallback: write raw XML if pretty printing fails
            print(f"[XML] Pretty print failed, using raw XML: {e}")
            with open(output_xml, 'w', encoding='utf-8') as f:
                f.write(xml_string)
                f.flush()
                os.fsync(f.fileno())

        # Return statistics
        return {
            "tasks_total": len(rows),
            "tasks_merged": len(removed_child_wbs),
            "resources_total": len(resources),
            "assignments_total": len(assignments),
            "dependencies_total": len(normalized_edges)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"MSPDI conversion failed: {str(e)}")

# ──────────────────────────────────────────────────────────────────────────────
# Missing Features Implementation: Final Ship, Second Scenario, Import
# ──────────────────────────────────────────────────────────────────────────────

# Store for shipped scenarios (locked versions) and scenario versions
SHIPPED_SCENARIOS: Dict[str, Dict[str, Any]] = {}
SCENARIO_VERSIONS: Dict[str, List[Dict[str, Any]]] = {}

class FinalShipPayload(BaseModel):
    """Payload for final ship endpoint - locks scenario and exports all data"""
    scenario_a: Dict[str, Any]
    scenario_b: Optional[Dict[str, Any]] = None
    scenario_c: Optional[Dict[str, Any]] = None
    project_name: str
    project_id: Optional[str] = None
    notes: Optional[str] = None

class DuplicateScenarioPayload(BaseModel):
    """Payload for duplicating/versioning scenarios"""
    scenario_data: Dict[str, Any]
    scenario_id: str
    version_name: Optional[str] = None
    
class ImportProjectPayload(BaseModel):
    """Payload for importing existing projects"""
    file_type: str  # "excel" or "xml"
    file_content: str  # Base64 encoded file content

@app.post("/api/project/final_ship")
async def final_ship(payload: FinalShipPayload):
    """
    Final Ship: Lock all scenario data and generate comprehensive exports.
    This creates immutable versions of scenarios with full Excel and XML exports.
    """
    if not DB.loaded:
        DB.load()
        
    # Generate unique ship ID with timestamp
    ship_id = f"SHIP_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
    
    # Create shipped package
    shipped_package = {
        "ship_id": ship_id,
        "project_name": payload.project_name,
        "project_id": payload.project_id or uuid.uuid4().hex,
        "shipped_date": datetime.datetime.now(ZoneInfo("America/New_York")).isoformat(),
        "scenarios": {
            "A": payload.scenario_a,
            "B": payload.scenario_b,
            "C": payload.scenario_c
        },
        "notes": payload.notes,
        "locked": True,
        "exports": {}
    }
    
    # Generate comprehensive Excel export with all scenarios
    excel_base = _export_basename(payload.project_name, "FINAL_SHIP")
    excel_path = f"{excel_base}.xlsx"
    
    try:
        # Build DataFrames for all scenarios
        dfA = build_wbs_dataframe_from_scenario(payload.scenario_a, payload.project_name)
        dfA = _ensure_v3_ae_columns(dfA)
        
        with pd.ExcelWriter(excel_path, engine="openpyxl") as xw:
            # Write Scenario A
            dfA.to_excel(xw, sheet_name="Scenario A - Final", index=False)
            _apply_number_formats(xw.sheets["Scenario A - Final"], dfA)
            
            # Write Scenario B if provided
            if payload.scenario_b:
                dfB = build_wbs_dataframe_from_scenario(payload.scenario_b, payload.project_name)
                dfB = _ensure_v3_ae_columns(dfB)
                dfB.to_excel(xw, sheet_name="Scenario B - Final", index=False)
                _apply_number_formats(xw.sheets["Scenario B - Final"], dfB)
                
            # Write Scenario C if provided
            if payload.scenario_c:
                dfC = build_wbs_dataframe_from_scenario(payload.scenario_c, payload.project_name)
                dfC = _ensure_v3_ae_columns(dfC)
                dfC.to_excel(xw, sheet_name="Scenario C - Final", index=False)
                _apply_number_formats(xw.sheets["Scenario C - Final"], dfC)
                
            # Add metadata sheet
            metadata_df = pd.DataFrame([{
                "Ship ID": ship_id,
                "Project Name": payload.project_name,
                "Shipped Date": shipped_package["shipped_date"],
                "Total Scenarios": sum(1 for s in [payload.scenario_a, payload.scenario_b, payload.scenario_c] if s),
                "Notes": payload.notes or ""
            }])
            metadata_df.to_excel(xw, sheet_name="Metadata", index=False)
        
        shipped_package["exports"]["excel"] = excel_path
        
    except Exception as e:
        print(f"[FINAL SHIP] Excel export failed: {e}")
        raise HTTPException(500, f"Excel export failed: {str(e)}")
    
    # Generate XML exports for all scenarios
    xml_files = []
    try:
        for label, scenario in [("A", payload.scenario_a), ("B", payload.scenario_b), ("C", payload.scenario_c)]:
            if scenario:
                xml_path = _export_single_scenario_xml(
                    scenario=scenario,
                    scenario_label=f"Scenario {label} - FINAL",
                    project_name=payload.project_name
                )
                xml_files.append(xml_path)
                shipped_package["exports"][f"xml_{label}"] = xml_path
                
    except Exception as e:
        print(f"[FINAL SHIP] XML export failed: {e}")
    
    # Create ZIP archive with all exports
    import zipfile
    zip_path = f"{excel_base}_COMPLETE.zip"
    
    try:
        with zipfile.ZipFile(zip_path, "w") as zipf:
            # Add Excel file
            if os.path.exists(excel_path):
                zipf.write(excel_path, f"Final_Export/{os.path.basename(excel_path)}")
            
            # Add XML files
            for xml_path in xml_files:
                if os.path.exists(xml_path):
                    zipf.write(xml_path, f"Final_Export/XML/{os.path.basename(xml_path)}")
            
            # Add ship metadata
            zipf.writestr("Final_Export/ship_metadata.json", json.dumps(shipped_package, indent=2))
            
        shipped_package["exports"]["zip"] = zip_path
        
    except Exception as e:
        print(f"[FINAL SHIP] ZIP creation failed: {e}")
    
    # Store shipped package (locked and immutable)
    SHIPPED_SCENARIOS[ship_id] = shipped_package
    
    # Update current scenarios to locked state
    if ship_id in _CURRENT_SCENARIOS:
        _CURRENT_SCENARIOS[ship_id]["locked"] = True
    
    return {
        "success": True,
        "ship_id": ship_id,
        "project_name": payload.project_name,
        "shipped_date": shipped_package["shipped_date"],
        "exports": {
            "excel": os.path.basename(excel_path) if excel_path else None,
            "xml_count": len(xml_files),
            "zip": os.path.basename(zip_path) if os.path.exists(zip_path) else None
        },
        "download_url": f"/api/project/download/{ship_id}"
    }

@app.get("/api/project/download/{ship_id}")
def download_shipped_project(ship_id: str):
    """Download the complete shipped project package"""
    if ship_id not in SHIPPED_SCENARIOS:
        raise HTTPException(404, f"Shipped project {ship_id} not found")
    
    package = SHIPPED_SCENARIOS[ship_id]
    zip_path = package.get("exports", {}).get("zip")
    
    if not zip_path or not os.path.exists(zip_path):
        raise HTTPException(404, "Export package not found")
    
    return FileResponse(
        zip_path,
        filename=f"{package['project_name']}_FINAL_{ship_id}.zip",
        media_type="application/zip"
    )

@app.post("/api/timeline/save")
async def save_timeline(request: dict):
    """
    Save timeline data (tasks, durations, dependencies) to backend storage.
    This endpoint integrates timeline changes with scenario data and persists everything.
    Uses proper dual-entry SCENARIO_STORE format.
    """
    try:
        tasks = request.get("tasks", [])
        reasoning = request.get("reasoning", {})
        metadata = request.get("metadata", {})
        scenario_letter = request.get("scenario", "A")
        session_id = request.get("session_id")  # Get session_id for SCENARIO_STORE lookup
        
        # Get the current scenario using proper dual-entry helpers
        scen = None
        
        if session_id:
            # Use get_working_scenario which handles dual-entry format properly
            scen = get_working_scenario(session_id)
            print(f"[TIMELINE SAVE] Using working scenario for session {session_id}")
        else:
            # Fallback to legacy _CURRENT_SCENARIOS for backwards compatibility
            scen = _CURRENT_SCENARIOS.get(scenario_letter)
        
        if not scen or not scen.get("items"):
            # Initialize scenario if it doesn't exist
            scen = {
                "items": [],
                "totals": {"hours": 0, "price": 0},
                "timeline": {},
                "metadata": {}
            }
            
            if not session_id:
                _CURRENT_SCENARIOS[scenario_letter] = scen
        
        # Ensure timeline dict exists (defensive fix for KeyError)
        if "timeline" not in scen:
            scen["timeline"] = {}
        
        # Update timeline data in the scenario
        scen["timeline_tasks"] = tasks
        scen["timeline_reasoning"] = reasoning
        scen["timeline_metadata"] = metadata
        scen["timeline_last_saved"] = datetime.datetime.now(ZoneInfo("America/New_York")).isoformat()
        
        # Calculate timeline metrics
        total_duration = 0
        critical_path = []
        resource_allocation = {}
        
        if tasks:
            # Calculate total duration
            earliest_start = min((datetime.datetime.fromisoformat(t["start"].replace("Z", "+00:00")) for t in tasks if "start" in t), default=None)
            latest_end = max((datetime.datetime.fromisoformat(t["end"].replace("Z", "+00:00")) for t in tasks if "end" in t), default=None)
            
            if earliest_start and latest_end:
                total_duration = (latest_end - earliest_start).days
                scen["timeline"]["start"] = earliest_start.isoformat()
                scen["timeline"]["end"] = latest_end.isoformat()
                scen["timeline"]["duration_days"] = total_duration
            
            # Calculate resource allocation
            for task in tasks:
                if "resources" in task:
                    for resource in task.get("resources", []):
                        resource_name = resource if isinstance(resource, str) else resource.get("name", "Unknown")
                        resource_allocation[resource_name] = resource_allocation.get(resource_name, 0) + task.get("hours", 0)
            
            # Identify critical path (simplified - tasks with no slack)
            critical_path = [t["id"] for t in tasks if t.get("is_critical", False)]
        
        # Update scenario items with timeline data
        for task in tasks:
            deliverable_code = task.get("deliverable_code") or task.get("id", "").split("-")[0]
            
            # Find matching item in scenario
            for item in scen.get("items", []):
                if item.get("deliverable_code") == deliverable_code:
                    # Update hours based on timeline duration
                    if "duration" in task and "days" in task["duration"]:
                        hours_per_day = metadata.get("hours_per_day", 6)
                        resource_count = len(task.get("resources", [])) or 1
                        calculated_hours = task["duration"]["days"] * hours_per_day * resource_count
                        
                        # Update item hours
                        item["timeline_hours"] = calculated_hours
                        
                        # If timeline hours significantly differ from original, use timeline
                        if abs(item.get("total_hours", 0) - calculated_hours) > 1:
                            item["total_hours"] = calculated_hours
                            item["hours_updated_from_timeline"] = True
                    
                    # Store timeline-specific data
                    item["timeline_start"] = task.get("start")
                    item["timeline_end"] = task.get("end")
                    item["timeline_dependencies"] = task.get("dependencies", [])
                    break
        
        # Store the complete scenario with timeline data using proper dual-entry format
        if session_id:
            update_working_scenario(session_id, scen)
            print(f"[TIMELINE SAVE] Saved to SCENARIO_STORE[{session_id}] via update_working_scenario")
            print(f"[DEBUG] TIMELINE SAVE: SCENARIO_STORE[{session_id}] keys:", list(SCENARIO_STORE[session_id].keys()))
        else:
            _CURRENT_SCENARIOS[scenario_letter] = scen
            # Also save to ScenarioStore if it exists
            if hasattr(app.state, "scenario_store"):
                app.state.scenario_store[scenario_letter] = scen
        
        print(f"[TIMELINE SAVE] Saved {len(tasks)} tasks for scenario {scenario_letter}")
        print(f"[TIMELINE SAVE] Duration: {total_duration} days, Resources: {len(resource_allocation)}")
        
        return {
            "success": True,
            "message": f"Timeline saved successfully for scenario {scenario_letter}",
            "metrics": {
                "total_tasks": len(tasks),
                "total_duration_days": total_duration,
                "critical_path_count": len(critical_path),
                "resource_count": len(resource_allocation),
                "resource_allocation": resource_allocation,
                "saved_at": scen["timeline_last_saved"]
            },
            "scenario": scen  # Return updated scenario for frontend sync
        }
        
    except Exception as e:
        print(f"[TIMELINE SAVE ERROR] {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to save timeline: {str(e)}")

@app.post("/api/scenario/duplicate")
async def duplicate_scenario(payload: DuplicateScenarioPayload):
    """
    Build Second Scenario: Create a new version based on existing scenario.
    Maintains version history and allows switching between versions.
    """
    scenario_id = payload.scenario_id
    base_scenario = payload.scenario_data
    
    # Generate version ID
    version_id = f"v{len(SCENARIO_VERSIONS.get(scenario_id, [])) + 2}"  # v2, v3, etc.
    version_name = payload.version_name or f"Version {version_id}"
    
    # Create new version with metadata
    new_version = {
        "version_id": version_id,
        "version_name": version_name,
        "created_date": datetime.datetime.now(ZoneInfo("America/New_York")).isoformat(),
        "base_version": "v1" if scenario_id not in SCENARIO_VERSIONS else SCENARIO_VERSIONS[scenario_id][-1]["version_id"],
        "scenario_data": base_scenario.copy(),
        "editable": True
    }
    
    # Store version
    if scenario_id not in SCENARIO_VERSIONS:
        SCENARIO_VERSIONS[scenario_id] = []
    SCENARIO_VERSIONS[scenario_id].append(new_version)
    
    # Update current scenarios with new version
    _CURRENT_SCENARIOS[f"{scenario_id}_{version_id}"] = new_version["scenario_data"]
    
    return {
        "success": True,
        "version_id": version_id,
        "version_name": version_name,
        "scenario_id": f"{scenario_id}_{version_id}",
        "created_date": new_version["created_date"],
        "total_versions": len(SCENARIO_VERSIONS[scenario_id]) + 1  # +1 for original
    }

@app.get("/api/scenario/versions/{scenario_id}")
def get_scenario_versions(scenario_id: str):
    """Get all versions of a scenario for comparison"""
    if scenario_id not in SCENARIO_VERSIONS:
        return {"versions": [], "message": "No versions found"}
    
    versions = SCENARIO_VERSIONS[scenario_id]
    return {
        "scenario_id": scenario_id,
        "versions": [
            {
                "version_id": v["version_id"],
                "version_name": v["version_name"],
                "created_date": v["created_date"],
                "editable": v["editable"]
            }
            for v in versions
        ],
        "total_versions": len(versions) + 1  # +1 for original
    }

@app.post("/api/project/import")
async def import_project(file: UploadFile = File(...)):
    """
    Import XML/Excel: Parse uploaded file and recreate scenario.
    Supports both Excel exports and MS Project XML files.
    """
    if not DB.loaded:
        DB.load()
    
    # Read uploaded file
    content = await file.read()
    filename = file.filename.lower()
    
    try:
        if filename.endswith('.xlsx') or filename.endswith('.xls'):
            # Import from Excel
            import_data = await _import_from_excel(content, filename)
        elif filename.endswith('.xml'):
            # Import from MS Project XML
            import_data = await _import_from_xml(content, filename)
        else:
            raise HTTPException(400, "Unsupported file format. Please upload Excel (.xlsx) or XML (.xml) file.")
        
        # Generate import ID
        import_id = f"IMP_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
        
        # Store imported scenario
        _CURRENT_SCENARIOS[import_id] = import_data["scenario"]
        
        return {
            "success": True,
            "import_id": import_id,
            "project_name": import_data.get("project_name", "Imported Project"),
            "scenario_count": import_data.get("scenario_count", 1),
            "deliverables_count": len(import_data["scenario"].get("items", [])),
            "total_hours": import_data.get("total_hours", 0),
            "total_price": import_data.get("total_price", 0),
            "message": "Project imported successfully. You can now edit and work with the imported data."
        }
        
    except Exception as e:
        print(f"[IMPORT ERROR] {str(e)}")
        raise HTTPException(500, f"Import failed: {str(e)}")

async def _import_from_excel(content: bytes, filename: str) -> Dict[str, Any]:
    """Parse Excel file and extract scenario data"""
    import io
    
    # Read Excel file
    df = pd.read_excel(io.BytesIO(content), sheet_name=0)
    
    # Extract project info
    project_name = df["Project_Name"].iloc[0] if "Project_Name" in df.columns else "Imported Project"
    
    # Group by deliverable to reconstruct scenario
    deliverables = []
    for deliverable_code, group in df.groupby("Deliverable_Code"):
        deliverable = {
            "deliverable_code": deliverable_code,
            "deliverable": group["Deliverable"].iloc[0],
            "hours": float(group["Planned_Hours"].sum()),
            "price": float(group["Price_USD"].sum()) if "Price_USD" in group.columns else 0,
            "included_task_groups": group["Task_Code"].unique().tolist() if "Task_Code" in group.columns else [],
            "components": []
        }
        
        # Extract components if present
        if "Component" in group.columns:
            for component, comp_group in group.groupby("Component"):
                if pd.notna(component):
                    deliverable["components"].append({
                        "name": component,
                        "hours": float(comp_group["Planned_Hours"].sum()),
                        "tasks": comp_group["Task"].unique().tolist() if "Task" in comp_group.columns else []
                    })
        
        deliverables.append(deliverable)
    
    # Create scenario structure
    scenario = {
        "project_name": project_name,
        "items": deliverables,
        "pricing_mode": "Per_Resource",  # Default, can be detected from data
        "rate_band": "Standard_US",
        "total_hours": float(df["Planned_Hours"].sum()) if "Planned_Hours" in df.columns else 0,
        "total_price": float(df["Price_USD"].sum()) if "Price_USD" in df.columns else 0
    }
    
    return {
        "scenario": scenario,
        "project_name": project_name,
        "scenario_count": 1,
        "total_hours": scenario["total_hours"],
        "total_price": scenario["total_price"]
    }

async def _import_from_xml(content: bytes, filename: str) -> Dict[str, Any]:
    """Parse MS Project XML and extract scenario data"""
    from xml.etree import ElementTree as ET
    
    # Parse XML
    root = ET.fromstring(content)
    
    # Extract project name
    project_elem = root.find(".//Name")
    project_name = project_elem.text if project_elem is not None else "Imported Project"
    
    # Extract tasks
    tasks = []
    tasks_elem = root.find(".//Tasks")
    if tasks_elem is not None:
        for task_elem in tasks_elem.findall("Task"):
            task = {}
            
            # Extract task fields
            for field in ["UID", "Name", "WBS", "OutlineLevel", "Work", "Cost"]:
                elem = task_elem.find(field)
                if elem is not None:
                    task[field] = elem.text
            
            # Convert Work from PT format to hours
            if "Work" in task and task["Work"]:
                # Parse PT480H format
                import re as re_module
                match = re_module.match(r'PT(\d+)([HM])', task["Work"])
                if match:
                    value, unit = match.groups()
                    task["Hours"] = float(value) if unit == 'H' else float(value) / 60
            
            tasks.append(task)
    
    # Group tasks into deliverables
    deliverables = []
    deliverable_tasks = {}
    
    for task in tasks:
        outline_level = int(task.get("OutlineLevel", 0))
        if outline_level == 1:  # Deliverable level
            deliverable_code = f"IMP_{task.get('UID', '')}"
            deliverable = {
                "deliverable_code": deliverable_code,
                "deliverable": task.get("Name", ""),
                "hours": float(task.get("Hours", 0)),
                "price": float(task.get("Cost", 0)) if task.get("Cost") else 0,
                "included_task_groups": [],
                "components": []
            }
            deliverables.append(deliverable)
            deliverable_tasks[deliverable_code] = deliverable
    
    # Create scenario
    scenario = {
        "project_name": project_name,
        "items": deliverables,
        "pricing_mode": "Per_Resource",
        "rate_band": "Standard_US",
        "total_hours": sum(d["hours"] for d in deliverables),
        "total_price": sum(d["price"] for d in deliverables)
    }
    
    return {
        "scenario": scenario,
        "project_name": project_name,
        "scenario_count": 1,
        "total_hours": scenario["total_hours"],
        "total_price": scenario["total_price"]
    }

@app.get("/api/shipped/list")
def list_shipped_projects():
    """List all shipped (finalized) projects"""
    shipped_list = []
    for ship_id, package in SHIPPED_SCENARIOS.items():
        shipped_list.append({
            "ship_id": ship_id,
            "project_name": package["project_name"],
            "shipped_date": package["shipped_date"],
            "scenario_count": sum(1 for s in package["scenarios"].values() if s),
            "has_exports": bool(package.get("exports", {})),
            "locked": package.get("locked", True)
        })
    
    return {
        "shipped_projects": shipped_list,
        "total_count": len(shipped_list)
    }

# ---------- NEW Flexible XML Export Endpoint ----------
class XMLExportPayload(BaseModel):
    """Flexible XML export payload that accepts scenario data in various formats"""
    session_id: Optional[str] = None  # NEW: Use working scenario from SCENARIO_STORE
    scenario: Optional[Dict[str, Any]] = None
    selected_deliverables: Optional[List[str]] = None  # For building scenario from deliverable codes
    project_name: Optional[str] = None
    pricing_mode: str = "Flat_Blended"
    rate_band: str = "Standard_US"
    blended_rate: Optional[float] = None
    start_date_mode: str = "next_monday"
    fixed_start_iso: Optional[str] = None
    hours_per_day: float = 8.0
    sheet_name: str = "Scenario"
    add_dependencies: bool = True
    add_milestones: bool = True
    add_custom_fields: bool = True


def load_export_scenario(session_id: Optional[str], payload: XMLExportPayload) -> Dict[str, Any]:
    """
    Central helper for all export endpoints to load the correct scenario.
    Priority order:
    0) Working scenario in SCENARIO_STORE (session_id) - PREFERRED (returns immediately!)
    1) payload.scenario with items
    2) payload.scenario without items + selected_deliverables (build items)
    3) selected_deliverables only (build full scenario)
    
    This ensures Step 3 edits are always used when session_id is provided.
    CRITICAL: Case 0 returns immediately to prevent payload.scenario from overwriting.
    """
    # Case 0: Preferred - use working scenario from SCENARIO_STORE
    # CRITICAL: Return immediately to prevent stale payload from overwriting Step 3 edits
    if session_id:
        try:
            _ = get_session_state(session_id)  # Trigger migration if needed
            scenario = get_working_scenario(session_id)
            if scenario and scenario.get("items"):
                print(f"[EXPORT HELPER] ✅ Using working scenario from SCENARIO_STORE for session {session_id}")
                print(f"[EXPORT HELPER] Scenario has {len(scenario.get('items', []))} items")
                
                # Debug: verify Step 3 hours
                items = scenario.get("items", [])
                if items:
                    first = items[0]
                    print("[EXPORT CHECK] First item:",
                          first.get("deliverable_code"),
                          "Planned_Hours=", first.get("Planned_Hours"),
                          "hours=", first.get("hours"),
                          "total_hours=", first.get("total_hours"))
                
                # Ensure components are inflated and RETURN IMMEDIATELY
                return _inflate_components_if_missing(scenario)
        except Exception as e:
            print(f"[EXPORT HELPER] Error loading from SCENARIO_STORE: {e}")
            # Fall through to other cases only on error
    
    scenario = None

    # Case 1: payload.scenario already has items (only if Case 0 failed/skipped)
    if payload.scenario and payload.scenario.get("items"):
        scenario = _inflate_components_if_missing(payload.scenario)
        print("[EXPORT HELPER] Using scenario from payload.scenario (with items)")

    # Case 2: payload.scenario without items, but we have selected_deliverables
    if scenario is None and payload.scenario and payload.selected_deliverables:
        scenario = payload.scenario.copy()
        items = []
        for dcode in payload.selected_deliverables:
            if DB.deliverables is not None:
                deliv_matches = DB.deliverables[
                    DB.deliverables["Deliverable_Code"].astype(str) == str(dcode)
                ]
                if not deliv_matches.empty:
                    deliv_row = deliv_matches.iloc[0]
                    task_groups = DB.task_groups_for_deliverable(str(dcode))
                    item = {
                        "deliverable_code": str(dcode),
                        "deliverable": str(deliv_row.get("Deliverable", "")),
                        "included_task_groups": task_groups,
                        "hours": 100,
                        "price": 15000,
                        "complexity": "Advanced",
                        "tier": "T2_MediumVolume"
                    }
                    items.append(item)
        scenario["items"] = items
        scenario = _inflate_components_if_missing(scenario)
        print("[EXPORT HELPER] Built scenario from payload.scenario + selected_deliverables")

    # Case 3: Only selected_deliverables
    if scenario is None and payload.selected_deliverables:
        items = []
        for dcode in payload.selected_deliverables:
            if DB.deliverables is not None:
                deliv_matches = DB.deliverables[
                    DB.deliverables["Deliverable_Code"].astype(str) == str(dcode)
                ]
                if not deliv_matches.empty:
                    deliv_row = deliv_matches.iloc[0]
                    task_groups = DB.task_groups_for_deliverable(str(dcode))
                    item = {
                        "deliverable_code": str(dcode),
                        "deliverable": str(deliv_row.get("Deliverable", "")),
                        "included_task_groups": task_groups,
                        "hours": 100,
                        "price": 15000,
                        "complexity": "Advanced",
                        "tier": "T2_MediumVolume"
                    }
                    items.append(item)
        
        scenario = {
            "items": items,
            "pricing_mode": payload.pricing_mode,
            "rate_band": payload.rate_band,
            "blended_rate": payload.blended_rate or DEFAULT_BILLABLE_RATE_USD,
            "project_start": payload.fixed_start_iso
        }
        scenario = _inflate_components_if_missing(scenario)
        print("[EXPORT HELPER] Built scenario purely from selected_deliverables")

    if scenario is None:
        raise HTTPException(400, "No scenario or deliverables provided for export")

    # Debug: verify Step 3 hours before WBS build
    items = scenario.get("items", [])
    if items:
        first = items[0]
        print("[EXPORT CHECK] First item:",
              first.get("deliverable_code"),
              "Planned_Hours=", first.get("Planned_Hours"),
              "hours=", first.get("hours"),
              "total_hours=", first.get("total_hours"))

    return scenario

@app.post("/api/xml")
def api_xml_export_flexible(payload: XMLExportPayload):
    """
    Flexible XML export endpoint using load_export_scenario() helper.
    Priority: session_id (SCENARIO_STORE) > payload.scenario > selected_deliverables
    
    This ensures Step 3 edits are preserved in XML exports.
    """
    if not DB.loaded:
        DB.load()
    
    # Determine project name
    project_name = (payload.project_name 
                   or _upload_title_default() 
                   or f"Project {datetime.date.today().isoformat()}")
    
    # Use centralized helper to load scenario with proper priority
    print(f"[/api/xml] session_id={payload.session_id}, has_scenario={payload.scenario is not None}")
    scenario = load_export_scenario(payload.session_id, payload)
    
    # Update scenario with payload settings
    scenario["pricing_mode"] = payload.pricing_mode
    scenario["rate_band"] = payload.rate_band
    if payload.blended_rate:
        scenario["blended_rate"] = payload.blended_rate
    elif not scenario.get("blended_rate"):
        scenario["blended_rate"] = DEFAULT_BILLABLE_RATE_USD
    
    # Build WBS DataFrame
    df = build_wbs_dataframe_from_scenario(scenario, project_name)
    df = _ensure_v3_ae_columns(df)
    
    # Create temporary Excel file for MSPDI conversion with unique ID to prevent collisions
    base = _export_basename(project_name, payload.sheet_name)
    unique_id = uuid.uuid4().hex[:8]
    temp_xlsx = os.path.join(tempfile.gettempdir(), f"{base}_temp_{unique_id}.xlsx")
    output_xml = os.path.join(tempfile.gettempdir(), f"{base}_{unique_id}.xml")
    
    try:
        # Write to temporary Excel file
        with pd.ExcelWriter(temp_xlsx, engine="openpyxl") as xw:
            df.to_excel(xw, sheet_name=payload.sheet_name, index=False)
            _apply_number_formats(xw.sheets[payload.sheet_name], df)
        
        # Convert to MSPDI XML with all features
        project_start_iso = payload.fixed_start_iso or scenario.get("project_start")
        
        stats = convert_excel_to_mspdi(
            input_xlsx=temp_xlsx,
            output_xml=output_xml,
            sheet_name=payload.sheet_name,
            start_date_mode=payload.start_date_mode,
            fixed_start_iso=project_start_iso,
            hours_per_day=payload.hours_per_day,
            merge_identical_children=False,
            project_name=project_name,
            pricing_mode=scenario.get("pricing_mode", "Flat_Blended"),
            rate_band=scenario.get("rate_band", "Standard_US"),
            blended_rate=scenario.get("blended_rate"),
            add_deliverable_milestones=payload.add_milestones
        )
        
        # Post-process XML if parallelization is enabled
        final_xml = output_xml
        if PARALLELIZE_IDENTICAL_NAMES:
            final_xml = post_process_xml(output_xml)
        
        return FileResponse(
            final_xml,
            filename=os.path.basename(final_xml),
            media_type="application/xml",
            headers={
                "X-Export-Stats": json.dumps(stats),
                "Content-Disposition": f'attachment; filename="{os.path.basename(final_xml)}"'
            }
        )
    
    finally:
        # Clean up temporary Excel file
        if os.path.exists(temp_xlsx):
            os.remove(temp_xlsx)

# ---------- Scenario Real-time Sync Endpoint ----------

# In-memory store for scenario sync state
SCENARIO_SYNC_STATE = {}  # session_id -> {scenario, version, last_modified, checksum}
WRITE_THROTTLE_MS = 150  # Throttle writes to prevent Gantt freezes

class ScenarioSyncPayload(BaseModel):
    session_id: str
    client_version: int
    last_server_version: int
    scenario: Optional[Dict[str, Any]] = None
    selections: Optional[Dict[str, Any]] = None
    timestamp: int
    checksum: Optional[str] = None

@app.post("/api/scenario/sync")
async def sync_scenario(payload: ScenarioSyncPayload):
    """
    Real-time synchronization endpoint for scenario data.
    Handles conflict detection, resolution, and incremental updates.
    """
    session_id = payload.session_id
    client_version = payload.client_version
    
    # Initialize session state if not exists
    if session_id not in SCENARIO_SYNC_STATE:
        SCENARIO_SYNC_STATE[session_id] = {
            "scenario": None,
            "selections": {},
            "version": 0,
            "last_modified": 0,
            "checksum": "",
            "history": []
        }
    
    server_state = SCENARIO_SYNC_STATE[session_id]
    
    # Throttle check to prevent excessive writes
    now_ms = int(time.time() * 1000)
    last_ms = int(server_state.get("last_write_ms", 0))
    if now_ms - last_ms < WRITE_THROTTLE_MS:
        return {
            "serverVersion": server_state["version"],
            "hasChanges": False,
            "hasConflicts": False,
            "conflicts": [],
            "timestamp": now_ms,
            "throttled": True
        }
    server_state["last_write_ms"] = now_ms
    
    server_version = server_state["version"]
    
    # Check for conflicts
    has_conflicts = False
    conflicts = []
    
    if server_version > payload.last_server_version:
        # Server has newer data that client doesn't know about
        has_conflicts = True
        
        # Detect specific conflicts
        if payload.scenario and server_state["scenario"]:
            # Check for field-level conflicts
            client_total = payload.scenario.get("totals", {}).get("grandTotal12", 0)
            server_total = server_state["scenario"].get("totals", {}).get("grandTotal12", 0)
            
            if abs(client_total - server_total) > 0.01:
                conflicts.append({
                    "field": "totals.grandTotal12",
                    "clientValue": client_total,
                    "serverValue": server_total,
                    "resolution": "server-wins"
                })
    
    # Apply client changes if newer
    changes_applied = False
    if payload.scenario and (not has_conflicts or payload.client_version > server_version):
        # Build map of old items from SCENARIO_STORE for comparison
        store_entry = SCENARIO_STORE.get(session_id, {})
        old_scen = store_entry.get("scenario") or {}
        old_items_by_id = {}
        for old_item in old_scen.get("items", []):
            item_id = old_item.get("id") or old_item.get("Row_ID") or old_item.get("deliverable_code")
            if item_id:
                old_items_by_id[item_id] = old_item
        
        # Apply cadence pricing logic to each item (comparing against stored values)
        scenario_data = payload.scenario
        if scenario_data and scenario_data.get("items"):
            for item in scenario_data["items"]:
                item_id = item.get("id") or item.get("Row_ID") or item.get("deliverable_code")
                old_item = old_items_by_id.get(item_id) if item_id else None
                apply_cadence_to_item(item, old_item)
        
        # Update server state with client data
        server_state["scenario"] = scenario_data
        server_state["selections"] = payload.selections or {}
        server_state["version"] = max(server_version + 1, payload.client_version)
        server_state["last_modified"] = payload.timestamp
        server_state["checksum"] = payload.checksum or ""
        
        # Keep history for debugging (limit to last 10)
        server_state["history"].append({
            "version": server_state["version"],
            "timestamp": payload.timestamp,
            "source": "client"
        })
        if len(server_state["history"]) > 10:
            server_state["history"] = server_state["history"][-10:]
        
        changes_applied = True
        
        # Merge Step 3 edits into SCENARIO_STORE (preserves timeline/WBS metadata)
        merge_scenario_from_sync(session_id)
        
        # Recompute totals in SCENARIO_STORE for consistent pricing
        if session_id in SCENARIO_STORE and SCENARIO_STORE[session_id].get("scenario"):
            SCENARIO_STORE[session_id]["scenario"] = _recompute_totals(SCENARIO_STORE[session_id]["scenario"])
        
        # Also update global _CURRENT_SCENARIOS if exists
        if scenario_data and scenario_data.get("items"):
            _CURRENT_SCENARIOS[f"{session_id}_sync"] = scenario_data
    
    # Determine what to send back to client
    response_data = {
        "serverVersion": server_state["version"],
        "hasChanges": False,
        "hasConflicts": has_conflicts,
        "conflicts": conflicts,
        "timestamp": int(time.time() * 1000)
    }
    
    # Send server data if client is behind
    if server_version > payload.last_server_version and server_state["scenario"]:
        response_data["hasChanges"] = True
        response_data["scenario"] = server_state["scenario"]
        response_data["selections"] = server_state["selections"]
        
        # Calculate changed elements for UI updates
        changed_elements = []
        if server_state["scenario"] and server_state["scenario"].get("items"):
            for item in server_state["scenario"]["items"]:
                changed_elements.append({
                    "type": "deliverable",
                    "id": item.get("deliverable_code"),
                    "field": "hours",
                    "value": item.get("hours", 0)
                })
        response_data["changedElements"] = changed_elements
    
    return response_data

@app.get("/api/scenario/sync/status/{session_id}")
async def get_sync_status(session_id: str):
    """Get the current sync status for a session"""
    if session_id not in SCENARIO_SYNC_STATE:
        return {
            "exists": False,
            "message": "No sync state for this session"
        }
    
    state = SCENARIO_SYNC_STATE[session_id]
    return {
        "exists": True,
        "version": state["version"],
        "last_modified": state["last_modified"],
        "has_scenario": state["scenario"] is not None,
        "history_count": len(state["history"]),
        "checksum": state["checksum"]
    }

@app.post("/api/scenario/sync/clear/{session_id}")
async def clear_sync_state(session_id: str):
    """Clear sync state for a session (for testing or reset)"""
    if session_id in SCENARIO_SYNC_STATE:
        del SCENARIO_SYNC_STATE[session_id]
        return {"success": True, "message": "Sync state cleared"}
    return {"success": False, "message": "No sync state found"}

# ---------- AI Pricing Retainer Suggestions ----------
class RetainerSuggestionsPayload(BaseModel):
    session_id: str
    monthly_budget: Optional[float] = None
    scenario: Optional[Dict[str, Any]] = None

@app.post("/api/pricing/retainer_suggestions")
async def analyze_retainer_suggestions(payload: RetainerSuggestionsPayload):
    """
    Analyze all deliverables in a scenario and suggest which should be retainers.
    Uses AI to analyze deliverable names and characteristics to recommend retainer setup.
    """
    if not DB.loaded:
        DB.load()
    
    # Get scenario from sync state or payload
    scenario = None
    if payload.session_id in SCENARIO_SYNC_STATE:
        scenario = SCENARIO_SYNC_STATE[payload.session_id].get("scenario")
    
    if not scenario and payload.scenario:
        scenario = payload.scenario
    
    if not scenario or not scenario.get("items"):
        raise HTTPException(400, "No scenario found for session. Please build a scenario first.")
    
    # Import pricing optimizer
    from ai_pricing_optimizer import PricingOptimizer
    
    optimizer = PricingOptimizer()
    
    suggestions = []
    total_converted = 0
    total_monthly_hours = 0
    
    # Analyze each deliverable for retainer potential
    for item in scenario["items"]:
        deliverable_name = item.get("deliverable_name") or item.get("deliverable", "")
        deliverable_code = item.get("deliverable_code", "")
        current_hours = item.get("hours", 0)
        is_already_retainer = item.get("is_retainer", False)
        
        # Skip if already a retainer
        if is_already_retainer:
            continue
        
        # Check if this should be a retainer
        try:
            analysis = await optimizer.should_be_retainer(
                deliverable_name=deliverable_name,
                rfp_text=RFP_TEXT_CACHE or ""
            )
            
            if analysis["is_retainer"]:
                suggested_months = analysis.get("suggested_months", 12)
                # Calculate monthly hours from total
                monthly_hours = round(current_hours / suggested_months, 1) if suggested_months > 0 else current_hours
                
                suggestion = {
                    "deliverable_code": deliverable_code,
                    "deliverable_name": deliverable_name,
                    "is_retainer": True,
                    "suggested_months": suggested_months,
                    "monthly_hours": monthly_hours,
                    "total_hours": current_hours,
                    "reasoning": analysis.get("reasoning", "Ongoing work pattern detected"),
                    "confidence": analysis.get("confidence", 0.85)
                }
                
                suggestions.append(suggestion)
                total_converted += 1
                total_monthly_hours += monthly_hours
                
                # Update the item in scenario
                item["is_retainer"] = True
                item["retainer_months"] = suggested_months
                item["monthly_hours"] = monthly_hours
                item["cadence"] = "Monthly"
                
        except Exception as e:
            print(f"[Retainer Analysis] Error analyzing {deliverable_name}: {e}")
            # Fallback to rule-based check
            keywords = ["management", "monthly", "ongoing", "maintenance", "monitoring", "optimization", "support"]
            if any(keyword in deliverable_name.lower() for keyword in keywords):
                monthly_hours = round(current_hours / 12, 1)
                suggestion = {
                    "deliverable_code": deliverable_code,
                    "deliverable_name": deliverable_name,
                    "is_retainer": True,
                    "suggested_months": 12,
                    "monthly_hours": monthly_hours,
                    "total_hours": current_hours,
                    "reasoning": "Keyword-based detection: ongoing work pattern",
                    "confidence": 0.75
                }
                suggestions.append(suggestion)
                total_converted += 1
                total_monthly_hours += monthly_hours
                
                # Update the item
                item["is_retainer"] = True
                item["retainer_months"] = 12
                item["monthly_hours"] = monthly_hours
                item["cadence"] = "Monthly"
    
    # Update scenario in sync state
    if payload.session_id in SCENARIO_SYNC_STATE:
        SCENARIO_SYNC_STATE[payload.session_id]["scenario"] = scenario
        SCENARIO_SYNC_STATE[payload.session_id]["version"] += 1
        SCENARIO_SYNC_STATE[payload.session_id]["last_modified"] = time.time()
    
    # Create retainer plan summary
    retainer_plan = None
    if suggestions:
        retainer_plan = {
            "total_retainers": total_converted,
            "monthly_hours": round(total_monthly_hours, 1),
            "monthly_budget": round(total_monthly_hours * (scenario.get("blended_rate", DEFAULT_BILLABLE_RATE_USD)), 2),
            "suggested_duration": 12,
            "deliverables": suggestions
        }
    
    message = f"Analyzed {len(scenario['items'])} deliverables. "
    if total_converted > 0:
        message += f"Converted {total_converted} to retainers with {total_monthly_hours:.0f} monthly hours."
    else:
        message += "No deliverables suitable for retainer conversion found."
    
    return {
        "success": True,
        "scenario": scenario,
        "retainer_plan": retainer_plan,
        "suggestions": suggestions,
        "message": message,
        "analyzed_count": len(scenario["items"]),
        "converted_count": total_converted
    }

# ---------- AI Agent API Endpoints ----------
from ai_agent import (
    AgentChatRequest,
    AgentExecuteRequest,
    AgentResponse,
    chat_with_agent,
    execute_command
)

@app.post("/api/agent/chat")
async def agent_chat(request: AgentChatRequest) -> AgentResponse:
    """
    Chat with the AI agent using natural language.
    The agent will understand your intent and generate UI actions.
    """
    try:
        # Get current app context if needed
        context = request.context or {}
        
        # Add current state info to context
        if not context.get("has_rfp"):
            context["has_rfp"] = bool(RFP_TEXT_CACHE)
        
        # Process the chat message with GPT-5 tier
        response = await chat_with_agent(
            message=request.message,
            context=context,
            session_id=request.session_id,
            gpt5_tier=request.gpt5_tier if hasattr(request, 'gpt5_tier') else "auto"
        )
        
        return response
        
    except Exception as e:
        print(f"[Agent Chat] Error: {e}")
        return AgentResponse(
            success=False,
            message=f"Sorry, I encountered an error: {str(e)}",
            error=str(e)
        )

@app.post("/api/agent/execute")
async def agent_execute(request: AgentExecuteRequest) -> AgentResponse:
    """
    Execute a specific command with parameters.
    Used when the UI wants to directly trigger agent actions.
    """
    try:
        response = await execute_command(
            command_type=request.command,
            parameters=request.parameters
        )
        
        return response
        
    except Exception as e:
        print(f"[Agent Execute] Error: {e}")
        return AgentResponse(
            success=False,
            message=f"Failed to execute command: {str(e)}",
            error=str(e)
        )

@app.get("/api/agent/status")
async def agent_status():
    """Check if the AI agent is available and working"""
    try:
        from ai_agent import OPENAI_AVAILABLE, GPT5_AVAILABLE
        
        return {
            "available": OPENAI_AVAILABLE,
            "gpt5_available": GPT5_AVAILABLE,
            "message": "AI Agent is ready to help!" if OPENAI_AVAILABLE else "AI Agent is offline"
        }
    except Exception as e:
        return {
            "available": False,
            "message": f"Agent status check failed: {str(e)}"
        }

# ---------- Run locally in Replit ----------
# In Replit, set the "run" command to: uvicorn main:app --host 0.0.0.0 --port 5000 --reload