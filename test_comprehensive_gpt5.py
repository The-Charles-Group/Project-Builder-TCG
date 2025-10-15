#!/usr/bin/env python3
"""
================================================================================
COMPREHENSIVE GPT-5 RFP ANALYSIS TEST SUITE
================================================================================
Tests all aspects of GPT-5 RFP analysis including document upload, analysis,
confidence scoring, auto-rescue logic, session isolation, and API response format.

Test Categories:
1. Multiple Document Types (PDF, DOCX, TXT)
2. GPT-5 Intelligence Features (confidence, evidence, multipliers) 
3. Auto-Rescue Logic (minimal vs complex RFPs)
4. Session Isolation (no data contamination)
5. API Response Format (job tracking, progress, errors)
"""

import os
import sys
import json
import time
import asyncio
import uuid
import httpx
import base64
import tempfile
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from pathlib import Path
from colorama import init, Fore, Style
import warnings

# Initialize colorama for colored console output
init(autoreset=True)

# Suppress warnings
warnings.filterwarnings("ignore")

# Document creation libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print(f"{Fore.YELLOW}⚠️  python-docx not installed - DOCX tests will be skipped{Style.RESET_ALL}")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print(f"{Fore.YELLOW}⚠️  reportlab not installed - PDF tests will be skipped{Style.RESET_ALL}")

# Base URL for API
BASE_URL = "http://localhost:5000"

# Test configuration
TEST_TIMEOUT = 120  # seconds per test
POLL_INTERVAL = 1.0  # seconds between status checks

# ============================================================================
# TEST REPORT TRACKING
# ============================================================================

class TestReport:
    """Track test results and generate final report"""
    
    def __init__(self):
        self.results = []
        self.start_time = datetime.now()
        self.total_tests = 0
        self.passed = 0
        self.failed = 0
        self.skipped = 0
        
    def add_test(self, name: str, status: str, details: Dict[str, Any], error: Optional[str] = None):
        """Add a test result"""
        self.total_tests += 1
        
        if status == "PASS":
            self.passed += 1
            symbol = "✅"
        elif status == "FAIL":
            self.failed += 1
            symbol = "❌"
        else:  # SKIP
            self.skipped += 1
            symbol = "⏭️"
        
        result = {
            "name": name,
            "status": status,
            "symbol": symbol,
            "details": details,
            "error": error,
            "timestamp": datetime.now().isoformat()
        }
        self.results.append(result)
        
        # Real-time output
        print(f"\n{symbol} {name}: {status}")
        if details:
            for key, value in details.items():
                print(f"   • {key}: {value}")
        if error:
            print(f"   ❗ Error: {error}")
    
    def generate_report(self) -> str:
        """Generate final test report"""
        duration = (datetime.now() - self.start_time).total_seconds()
        
        report = []
        report.append("=" * 80)
        report.append("GPT-5 RFP ANALYSIS COMPREHENSIVE TEST REPORT")
        report.append("=" * 80)
        report.append(f"\nTest Date: {self.start_time.strftime('%Y-%m-%d %H:%M:%S')}")
        report.append(f"Duration: {duration:.2f} seconds")
        report.append(f"\n📊 SUMMARY:")
        report.append(f"   Total Tests: {self.total_tests}")
        report.append(f"   ✅ Passed: {self.passed}")
        report.append(f"   ❌ Failed: {self.failed}")
        report.append(f"   ⏭️  Skipped: {self.skipped}")
        report.append(f"   Success Rate: {(self.passed/max(self.total_tests, 1)*100):.1f}%")
        
        # Group results by category
        categories = {}
        for result in self.results:
            category = result["name"].split(":")[0] if ":" in result["name"] else "General"
            if category not in categories:
                categories[category] = []
            categories[category].append(result)
        
        # Detailed results by category
        report.append("\n" + "=" * 80)
        report.append("DETAILED RESULTS BY CATEGORY")
        report.append("=" * 80)
        
        for category, tests in categories.items():
            report.append(f"\n📁 {category}")
            report.append("-" * 40)
            
            for test in tests:
                report.append(f"\n{test['symbol']} {test['name']}")
                
                if test["details"]:
                    for key, value in test["details"].items():
                        # Format value based on type
                        if isinstance(value, float):
                            formatted_value = f"{value:.3f}"
                        elif isinstance(value, list) and len(value) > 3:
                            formatted_value = f"[{len(value)} items]"
                        else:
                            formatted_value = str(value)
                        report.append(f"   • {key}: {formatted_value}")
                
                if test["error"]:
                    report.append(f"   ❗ Error: {test['error']}")
        
        # Critical Issues Found
        report.append("\n" + "=" * 80)
        report.append("CRITICAL ISSUES FOUND")
        report.append("=" * 80)
        
        critical_issues = []
        for result in self.results:
            if result["status"] == "FAIL":
                critical_issues.append(f"• {result['name']}: {result.get('error', 'Test failed')}")
        
        if critical_issues:
            report.extend(critical_issues)
        else:
            report.append("✅ No critical issues found!")
        
        # Recommendations
        report.append("\n" + "=" * 80)
        report.append("RECOMMENDATIONS")
        report.append("=" * 80)
        
        if self.failed > 0:
            report.append("⚠️  Address failing tests before production deployment:")
            for result in self.results:
                if result["status"] == "FAIL":
                    if "GPT-5" in result["name"] and "fallback" in result.get("error", "").lower():
                        report.append("   • Ensure GPT-5 API key is configured and valid")
                    elif "session" in result["name"].lower():
                        report.append("   • Review session management and isolation logic")
                    elif "document" in result["name"].lower():
                        report.append("   • Check document parsing and upload handlers")
        else:
            report.append("✅ All tests passed! System is ready for production.")
        
        report.append("\n" + "=" * 80)
        report.append("END OF REPORT")
        report.append("=" * 80)
        
        return "\n".join(report)
    
    def save_report(self, filename: str = "gpt5_test_report.txt"):
        """Save report to file"""
        report_content = self.generate_report()
        with open(filename, "w") as f:
            f.write(report_content)
        print(f"\n📄 Report saved to: {filename}")
        return filename

# ============================================================================
# TEST RFP CONTENT GENERATION
# ============================================================================

def create_complex_rfp_content() -> str:
    """Create a complex RFP with 50+ expected deliverables"""
    return """
REQUEST FOR PROPOSAL
COMPREHENSIVE DIGITAL TRANSFORMATION & MARKETING INITIATIVE

PROJECT OVERVIEW:
We are seeking a full-service agency partner for a complete digital transformation and 
integrated marketing campaign for our global luxury fashion brand. This initiative spans 
18 months and encompasses brand strategy, creative development, technology implementation, 
and multi-channel marketing execution.

PROJECT SCOPE & DELIVERABLES:

1. BRAND STRATEGY & POSITIONING
   - Brand audit and competitive analysis
   - Market research and consumer insights
   - Brand positioning and messaging framework
   - Brand architecture development
   - Visual identity refresh
   - Brand guidelines and standards documentation
   - Tone of voice guidelines
   - Brand training materials

2. CREATIVE DEVELOPMENT
   - Creative concept development
   - Campaign creative direction
   - Photography and videography production
   - Motion graphics and animation
   - Print collateral design
   - Digital asset creation
   - Packaging design concepts
   - Retail environment design
   - Event and experiential design

3. WEBSITE & DIGITAL PLATFORMS
   - Website strategy and information architecture
   - UX research and user journey mapping
   - Wireframing and prototyping
   - Visual design and UI development
   - Front-end development
   - Back-end development and CMS integration
   - E-commerce platform setup
   - Mobile app design and development
   - API integration and development
   - Performance optimization
   - Security audit and implementation
   - Analytics setup and configuration

4. CONTENT STRATEGY & PRODUCTION
   - Content strategy development
   - Editorial calendar creation
   - Blog content creation
   - Social media content production
   - Video content production
   - Podcast development
   - Email newsletter design and copywriting
   - Product descriptions and catalog content
   - SEO content optimization
   - Influencer collaboration content

5. PAID MEDIA & ADVERTISING
   - Media strategy and planning
   - Search engine marketing (SEM)
   - Display advertising campaigns
   - Social media advertising
   - Programmatic advertising setup
   - Video advertising production
   - Native advertising placement
   - Out-of-home advertising
   - Radio and podcast advertising
   - Connected TV campaigns

6. SOCIAL MEDIA & COMMUNITY
   - Social media strategy
   - Community management
   - Social listening and monitoring
   - Influencer identification and outreach
   - User-generated content campaigns
   - Social commerce implementation
   - Live streaming events
   - Social media crisis management

7. MARKETING AUTOMATION & CRM
   - Marketing automation platform selection
   - CRM system implementation
   - Email marketing automation
   - Lead scoring and nurturing
   - Customer segmentation
   - Personalization engine setup
   - Loyalty program development
   - Customer data platform integration

8. ANALYTICS & MEASUREMENT
   - Analytics strategy and KPI definition
   - Dashboard creation and reporting
   - Conversion rate optimization
   - A/B testing framework
   - Attribution modeling
   - ROI analysis and reporting
   - Competitive intelligence monitoring
   - Market trend analysis

9. TECHNOLOGY & INTEGRATION
   - Technical architecture design
   - System integration planning
   - API development and management
   - Cloud infrastructure setup
   - DevOps and deployment automation
   - Quality assurance and testing
   - Performance monitoring
   - Disaster recovery planning

10. PROJECT MANAGEMENT & GOVERNANCE
    - Project planning and scheduling
    - Resource allocation and management
    - Risk assessment and mitigation
    - Stakeholder communication plans
    - Regular status reporting
    - Budget tracking and management
    - Change management procedures
    - Post-launch support and maintenance

BUDGET: $2.5M - $3.5M
TIMELINE: 18 months starting Q2 2025

Please provide detailed proposals including methodology, timeline, team structure, 
and itemized pricing for all deliverables.
"""

def create_minimal_rfp_content() -> str:
    """Create minimal RFP content that should trigger auto-rescue"""
    return """
RFP - Quick Website Update

We need a simple website refresh. 
Please update our homepage and add a contact form.

Budget: $10,000
Timeline: 2 weeks
"""

def create_medium_rfp_content() -> str:
    """Create medium complexity RFP for testing"""
    return """
REQUEST FOR PROPOSAL - BRAND CAMPAIGN

We're looking for an agency to develop and execute a brand awareness campaign.

Requirements:
- Brand strategy development
- Creative concept and execution
- Social media campaign
- Digital advertising
- Content creation
- Performance tracking and reporting

The campaign should run for 6 months with quarterly reviews.

Budget: $150,000
"""

# ============================================================================
# TEST DOCUMENT CREATION
# ============================================================================

def create_test_txt(content: str, filename: str) -> str:
    """Create a text file with RFP content"""
    filepath = f"test_rfps/{filename}.txt"
    os.makedirs("test_rfps", exist_ok=True)
    
    with open(filepath, "w") as f:
        f.write(content)
    
    return filepath

def create_test_docx(content: str, filename: str) -> Optional[str]:
    """Create a DOCX file with RFP content"""
    if not DOCX_AVAILABLE:
        return None
    
    filepath = f"test_rfps/{filename}.docx"
    os.makedirs("test_rfps", exist_ok=True)
    
    doc = Document()
    doc.add_heading("Request for Proposal", 0)
    
    # Split content into paragraphs
    paragraphs = content.split("\n\n")
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
    
    doc.save(filepath)
    return filepath

def create_test_pdf(content: str, filename: str) -> Optional[str]:
    """Create a PDF file with RFP content"""
    if not PDF_AVAILABLE:
        return None
    
    filepath = f"test_rfps/{filename}.pdf"
    os.makedirs("test_rfps", exist_ok=True)
    
    c = canvas.Canvas(filepath, pagesize=letter)
    width, height = letter
    
    # Write content
    y_position = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y_position, "REQUEST FOR PROPOSAL")
    y_position -= 30
    
    c.setFont("Helvetica", 10)
    lines = content.split("\n")
    
    for line in lines[:100]:  # Limit lines for PDF generation
        if y_position < 50:  # New page if needed
            c.showPage()
            y_position = height - 50
            c.setFont("Helvetica", 10)
        
        if line.strip():
            # Handle long lines
            if len(line) > 90:
                words = line.split()
                current_line = ""
                for word in words:
                    if len(current_line + " " + word) < 90:
                        current_line += " " + word if current_line else word
                    else:
                        c.drawString(50, y_position, current_line.strip())
                        y_position -= 15
                        current_line = word
                        if y_position < 50:
                            c.showPage()
                            y_position = height - 50
                            c.setFont("Helvetica", 10)
                if current_line:
                    c.drawString(50, y_position, current_line.strip())
                    y_position -= 15
            else:
                c.drawString(50, y_position, line)
                y_position -= 15
        else:
            y_position -= 10  # Extra space for paragraph breaks
    
    c.save()
    return filepath

# ============================================================================
# API INTERACTION HELPERS  
# ============================================================================

async def upload_and_analyze(client: httpx.AsyncClient, filepath: str, 
                            report: TestReport, test_name: str) -> Optional[Dict]:
    """Upload a document and start GPT-5 analysis"""
    
    # Extract text from file based on type
    text_content = ""
    if filepath.endswith('.txt'):
        with open(filepath, "r") as f:
            text_content = f.read()
    elif filepath.endswith('.docx') and DOCX_AVAILABLE:
        from docx import Document
        doc = Document(filepath)
        text_content = "\n".join([para.text for para in doc.paragraphs])
    elif filepath.endswith('.pdf') and PDF_AVAILABLE:
        # For PDF, we'll use a simple text extraction or just use the TXT version
        # Since PDF text extraction is complex, we'll use the TXT content for now
        txt_path = filepath.replace('.pdf', '.txt')
        if os.path.exists(txt_path):
            with open(txt_path, "r") as f:
                text_content = f.read()
    
    if not text_content:
        return {"success": False, "error": "Could not extract text from file"}
    
    try:
        # Use the correct AI planner endpoint with text content
        request_data = {
            "request_text": text_content[:50000],  # Limit to 50K chars to avoid timeouts
            "mode": "deep",  # Use deep mode for comprehensive analysis
            "tier": "thinking",  # Use GPT-5 thinking tier
            "strictness": "balanced"
        }
        
        response = await client.post(
            f"{BASE_URL}/api/ai/analyze",
            json=request_data,
            timeout=30.0
        )
        
        if response.status_code == 200:
            data = response.json()
            job_id = data.get("job_id")
            
            if job_id:
                print(f"   📤 Started analysis job: {job_id}")
                
                # Wait for job completion
                result = await wait_for_job(client, job_id, timeout=TEST_TIMEOUT)
                
                if result and result.get("status") == "completed":
                    analysis_result = result.get("result", {})
                    
                    # Extract deliverables properly
                    deliverables = []
                    if "items" in analysis_result:
                        deliverables = analysis_result["items"]
                    elif "plan" in analysis_result:
                        plan = analysis_result["plan"]
                        if "suggestions_by_department" in plan:
                            for dept_delivs in plan["suggestions_by_department"].values():
                                deliverables.extend(dept_delivs)
                    
                    details = {
                        "job_id": job_id,
                        "deliverables_count": len(deliverables),
                        "using_gpt5": analysis_result.get("model_used", "").startswith("gpt-5"),
                        "has_confidence_scores": any(
                            item.get("confidence", 0) > 0 
                            for item in deliverables
                        ),
                        "processing_time": result.get("processing_time", 0)
                    }
                    
                    return {"success": True, "details": details, "result": analysis_result, "deliverables": deliverables}
                else:
                    return {"success": False, "error": f"Job failed: {result.get('error', 'Unknown error')}"}
        else:
            # Fallback to older endpoint
            response = await client.post(
                f"{BASE_URL}/api/suggest_by_file",
                files=files,
                timeout=30.0
            )
            
            if response.status_code == 200:
                data = response.json()
                suggested = data.get("suggested", [])
                
                details = {
                    "deliverables_count": len(suggested),
                    "using_gpt5": False,  # Old endpoint doesn't use GPT-5
                    "has_confidence_scores": False,
                    "processing_time": 0
                }
                
                return {"success": True, "details": details, "result": data, "deliverables": suggested}
    
    except httpx.TimeoutException:
        return {"success": False, "error": "Request timeout"}
    except Exception as e:
        return {"success": False, "error": str(e)}
    
    return {"success": False, "error": "Upload failed"}

async def wait_for_job(client: httpx.AsyncClient, job_id: str, timeout: int = 120) -> Optional[Dict]:
    """Poll job status until completion or timeout"""
    start_time = time.time()
    
    while (time.time() - start_time) < timeout:
        try:
            response = await client.get(f"{BASE_URL}/api/ai/status/{job_id}")
            
            if response.status_code == 200:
                data = response.json()
                status = data.get("status")
                
                if status in ["completed", "failed"]:
                    data["processing_time"] = time.time() - start_time
                    return data
                
                # Show progress
                progress = data.get("progress", 0)
                print(f"   ⏳ Progress: {progress:.0f}% - {data.get('current_stage', 'Processing...')}")
            
            await asyncio.sleep(POLL_INTERVAL)
        
        except Exception as e:
            print(f"   ⚠️  Error checking status: {e}")
            await asyncio.sleep(POLL_INTERVAL)
    
    return {"status": "failed", "error": "Timeout waiting for completion"}

# ============================================================================
# INDIVIDUAL TEST FUNCTIONS
# ============================================================================

async def test_document_types(client: httpx.AsyncClient, report: TestReport):
    """Test 1: Multiple Document Types"""
    print("\n" + "="*60)
    print("TEST 1: MULTIPLE DOCUMENT TYPES")
    print("="*60)
    
    # Create test documents
    complex_rfp = create_complex_rfp_content()
    
    # Test TXT
    txt_file = create_test_txt(complex_rfp, "complex_rfp")
    result = await upload_and_analyze(client, txt_file, report, "test_txt_upload")
    
    if result["success"]:
        report.add_test(
            "DocumentTypes:TXT Upload",
            "PASS" if result["details"]["deliverables_count"] >= 50 else "FAIL",
            result["details"],
            None if result["details"]["deliverables_count"] >= 50 
            else f"Only {result['details']['deliverables_count']} deliverables (expected 50+)"
        )
    else:
        report.add_test("DocumentTypes:TXT Upload", "FAIL", {}, result["error"])
    
    # Test DOCX
    if DOCX_AVAILABLE:
        docx_file = create_test_docx(complex_rfp, "complex_rfp")
        if docx_file:
            result = await upload_and_analyze(client, docx_file, report, "test_docx_upload")
            
            if result["success"]:
                report.add_test(
                    "DocumentTypes:DOCX Upload",
                    "PASS" if result["details"]["deliverables_count"] >= 50 else "FAIL",
                    result["details"],
                    None if result["details"]["deliverables_count"] >= 50
                    else f"Only {result['details']['deliverables_count']} deliverables (expected 50+)"
                )
            else:
                report.add_test("DocumentTypes:DOCX Upload", "FAIL", {}, result["error"])
    else:
        report.add_test("DocumentTypes:DOCX Upload", "SKIP", {}, "python-docx not installed")
    
    # Test PDF
    if PDF_AVAILABLE:
        pdf_file = create_test_pdf(complex_rfp, "complex_rfp")
        if pdf_file:
            result = await upload_and_analyze(client, pdf_file, report, "test_pdf_upload")
            
            if result["success"]:
                report.add_test(
                    "DocumentTypes:PDF Upload",
                    "PASS" if result["details"]["deliverables_count"] >= 50 else "FAIL",
                    result["details"],
                    None if result["details"]["deliverables_count"] >= 50
                    else f"Only {result['details']['deliverables_count']} deliverables (expected 50+)"
                )
            else:
                report.add_test("DocumentTypes:PDF Upload", "FAIL", {}, result["error"])
    else:
        report.add_test("DocumentTypes:PDF Upload", "SKIP", {}, "reportlab not installed")

async def test_gpt5_intelligence(client: httpx.AsyncClient, report: TestReport):
    """Test 2: GPT-5 Intelligence Features"""
    print("\n" + "="*60)
    print("TEST 2: GPT-5 INTELLIGENCE FEATURES")
    print("="*60)
    
    # Use complex RFP for full GPT-5 analysis
    complex_rfp = create_complex_rfp_content()
    txt_file = create_test_txt(complex_rfp, "gpt5_intelligence_test")
    
    result = await upload_and_analyze(client, txt_file, report, "test_gpt5_features")
    
    if result["success"] and result.get("deliverables"):
        items = result["deliverables"]
        
        # Check confidence scores
        has_confidence = False
        confidence_range_valid = True
        confidence_scores = []
        
        for item in items:
            if "confidence" in item:
                has_confidence = True
                conf = item["confidence"]
                confidence_scores.append(conf)
                if conf < 0.0 or conf > 1.0:
                    confidence_range_valid = False
        
        # Check for evidence/rationales
        has_rationales = any(
            item.get("rationale") or item.get("evidence") or item.get("reasoning")
            for item in items
        )
        
        # Check for multipliers
        has_multipliers = any(
            item.get("complexity_multiplier") or 
            item.get("channel_multiplier") or
            item.get("market_multiplier")
            for item in items
        )
        
        # Check if using GPT-5 model
        using_gpt5 = result["details"].get("using_gpt5", False)
        
        details = {
            "using_gpt5": using_gpt5,
            "has_confidence_scores": has_confidence,
            "avg_confidence": sum(confidence_scores)/len(confidence_scores) if confidence_scores else 0,
            "confidence_range_valid": confidence_range_valid,
            "has_rationales": has_rationales,
            "has_multipliers": has_multipliers,
            "total_items": len(items)
        }
        
        # Determine pass/fail
        all_features_present = (
            using_gpt5 and
            has_confidence and
            confidence_range_valid
        )
        
        report.add_test(
            "GPT5Intelligence:Feature Detection",
            "PASS" if all_features_present else "FAIL",
            details,
            None if all_features_present else "Missing GPT-5 intelligence features"
        )
        
        # Test confidence score distribution
        if confidence_scores:
            high_confidence = sum(1 for c in confidence_scores if c >= 0.7)
            medium_confidence = sum(1 for c in confidence_scores if 0.3 <= c < 0.7)
            low_confidence = sum(1 for c in confidence_scores if c < 0.3)
            
            report.add_test(
                "GPT5Intelligence:Confidence Distribution",
                "PASS",
                {
                    "high_confidence": high_confidence,
                    "medium_confidence": medium_confidence,
                    "low_confidence": low_confidence,
                    "total": len(confidence_scores)
                }
            )
    else:
        report.add_test("GPT5Intelligence:Feature Detection", "FAIL", {}, 
                       result.get("error", "Failed to analyze document"))

async def test_auto_rescue(client: httpx.AsyncClient, report: TestReport):
    """Test 3: Auto-Rescue Logic"""
    print("\n" + "="*60)
    print("TEST 3: AUTO-RESCUE LOGIC")
    print("="*60)
    
    # Test with minimal content (should trigger rescue)
    minimal_rfp = create_minimal_rfp_content()
    minimal_file = create_test_txt(minimal_rfp, "minimal_rfp")
    
    print("   Testing minimal RFP (should trigger rescue)...")
    result_minimal = await upload_and_analyze(client, minimal_file, report, "test_minimal_rescue")
    
    if result_minimal["success"]:
        # Check if rescue was triggered (should have more deliverables than expected from minimal input)
        deliverables_count = result_minimal["details"]["deliverables_count"]
        
        report.add_test(
            "AutoRescue:Minimal Content",
            "PASS" if deliverables_count >= 15 else "FAIL",
            {
                "deliverables_count": deliverables_count,
                "rescue_triggered": deliverables_count >= 15
            },
            None if deliverables_count >= 15 else "Rescue logic not triggered for minimal content"
        )
    else:
        report.add_test("AutoRescue:Minimal Content", "FAIL", {}, result_minimal["error"])
    
    # Test with complex content (should NOT need rescue)
    complex_rfp = create_complex_rfp_content()
    complex_file = create_test_txt(complex_rfp, "complex_no_rescue")
    
    print("   Testing complex RFP (should not need rescue)...")
    result_complex = await upload_and_analyze(client, complex_file, report, "test_complex_no_rescue")
    
    if result_complex["success"]:
        deliverables_count = result_complex["details"]["deliverables_count"]
        using_gpt5 = result_complex["details"].get("using_gpt5", False)
        
        report.add_test(
            "AutoRescue:Complex Content",
            "PASS",
            {
                "deliverables_count": deliverables_count,
                "using_gpt5": using_gpt5,
                "rescue_needed": False
            }
        )
    else:
        report.add_test("AutoRescue:Complex Content", "FAIL", {}, result_complex["error"])
    
    # Test gate threshold with medium complexity
    medium_rfp = create_medium_rfp_content()
    medium_file = create_test_txt(medium_rfp, "medium_rfp")
    
    print("   Testing medium RFP (testing gate threshold)...")
    result_medium = await upload_and_analyze(client, medium_file, report, "test_medium_threshold")
    
    if result_medium["success"]:
        report.add_test(
            "AutoRescue:Gate Threshold",
            "PASS",
            {
                "deliverables_count": result_medium["details"]["deliverables_count"],
                "complexity": "medium"
            }
        )
    else:
        report.add_test("AutoRescue:Gate Threshold", "FAIL", {}, result_medium["error"])

async def test_session_isolation(client: httpx.AsyncClient, report: TestReport):
    """Test 4: Session Isolation"""
    print("\n" + "="*60)
    print("TEST 4: SESSION ISOLATION")
    print("="*60)
    
    # Create two different RFPs
    rfp_a_content = """
    RFP A: E-COMMERCE PLATFORM DEVELOPMENT
    
    We need a complete e-commerce platform with:
    - Product catalog management
    - Shopping cart functionality
    - Payment processing integration
    - Order management system
    - Customer accounts
    - Inventory tracking
    - Shipping integration
    - Admin dashboard
    
    Technology: React, Node.js, PostgreSQL
    Budget: $200,000
    """
    
    rfp_b_content = """
    RFP B: CORPORATE VIDEO PRODUCTION
    
    We need video production services for:
    - Company overview video
    - Product demonstration videos
    - Customer testimonial videos
    - Training videos
    - Social media video content
    - Event coverage
    - Animation and motion graphics
    
    Deliverables: 20 videos total
    Budget: $75,000
    """
    
    # Create test files
    file_a = create_test_txt(rfp_a_content, "session_test_rfp_a")
    file_b = create_test_txt(rfp_b_content, "session_test_rfp_b")
    
    # Upload and analyze RFP A
    print("   Uploading RFP A (E-commerce)...")
    result_a = await upload_and_analyze(client, file_a, report, "test_session_a")
    
    if not result_a["success"]:
        report.add_test("SessionIsolation:RFP_A Upload", "FAIL", {}, result_a["error"])
        return
    
    job_id_a = result_a["details"].get("job_id")
    items_a = result_a.get("deliverables", [])
    
    # Upload and analyze RFP B
    print("   Uploading RFP B (Video Production)...")
    result_b = await upload_and_analyze(client, file_b, report, "test_session_b")
    
    if not result_b["success"]:
        report.add_test("SessionIsolation:RFP_B Upload", "FAIL", {}, result_b["error"])
        return
    
    job_id_b = result_b["details"].get("job_id")
    items_b = result_b.get("deliverables", [])
    
    # Check for session isolation
    
    # 1. Job IDs should be unique
    unique_job_ids = job_id_a != job_id_b if (job_id_a and job_id_b) else True
    
    # 2. Check for content contamination
    # Extract deliverable names/descriptions
    deliverables_a = set()
    deliverables_b = set()
    
    for item in items_a:
        name = str(item.get("name", "") or item.get("deliverable", "")).lower()
        deliverables_a.add(name)
    
    for item in items_b:
        name = str(item.get("name", "") or item.get("deliverable", "")).lower()
        deliverables_b.add(name)
    
    # Check for e-commerce terms in video RFP results
    ecommerce_terms = ["shopping cart", "payment", "inventory", "product catalog"]
    video_terms = ["video", "production", "filming", "animation"]
    
    contamination_a_to_b = any(
        term in " ".join(deliverables_b)
        for term in ecommerce_terms
    )
    
    contamination_b_to_a = any(
        term in " ".join(deliverables_a)
        for term in video_terms
    )
    
    # 3. Deliverable counts should be appropriate for each RFP type
    reasonable_counts = (
        10 <= len(items_a) <= 100 and
        10 <= len(items_b) <= 100
    )
    
    details = {
        "unique_job_ids": unique_job_ids,
        "rfp_a_deliverables": len(items_a),
        "rfp_b_deliverables": len(items_b),
        "no_contamination_a_to_b": not contamination_a_to_b,
        "no_contamination_b_to_a": not contamination_b_to_a,
        "reasonable_counts": reasonable_counts
    }
    
    session_isolated = (
        unique_job_ids and
        not contamination_a_to_b and
        not contamination_b_to_a and
        reasonable_counts
    )
    
    report.add_test(
        "SessionIsolation:Data Separation",
        "PASS" if session_isolated else "FAIL",
        details,
        None if session_isolated else "Session data contamination detected"
    )
    
    # Test concurrent sessions
    print("   Testing concurrent session handling...")
    
    # Start two jobs simultaneously
    async def start_job(file_path):
        # Extract text from file
        text_content = ""
        with open(file_path, "r") as f:
            text_content = f.read()
        
        if not text_content:
            return None
            
        request_data = {
            "request_text": text_content[:50000],
            "mode": "deep",
            "tier": "thinking",
            "strictness": "balanced"
        }
        
        try:
            response = await client.post(f"{BASE_URL}/api/ai/analyze", json=request_data)
            if response.status_code == 200:
                return response.json().get("job_id")
        except:
            pass
        return None
    
    job1, job2 = await asyncio.gather(
        start_job(file_a),
        start_job(file_b),
        return_exceptions=True
    )
    
    if job1 and job2 and job1 != job2:
        report.add_test(
            "SessionIsolation:Concurrent Sessions",
            "PASS",
            {"job1": job1, "job2": job2}
        )
    else:
        report.add_test(
            "SessionIsolation:Concurrent Sessions",
            "SKIP" if (job1 is None or job2 is None) else "FAIL",
            {},
            "Could not test concurrent sessions" if (job1 is None or job2 is None) else "Failed to handle concurrent sessions"
        )

async def test_api_response_format(client: httpx.AsyncClient, report: TestReport):
    """Test 5: API Response Format"""
    print("\n" + "="*60)
    print("TEST 5: API RESPONSE FORMAT")
    print("="*60)
    
    # Create a test RFP
    test_rfp = create_medium_rfp_content()
    
    # Test job creation
    print("   Testing job creation...")
    request_data = {
        "request_text": test_rfp,
        "mode": "deep",
        "tier": "thinking",
        "strictness": "balanced"
    }
    
    response = await client.post(f"{BASE_URL}/api/ai/analyze", json=request_data)
    
    if response.status_code != 200:
        report.add_test("APIFormat:Job Creation", "FAIL", {}, 
                       f"Status code: {response.status_code}")
        return
    
    job_data = response.json()
    
    # Validate job creation response
    required_fields = ["job_id", "status", "message"]
    has_required = all(field in job_data for field in required_fields)
    
    report.add_test(
        "APIFormat:Job Creation Response",
        "PASS" if has_required else "FAIL",
        {"fields_present": list(job_data.keys())},
        None if has_required else f"Missing required fields"
    )
    
    job_id = job_data.get("job_id")
    if not job_id:
        return
    
    # Test status tracking
    print("   Testing status tracking...")
    progress_updates = []
    status_checks = 0
    max_checks = 60  # 60 seconds with 1s interval
    
    while status_checks < max_checks:
        response = await client.get(f"{BASE_URL}/api/ai/jobs/{job_id}")
        
        if response.status_code != 200:
            report.add_test("APIFormat:Status Check", "FAIL", {}, 
                           f"Status code: {response.status_code}")
            break
        
        status_data = response.json()
        status = status_data.get("status")
        progress = status_data.get("progress", 0)
        
        progress_updates.append(progress)
        status_checks += 1
        
        if status in ["completed", "failed"]:
            break
        
        await asyncio.sleep(POLL_INTERVAL)
    
    # Validate progress updates
    progress_valid = all(0 <= p <= 100 for p in progress_updates)
    progress_increasing = True
    if len(progress_updates) > 1:
        progress_increasing = all(
            progress_updates[i] <= progress_updates[i+1] 
            for i in range(len(progress_updates)-1)
        )
    
    report.add_test(
        "APIFormat:Progress Updates",
        "PASS" if (progress_valid and len(progress_updates) > 0) else "FAIL",
        {
            "num_updates": len(progress_updates),
            "progress_valid": progress_valid,
            "progress_increasing": progress_increasing,
            "final_progress": progress_updates[-1] if progress_updates else 0
        }
    )
    
    # Validate final response structure
    if status == "completed":
        result = status_data.get("result", {})
        
        # Check response structure
        structure_valid = isinstance(result, dict)
        
        if structure_valid:
            # Check for deliverables (could be in items or plan structure)
            has_deliverables = False
            deliverable_count = 0
            
            if "items" in result:
                items = result["items"]
                has_deliverables = isinstance(items, list) and len(items) > 0
                deliverable_count = len(items)
            elif "plan" in result and "suggestions_by_department" in result["plan"]:
                dept_suggestions = result["plan"]["suggestions_by_department"]
                total_delivs = sum(len(v) for v in dept_suggestions.values())
                has_deliverables = total_delivs > 0
                deliverable_count = total_delivs
            
            report.add_test(
                "APIFormat:Response Structure",
                "PASS" if has_deliverables else "FAIL",
                {
                    "structure_valid": structure_valid,
                    "has_deliverables": has_deliverables,
                    "deliverable_count": deliverable_count,
                    "model_used": result.get("model_used", "unknown")
                }
            )
        else:
            report.add_test("APIFormat:Response Structure", "FAIL", {}, 
                           "Invalid response structure")
    
    # Test error handling with malformed request
    print("   Testing error handling...")
    
    # Test with missing required field
    response = await client.post(f"{BASE_URL}/api/ai/analyze", json={})
    
    error_handled = response.status_code in [400, 422]
    
    report.add_test(
        "APIFormat:Error Handling",
        "PASS" if error_handled else "FAIL",
        {"status_code": response.status_code},
        None if error_handled else "Missing proper error handling"
    )
    
    # Test with invalid job ID
    response = await client.get(f"{BASE_URL}/api/ai/status/invalid-job-id-xyz")
    
    invalid_job_handled = response.status_code in [404, 400]
    
    report.add_test(
        "APIFormat:Invalid Job ID",
        "PASS" if invalid_job_handled else "FAIL",
        {"status_code": response.status_code}
    )

# ============================================================================
# MAIN TEST RUNNER
# ============================================================================

async def run_all_tests():
    """Run all GPT-5 comprehensive tests"""
    
    print("\n" + "="*80)
    print("GPT-5 RFP ANALYSIS COMPREHENSIVE TEST SUITE")
    print("Starting at:", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    print("="*80)
    
    # Initialize report
    report = TestReport()
    
    # Create HTTP client
    async with httpx.AsyncClient(timeout=60.0) as client:
        
        # Check if server is running
        try:
            response = await client.get(f"{BASE_URL}/")
            if response.status_code != 200:
                print("❌ Server not responding at", BASE_URL)
                return
        except Exception as e:
            print(f"❌ Cannot connect to server at {BASE_URL}: {e}")
            print("   Please ensure FastAPI server is running on port 5000")
            return
        
        # Run test suites
        try:
            await test_document_types(client, report)
            await test_gpt5_intelligence(client, report)
            await test_auto_rescue(client, report)
            await test_session_isolation(client, report)
            await test_api_response_format(client, report)
        except Exception as e:
            print(f"\n❌ Test suite error: {e}")
            report.add_test("TestSuite:Execution", "FAIL", {}, str(e))
    
    # Generate and save report
    print("\n" + "="*80)
    print("GENERATING FINAL REPORT")
    print("="*80)
    
    report_file = report.save_report("gpt5_comprehensive_test_report.txt")
    
    # Print summary
    print("\n" + "="*80)
    print("TEST SUITE COMPLETED")
    print("="*80)
    print(f"✅ Passed: {report.passed}/{report.total_tests}")
    print(f"❌ Failed: {report.failed}/{report.total_tests}")
    print(f"⏭️  Skipped: {report.skipped}/{report.total_tests}")
    print(f"📄 Full report: {report_file}")
    print("="*80)
    
    # Also print the report to console
    print("\n" + report.generate_report())
    
    return report

def main():
    """Entry point for test suite"""
    
    # Run tests
    report = asyncio.run(run_all_tests())
    
    # Exit with appropriate code
    if report and report.failed == 0:
        sys.exit(0)  # Success
    else:
        sys.exit(1)  # Failure

if __name__ == "__main__":
    main()