#!/usr/bin/env python3
"""Generate test documents in PDF and DOCX formats from text files"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY

def txt_to_docx(txt_file, docx_file):
    """Convert text file to DOCX format"""
    doc = Document()
    
    # Add title style
    with open(txt_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Process lines to create formatted document
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        # First line is title
        if i == 0:
            heading = doc.add_heading(line, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Section headers (all caps lines)
        elif line.isupper() and len(line) > 5:
            doc.add_heading(line, 1)
        # Subsection headers (lines ending with colon)
        elif line.endswith(':') and len(line) < 50:
            doc.add_heading(line[:-1], 2)
        # Regular content
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Save document
    doc.save(docx_file)
    print(f"Created DOCX: {docx_file}")

def txt_to_pdf(txt_file, pdf_file):
    """Convert text file to PDF format"""
    # Create PDF document
    doc = SimpleDocTemplate(pdf_file, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)
    
    # Container for page elements
    story = []
    
    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='black',
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor='black',
        spaceAfter=12,
        spaceBefore=12
    )
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        spaceAfter=12
    )
    
    # Read and process text file
    with open(txt_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Process lines and create PDF elements
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.2*inch))
            continue
        
        # Escape special characters for ReportLab
        line = line.replace('&', '&amp;')
        line = line.replace('<', '&lt;')
        line = line.replace('>', '&gt;')
        
        # First line is title
        if i == 0:
            story.append(Paragraph(line, title_style))
            story.append(Spacer(1, 0.3*inch))
        # Section headers (all caps lines)
        elif line.isupper() and len(line) > 5:
            story.append(Spacer(1, 0.2*inch))
            story.append(Paragraph(line, heading_style))
        # Subsection headers (lines ending with colon or starting with number)
        elif (line.endswith(':') and len(line) < 50) or (len(line) > 0 and line[0].isdigit() and '.' in line[:3]):
            story.append(Paragraph(line, heading_style))
        # Regular content
        else:
            story.append(Paragraph(line, body_style))
    
    # Build PDF
    doc.build(story)
    print(f"Created PDF: {pdf_file}")

def create_empty_pdf(pdf_file):
    """Create an empty PDF for testing"""
    doc = SimpleDocTemplate(pdf_file, pagesize=letter)
    story = []
    doc.build(story)
    print(f"Created empty PDF: {pdf_file}")

def create_corrupted_file(file_path):
    """Create a corrupted file for testing"""
    with open(file_path, 'wb') as f:
        # Write random bytes that don't form a valid file
        f.write(b'%PDF-1.4\n%\xE2\xE3\xCF\xD3\n')
        f.write(b'This is not a valid PDF file structure')
        f.write(b'\x00\x01\x02\x03' * 100)
    print(f"Created corrupted file: {file_path}")

def main():
    """Generate all test documents"""
    # Ensure test_docs directory exists
    os.makedirs('test_docs', exist_ok=True)
    
    # Convert existing TXT files to DOCX and PDF
    txt_files = [
        ('test_docs/luxury_fashion_rfp_small.txt', 'test_docs/luxury_fashion_rfp_small'),
        ('test_docs/luxury_fashion_rfp_large.txt', 'test_docs/luxury_fashion_rfp_large')
    ]
    
    for txt_file, base_name in txt_files:
        if os.path.exists(txt_file):
            # Convert to DOCX
            docx_file = f"{base_name}.docx"
            txt_to_docx(txt_file, docx_file)
            
            # Convert to PDF
            pdf_file = f"{base_name}.pdf"
            txt_to_pdf(txt_file, pdf_file)
    
    # Create empty documents for testing
    create_empty_pdf('test_docs/empty.pdf')
    
    # Create empty DOCX
    doc = Document()
    doc.save('test_docs/empty.docx')
    print("Created empty DOCX: test_docs/empty.docx")
    
    # Create empty TXT
    with open('test_docs/empty.txt', 'w') as f:
        f.write('')
    print("Created empty TXT: test_docs/empty.txt")
    
    # Create corrupted files for testing
    create_corrupted_file('test_docs/corrupted.pdf')
    create_corrupted_file('test_docs/corrupted.docx')
    
    print("\nAll test documents generated successfully!")
    print("\nGenerated files:")
    for file in os.listdir('test_docs'):
        file_path = os.path.join('test_docs', file)
        size = os.path.getsize(file_path)
        print(f"  - {file} ({size:,} bytes)")

if __name__ == '__main__':
    main()